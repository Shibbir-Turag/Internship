from __future__ import annotations

"""
SCRIPT 05 — THIRD TRIAL RGB GROWTH-RATE AND TREATMENT COMPARISON

Purpose
-------
Use Script 04 visible-emergence and RGB green-cover outputs to calculate
Trial 3 tray-level and treatment-level RGB growth metrics.

This script DOES:
- read Script 04 tray and cell results
- preserve observed Day 7 RGB green-cover values
- create adjusted/imputed Day 7 RGB green-cover values for likely bug-eaten cells
- calculate visible-emergence rate and RGB green-cover growth rate using real elapsed days
- compare Microbes vs No Microbes
- compare Ideal vs Heat vs Moisture
- compare Microbes x Treatment groups
- compare fixed Inside/Outside placement for Ideal and Moisture trays
- exclude Heat trays from fixed Inside/Outside comparison
- generate CSV outputs
- generate Excel workbook
- generate charts
- generate PDF summary
- generate expanded Word report

Important fix in this version
-----------------------------
The Inside/Outside comparison now uses tray-number fallback metadata directly:

Ideal Inside     = Tray 1, Tray 9
Ideal Outside    = Tray 2, Tray 7
Moisture Inside  = Tray 5, Tray 12
Moisture Outside = Tray 3, Tray 10

This avoids blank charts caused by inconsistent or missing label_environment values
from earlier CSV files.
"""

import argparse
import json
import math
import re
from pathlib import Path

import matplotlib
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(r"C:\Users\tshib\OneDrive\Desktop\Internship")

SCRIPT04_REPORTS = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "04_Visible_Emergence"
    / "_reports"
)

TRAY_SUMMARY_CSV = SCRIPT04_REPORTS / "visible_emergence_tray_summary.csv"
CELL_RESULTS_CSV = SCRIPT04_REPORTS / "visible_emergence_cell_results.csv"
FIRST_EMERGENCE_CSV = SCRIPT04_REPORTS / "first_emergence_summary.csv"

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "05_RGB_Growth_Rate_Treatment_Comparison"
)

CHARTS_ROOT = OUTPUT_ROOT / "charts"
REPORTS_ROOT = OUTPUT_ROOT / "_reports"
CONFIG_ROOT = OUTPUT_ROOT / "_config"


# ============================================================
# 2) SETTINGS
# ============================================================

EXPECTED_CELLS = 70
EXPECTED_OBSERVATION_DAYS = [1, 2, 3, 4, 5, 6, 7]

DATE_MAP = {
    1: "2026-06-29",
    2: "2026-06-30",
    3: "2026-07-01",
    4: "2026-07-02",
    5: "2026-07-03",
    6: "2026-07-04",
    7: "2026-07-07",
}

ELAPSED_DAY_MAP = {
    1: 0,
    2: 1,
    3: 2,
    4: 3,
    5: 4,
    6: 5,
    7: 8,
}

PREVIOUS_PHOTO_INTERVAL_MAP = {
    1: "",
    2: 1,
    3: 1,
    4: 1,
    5: 1,
    6: 1,
    7: 3,
}

MICROBE_ORDER = ["No Microbes", "Microbes"]
TREATMENT_ORDER = ["Ideal", "Heat", "Moisture"]
FIXED_ENVIRONMENT_TREATMENTS = ["Ideal", "Moisture"]

# Correct Trial 3 tray metadata.
TRAY_FALLBACK = {
    1:  {"microbe_status": "No Microbes", "treatment": "Ideal",    "fixed_environment": "Inside"},
    2:  {"microbe_status": "No Microbes", "treatment": "Ideal",    "fixed_environment": "Outside"},
    3:  {"microbe_status": "No Microbes", "treatment": "Moisture", "fixed_environment": "Outside"},
    4:  {"microbe_status": "No Microbes", "treatment": "Heat",     "fixed_environment": "Dynamic Heat"},
    5:  {"microbe_status": "Microbes",    "treatment": "Moisture", "fixed_environment": "Inside"},
    6:  {"microbe_status": "Microbes",    "treatment": "Heat",     "fixed_environment": "Dynamic Heat"},
    7:  {"microbe_status": "Microbes",    "treatment": "Ideal",    "fixed_environment": "Outside"},
    8:  {"microbe_status": "Microbes",    "treatment": "Heat",     "fixed_environment": "Dynamic Heat"},
    9:  {"microbe_status": "Microbes",    "treatment": "Ideal",    "fixed_environment": "Inside"},
    10: {"microbe_status": "Microbes",    "treatment": "Moisture", "fixed_environment": "Outside"},
    11: {"microbe_status": "No Microbes", "treatment": "Heat",     "fixed_environment": "Dynamic Heat"},
    12: {"microbe_status": "No Microbes", "treatment": "Moisture", "fixed_environment": "Inside"},
}

PERFORMANCE_COMPONENTS_OBSERVED = [
    "final_tracked_emergence_percent",
    "final_observed_green_cover_percent",
    "observed_green_cover_rate_day1_to_final_pp_per_day",
    "emergence_rate_day1_to_final_pp_per_day",
]

PERFORMANCE_COMPONENTS_ADJUSTED = [
    "final_tracked_emergence_percent",
    "final_adjusted_green_cover_percent",
    "adjusted_green_cover_rate_day1_to_final_pp_per_day",
    "emergence_rate_day1_to_final_pp_per_day",
]


# ============================================================
# 3) OPTIONAL REPORT IMPORTS
# ============================================================

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image as PDFImage,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 4) GENERAL HELPERS
# ============================================================

def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r"(\d+)", str(value))
    ]


def normalise_text(value: object) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value).strip().casefold())


def parse_yes(value: object) -> bool:
    return str(value).strip().casefold() in {"yes", "y", "true", "1", "p"}


def yes_no(value: bool) -> str:
    return "Yes" if value else "No"


def require_columns(dataframe: pd.DataFrame, required_columns: list[str], source_name: str) -> None:
    missing = [column for column in required_columns if column not in dataframe.columns]

    if missing:
        raise ValueError(
            f"{source_name} is missing required column(s): "
            + ", ".join(missing)
        )


def safe_numeric(dataframe: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    dataframe = dataframe.copy()

    for column in columns:
        if column in dataframe.columns:
            dataframe[column] = pd.to_numeric(dataframe[column], errors="coerce")

    return dataframe


def safe_round_dataframe(dataframe: pd.DataFrame, decimals: int = 4) -> pd.DataFrame:
    output = dataframe.copy()
    numeric_columns = output.select_dtypes(include=["number"]).columns
    output[numeric_columns] = output[numeric_columns].round(decimals)
    return output


def minmax_score(series: pd.Series) -> pd.Series:
    values = pd.to_numeric(series, errors="coerce")
    valid = values.dropna()

    if valid.empty:
        return pd.Series([math.nan] * len(values), index=values.index)

    minimum = valid.min()
    maximum = valid.max()

    if maximum == minimum:
        return pd.Series(
            [50.0 if pd.notna(value) else math.nan for value in values],
            index=values.index,
        )

    return (values - minimum) / (maximum - minimum) * 100.0


def inverse_minmax_score(series: pd.Series) -> pd.Series:
    values = pd.to_numeric(series, errors="coerce")
    valid = values.dropna()

    if valid.empty:
        return pd.Series([math.nan] * len(values), index=values.index)

    minimum = valid.min()
    maximum = valid.max()

    if maximum == minimum:
        return pd.Series(
            [50.0 if pd.notna(value) else math.nan for value in values],
            index=values.index,
        )

    return (maximum - values) / (maximum - minimum) * 100.0


def mean_or_nan(dataframe: pd.DataFrame, column: str):
    if dataframe.empty or column not in dataframe.columns:
        return math.nan

    return pd.to_numeric(dataframe[column], errors="coerce").mean()


def count_unique_or_zero(dataframe: pd.DataFrame, column: str) -> int:
    if dataframe.empty or column not in dataframe.columns:
        return 0

    return int(dataframe[column].nunique())


def format_number(value: object, decimals: int = 2) -> str:
    try:
        if pd.isna(value):
            return "N/A"
        return f"{float(value):.{decimals}f}"
    except Exception:
        return "N/A"


def standardise_treatment(value: object, tray_no: int | None = None) -> str:
    text = normalise_text(value)

    if text in {"ideal", "control"}:
        return "Ideal"

    if text in {"heat", "heattreatment"}:
        return "Heat"

    if text in {"moist", "moisture", "moisturetreatment"}:
        return "Moisture"

    if tray_no in TRAY_FALLBACK:
        return TRAY_FALLBACK[tray_no]["treatment"]

    return str(value).strip()


def standardise_microbe(value: object, tray_no: int | None = None) -> str:
    text = normalise_text(value)

    if text in {"m", "microbe", "microbes", "withmicrobes"}:
        return "Microbes"

    if text in {
        "nm",
        "n/m",
        "nomicrobe",
        "nomicrobes",
        "withoutmicrobes",
        "no",
    }:
        return "No Microbes"

    if tray_no in TRAY_FALLBACK:
        return TRAY_FALLBACK[tray_no]["microbe_status"]

    return str(value).strip()


def standardise_environment(
    value: object,
    tray_no: int | None = None,
    treatment: str | None = None,
) -> str:
    text = normalise_text(value)

    if text in {"in", "inside", "insidegreenhouse", "greenhouse"}:
        return "Inside"

    if text in {"out", "outside", "outsideopen", "open", "outsideinopen"}:
        return "Outside"

    if text in {"dynamic", "dynamicheat", "heat"}:
        return "Dynamic Heat"

    if treatment and str(treatment).casefold() == "heat":
        return "Dynamic Heat"

    if tray_no in TRAY_FALLBACK:
        return TRAY_FALLBACK[tray_no]["fixed_environment"]

    return str(value).strip()


def tray_no_as_int(value) -> int | None:
    try:
        if pd.isna(value):
            return None
        return int(value)
    except Exception:
        return None


def add_standardised_metadata(frame: pd.DataFrame) -> pd.DataFrame:
    frame = frame.copy()

    if "treatment" not in frame.columns and "treatment_type" in frame.columns:
        frame["treatment"] = frame["treatment_type"]

    if "label_environment" not in frame.columns and "base_environment" in frame.columns:
        frame["label_environment"] = frame["base_environment"]

    if "label_environment" not in frame.columns and "environment" in frame.columns:
        frame["label_environment"] = frame["environment"]

    if "microbe_status" not in frame.columns and "microbes_status" in frame.columns:
        frame["microbe_status"] = frame["microbes_status"]

    if "microbe_status" not in frame.columns:
        frame["microbe_status"] = ""

    if "treatment" not in frame.columns:
        frame["treatment"] = ""

    if "label_environment" not in frame.columns:
        frame["label_environment"] = ""

    if "observed_environment" not in frame.columns:
        frame["observed_environment"] = frame["label_environment"]

    frame["tray_no"] = pd.to_numeric(frame["tray_no"], errors="coerce")

    frame["treatment"] = frame.apply(
        lambda row: standardise_treatment(
            row.get("treatment", ""),
            tray_no_as_int(row.get("tray_no")),
        ),
        axis=1,
    )

    frame["microbe_status"] = frame.apply(
        lambda row: standardise_microbe(
            row.get("microbe_status", ""),
            tray_no_as_int(row.get("tray_no")),
        ),
        axis=1,
    )

    # Force fixed_environment from tray number where possible.
    frame["fixed_environment"] = frame.apply(
        lambda row: (
            TRAY_FALLBACK[tray_no_as_int(row.get("tray_no"))]["fixed_environment"]
            if tray_no_as_int(row.get("tray_no")) in TRAY_FALLBACK
            else standardise_environment(
                row.get("label_environment", ""),
                tray_no_as_int(row.get("tray_no")),
                row.get("treatment", ""),
            )
        ),
        axis=1,
    )

    frame["label_environment"] = frame["fixed_environment"]

    frame["observed_environment"] = frame.apply(
        lambda row: (
            standardise_environment(
                row.get("observed_environment", ""),
                tray_no_as_int(row.get("tray_no")),
                row.get("treatment", ""),
            )
            if row.get("treatment", "") == "Heat"
            else row.get("fixed_environment", "")
        ),
        axis=1,
    )

    frame["environment_group"] = np.where(
        frame["treatment"].astype(str).str.casefold().eq("heat"),
        "Dynamic Heat",
        frame["fixed_environment"],
    )

    frame["microbe_treatment"] = frame["microbe_status"] + " | " + frame["treatment"]
    frame["treatment_environment"] = frame["treatment"] + " | " + frame["environment_group"]
    frame["microbe_environment"] = frame["microbe_status"] + " | " + frame["environment_group"]

    fill_string_columns = [
        "microbe_status",
        "treatment",
        "label_environment",
        "fixed_environment",
        "observed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
    ]

    for column in fill_string_columns:
        if column in frame.columns:
            frame[column] = frame[column].fillna("").astype(str)

    return frame


# ============================================================
# 5) LOAD SCRIPT 04 OUTPUTS
# ============================================================

def load_script04_outputs() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if not TRAY_SUMMARY_CSV.exists():
        raise FileNotFoundError(f"Missing Script 04 tray summary:\n{TRAY_SUMMARY_CSV}")

    if not CELL_RESULTS_CSV.exists():
        raise FileNotFoundError(f"Missing Script 04 cell results:\n{CELL_RESULTS_CSV}")

    tray = pd.read_csv(TRAY_SUMMARY_CSV)
    cell = pd.read_csv(CELL_RESULTS_CSV)

    first = pd.read_csv(FIRST_EMERGENCE_CSV) if FIRST_EMERGENCE_CSV.exists() else pd.DataFrame()

    if "treatment" not in tray.columns and "treatment_type" in tray.columns:
        tray["treatment"] = tray["treatment_type"]

    if "treatment" not in cell.columns and "treatment_type" in cell.columns:
        cell["treatment"] = cell["treatment_type"]

    if "label_environment" not in tray.columns and "base_environment" in tray.columns:
        tray["label_environment"] = tray["base_environment"]

    if "label_environment" not in cell.columns and "base_environment" in cell.columns:
        cell["label_environment"] = cell["base_environment"]

    if "mean_rgb_green_cover_percent" not in tray.columns and "mean_green_area_percent" in tray.columns:
        tray["mean_rgb_green_cover_percent"] = tray["mean_green_area_percent"]

    if "rgb_green_cover_percent" not in cell.columns and "green_area_percent" in cell.columns:
        cell["rgb_green_cover_percent"] = cell["green_area_percent"]

    require_columns(
        tray,
        [
            "day_order",
            "day",
            "calendar_date",
            "days_since_day1",
            "tray",
            "tray_no",
            "microbe_status",
            "treatment",
            "label_environment",
            "tracked_emergence_percent",
            "mean_rgb_green_cover_percent",
            "status",
            "capture_id",
        ],
        "visible_emergence_tray_summary.csv",
    )

    require_columns(
        cell,
        [
            "day_order",
            "day",
            "calendar_date",
            "days_since_day1",
            "tray",
            "tray_no",
            "cell_id",
            "raw_current_green_evidence",
            "tracked_visible_emerged",
            "green_area_percent",
            "microbe_status",
            "treatment",
            "label_environment",
            "capture_id",
        ],
        "visible_emergence_cell_results.csv",
    )

    tray = safe_numeric(
        tray,
        [
            "day_order",
            "days_since_planting",
            "days_since_day1",
            "days_since_previous_photo",
            "tray_no",
            "raw_green_cells",
            "raw_green_percent",
            "tracked_emerged_cells",
            "tracked_emergence_percent",
            "newly_emerged_today",
            "newly_emerged_percent",
            "carried_forward_cells",
            "mean_green_area_percent",
            "mean_rgb_green_cover_percent",
            "possible_day7_bug_eaten_cells",
        ],
    )

    cell = safe_numeric(
        cell,
        [
            "day_order",
            "days_since_planting",
            "days_since_day1",
            "days_since_previous_photo",
            "tray_no",
            "cell_id",
            "row",
            "column",
            "green_pixels",
            "largest_green_component",
            "green_area_percent",
            "rgb_green_cover_percent",
            "zone_pixels",
        ],
    )

    tray["status"] = tray["status"].astype(str).str.upper()

    pass_tray = tray.loc[tray["status"].eq("PASS")].copy()

    if not pass_tray.empty and "capture_id" in cell.columns:
        pass_keys = pass_tray[
            ["day_order", "tray_no", "capture_id"]
        ].drop_duplicates()

        cell = cell.merge(
            pass_keys,
            on=["day_order", "tray_no", "capture_id"],
            how="inner",
        )

    for frame in [pass_tray, cell]:
        frame["day_order"] = pd.to_numeric(frame["day_order"], errors="coerce").astype(int)

        frame["calendar_date"] = frame.apply(
            lambda row: (
                DATE_MAP.get(int(row["day_order"]), "")
                if pd.isna(row.get("calendar_date"))
                or str(row.get("calendar_date")).strip() == ""
                else str(row.get("calendar_date")).strip()
            ),
            axis=1,
        )

        frame["days_since_day1"] = pd.to_numeric(frame["days_since_day1"], errors="coerce")

        frame["days_since_day1"] = frame.apply(
            lambda row: (
                ELAPSED_DAY_MAP.get(int(row["day_order"]), math.nan)
                if pd.isna(row["days_since_day1"])
                else row["days_since_day1"]
            ),
            axis=1,
        )

        if "days_since_previous_photo" not in frame.columns:
            frame["days_since_previous_photo"] = math.nan

        frame["days_since_previous_photo"] = pd.to_numeric(
            frame["days_since_previous_photo"],
            errors="coerce",
        )

        frame["days_since_previous_photo"] = frame.apply(
            lambda row: (
                PREVIOUS_PHOTO_INTERVAL_MAP.get(int(row["day_order"]), math.nan)
                if pd.isna(row["days_since_previous_photo"])
                else row["days_since_previous_photo"]
            ),
            axis=1,
        )

    pass_tray = add_standardised_metadata(pass_tray)
    cell = add_standardised_metadata(cell)

    return pass_tray, cell, first


# ============================================================
# 6) DAY 7 ADJUSTMENT / IMPUTATION
# ============================================================

def estimate_adjusted_day7_for_group(group: pd.DataFrame) -> pd.DataFrame:
    group = group.sort_values("day_order").copy()

    group["observed_green_area_percent"] = group["green_area_percent"]
    group["adjusted_green_area_percent"] = group["green_area_percent"]
    group["day7_imputed"] = "No"
    group["imputation_reason"] = ""
    group["imputation_method"] = ""
    group["previous_growth_rate_pp_per_day"] = math.nan
    group["previous_observed_green_percent"] = math.nan
    group["previous_observed_day_order"] = math.nan
    group["previous_observed_days_since_day1"] = math.nan

    if 7 not in set(group["day_order"]):
        return group

    day7_index = group.index[group["day_order"].eq(7)][0]
    day7_row = group.loc[day7_index]

    raw_day7 = parse_yes(day7_row.get("raw_current_green_evidence", ""))
    possible_flag = parse_yes(day7_row.get("possible_day7_bug_eaten", ""))

    prior_rows = group.loc[group["day_order"] < 7].copy()

    prior_raw_positive = prior_rows.loc[
        prior_rows["raw_current_green_evidence"].apply(parse_yes)
    ].copy()

    had_prior_visible_crop = not prior_raw_positive.empty
    should_impute = possible_flag or (had_prior_visible_crop and not raw_day7)

    if not should_impute:
        return group

    prior_positive_green = prior_rows.loc[
        pd.to_numeric(prior_rows["green_area_percent"], errors="coerce").fillna(0) > 0
    ].copy()

    if prior_positive_green.empty:
        return group

    prior_positive_green = prior_positive_green.sort_values("day_order")

    last_row = prior_positive_green.iloc[-1]
    last_green = float(last_row["green_area_percent"])
    last_elapsed = float(last_row["days_since_day1"])
    last_day_order = int(last_row["day_order"])

    slope = 0.0

    if len(prior_positive_green) >= 2:
        previous_row = prior_positive_green.iloc[-2]
        previous_green = float(previous_row["green_area_percent"])
        previous_elapsed = float(previous_row["days_since_day1"])

        elapsed_difference = max(last_elapsed - previous_elapsed, 1e-9)
        raw_slope = (last_green - previous_green) / elapsed_difference
        slope = max(0.0, raw_slope)

    day7_elapsed = float(day7_row["days_since_day1"])
    elapsed_to_day7 = max(day7_elapsed - last_elapsed, 0.0)

    estimated = last_green + slope * elapsed_to_day7
    observed_day7 = float(day7_row["green_area_percent"])

    adjusted = min(100.0, max(observed_day7, estimated, last_green))

    group.loc[day7_index, "adjusted_green_area_percent"] = adjusted
    group.loc[day7_index, "day7_imputed"] = "Yes"
    group.loc[day7_index, "imputation_reason"] = (
        "Day 7 cell had no current green evidence but showed visible green "
        "evidence earlier; treated as possible bug-eaten/missing crop."
    )
    group.loc[day7_index, "imputation_method"] = (
        "Estimated from latest prior observed green cover and previous positive "
        "growth rate. Negative previous growth was not projected."
    )
    group.loc[day7_index, "previous_growth_rate_pp_per_day"] = slope
    group.loc[day7_index, "previous_observed_green_percent"] = last_green
    group.loc[day7_index, "previous_observed_day_order"] = last_day_order
    group.loc[day7_index, "previous_observed_days_since_day1"] = last_elapsed

    return group


def create_adjusted_cell_table(cell: pd.DataFrame) -> pd.DataFrame:
    adjusted_groups = []

    for (_tray_no, _cell_id), group in cell.groupby(["tray_no", "cell_id"], sort=True, dropna=False):
        adjusted_groups.append(estimate_adjusted_day7_for_group(group))

    if not adjusted_groups:
        return pd.DataFrame()

    adjusted = pd.concat(adjusted_groups, ignore_index=True)

    adjusted["adjusted_raw_current_green_evidence"] = np.where(
        pd.to_numeric(adjusted["adjusted_green_area_percent"], errors="coerce").fillna(0) > 0,
        "Yes",
        "No",
    )

    adjusted["adjustment_difference_pp"] = (
        adjusted["adjusted_green_area_percent"]
        - adjusted["observed_green_area_percent"]
    )

    return adjusted.sort_values(
        ["day_order", "tray_no", "cell_id"]
    ).reset_index(drop=True)


# ============================================================
# 7) TRAY DAILY METRICS
# ============================================================

def create_tray_daily_metrics(adjusted_cell: pd.DataFrame) -> pd.DataFrame:
    adjusted_cell = adjusted_cell.copy()

    adjusted_cell["raw_green_bool"] = adjusted_cell["raw_current_green_evidence"].apply(parse_yes)
    adjusted_cell["tracked_bool"] = adjusted_cell["tracked_visible_emerged"].apply(parse_yes)

    if "newly_emerged_today" in adjusted_cell.columns:
        adjusted_cell["newly_emerged_bool"] = adjusted_cell["newly_emerged_today"].apply(parse_yes)
    else:
        adjusted_cell["newly_emerged_bool"] = False

    group_columns = [
        "day_order",
        "day",
        "calendar_date",
        "days_since_day1",
        "days_since_previous_photo",
        "tray",
        "tray_no",
        "microbe_status",
        "treatment",
        "label_environment",
        "fixed_environment",
        "observed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
        "heat_phase",
        "moisture_phase",
        "moisture_watered_today",
        "capture_id",
    ]

    group_columns = [column for column in group_columns if column in adjusted_cell.columns]

    for column in group_columns:
        if adjusted_cell[column].dtype == "O":
            adjusted_cell[column] = adjusted_cell[column].fillna("")

    daily = (
        adjusted_cell.groupby(group_columns, as_index=False, dropna=False)
        .agg(
            cells_processed=("cell_id", "count"),
            raw_green_cells=("raw_green_bool", "sum"),
            tracked_emerged_cells=("tracked_bool", "sum"),
            newly_emerged_cells=("newly_emerged_bool", "sum"),
            observed_green_cover_percent=("observed_green_area_percent", "mean"),
            adjusted_green_cover_percent=("adjusted_green_area_percent", "mean"),
            day7_imputed_cells=("day7_imputed", lambda s: int((s == "Yes").sum())),
            mean_adjustment_difference_pp=("adjustment_difference_pp", "mean"),
        )
    )

    invalid = daily.loc[daily["cells_processed"] != EXPECTED_CELLS]

    if not invalid.empty:
        print("\nWARNING: Some tray/day records do not have exactly 70 cells.")
        print(invalid[["day_order", "tray_no", "cells_processed"]].to_string(index=False))

    daily["raw_green_percent"] = daily["raw_green_cells"] / EXPECTED_CELLS * 100.0
    daily["tracked_emergence_percent"] = daily["tracked_emerged_cells"] / EXPECTED_CELLS * 100.0
    daily["newly_emerged_percent"] = daily["newly_emerged_cells"] / EXPECTED_CELLS * 100.0

    return daily.sort_values(["day_order", "tray_no"]).reset_index(drop=True)


# ============================================================
# 8) TRAY GROWTH METRICS
# ============================================================

def value_for_day(rows_by_day: dict[int, pd.Series], day_order: int, column: str):
    row = rows_by_day.get(day_order)
    if row is None:
        return math.nan
    return row.get(column, math.nan)


def elapsed_for_day(rows_by_day: dict[int, pd.Series], day_order: int):
    row = rows_by_day.get(day_order)
    if row is None:
        return math.nan
    return pd.to_numeric(row.get("days_since_day1", math.nan), errors="coerce")


def first_non_nan_day(rows_by_day: dict[int, pd.Series], preferred_day: int, fallback_day: int, column: str):
    preferred = value_for_day(rows_by_day, preferred_day, column)
    if pd.notna(preferred):
        return preferred
    return value_for_day(rows_by_day, fallback_day, column)


def rate_between_days(rows_by_day: dict[int, pd.Series], start_day: int, end_day: int, column: str):
    start_value = value_for_day(rows_by_day, start_day, column)
    end_value = value_for_day(rows_by_day, end_day, column)
    start_elapsed = elapsed_for_day(rows_by_day, start_day)
    end_elapsed = elapsed_for_day(rows_by_day, end_day)

    if (
        pd.isna(start_value)
        or pd.isna(end_value)
        or pd.isna(start_elapsed)
        or pd.isna(end_elapsed)
        or end_elapsed <= start_elapsed
    ):
        return math.nan

    return (end_value - start_value) / (end_elapsed - start_elapsed)


def t50_elapsed_days(group: pd.DataFrame, emergence_column: str = "tracked_emergence_percent"):
    series = group.sort_values("days_since_day1")[
        ["day_order", "days_since_day1", emergence_column]
    ].dropna()

    if series.empty or series[emergence_column].max() < 50:
        return math.nan

    previous_elapsed = None
    previous_value = None

    for row in series.itertuples(index=False):
        elapsed = float(row.days_since_day1)
        value = float(getattr(row, emergence_column))

        if value >= 50:
            if previous_elapsed is None:
                return elapsed
            if value == previous_value:
                return elapsed

            fraction = (50.0 - previous_value) / (value - previous_value)
            return previous_elapsed + fraction * (elapsed - previous_elapsed)

        previous_elapsed = elapsed
        previous_value = value

    return math.nan


def create_first_emergence_metrics(adjusted_cell: pd.DataFrame) -> pd.DataFrame:
    if "first_visible_emergence_day_order" not in adjusted_cell.columns:
        return pd.DataFrame()

    first = adjusted_cell[
        [
            "tray_no",
            "cell_id",
            "first_visible_emergence_day_order",
            "first_visible_emergence_date",
        ]
    ].drop_duplicates(subset=["tray_no", "cell_id"]).copy()

    first["first_visible_emergence_day_order"] = pd.to_numeric(
        first["first_visible_emergence_day_order"],
        errors="coerce",
    )

    elapsed_lookup = (
        adjusted_cell[["day_order", "days_since_day1"]]
        .drop_duplicates()
        .set_index("day_order")["days_since_day1"]
        .to_dict()
    )

    first["first_visible_emergence_elapsed_days"] = first[
        "first_visible_emergence_day_order"
    ].map(elapsed_lookup)

    first = first.dropna(subset=["first_visible_emergence_day_order"])

    if first.empty:
        return pd.DataFrame()

    return (
        first.groupby("tray_no", as_index=False, dropna=False)
        .agg(
            mean_first_visible_emergence_day_order=("first_visible_emergence_day_order", "mean"),
            median_first_visible_emergence_day_order=("first_visible_emergence_day_order", "median"),
            mean_first_visible_emergence_elapsed_days=("first_visible_emergence_elapsed_days", "mean"),
            median_first_visible_emergence_elapsed_days=("first_visible_emergence_elapsed_days", "median"),
        )
    )


def create_tray_growth_metrics(tray_daily: pd.DataFrame, adjusted_cell: pd.DataFrame) -> pd.DataFrame:
    records = []

    for tray_no, group in tray_daily.groupby("tray_no", dropna=False):
        group = group.sort_values("day_order")
        rows_by_day = {int(row["day_order"]): row for _, row in group.iterrows()}
        reference = group.iloc[0]

        final_day = int(group["day_order"].max())
        preferred_final_day = 7 if 7 in rows_by_day else final_day
        final_row = rows_by_day[preferred_final_day]

        record = {
            "tray_no": int(tray_no),
            "tray": str(reference["tray"]),
            "microbe_status": str(reference["microbe_status"]),
            "treatment": str(reference["treatment"]),
            "label_environment": str(reference["label_environment"]),
            "fixed_environment": str(reference["fixed_environment"]),
            "environment_group": str(reference["environment_group"]),
            "microbe_treatment": str(reference["microbe_treatment"]),
            "treatment_environment": str(reference["treatment_environment"]),
            "microbe_environment": str(reference["microbe_environment"]),
            "available_day_count": int(group["day_order"].nunique()),
            "final_day_order": preferred_final_day,
            "final_calendar_date": str(final_row["calendar_date"]),
            "final_days_since_day1": float(final_row["days_since_day1"]),
            "day1_tracked_emergence_percent": value_for_day(rows_by_day, 1, "tracked_emergence_percent"),
            "final_tracked_emergence_percent": value_for_day(rows_by_day, preferred_final_day, "tracked_emergence_percent"),
            "day1_observed_green_cover_percent": value_for_day(rows_by_day, 1, "observed_green_cover_percent"),
            "final_observed_green_cover_percent": value_for_day(rows_by_day, preferred_final_day, "observed_green_cover_percent"),
            "day1_adjusted_green_cover_percent": value_for_day(rows_by_day, 1, "adjusted_green_cover_percent"),
            "final_adjusted_green_cover_percent": value_for_day(rows_by_day, preferred_final_day, "adjusted_green_cover_percent"),
            "day7_imputed_cells": value_for_day(rows_by_day, 7, "day7_imputed_cells"),
            "day7_mean_adjustment_difference_pp": value_for_day(rows_by_day, 7, "mean_adjustment_difference_pp"),
            "emergence_rate_day1_to_final_pp_per_day": rate_between_days(rows_by_day, 1, preferred_final_day, "tracked_emergence_percent"),
            "observed_green_cover_rate_day1_to_final_pp_per_day": rate_between_days(rows_by_day, 1, preferred_final_day, "observed_green_cover_percent"),
            "adjusted_green_cover_rate_day1_to_final_pp_per_day": rate_between_days(rows_by_day, 1, preferred_final_day, "adjusted_green_cover_percent"),
            "observed_green_cover_rate_day1_to_day6_pp_per_day": rate_between_days(rows_by_day, 1, 6, "observed_green_cover_percent"),
            "adjusted_green_cover_rate_day1_to_day6_pp_per_day": rate_between_days(rows_by_day, 1, 6, "adjusted_green_cover_percent"),
            "observed_green_cover_rate_day6_to_day7_pp_per_day": rate_between_days(rows_by_day, 6, 7, "observed_green_cover_percent"),
            "adjusted_green_cover_rate_day6_to_day7_pp_per_day": rate_between_days(rows_by_day, 6, 7, "adjusted_green_cover_percent"),
            "observed_day7_change_from_day6_pp": (
                value_for_day(rows_by_day, 7, "observed_green_cover_percent")
                - value_for_day(rows_by_day, 6, "observed_green_cover_percent")
            ),
            "adjusted_day7_change_from_day6_pp": (
                value_for_day(rows_by_day, 7, "adjusted_green_cover_percent")
                - value_for_day(rows_by_day, 6, "adjusted_green_cover_percent")
            ),
            "t50_elapsed_days": t50_elapsed_days(group),
        }

        # Backward-compatible column names.
        record["emergence_rate_day1_to_day7_pp_per_day"] = record["emergence_rate_day1_to_final_pp_per_day"]
        record["observed_green_cover_rate_day1_to_day7_pp_per_day"] = record["observed_green_cover_rate_day1_to_final_pp_per_day"]
        record["adjusted_green_cover_rate_day1_to_day7_pp_per_day"] = record["adjusted_green_cover_rate_day1_to_final_pp_per_day"]

        records.append(record)

    metrics = pd.DataFrame(records)

    first_metrics = create_first_emergence_metrics(adjusted_cell)

    if not first_metrics.empty:
        metrics = metrics.merge(
            first_metrics,
            on="tray_no",
            how="left",
            validate="one_to_one",
        )

    for component in PERFORMANCE_COMPONENTS_OBSERVED:
        metrics[f"{component}_observed_score"] = (
            minmax_score(metrics[component]) if component in metrics.columns else math.nan
        )

    observed_score_columns = [
        f"{component}_observed_score"
        for component in PERFORMANCE_COMPONENTS_OBSERVED
        if f"{component}_observed_score" in metrics.columns
    ]

    if "t50_elapsed_days" in metrics.columns:
        metrics["t50_observed_score"] = inverse_minmax_score(metrics["t50_elapsed_days"])
        observed_score_columns.append("t50_observed_score")

    metrics["overall_observed_rgb_score"] = metrics[observed_score_columns].mean(axis=1, skipna=True)
    metrics["overall_observed_rgb_rank"] = metrics["overall_observed_rgb_score"].rank(
        ascending=False,
        method="min",
    ).astype("Int64")

    for component in PERFORMANCE_COMPONENTS_ADJUSTED:
        metrics[f"{component}_adjusted_score"] = (
            minmax_score(metrics[component]) if component in metrics.columns else math.nan
        )

    adjusted_score_columns = [
        f"{component}_adjusted_score"
        for component in PERFORMANCE_COMPONENTS_ADJUSTED
        if f"{component}_adjusted_score" in metrics.columns
    ]

    if "t50_elapsed_days" in metrics.columns:
        metrics["t50_adjusted_score"] = inverse_minmax_score(metrics["t50_elapsed_days"])
        adjusted_score_columns.append("t50_adjusted_score")

    metrics["overall_adjusted_rgb_score"] = metrics[adjusted_score_columns].mean(axis=1, skipna=True)
    metrics["overall_adjusted_rgb_rank"] = metrics["overall_adjusted_rgb_score"].rank(
        ascending=False,
        method="min",
    ).astype("Int64")

    return metrics.sort_values(
        ["overall_adjusted_rgb_rank", "tray_no"],
        na_position="last",
    ).reset_index(drop=True)


# ============================================================
# 9) GROUP TABLES
# ============================================================

def create_group_daily(tray_daily: pd.DataFrame, group_column: str, group_type: str) -> pd.DataFrame:
    result = (
        tray_daily.groupby(
            [
                group_column,
                "day_order",
                "day",
                "calendar_date",
                "days_since_day1",
            ],
            as_index=False,
            dropna=False,
        )
        .agg(
            tray_count=("tray_no", "nunique"),
            mean_tracked_emergence_percent=("tracked_emergence_percent", "mean"),
            sd_tracked_emergence_percent=("tracked_emergence_percent", "std"),
            mean_observed_green_cover_percent=("observed_green_cover_percent", "mean"),
            sd_observed_green_cover_percent=("observed_green_cover_percent", "std"),
            mean_adjusted_green_cover_percent=("adjusted_green_cover_percent", "mean"),
            sd_adjusted_green_cover_percent=("adjusted_green_cover_percent", "std"),
            mean_newly_emerged_cells=("newly_emerged_cells", "mean"),
            mean_day7_imputed_cells=("day7_imputed_cells", "mean"),
        )
        .rename(columns={group_column: "group"})
    )

    result["group_type"] = group_type

    for column in [
        "sd_tracked_emergence_percent",
        "sd_observed_green_cover_percent",
        "sd_adjusted_green_cover_percent",
    ]:
        result[column] = result[column].fillna(0.0)

    return result


def create_group_growth(tray_metrics: pd.DataFrame, group_column: str, group_type: str) -> pd.DataFrame:
    metric_columns = [
        "final_tracked_emergence_percent",
        "final_observed_green_cover_percent",
        "final_adjusted_green_cover_percent",
        "day7_imputed_cells",
        "day7_mean_adjustment_difference_pp",
        "emergence_rate_day1_to_final_pp_per_day",
        "observed_green_cover_rate_day1_to_final_pp_per_day",
        "adjusted_green_cover_rate_day1_to_final_pp_per_day",
        "observed_green_cover_rate_day6_to_day7_pp_per_day",
        "adjusted_green_cover_rate_day6_to_day7_pp_per_day",
        "t50_elapsed_days",
        "mean_first_visible_emergence_elapsed_days",
        "overall_observed_rgb_score",
        "overall_adjusted_rgb_score",
    ]

    aggregations = {"tray_count": ("tray_no", "nunique")}

    for metric in metric_columns:
        if metric in tray_metrics.columns:
            aggregations[f"mean_{metric}"] = (metric, "mean")
            aggregations[f"sd_{metric}"] = (metric, "std")

    output = (
        tray_metrics.groupby(group_column, as_index=False, dropna=False)
        .agg(**aggregations)
        .rename(columns={group_column: "group"})
    )

    output["group_type"] = group_type

    for column in output.columns:
        if column.startswith("sd_"):
            output[column] = output[column].fillna(0.0)

    # Add backward-compatible group-growth names.
    rename_map = {
        "mean_emergence_rate_day1_to_final_pp_per_day": "mean_emergence_rate_day1_to_day7_pp_per_day",
        "mean_observed_green_cover_rate_day1_to_final_pp_per_day": "mean_observed_green_cover_rate_day1_to_day7_pp_per_day",
        "mean_adjusted_green_cover_rate_day1_to_final_pp_per_day": "mean_adjusted_green_cover_rate_day1_to_day7_pp_per_day",
    }

    for source, target in rename_map.items():
        if source in output.columns and target not in output.columns:
            output[target] = output[source]

    return output


def create_all_group_tables(
    tray_daily: pd.DataFrame,
    tray_metrics: pd.DataFrame,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    group_specs = [
        ("microbe_status", "Microbe Status"),
        ("treatment", "Treatment Type"),
        ("microbe_treatment", "Microbe x Treatment"),
        ("treatment_environment", "Treatment x Environment"),
        ("microbe_environment", "Microbe x Environment"),
    ]

    daily_parts = []
    growth_parts = []

    for group_column, group_type in group_specs:
        daily_parts.append(create_group_daily(tray_daily, group_column, group_type))
        growth_parts.append(create_group_growth(tray_metrics, group_column, group_type))

    group_daily = pd.concat(daily_parts, ignore_index=True)
    group_growth = pd.concat(growth_parts, ignore_index=True)

    return group_daily, group_growth


# ============================================================
# 10) INSIDE VS OUTSIDE COMPARISON
# ============================================================

def create_inside_outside_comparison(tray_metrics: pd.DataFrame) -> pd.DataFrame:
    rows = []

    frame = tray_metrics.copy()

    frame["tray_no"] = pd.to_numeric(frame["tray_no"], errors="coerce")
    frame["treatment"] = frame.apply(
        lambda row: TRAY_FALLBACK.get(int(row["tray_no"]), {}).get(
            "treatment",
            standardise_treatment(row.get("treatment", ""), tray_no_as_int(row.get("tray_no"))),
        ),
        axis=1,
    )
    frame["fixed_environment"] = frame.apply(
        lambda row: TRAY_FALLBACK.get(int(row["tray_no"]), {}).get(
            "fixed_environment",
            standardise_environment(
                row.get("fixed_environment", row.get("label_environment", "")),
                tray_no_as_int(row.get("tray_no")),
                row.get("treatment", ""),
            ),
        ),
        axis=1,
    )

    for treatment in FIXED_ENVIRONMENT_TREATMENTS:
        subset = frame.loc[
            frame["treatment"].astype(str).str.casefold().eq(treatment.casefold())
        ].copy()

        inside = subset.loc[
            subset["fixed_environment"].astype(str).str.casefold().eq("inside")
        ].copy()

        outside = subset.loc[
            subset["fixed_environment"].astype(str).str.casefold().eq("outside")
        ].copy()

        row = {
            "treatment": treatment,
            "inside_tray_count": count_unique_or_zero(inside, "tray_no"),
            "outside_tray_count": count_unique_or_zero(outside, "tray_no"),
            "inside_trays": ", ".join(inside["tray"].astype(str).tolist()),
            "outside_trays": ", ".join(outside["tray"].astype(str).tolist()),
            "inside_final_emergence_percent": mean_or_nan(inside, "final_tracked_emergence_percent"),
            "outside_final_emergence_percent": mean_or_nan(outside, "final_tracked_emergence_percent"),
            "inside_final_observed_green_cover_percent": mean_or_nan(inside, "final_observed_green_cover_percent"),
            "outside_final_observed_green_cover_percent": mean_or_nan(outside, "final_observed_green_cover_percent"),
            "inside_final_adjusted_green_cover_percent": mean_or_nan(inside, "final_adjusted_green_cover_percent"),
            "outside_final_adjusted_green_cover_percent": mean_or_nan(outside, "final_adjusted_green_cover_percent"),
            "inside_emergence_rate_pp_per_day": mean_or_nan(inside, "emergence_rate_day1_to_final_pp_per_day"),
            "outside_emergence_rate_pp_per_day": mean_or_nan(outside, "emergence_rate_day1_to_final_pp_per_day"),
            "inside_adjusted_green_cover_rate_pp_per_day": mean_or_nan(inside, "adjusted_green_cover_rate_day1_to_final_pp_per_day"),
            "outside_adjusted_green_cover_rate_pp_per_day": mean_or_nan(outside, "adjusted_green_cover_rate_day1_to_final_pp_per_day"),
            "inside_t50_elapsed_days": mean_or_nan(inside, "t50_elapsed_days"),
            "outside_t50_elapsed_days": mean_or_nan(outside, "t50_elapsed_days"),
            "inside_overall_adjusted_rgb_score": mean_or_nan(inside, "overall_adjusted_rgb_score"),
            "outside_overall_adjusted_rgb_score": mean_or_nan(outside, "overall_adjusted_rgb_score"),
        }

        row["inside_minus_outside_final_emergence_pp"] = (
            row["inside_final_emergence_percent"]
            - row["outside_final_emergence_percent"]
        )

        row["inside_minus_outside_adjusted_green_cover_pp"] = (
            row["inside_final_adjusted_green_cover_percent"]
            - row["outside_final_adjusted_green_cover_percent"]
        )

        row["inside_minus_outside_adjusted_growth_rate_pp_per_day"] = (
            row["inside_adjusted_green_cover_rate_pp_per_day"]
            - row["outside_adjusted_green_cover_rate_pp_per_day"]
        )

        row["inside_minus_outside_overall_score"] = (
            row["inside_overall_adjusted_rgb_score"]
            - row["outside_overall_adjusted_rgb_score"]
        )

        if row["inside_tray_count"] == 0 or row["outside_tray_count"] == 0:
            interpretation = "Insufficient Inside or Outside tray data for this comparison."
        elif pd.isna(row["inside_minus_outside_adjusted_green_cover_pp"]):
            interpretation = "Insufficient numeric data for adjusted green-cover comparison."
        elif row["inside_minus_outside_adjusted_green_cover_pp"] > 0:
            interpretation = "Inside performed higher than Outside for adjusted Day 7 RGB green-cover."
        elif row["inside_minus_outside_adjusted_green_cover_pp"] < 0:
            interpretation = "Outside performed higher than Inside for adjusted Day 7 RGB green-cover."
        else:
            interpretation = "Inside and Outside were equal for adjusted Day 7 RGB green-cover."

        row["interpretation"] = interpretation
        rows.append(row)

    return pd.DataFrame(rows)


# ============================================================
# 11) PHASE TABLES
# ============================================================

def create_phase_tables(tray_daily: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    heat = tray_daily.loc[
        tray_daily["treatment"].astype(str).str.casefold().eq("heat")
    ].copy()

    moisture = tray_daily.loc[
        tray_daily["treatment"].astype(str).str.casefold().eq("moisture")
    ].copy()

    if not heat.empty and "heat_phase" in heat.columns:
        heat_phase = (
            heat.groupby(
                [
                    "heat_phase",
                    "microbe_status",
                    "day_order",
                    "day",
                    "days_since_day1",
                ],
                as_index=False,
                dropna=False,
            )
            .agg(
                tray_count=("tray_no", "nunique"),
                mean_tracked_emergence_percent=("tracked_emergence_percent", "mean"),
                mean_observed_green_cover_percent=("observed_green_cover_percent", "mean"),
                mean_adjusted_green_cover_percent=("adjusted_green_cover_percent", "mean"),
                mean_day7_imputed_cells=("day7_imputed_cells", "mean"),
            )
            .sort_values(["day_order", "microbe_status"])
            .reset_index(drop=True)
        )
    else:
        heat_phase = pd.DataFrame()

    if not moisture.empty and "moisture_phase" in moisture.columns:
        moisture_phase = (
            moisture.groupby(
                [
                    "moisture_phase",
                    "moisture_watered_today",
                    "microbe_status",
                    "fixed_environment",
                    "day_order",
                    "day",
                    "days_since_day1",
                ],
                as_index=False,
                dropna=False,
            )
            .agg(
                tray_count=("tray_no", "nunique"),
                mean_tracked_emergence_percent=("tracked_emergence_percent", "mean"),
                mean_observed_green_cover_percent=("observed_green_cover_percent", "mean"),
                mean_adjusted_green_cover_percent=("adjusted_green_cover_percent", "mean"),
                mean_day7_imputed_cells=("day7_imputed_cells", "mean"),
            )
            .sort_values(["day_order", "microbe_status", "fixed_environment"])
            .reset_index(drop=True)
        )
    else:
        moisture_phase = pd.DataFrame()

    return heat_phase, moisture_phase


# ============================================================
# 12) CHARTS
# ============================================================

def no_data_chart(title: str, output_path: Path, message: str = "No valid numeric data available") -> None:
    figure, axis = plt.subplots(figsize=(9, 5.5))
    axis.text(0.5, 0.5, message, ha="center", va="center", fontsize=12)
    axis.set_title(title)
    axis.axis("off")
    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_trend_chart(
    daily: pd.DataFrame,
    group_type: str,
    groups: list[str],
    value_column: str,
    sd_column: str,
    title: str,
    y_label: str,
    output_path: Path,
) -> None:
    subset = daily.loc[daily["group_type"].eq(group_type)].copy()

    if subset.empty or value_column not in subset.columns:
        no_data_chart(title, output_path)
        return

    figure, axis = plt.subplots(figsize=(11.5, 6.8))
    plotted = False

    for group in groups:
        series = subset.loc[subset["group"].eq(group)].sort_values("days_since_day1")

        if series.empty:
            continue

        y = pd.to_numeric(series[value_column], errors="coerce")

        if y.notna().sum() == 0:
            continue

        axis.errorbar(
            series["days_since_day1"],
            y,
            yerr=series[sd_column] if sd_column in series.columns else None,
            marker="o",
            linewidth=2,
            capsize=4,
            label=f"{group} (n={int(series['tray_count'].max())})",
        )

        plotted = True

    if not plotted:
        plt.close(figure)
        no_data_chart(title, output_path)
        return

    ticks = subset[["day", "days_since_day1"]].drop_duplicates().sort_values("days_since_day1")

    axis.set_title(title)
    axis.set_xlabel("Elapsed days since Day 1 image")
    axis.set_ylabel(y_label)
    axis.set_xticks(ticks["days_since_day1"])
    axis.set_xticklabels(ticks["day"])
    axis.grid(True, axis="y", alpha=0.30)
    axis.legend(loc="best")

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_day7_observed_adjusted_chart(tray_metrics: pd.DataFrame, output_path: Path) -> None:
    frame = tray_metrics.sort_values("tray_no").copy()

    if frame.empty:
        no_data_chart("Day 7 RGB green-cover: observed vs adjusted", output_path)
        return

    x = np.arange(len(frame))
    width = 0.38

    figure, axis = plt.subplots(figsize=(12, 6.5))

    axis.bar(
        x - width / 2,
        frame["final_observed_green_cover_percent"],
        width,
        label="Observed final",
    )

    axis.bar(
        x + width / 2,
        frame["final_adjusted_green_cover_percent"],
        width,
        label="Adjusted final",
    )

    axis.set_title("Final RGB green-cover: observed vs adjusted")
    axis.set_xlabel("Tray")
    axis.set_ylabel("Mean RGB green-cover (%)")
    axis.set_xticks(x)
    axis.set_xticklabels(frame["tray"], rotation=30, ha="right")
    axis.grid(True, axis="y", alpha=0.30)
    axis.legend(loc="best")

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_tray_ranking_chart(tray_metrics: pd.DataFrame, output_path: Path) -> None:
    frame = tray_metrics.dropna(subset=["overall_adjusted_rgb_score"]).sort_values(
        "overall_adjusted_rgb_score"
    ).copy()

    if frame.empty:
        no_data_chart("Trial 3 tray ranking by adjusted RGB performance score", output_path)
        return

    frame["label"] = (
        frame["tray"].astype(str)
        + " — "
        + frame["microbe_status"].astype(str)
        + " | "
        + frame["treatment"].astype(str)
        + " | "
        + frame["environment_group"].astype(str)
    )

    figure, axis = plt.subplots(figsize=(12.5, 7.2))

    axis.barh(frame["label"], frame["overall_adjusted_rgb_score"])
    axis.set_title("Trial 3 tray ranking by adjusted RGB performance score")
    axis.set_xlabel("Adjusted RGB performance score")
    axis.set_ylabel("Tray")
    axis.grid(True, axis="x", alpha=0.30)

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_bug_imputed_cells_chart(tray_metrics: pd.DataFrame, output_path: Path) -> None:
    frame = tray_metrics.sort_values("tray_no").copy()

    if frame.empty or "day7_imputed_cells" not in frame.columns:
        no_data_chart("Possible Day 7 bug-eaten/missing cells flagged by tray", output_path)
        return

    figure, axis = plt.subplots(figsize=(11.5, 6.2))

    axis.bar(frame["tray"], frame["day7_imputed_cells"].fillna(0))
    axis.set_title("Possible Day 7 bug-eaten/missing cells flagged by tray")
    axis.set_xlabel("Tray")
    axis.set_ylabel("Cells flagged")
    axis.grid(True, axis="y", alpha=0.30)
    axis.tick_params(axis="x", rotation=30)

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_inside_outside_green_chart(comparison: pd.DataFrame, output_path: Path) -> None:
    title = "Ideal and Moisture: Inside vs Outside adjusted final RGB green-cover"

    if comparison.empty:
        no_data_chart(title, output_path)
        return

    frame = comparison.copy()

    inside_values = pd.to_numeric(
        frame["inside_final_adjusted_green_cover_percent"],
        errors="coerce",
    )
    outside_values = pd.to_numeric(
        frame["outside_final_adjusted_green_cover_percent"],
        errors="coerce",
    )

    if inside_values.notna().sum() == 0 and outside_values.notna().sum() == 0:
        no_data_chart(title, output_path, "Inside/Outside values could not be calculated")
        return

    x = np.arange(len(frame))
    width = 0.36

    figure, axis = plt.subplots(figsize=(9.5, 5.8))

    axis.bar(x - width / 2, inside_values, width, label="Inside")
    axis.bar(x + width / 2, outside_values, width, label="Outside")

    for index, value in enumerate(inside_values):
        if pd.notna(value):
            axis.text(index - width / 2, value, f"{value:.2f}", ha="center", va="bottom", fontsize=8)

    for index, value in enumerate(outside_values):
        if pd.notna(value):
            axis.text(index + width / 2, value, f"{value:.2f}", ha="center", va="bottom", fontsize=8)

    axis.set_title(title)
    axis.set_xlabel("Treatment")
    axis.set_ylabel("Adjusted final RGB green-cover (%)")
    axis.set_xticks(x)
    axis.set_xticklabels(frame["treatment"])
    axis.grid(True, axis="y", alpha=0.30)
    axis.legend(loc="best")

    y_max = np.nanmax([inside_values.max(), outside_values.max()])
    if pd.notna(y_max):
        axis.set_ylim(0, max(y_max * 1.18, 1.0))

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def save_inside_outside_emergence_chart(comparison: pd.DataFrame, output_path: Path) -> None:
    title = "Ideal and Moisture: Inside vs Outside final visible emergence"

    if comparison.empty:
        no_data_chart(title, output_path)
        return

    frame = comparison.copy()

    inside_values = pd.to_numeric(
        frame["inside_final_emergence_percent"],
        errors="coerce",
    )
    outside_values = pd.to_numeric(
        frame["outside_final_emergence_percent"],
        errors="coerce",
    )

    if inside_values.notna().sum() == 0 and outside_values.notna().sum() == 0:
        no_data_chart(title, output_path, "Inside/Outside values could not be calculated")
        return

    x = np.arange(len(frame))
    width = 0.36

    figure, axis = plt.subplots(figsize=(9.5, 5.8))

    axis.bar(x - width / 2, inside_values, width, label="Inside")
    axis.bar(x + width / 2, outside_values, width, label="Outside")

    for index, value in enumerate(inside_values):
        if pd.notna(value):
            axis.text(index - width / 2, value, f"{value:.2f}", ha="center", va="bottom", fontsize=8)

    for index, value in enumerate(outside_values):
        if pd.notna(value):
            axis.text(index + width / 2, value, f"{value:.2f}", ha="center", va="bottom", fontsize=8)

    axis.set_title(title)
    axis.set_xlabel("Treatment")
    axis.set_ylabel("Final tracked visible emergence (%)")
    axis.set_xticks(x)
    axis.set_xticklabels(frame["treatment"])
    axis.grid(True, axis="y", alpha=0.30)
    axis.legend(loc="best")

    y_max = np.nanmax([inside_values.max(), outside_values.max()])
    if pd.notna(y_max):
        axis.set_ylim(0, max(y_max * 1.18, 1.0))

    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220)
    plt.close(figure)


def create_charts(
    group_daily: pd.DataFrame,
    tray_metrics: pd.DataFrame,
    inside_outside_comparison: pd.DataFrame,
) -> dict[str, Path]:
    CHARTS_ROOT.mkdir(parents=True, exist_ok=True)

    charts = {}

    path = CHARTS_ROOT / "01_microbes_visible_emergence_trend.png"
    save_trend_chart(
        group_daily,
        "Microbe Status",
        MICROBE_ORDER,
        "mean_tracked_emergence_percent",
        "sd_tracked_emergence_percent",
        "Tracked visible emergence: Microbes vs No Microbes",
        "Mean tracked visible emergence (%)",
        path,
    )
    charts["microbes_emergence"] = path

    path = CHARTS_ROOT / "02_microbes_adjusted_green_cover_trend.png"
    save_trend_chart(
        group_daily,
        "Microbe Status",
        MICROBE_ORDER,
        "mean_adjusted_green_cover_percent",
        "sd_adjusted_green_cover_percent",
        "Adjusted RGB green-cover: Microbes vs No Microbes",
        "Mean adjusted RGB green-cover (%)",
        path,
    )
    charts["microbes_green"] = path

    path = CHARTS_ROOT / "03_treatment_visible_emergence_trend.png"
    save_trend_chart(
        group_daily,
        "Treatment Type",
        TREATMENT_ORDER,
        "mean_tracked_emergence_percent",
        "sd_tracked_emergence_percent",
        "Tracked visible emergence by treatment type",
        "Mean tracked visible emergence (%)",
        path,
    )
    charts["treatment_emergence"] = path

    path = CHARTS_ROOT / "04_treatment_adjusted_green_cover_trend.png"
    save_trend_chart(
        group_daily,
        "Treatment Type",
        TREATMENT_ORDER,
        "mean_adjusted_green_cover_percent",
        "sd_adjusted_green_cover_percent",
        "Adjusted RGB green-cover by treatment type",
        "Mean adjusted RGB green-cover (%)",
        path,
    )
    charts["treatment_green"] = path

    path = CHARTS_ROOT / "05_day7_observed_vs_adjusted_green_cover_by_tray.png"
    save_day7_observed_adjusted_chart(tray_metrics, path)
    charts["day7_observed_adjusted"] = path

    path = CHARTS_ROOT / "06_adjusted_rgb_tray_ranking.png"
    save_tray_ranking_chart(tray_metrics, path)
    charts["tray_ranking"] = path

    path = CHARTS_ROOT / "07_day7_possible_bug_eaten_cells_by_tray.png"
    save_bug_imputed_cells_chart(tray_metrics, path)
    charts["bug_cells"] = path

    path = CHARTS_ROOT / "08_inside_outside_adjusted_green_cover_ideal_moisture.png"
    save_inside_outside_green_chart(inside_outside_comparison, path)
    charts["inside_outside_green"] = path

    path = CHARTS_ROOT / "09_inside_outside_emergence_ideal_moisture.png"
    save_inside_outside_emergence_chart(inside_outside_comparison, path)
    charts["inside_outside_emergence"] = path

    return charts


# ============================================================
# 13) SAVE CSV AND EXCEL
# ============================================================

def style_workbook(path: Path) -> None:
    workbook = load_workbook(path)

    header_fill = PatternFill("solid", fgColor="1F4E78")

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.row_dimensions[1].height = 34

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = column_cells[0].column_letter
            longest = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in column_cells
            )
            worksheet.column_dimensions[letter].width = min(max(12, longest + 2), 58)

    workbook.save(path)


def save_tables(
    adjusted_cell: pd.DataFrame,
    tray_daily: pd.DataFrame,
    tray_metrics: pd.DataFrame,
    group_daily: pd.DataFrame,
    group_growth: pd.DataFrame,
    heat_phase: pd.DataFrame,
    moisture_phase: pd.DataFrame,
    inside_outside_comparison: pd.DataFrame,
) -> dict[str, Path]:
    REPORTS_ROOT.mkdir(parents=True, exist_ok=True)

    paths = {
        "adjusted_cell": REPORTS_ROOT / "cell_growth_with_day7_adjustment.csv",
        "tray_daily": REPORTS_ROOT / "tray_daily_rgb_metrics.csv",
        "tray_metrics": REPORTS_ROOT / "tray_growth_metrics.csv",
        "group_daily": REPORTS_ROOT / "group_daily_metrics.csv",
        "group_growth": REPORTS_ROOT / "group_growth_metrics.csv",
        "heat_phase": REPORTS_ROOT / "heat_phase_response.csv",
        "moisture_phase": REPORTS_ROOT / "moisture_phase_response.csv",
        "inside_outside": REPORTS_ROOT / "inside_outside_comparison_ideal_moisture.csv",
        "environment_debug": REPORTS_ROOT / "environment_metadata_debug.csv",
        "day7_imputed": REPORTS_ROOT / "possible_day7_bug_eaten_cells.csv",
        "excel": REPORTS_ROOT / "rgb_growth_treatment_report.xlsx",
    }

    adjusted_cell.to_csv(paths["adjusted_cell"], index=False)
    tray_daily.to_csv(paths["tray_daily"], index=False)
    tray_metrics.to_csv(paths["tray_metrics"], index=False)
    group_daily.to_csv(paths["group_daily"], index=False)
    group_growth.to_csv(paths["group_growth"], index=False)
    heat_phase.to_csv(paths["heat_phase"], index=False)
    moisture_phase.to_csv(paths["moisture_phase"], index=False)
    inside_outside_comparison.to_csv(paths["inside_outside"], index=False)

    environment_debug_columns = [
        "tray_no",
        "tray",
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
        "final_tracked_emergence_percent",
        "final_adjusted_green_cover_percent",
        "overall_adjusted_rgb_score",
    ]

    environment_debug_columns = [
        column for column in environment_debug_columns if column in tray_metrics.columns
    ]

    tray_metrics[environment_debug_columns].to_csv(paths["environment_debug"], index=False)

    day7_imputed = adjusted_cell.loc[adjusted_cell["day7_imputed"].eq("Yes")].copy()
    day7_imputed.to_csv(paths["day7_imputed"], index=False)

    readme = pd.DataFrame(
        {
            "Notes": [
                "This workbook summarises Trial 3 RGB visible-emergence and green-cover growth metrics.",
                "Observed Day 7 values come directly from Script 04 and are not overwritten.",
                "Adjusted Day 7 values are created only for cells visible before Day 7 but missing on Day 7.",
                "day7_imputed = Yes identifies adjusted records.",
                "Growth rates use real elapsed days since the Day 1 image.",
                "Day 7 is 8 elapsed days after Day 1 because images were skipped on 05/07/2026 and 06/07/2026.",
                "Inside vs Outside comparison uses tray-number fallback metadata for Ideal and Moisture trays.",
                "Heat trays are excluded from direct Inside vs Outside comparison because they moved between environments.",
                "These are descriptive image-derived results, not formal statistical proof.",
                "NDVI/NDRE are not included here; they are handled by later multispectral scripts.",
            ]
        }
    )

    with pd.ExcelWriter(paths["excel"], engine="openpyxl") as writer:
        safe_round_dataframe(tray_metrics).to_excel(
            writer,
            sheet_name="Tray Growth Metrics",
            index=False,
        )

        safe_round_dataframe(tray_daily).to_excel(
            writer,
            sheet_name="Tray Daily Metrics",
            index=False,
        )

        safe_round_dataframe(group_daily).to_excel(
            writer,
            sheet_name="Group Daily Metrics",
            index=False,
        )

        safe_round_dataframe(group_growth).to_excel(
            writer,
            sheet_name="Group Growth Metrics",
            index=False,
        )

        safe_round_dataframe(inside_outside_comparison).to_excel(
            writer,
            sheet_name="Inside Outside Compare",
            index=False,
        )

        safe_round_dataframe(day7_imputed).to_excel(
            writer,
            sheet_name="Day7 Imputed Cells",
            index=False,
        )

        safe_round_dataframe(heat_phase).to_excel(
            writer,
            sheet_name="Heat Phase Response",
            index=False,
        )

        safe_round_dataframe(moisture_phase).to_excel(
            writer,
            sheet_name="Moisture Phase Response",
            index=False,
        )

        safe_round_dataframe(tray_metrics[environment_debug_columns]).to_excel(
            writer,
            sheet_name="Environment Debug",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_workbook(paths["excel"])

    return paths


# ============================================================
# 14) WORD REPORT
# ============================================================

def add_docx_table(document, dataframe: pd.DataFrame, columns: list[str], max_rows: int = 12):
    if dataframe.empty:
        document.add_paragraph("No table data was available.")
        return

    columns = [column for column in columns if column in dataframe.columns]

    if not columns:
        document.add_paragraph("No valid columns were available for this table.")
        return

    frame = dataframe[columns].head(max_rows).copy()
    frame = safe_round_dataframe(frame, 3)

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column.replace("_", " ").title()

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row[column])


def add_picture_if_exists(document, path: Path | None, width_inches: float = 6.3):
    if path is None or not Path(path).exists():
        return

    document.add_picture(str(path), width=Inches(width_inches))


def best_row_by_metric(dataframe: pd.DataFrame, metric: str):
    if dataframe.empty or metric not in dataframe.columns:
        return None

    frame = dataframe.dropna(subset=[metric])

    if frame.empty:
        return None

    return frame.loc[frame[metric].idxmax()]


def group_summary_sentence(group_growth: pd.DataFrame, group_type: str, metric: str, metric_label: str) -> str:
    subset = group_growth.loc[group_growth["group_type"].eq(group_type)].copy()

    if subset.empty or metric not in subset.columns:
        return f"No valid {group_type.lower()} summary was available for {metric_label}."

    subset = subset.dropna(subset=[metric])

    if subset.empty:
        return f"No valid {group_type.lower()} summary was available for {metric_label}."

    best = subset.loc[subset[metric].idxmax()]

    return (
        f"For {metric_label}, the highest mean value in the {group_type.lower()} "
        f"comparison was recorded by {best['group']} ({format_number(best[metric])})."
    )


def describe_output_file(document, filename: str, description: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(filename).bold = True
    paragraph.add_run(f": {description}")


def create_word_report(
    output_path: Path,
    tray_metrics: pd.DataFrame,
    group_growth: pd.DataFrame,
    inside_outside_comparison: pd.DataFrame,
    adjusted_cell: pd.DataFrame,
    charts: dict[str, Path],
):
    if not DOCX_AVAILABLE:
        print("WARNING: python-docx is not installed. Word report was skipped.")
        print("Install it with: pip install python-docx")
        return None

    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    title = document.add_heading(
        "Trial 3 RGB Growth-Rate and Treatment Comparison Report",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This report summarises the Trial 3 RGB visible-emergence and green-cover "
        "analysis. The purpose of this script was to convert the cell-level outputs "
        "from Script 04 into tray-level and treatment-level growth comparisons. "
        "The report also documents the generated charts, CSV files, Excel sheets, "
        "and interpretation notes for final internship reporting."
    )

    document.add_paragraph(
        "The results are based on visible seedling evidence detected from RGB images. "
        "They should be interpreted as visible emergence and image-derived green-cover "
        "measurements, not direct biological germination below the soil surface."
    )

    document.add_heading("1. Data processed", level=1)

    tray_count = int(tray_metrics["tray_no"].nunique()) if not tray_metrics.empty else 0
    cell_record_count = len(adjusted_cell)
    imputed_count = int(adjusted_cell["day7_imputed"].eq("Yes").sum()) if "day7_imputed" in adjusted_cell.columns else 0

    document.add_paragraph(
        f"The analysis used {tray_count} trays, with 70 cells per tray. "
        f"The adjusted cell table contained {cell_record_count} cell-day records. "
        f"The corrected observation period was Day 1 (29/06/2026) to Day 7 "
        f"(07/07/2026). A total of {imputed_count} Day 7 cell records were flagged "
        f"for adjusted estimation."
    )

    document.add_paragraph(
        "Observed Day 7 values were preserved exactly. Adjusted Day 7 values were "
        "created separately only where a cell had visible green evidence before Day 7 "
        "but was missing on Day 7."
    )

    document.add_heading("2. Tray-level performance", level=1)

    best = best_row_by_metric(tray_metrics, "overall_adjusted_rgb_score")

    if best is not None:
        document.add_paragraph(
            f"The highest adjusted RGB performance score was recorded by "
            f"{best['tray']} ({best['microbe_status']} | {best['treatment']} | "
            f"{best['environment_group']}). Its adjusted RGB score was "
            f"{format_number(best['overall_adjusted_rgb_score'])}."
        )

    tray_columns = [
        "tray",
        "microbe_status",
        "treatment",
        "environment_group",
        "final_tracked_emergence_percent",
        "final_adjusted_green_cover_percent",
        "day7_imputed_cells",
        "overall_adjusted_rgb_score",
        "overall_adjusted_rgb_rank",
    ]

    add_docx_table(
        document,
        tray_metrics.sort_values("overall_adjusted_rgb_rank", na_position="last"),
        tray_columns,
        max_rows=12,
    )

    add_picture_if_exists(document, charts.get("tray_ranking"))

    document.add_heading("3. Microbes vs No Microbes", level=1)

    document.add_paragraph(
        group_summary_sentence(
            group_growth,
            "Microbe Status",
            "mean_final_adjusted_green_cover_percent",
            "final adjusted RGB green-cover",
        )
    )

    document.add_paragraph(
        group_summary_sentence(
            group_growth,
            "Microbe Status",
            "mean_overall_adjusted_rgb_score",
            "overall adjusted RGB performance score",
        )
    )

    microbe_group = group_growth.loc[group_growth["group_type"].eq("Microbe Status")].copy()

    microbe_columns = [
        "group",
        "tray_count",
        "mean_final_tracked_emergence_percent",
        "mean_final_adjusted_green_cover_percent",
        "mean_adjusted_green_cover_rate_day1_to_final_pp_per_day",
        "mean_overall_adjusted_rgb_score",
    ]

    add_docx_table(document, microbe_group, microbe_columns, max_rows=5)
    add_picture_if_exists(document, charts.get("microbes_green"))

    document.add_paragraph(
        "The Microbes vs No Microbes trend chart shows how adjusted RGB green-cover "
        "changed across observation days for both groups."
    )

    document.add_heading("4. Ideal vs Heat vs Moisture", level=1)

    document.add_paragraph(
        group_summary_sentence(
            group_growth,
            "Treatment Type",
            "mean_final_adjusted_green_cover_percent",
            "final adjusted RGB green-cover",
        )
    )

    document.add_paragraph(
        group_summary_sentence(
            group_growth,
            "Treatment Type",
            "mean_overall_adjusted_rgb_score",
            "overall adjusted RGB performance score",
        )
    )

    treatment_group = group_growth.loc[group_growth["group_type"].eq("Treatment Type")].copy()

    treatment_columns = [
        "group",
        "tray_count",
        "mean_final_tracked_emergence_percent",
        "mean_final_adjusted_green_cover_percent",
        "mean_adjusted_green_cover_rate_day1_to_final_pp_per_day",
        "mean_overall_adjusted_rgb_score",
    ]

    add_docx_table(document, treatment_group, treatment_columns, max_rows=10)
    add_picture_if_exists(document, charts.get("treatment_green"))

    document.add_heading("5. Ideal and Moisture Inside vs Outside comparison", level=1)

    document.add_paragraph(
        "This section directly compares fixed Inside and fixed Outside trays for "
        "Ideal and Moisture treatments. Heat trays are excluded because they were "
        "moved between environments during the trial."
    )

    inside_columns = [
        "treatment",
        "inside_tray_count",
        "outside_tray_count",
        "inside_trays",
        "outside_trays",
        "inside_final_emergence_percent",
        "outside_final_emergence_percent",
        "inside_minus_outside_final_emergence_pp",
        "inside_final_adjusted_green_cover_percent",
        "outside_final_adjusted_green_cover_percent",
        "inside_minus_outside_adjusted_green_cover_pp",
        "inside_minus_outside_adjusted_growth_rate_pp_per_day",
        "interpretation",
    ]

    add_docx_table(document, inside_outside_comparison, inside_columns, max_rows=5)

    for _, row in inside_outside_comparison.iterrows():
        document.add_paragraph(
            f"For {row['treatment']} trays, the inside-minus-outside difference "
            f"in adjusted final RGB green-cover was "
            f"{format_number(row.get('inside_minus_outside_adjusted_green_cover_pp'))} "
            f"percentage points. Interpretation: {row.get('interpretation', '')}"
        )

    add_picture_if_exists(document, charts.get("inside_outside_green"))
    add_picture_if_exists(document, charts.get("inside_outside_emergence"))

    document.add_paragraph(
        "The Inside vs Outside green-cover chart compares final adjusted green-cover. "
        "The emergence chart compares final tracked visible emergence. These are "
        "different because emergence measures how many cells had visible seedlings, "
        "while green-cover measures the amount of green plant area within cell zones."
    )

    document.add_heading("6. Observed vs adjusted Day 7", level=1)

    document.add_paragraph(
        "Observed Day 7 results represent what was visible in the final images. "
        "Adjusted Day 7 values are a separate scenario created for growth-rate "
        "comparison when a cell had earlier crop evidence but no visible green "
        "evidence on Day 7."
    )

    add_picture_if_exists(document, charts.get("day7_observed_adjusted"))
    add_picture_if_exists(document, charts.get("bug_cells"))

    document.add_heading("7. Description of generated CSV files", level=1)

    describe_output_file(
        document,
        "cell_growth_with_day7_adjustment.csv",
        "Full cell-level dataset after Day 7 adjustment processing. It preserves observed values and adds adjusted green-cover columns and imputation flags.",
    )

    describe_output_file(
        document,
        "tray_daily_rgb_metrics.csv",
        "Tray-by-day summary containing daily visible emergence, observed green-cover, adjusted green-cover, newly emerged cells, and Day 7 imputed cell counts.",
    )

    describe_output_file(
        document,
        "tray_growth_metrics.csv",
        "Tray-level growth summary containing final emergence, final observed and adjusted green-cover, growth rates, T50 timing, and RGB performance scores.",
    )

    describe_output_file(
        document,
        "group_daily_metrics.csv",
        "Group-by-day table used for trend charts across Microbes, treatments, and environment combinations.",
    )

    describe_output_file(
        document,
        "group_growth_metrics.csv",
        "Final group-level comparison table for treatment and environment groups.",
    )

    describe_output_file(
        document,
        "inside_outside_comparison_ideal_moisture.csv",
        "Dedicated Ideal and Moisture Inside vs Outside comparison table.",
    )

    describe_output_file(
        document,
        "environment_metadata_debug.csv",
        "Debug table showing each tray's standardised treatment and fixed environment assignment.",
    )

    describe_output_file(
        document,
        "possible_day7_bug_eaten_cells.csv",
        "Filtered table of cells where adjusted Day 7 values were estimated.",
    )

    describe_output_file(
        document,
        "heat_phase_response.csv",
        "Descriptive summary of Heat tray performance across movement phases.",
    )

    describe_output_file(
        document,
        "moisture_phase_response.csv",
        "Descriptive summary of Moisture tray performance across watering and dry phases.",
    )

    document.add_heading("8. Description of Excel workbook", level=1)

    document.add_paragraph(
        "The Excel workbook rgb_growth_treatment_report.xlsx combines the main CSV "
        "outputs into one inspection file. The most important sheets are Tray Growth "
        "Metrics, Group Growth Metrics, Inside Outside Compare, Day7 Imputed Cells, "
        "and Environment Debug."
    )

    document.add_heading("9. Description of generated charts", level=1)

    chart_descriptions = [
        ("01_microbes_visible_emergence_trend.png", "Tracked visible emergence trend for Microbes and No Microbes."),
        ("02_microbes_adjusted_green_cover_trend.png", "Adjusted RGB green-cover trend for Microbes and No Microbes."),
        ("03_treatment_visible_emergence_trend.png", "Tracked visible emergence trend for Ideal, Heat, and Moisture."),
        ("04_treatment_adjusted_green_cover_trend.png", "Adjusted RGB green-cover trend for Ideal, Heat, and Moisture."),
        ("05_day7_observed_vs_adjusted_green_cover_by_tray.png", "Observed vs adjusted final green-cover for each tray."),
        ("06_adjusted_rgb_tray_ranking.png", "Tray ranking by adjusted RGB performance score."),
        ("07_day7_possible_bug_eaten_cells_by_tray.png", "Number of possible Day 7 missing or bug-eaten cells by tray."),
        ("08_inside_outside_adjusted_green_cover_ideal_moisture.png", "Inside vs Outside adjusted final green-cover for Ideal and Moisture trays."),
        ("09_inside_outside_emergence_ideal_moisture.png", "Inside vs Outside final visible emergence for Ideal and Moisture trays."),
    ]

    for filename, description in chart_descriptions:
        describe_output_file(document, filename, description)

    document.add_heading("10. Interpretation and limitations", level=1)

    document.add_paragraph(
        "These results are suitable for identifying treatment trends, but they are "
        "descriptive image-derived results, not formal statistical proof. Lighting, "
        "image quality, green-pixel thresholding, crop overlap, and cell visibility "
        "can affect RGB green-cover estimates."
    )

    document.add_paragraph(
        "The Day 7 adjustment should be used only for adjusted growth-rate comparison. "
        "Observed Day 7 values should remain the source for reporting what was actually "
        "visible in the images."
    )

    document.add_paragraph(
        "The Inside vs Outside comparison is restricted to Ideal and Moisture trays. "
        "Heat trays are excluded because they were moved between environments. Later "
        "multispectral scripts should be used to check whether NDVI and NDRE support "
        "the RGB findings."
    )

    document.save(output_path)

    return output_path


# ============================================================
# 15) PDF REPORT
# ============================================================

def make_pdf_table(dataframe: pd.DataFrame, columns: list[str], max_rows: int = 20):
    if not REPORTLAB_AVAILABLE:
        return None

    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        "HeaderCell",
        parent=styles["BodyText"],
        fontSize=7,
        leading=8,
        textColor=colors.white,
        alignment=1,
        wordWrap="CJK",
    )

    cell_style = ParagraphStyle(
        "BodyCell",
        parent=styles["BodyText"],
        fontSize=7,
        leading=8,
        wordWrap="CJK",
    )

    columns = [column for column in columns if column in dataframe.columns]

    if not columns:
        return Paragraph("No valid table columns available.", styles["BodyText"])

    frame = dataframe[columns].head(max_rows).copy()
    frame = safe_round_dataframe(frame, 3)

    data = [[Paragraph(str(column), header_style) for column in columns]]

    for _, row in frame.iterrows():
        data.append([Paragraph(str(row[column]), cell_style) for column in columns])

    table = Table(data, repeatRows=1)

    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
            ]
        )
    )

    return table


def create_pdf_report(
    output_path: Path,
    tray_metrics: pd.DataFrame,
    group_growth: pd.DataFrame,
    inside_outside_comparison: pd.DataFrame,
    charts: dict[str, Path],
):
    if not REPORTLAB_AVAILABLE:
        print("WARNING: reportlab is not installed. PDF report was skipped.")
        return None

    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=landscape(A4),
        rightMargin=0.35 * inch,
        leftMargin=0.35 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )

    styles = getSampleStyleSheet()
    story = []

    story.append(
        Paragraph(
            "Trial 3 RGB Growth-Rate and Treatment Comparison Report",
            styles["Title"],
        )
    )

    story.append(Spacer(1, 10))

    story.append(
        Paragraph(
            "This report summarises Trial 3 visible-emergence and RGB green-cover "
            "growth results. Observed Day 7 values are preserved, while adjusted "
            "Day 7 values are reported separately for likely missing or bug-eaten seedlings.",
            styles["BodyText"],
        )
    )

    story.append(Spacer(1, 12))

    top_columns = [
        "tray",
        "microbe_status",
        "treatment",
        "environment_group",
        "final_tracked_emergence_percent",
        "final_observed_green_cover_percent",
        "final_adjusted_green_cover_percent",
        "day7_imputed_cells",
        "overall_adjusted_rgb_score",
        "overall_adjusted_rgb_rank",
    ]

    story.append(Paragraph("Tray-level adjusted RGB ranking", styles["Heading2"]))
    story.append(
        make_pdf_table(
            tray_metrics.sort_values("overall_adjusted_rgb_rank", na_position="last"),
            top_columns,
            max_rows=12,
        )
    )

    story.append(PageBreak())

    story.append(Paragraph("Ideal and Moisture Inside vs Outside comparison", styles["Heading2"]))

    inside_columns = [
        "treatment",
        "inside_tray_count",
        "outside_tray_count",
        "inside_final_emergence_percent",
        "outside_final_emergence_percent",
        "inside_minus_outside_final_emergence_pp",
        "inside_final_adjusted_green_cover_percent",
        "outside_final_adjusted_green_cover_percent",
        "inside_minus_outside_adjusted_green_cover_pp",
        "interpretation",
    ]

    story.append(
        make_pdf_table(
            inside_outside_comparison,
            inside_columns,
            max_rows=5,
        )
    )

    figure_items = [
        ("Microbes adjusted green-cover trend", charts.get("microbes_green")),
        ("Treatment adjusted green-cover trend", charts.get("treatment_green")),
        ("Inside vs Outside adjusted green-cover", charts.get("inside_outside_green")),
        ("Inside vs Outside visible emergence", charts.get("inside_outside_emergence")),
        ("Day 7 observed vs adjusted green-cover", charts.get("day7_observed_adjusted")),
        ("Adjusted RGB tray ranking", charts.get("tray_ranking")),
        ("Possible Day 7 bug-eaten cells", charts.get("bug_cells")),
    ]

    for title, path in figure_items:
        if path is None or not Path(path).exists():
            continue

        story.append(PageBreak())
        story.append(Paragraph(title, styles["Heading2"]))
        story.append(Spacer(1, 6))
        story.append(PDFImage(str(path), width=10.5 * inch, height=5.9 * inch))

    document.build(story)

    return output_path


# ============================================================
# 16) SETTINGS
# ============================================================

def save_settings(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)

    settings = {
        "purpose": "Trial 3 RGB growth-rate and treatment comparison",
        "script04_tray_summary_csv": str(TRAY_SUMMARY_CSV),
        "script04_cell_results_csv": str(CELL_RESULTS_CSV),
        "output_root": str(OUTPUT_ROOT),
        "expected_cells_per_tray": EXPECTED_CELLS,
        "expected_observation_days": EXPECTED_OBSERVATION_DAYS,
        "corrected_day1_photo_date": "2026-06-29",
        "day7_photo_date": "2026-07-07",
        "inside_outside_comparison_fix": (
            "Inside/Outside comparison uses tray-number fallback metadata from Trial 3 tray design. "
            "This prevents In/Out, Inside/Outside, blanks, or inconsistent labels from producing blank charts."
        ),
        "inside_outside_comparison_policy": (
            "Direct Inside vs Outside comparison is only performed for Ideal and Moisture trays. "
            "Heat trays are excluded because their environment changed during the trial."
        ),
        "observed_day7_policy": "Observed Day 7 values from Script 04 are preserved exactly.",
        "adjusted_day7_policy": (
            "If a cell had prior green evidence but no current green evidence on Day 7, "
            "an adjusted Day 7 green-cover value is estimated from the most recent prior "
            "green-cover value and previous positive growth slope."
        ),
        "word_report": "Expanded Word report is generated for this comparison script.",
        "statistical_warning": (
            "This script produces descriptive image-derived comparisons only; "
            "it is not a formal statistical test."
        ),
    }

    with path.open("w", encoding="utf-8") as file:
        json.dump(settings, file, indent=2)

    return path


# ============================================================
# 17) MAIN
# ============================================================

def run_analysis(args) -> int:
    print("\nSCRIPT 05 — THIRD TRIAL RGB GROWTH-RATE AND TREATMENT COMPARISON")
    print("=" * 78)
    print(f"Script 04 reports folder:\n{SCRIPT04_REPORTS}")
    print(f"\nOutput folder:\n{OUTPUT_ROOT}")

    tray_summary, cell_results, first_summary = load_script04_outputs()

    print(f"\nLoaded Script 04 PASS tray rows: {len(tray_summary)}")
    print(f"Loaded Script 04 PASS cell rows: {len(cell_results)}")

    day_counts = (
        tray_summary.groupby("day_order", dropna=False)["tray_no"]
        .nunique()
        .reset_index(name="tray_count")
    )

    print("\nTray count by observation day:")
    print(day_counts.to_string(index=False))

    missing_days = set(EXPECTED_OBSERVATION_DAYS) - set(
        tray_summary["day_order"].astype(int).unique()
    )

    if missing_days:
        print(
            "\nWARNING: Missing expected observation days: "
            + ", ".join(f"Day {day}" for day in sorted(missing_days))
        )

    adjusted_cell = create_adjusted_cell_table(cell_results)

    if adjusted_cell.empty:
        raise RuntimeError("Adjusted cell table is empty. Check Script 04 cell results.")

    tray_daily = create_tray_daily_metrics(adjusted_cell)
    tray_metrics = create_tray_growth_metrics(tray_daily, adjusted_cell)
    group_daily, group_growth = create_all_group_tables(tray_daily, tray_metrics)
    heat_phase, moisture_phase = create_phase_tables(tray_daily)
    inside_outside_comparison = create_inside_outside_comparison(tray_metrics)

    print("\nEnvironment metadata debug:")
    print(
        tray_metrics[
            [
                "tray_no",
                "tray",
                "microbe_status",
                "treatment",
                "fixed_environment",
                "final_tracked_emergence_percent",
                "final_adjusted_green_cover_percent",
            ]
        ].sort_values("tray_no").to_string(index=False)
    )

    print("\nInside vs Outside comparison check:")
    print(
        inside_outside_comparison[
            [
                "treatment",
                "inside_tray_count",
                "outside_tray_count",
                "inside_trays",
                "outside_trays",
                "inside_final_adjusted_green_cover_percent",
                "outside_final_adjusted_green_cover_percent",
                "inside_minus_outside_adjusted_green_cover_pp",
            ]
        ].to_string(index=False)
    )

    charts = create_charts(
        group_daily,
        tray_metrics,
        inside_outside_comparison,
    )

    table_paths = save_tables(
        adjusted_cell,
        tray_daily,
        tray_metrics,
        group_daily,
        group_growth,
        heat_phase,
        moisture_phase,
        inside_outside_comparison,
    )

    pdf_path = REPORTS_ROOT / "rgb_growth_treatment_summary.pdf"

    create_pdf_report(
        pdf_path,
        tray_metrics,
        group_growth,
        inside_outside_comparison,
        charts,
    )

    docx_path = REPORTS_ROOT / "rgb_growth_treatment_short_report.docx"

    create_word_report(
        docx_path,
        tray_metrics,
        group_growth,
        inside_outside_comparison,
        adjusted_cell,
        charts,
    )

    settings_path = save_settings(
        CONFIG_ROOT / "rgb_growth_treatment_settings.json"
    )

    print("\n" + "=" * 78)
    print("SCRIPT 05 FINISHED")
    print("=" * 78)

    day7_imputed_cells = int(adjusted_cell["day7_imputed"].eq("Yes").sum())

    print(f"Day 7 imputed cell records: {day7_imputed_cells}")

    if not tray_metrics.empty:
        best = tray_metrics.sort_values(
            "overall_adjusted_rgb_rank",
            na_position="last",
        ).iloc[0]

        print(
            "\nTop adjusted RGB tray:"
            f"\n  {best['tray']} — {best['microbe_status']} | "
            f"{best['treatment']} | {best['environment_group']}"
            f"\n  Adjusted RGB score: {best['overall_adjusted_rgb_score']:.2f}"
        )

    print("\nMain output files:")
    print(f"Adjusted cell table:\n{table_paths['adjusted_cell']}")
    print(f"\nTray daily metrics:\n{table_paths['tray_daily']}")
    print(f"\nTray growth metrics:\n{table_paths['tray_metrics']}")
    print(f"\nGroup daily metrics:\n{table_paths['group_daily']}")
    print(f"\nGroup growth metrics:\n{table_paths['group_growth']}")
    print(f"\nInside vs Outside comparison:\n{table_paths['inside_outside']}")
    print(f"\nEnvironment debug:\n{table_paths['environment_debug']}")
    print(f"\nDay 7 imputed cells:\n{table_paths['day7_imputed']}")
    print(f"\nExcel report:\n{table_paths['excel']}")

    if REPORTLAB_AVAILABLE:
        print(f"\nPDF report:\n{pdf_path}")
    else:
        print("\nPDF report skipped because reportlab is not installed.")

    if DOCX_AVAILABLE:
        print(f"\nWord report:\n{docx_path}")
    else:
        print("\nWord report skipped because python-docx is not installed.")

    print(f"\nSettings:\n{settings_path}")
    print(f"\nCharts folder:\n{CHARTS_ROOT}")

    print("\nImportant charts to re-check:")
    print(CHARTS_ROOT / "08_inside_outside_adjusted_green_cover_ideal_moisture.png")
    print(CHARTS_ROOT / "09_inside_outside_emergence_ideal_moisture.png")

    return 0


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Trial 3 Script 05: RGB growth-rate and treatment comparison."
    )

    args = parser.parse_args()

    return run_analysis(args)


if __name__ == "__main__":
    raise SystemExit(main())