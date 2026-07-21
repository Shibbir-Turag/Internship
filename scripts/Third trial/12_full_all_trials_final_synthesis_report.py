from __future__ import annotations

"""
SCRIPT 12 — FULL FIRST + SECOND + THIRD TRIAL FINAL SYNTHESIS REPORT

Purpose
-------
Create one final overall synthesis report across the internship germination trials.

This script does not recalculate image-processing outputs. It reads the final
tables, reports and packages already produced by the First, Second and Third
Trial workflows, then writes a final Word report.

Main output
-----------
outputs/All Trials Final Synthesis/
    _reports/
        all_trials_final_synthesis_report.docx
        all_trials_key_findings.csv
        all_trials_input_inventory.csv
        all_trials_synthesis_workbook.xlsx
    _config/
        all_trials_synthesis_settings.json

The report covers:
1. Overall internship workflow
2. First Trial role as pilot / baseline development stage
3. Second Trial RGB + MS workflow and treatment findings
4. Third Trial RGB + MS workflow and treatment findings
5. Cross-trial methodological development
6. Treatment-level interpretation across trials
7. Observed vs adjusted Day 7 handling
8. Limitations
9. Final output guide

Notes
-----
The script is intentionally defensive. If a First or Second Trial file has a
different name, it tries to locate the closest matching CSV/DOCX/XLSX files
using folder scanning. Missing inputs are recorded in the inventory instead of
crashing the whole report.
"""

import json
import math
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

OUTPUTS_ROOT = PROJECT_ROOT / "outputs"

FIRST_OUTPUT_CANDIDATES = [
    OUTPUTS_ROOT / "First Trial",
    OUTPUTS_ROOT / "First Trial (Two Trays)",
    OUTPUTS_ROOT / "First trial",
    OUTPUTS_ROOT / "First Trial Two Trays",
]

SECOND_OUTPUT_ROOT = OUTPUTS_ROOT / "Second Trial"

THIRD_OUTPUT_ROOT = OUTPUTS_ROOT / "Third trial"

FINAL_OUTPUT_ROOT = OUTPUTS_ROOT / "All Trials Final Synthesis"

REPORTS_ROOT = FINAL_OUTPUT_ROOT / "_reports"

CONFIG_ROOT = FINAL_OUTPUT_ROOT / "_config"

CHARTS_USED_ROOT = FINAL_OUTPUT_ROOT / "charts_used"

WORD_REPORT = REPORTS_ROOT / "all_trials_final_synthesis_report.docx"

KEY_FINDINGS_CSV = REPORTS_ROOT / "all_trials_key_findings.csv"

INPUT_INVENTORY_CSV = REPORTS_ROOT / "all_trials_input_inventory.csv"

SYNTHESIS_WORKBOOK = REPORTS_ROOT / "all_trials_synthesis_workbook.xlsx"

SETTINGS_JSON = CONFIG_ROOT / "all_trials_synthesis_settings.json"


# ============================================================
# 2) OPTIONAL WORD SUPPORT
# ============================================================

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True

except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 3) TRIAL-SPECIFIC LIKELY FILES
# ============================================================

THIRD_FILES = {
    "combined_tray": THIRD_OUTPUT_ROOT
    / "09_Final_RGB_MS_Synthesis"
    / "_reports"
    / "final_combined_tray_synthesis.csv",

    "group_synthesis": THIRD_OUTPUT_ROOT
    / "09_Final_RGB_MS_Synthesis"
    / "_reports"
    / "final_group_synthesis.csv",

    "inside_outside": THIRD_OUTPUT_ROOT
    / "09_Final_RGB_MS_Synthesis"
    / "_reports"
    / "final_inside_outside_synthesis.csv",

    "summary_report": THIRD_OUTPUT_ROOT
    / "11_Third_Trial_Short_Final_Summary_Report"
    / "_reports"
    / "third_trial_short_final_summary_report.docx",

    "package_zip": THIRD_OUTPUT_ROOT
    / "10_Final_Figure_Report_Package"
    / "Third_Trial_Final_Figure_Report_Package.zip",
}

SECOND_FILE_CANDIDATES = {
    "combined_tray": [
        SECOND_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis" / "_reports" / "final_combined_tray_synthesis.csv",
        SECOND_OUTPUT_ROOT / "09_Final_Second_Trial_Synthesis" / "_reports" / "final_combined_tray_synthesis.csv",
        SECOND_OUTPUT_ROOT / "09_Final_Synthesis" / "_reports" / "final_combined_tray_synthesis.csv",
        SECOND_OUTPUT_ROOT / "09_Final_Second_Trial_Synthesis" / "_reports" / "second_trial_combined_tray_synthesis.csv",
        SECOND_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis" / "_reports" / "second_trial_combined_tray_synthesis.csv",
    ],

    "group_synthesis": [
        SECOND_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis" / "_reports" / "final_group_synthesis.csv",
        SECOND_OUTPUT_ROOT / "09_Final_Second_Trial_Synthesis" / "_reports" / "final_group_synthesis.csv",
        SECOND_OUTPUT_ROOT / "08_MS_Treatment_Comparison" / "_reports" / "ms_group_growth_metrics.csv",
        SECOND_OUTPUT_ROOT / "05_RGB_Growth_Rate_Treatment_Comparison" / "_reports" / "group_growth_metrics.csv",
        SECOND_OUTPUT_ROOT / "05_Growth_Rate_Treatment_Comparison" / "_reports" / "group_growth_metrics.csv",
    ],

    "inside_outside": [
        SECOND_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis" / "_reports" / "final_inside_outside_synthesis.csv",
        SECOND_OUTPUT_ROOT / "08_MS_Treatment_Comparison" / "_reports" / "inside_outside_comparison_ideal_moisture.csv",
        SECOND_OUTPUT_ROOT / "05_RGB_Growth_Rate_Treatment_Comparison" / "_reports" / "inside_outside_comparison_ideal_moisture.csv",
    ],

    "summary_report": [
        SECOND_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis" / "_reports" / "final_rgb_ms_synthesis_report.docx",
        SECOND_OUTPUT_ROOT / "09_Final_Second_Trial_Synthesis" / "_reports" / "final_second_trial_synthesis_report.docx",
        SECOND_OUTPUT_ROOT / "10_Report_Figure_Package" / "_reports" / "second_trial_final_package_handover.docx",
    ],

    "package_zip": [
        SECOND_OUTPUT_ROOT / "10_Report_Figure_Package" / "Second_Trial_Final_Figure_Report_Package.zip",
        SECOND_OUTPUT_ROOT / "10_Final_Figure_Report_Package" / "Second_Trial_Final_Figure_Report_Package.zip",
    ],
}

FIRST_FILE_NAME_HINTS = {
    "combined_tray": [
        "final",
        "synthesis",
        "tray",
        ".csv",
    ],
    "group_synthesis": [
        "group",
        "summary",
        ".csv",
    ],
    "summary_report": [
        "report",
        ".docx",
    ],
}


# ============================================================
# 4) COLUMN CANDIDATES
# ============================================================

TRAY_SCORE_COLUMNS = [
    "combined_rgb_ms_score",
    "overall_adjusted_ms_score",
    "overall_adjusted_rgb_score",
    "overall_ms_score",
    "overall_rgb_score",
    "overall_adjusted_growth_score",
    "overall_score",
    "adjusted_ms_score",
    "adjusted_rgb_score",
    "score",
]

TRAY_RANK_COLUMNS = [
    "combined_rank",
    "overall_adjusted_ms_rank",
    "overall_adjusted_rgb_rank",
    "rank",
]

GROUP_SCORE_COLUMNS = [
    "mean_combined_rgb_ms_score",
    "mean_overall_adjusted_ms_score",
    "mean_overall_adjusted_rgb_score",
    "mean_overall_score",
    "mean_adjusted_ms_score",
    "mean_adjusted_rgb_score",
    "mean_score",
]


# ============================================================
# 5) DATA CLASSES
# ============================================================

@dataclass
class TrialBundle:
    name: str
    output_root: Path | None
    combined_tray_path: Path | None
    group_synthesis_path: Path | None
    inside_outside_path: Path | None
    summary_report_path: Path | None
    package_zip_path: Path | None
    combined_tray: pd.DataFrame
    group_synthesis: pd.DataFrame
    inside_outside: pd.DataFrame
    notes: list[str]


# ============================================================
# 6) GENERAL HELPERS
# ============================================================

def ensure_output_folders() -> None:
    for folder in [
        FINAL_OUTPUT_ROOT,
        REPORTS_ROOT,
        CONFIG_ROOT,
        CHARTS_USED_ROOT,
    ]:
        folder.mkdir(
            parents=True,
            exist_ok=True,
        )


def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(
            r"(\d+)",
            str(value),
        )
    ]


def safe_float(value: object, default: float = math.nan) -> float:
    try:
        if pd.isna(value):
            return default

        return float(value)

    except Exception:
        return default


def format_number(value: object, decimals: int = 2) -> str:
    value = safe_float(value)

    if not math.isfinite(value):
        return "N/A"

    return f"{value:.{decimals}f}"


def clean_text(value: object, default: str = "N/A") -> str:
    text = str(value).strip()

    if text.casefold() in {
        "",
        "nan",
        "none",
        "unknown",
    }:
        return default

    return text


def first_existing(paths: list[Path]) -> Path | None:
    for path in paths:
        if path.exists():
            return path

    return None


def read_csv_optional(path: Path | None) -> pd.DataFrame:
    if path is None or not path.exists():
        return pd.DataFrame()

    try:
        return pd.read_csv(path)

    except Exception:
        return pd.DataFrame()


def pick_column(frame: pd.DataFrame, candidates: list[str]) -> str | None:
    for column in candidates:
        if column in frame.columns:
            return column

    return None


def numeric_series(frame: pd.DataFrame, column: str | None) -> pd.Series:
    if column is None or column not in frame.columns:
        return pd.Series(
            [math.nan] * len(frame),
            index=frame.index,
            dtype=float,
        )

    return pd.to_numeric(
        frame[column],
        errors="coerce",
    )


def derive_tray_no(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()

    if "tray_no" not in output.columns:
        if "tray" in output.columns:
            output["tray_no"] = (
                output["tray"]
                .astype(str)
                .str.extract(r"(\d+)")
                .iloc[:, 0]
            )

    if "tray_no" in output.columns:
        output["tray_no"] = pd.to_numeric(
            output["tray_no"],
            errors="coerce",
        )

    if "tray" not in output.columns and "tray_no" in output.columns:
        output["tray"] = output["tray_no"].apply(
            lambda value: f"Tray {int(value)}" if pd.notna(value) else "Tray"
        )

    return output


def best_row_by_score(frame: pd.DataFrame) -> pd.Series | None:
    if frame.empty:
        return None

    score_column = pick_column(
        frame,
        TRAY_SCORE_COLUMNS,
    )

    if score_column is None:
        return None

    temp = frame.copy()

    temp["_score_numeric"] = pd.to_numeric(
        temp[score_column],
        errors="coerce",
    )

    temp = temp.dropna(
        subset=["_score_numeric"]
    )

    if temp.empty:
        return None

    return temp.loc[
        temp["_score_numeric"].idxmax()
    ]


def best_group_by_score(frame: pd.DataFrame, group_type: str | None = None) -> pd.Series | None:
    if frame.empty:
        return None

    temp = frame.copy()

    if group_type and "group_type" in temp.columns:
        temp = temp.loc[
            temp["group_type"].astype(str).eq(group_type)
        ].copy()

    if temp.empty:
        return None

    score_column = pick_column(
        temp,
        GROUP_SCORE_COLUMNS,
    )

    if score_column is None:
        return None

    temp["_score_numeric"] = pd.to_numeric(
        temp[score_column],
        errors="coerce",
    )

    temp = temp.dropna(
        subset=["_score_numeric"]
    )

    if temp.empty:
        return None

    return temp.loc[
        temp["_score_numeric"].idxmax()
    ]


def safe_copy_chart(path: Path | None) -> Path | None:
    if path is None or not path.exists():
        return None

    destination = CHARTS_USED_ROOT / path.name

    counter = 2

    while destination.exists():
        destination = CHARTS_USED_ROOT / f"{destination.stem}_{counter}{destination.suffix}"
        counter += 1

    shutil.copy2(
        path,
        destination,
    )

    return destination


def find_by_hints(root: Path | None, hints: list[str], suffixes: set[str]) -> Path | None:
    if root is None or not root.exists():
        return None

    matches = []

    for file in root.rglob("*"):
        if not file.is_file():
            continue

        name = file.name.casefold()

        if file.suffix.casefold() not in suffixes:
            continue

        if all(hint.casefold() in name for hint in hints if hint != ".csv" and hint != ".docx"):
            matches.append(file)

    if not matches:
        return None

    matches = sorted(
        matches,
        key=lambda path: (
            0 if "_reports" in str(path).casefold() else 1,
            len(str(path)),
            str(path).casefold(),
        ),
    )

    return matches[0]


def find_first_trial_root() -> Path | None:
    for root in FIRST_OUTPUT_CANDIDATES:
        if root.exists():
            return root

    possible = [
        folder
        for folder in OUTPUTS_ROOT.glob("*")
        if folder.is_dir()
        and "first" in folder.name.casefold()
        and "trial" in folder.name.casefold()
    ]

    if possible:
        return sorted(
            possible,
            key=lambda path: str(path).casefold(),
        )[0]

    return None


def find_generic_final_files(root: Path | None) -> dict[str, Path | None]:
    if root is None or not root.exists():
        return {
            "combined_tray": None,
            "group_synthesis": None,
            "inside_outside": None,
            "summary_report": None,
            "package_zip": None,
        }

    csv_files = sorted(
        root.rglob("*.csv"),
        key=lambda path: str(path).casefold(),
    )

    docx_files = sorted(
        root.rglob("*.docx"),
        key=lambda path: str(path).casefold(),
    )

    zip_files = sorted(
        root.rglob("*.zip"),
        key=lambda path: str(path).casefold(),
    )

    combined = None
    group = None
    inside_outside = None

    for file in csv_files:
        name = file.name.casefold()

        if combined is None and (
            (
                "combined" in name
                and "tray" in name
            )
            or (
                "synthesis" in name
                and "tray" in name
            )
            or (
                "tray_growth" in name
            )
            or (
                "tray" in name
                and "metric" in name
            )
        ):
            combined = file

        if group is None and (
            "group" in name
            and (
                "synthesis" in name
                or "growth" in name
                or "metric" in name
                or "summary" in name
            )
        ):
            group = file

        if inside_outside is None and (
            "inside" in name
            and "outside" in name
        ):
            inside_outside = file

    summary_report = None

    for file in docx_files:
        name = file.name.casefold()

        if "report" in name or "summary" in name or "synthesis" in name:
            summary_report = file
            break

    package_zip = zip_files[0] if zip_files else None

    return {
        "combined_tray": combined,
        "group_synthesis": group,
        "inside_outside": inside_outside,
        "summary_report": summary_report,
        "package_zip": package_zip,
    }


# ============================================================
# 7) LOAD TRIAL BUNDLES
# ============================================================

def load_first_trial_bundle() -> TrialBundle:
    root = find_first_trial_root()

    found = find_generic_final_files(root)

    combined = read_csv_optional(
        found["combined_tray"]
    )

    group = read_csv_optional(
        found["group_synthesis"]
    )

    inside_outside = read_csv_optional(
        found["inside_outside"]
    )

    notes = []

    if root is None:
        notes.append(
            "First Trial output folder was not found automatically."
        )

    if combined.empty:
        notes.append(
            "First Trial tray-level synthesis table was not found. The report treats the First Trial mainly as the pilot workflow stage."
        )

    return TrialBundle(
        name="First Trial",
        output_root=root,
        combined_tray_path=found["combined_tray"],
        group_synthesis_path=found["group_synthesis"],
        inside_outside_path=found["inside_outside"],
        summary_report_path=found["summary_report"],
        package_zip_path=found["package_zip"],
        combined_tray=derive_tray_no(combined) if not combined.empty else combined,
        group_synthesis=group,
        inside_outside=inside_outside,
        notes=notes,
    )


def load_second_trial_bundle() -> TrialBundle:
    generic = find_generic_final_files(
        SECOND_OUTPUT_ROOT
    )

    combined_path = first_existing(
        SECOND_FILE_CANDIDATES["combined_tray"]
    ) or generic["combined_tray"]

    group_path = first_existing(
        SECOND_FILE_CANDIDATES["group_synthesis"]
    ) or generic["group_synthesis"]

    inside_outside_path = first_existing(
        SECOND_FILE_CANDIDATES["inside_outside"]
    ) or generic["inside_outside"]

    summary_report_path = first_existing(
        SECOND_FILE_CANDIDATES["summary_report"]
    ) or generic["summary_report"]

    package_zip_path = first_existing(
        SECOND_FILE_CANDIDATES["package_zip"]
    ) or generic["package_zip"]

    combined = read_csv_optional(
        combined_path
    )

    group = read_csv_optional(
        group_path
    )

    inside_outside = read_csv_optional(
        inside_outside_path
    )

    notes = []

    if not SECOND_OUTPUT_ROOT.exists():
        notes.append(
            "Second Trial output folder was not found."
        )

    if combined.empty:
        notes.append(
            "A direct Second Trial final tray synthesis table was not found. The report still includes the confirmed workflow conclusion that the Second Trial matured the 7 × 10 grid workflow and showed strongest performance in the Microbes Inside trays."
        )

    return TrialBundle(
        name="Second Trial",
        output_root=SECOND_OUTPUT_ROOT if SECOND_OUTPUT_ROOT.exists() else None,
        combined_tray_path=combined_path,
        group_synthesis_path=group_path,
        inside_outside_path=inside_outside_path,
        summary_report_path=summary_report_path,
        package_zip_path=package_zip_path,
        combined_tray=derive_tray_no(combined) if not combined.empty else combined,
        group_synthesis=group,
        inside_outside=inside_outside,
        notes=notes,
    )


def load_third_trial_bundle() -> TrialBundle:
    combined = read_csv_optional(
        THIRD_FILES["combined_tray"]
    )

    group = read_csv_optional(
        THIRD_FILES["group_synthesis"]
    )

    inside_outside = read_csv_optional(
        THIRD_FILES["inside_outside"]
    )

    notes = []

    if combined.empty:
        notes.append(
            "Third Trial final combined tray synthesis table was not found. Run Script 09 before using this final report."
        )

    return TrialBundle(
        name="Third Trial",
        output_root=THIRD_OUTPUT_ROOT if THIRD_OUTPUT_ROOT.exists() else None,
        combined_tray_path=THIRD_FILES["combined_tray"] if THIRD_FILES["combined_tray"].exists() else None,
        group_synthesis_path=THIRD_FILES["group_synthesis"] if THIRD_FILES["group_synthesis"].exists() else None,
        inside_outside_path=THIRD_FILES["inside_outside"] if THIRD_FILES["inside_outside"].exists() else None,
        summary_report_path=THIRD_FILES["summary_report"] if THIRD_FILES["summary_report"].exists() else None,
        package_zip_path=THIRD_FILES["package_zip"] if THIRD_FILES["package_zip"].exists() else None,
        combined_tray=derive_tray_no(combined) if not combined.empty else combined,
        group_synthesis=group,
        inside_outside=inside_outside,
        notes=notes,
    )


# ============================================================
# 8) INVENTORY AND FINDINGS
# ============================================================

def create_input_inventory(bundles: list[TrialBundle]) -> pd.DataFrame:
    rows = []

    for bundle in bundles:
        file_items = [
            (
                "Output root",
                bundle.output_root,
                "Main output folder for the trial.",
            ),
            (
                "Combined tray synthesis",
                bundle.combined_tray_path,
                "Tray-level final or closest available synthesis table.",
            ),
            (
                "Group synthesis",
                bundle.group_synthesis_path,
                "Treatment/group-level comparison table.",
            ),
            (
                "Inside vs Outside comparison",
                bundle.inside_outside_path,
                "Fixed-environment comparison table where available.",
            ),
            (
                "Summary report",
                bundle.summary_report_path,
                "Existing Word report from the trial workflow.",
            ),
            (
                "Package ZIP",
                bundle.package_zip_path,
                "Final packaged outputs for the trial, if available.",
            ),
        ]

        for label, path, description in file_items:
            exists = bool(path is not None and Path(path).exists())

            rows.append(
                {
                    "trial": bundle.name,
                    "item": label,
                    "path": str(path) if path else "",
                    "exists": "Yes" if exists else "No",
                    "description": description,
                }
            )

    return pd.DataFrame(rows)


def create_key_findings(bundles: list[TrialBundle]) -> pd.DataFrame:
    rows = []

    for bundle in bundles:
        best_tray = best_row_by_score(
            bundle.combined_tray
        )

        if best_tray is not None:
            score_column = pick_column(
                bundle.combined_tray,
                TRAY_SCORE_COLUMNS,
            )

            rows.append(
                {
                    "trial": bundle.name,
                    "topic": "Top tray",
                    "finding": (
                        f"{clean_text(best_tray.get('tray', 'N/A'))} had the highest available tray-level score."
                    ),
                    "value": format_number(
                        best_tray.get(score_column, math.nan),
                        2,
                    ),
                    "source": Path(bundle.combined_tray_path).name if bundle.combined_tray_path else "",
                }
            )

        elif bundle.name == "Second Trial":
            rows.append(
                {
                    "trial": bundle.name,
                    "topic": "Top treatment pattern",
                    "finding": (
                        "The Second Trial workflow showed Microbes Inside as the strongest treatment pattern, with Tray 8 and Tray 5 identified as the leading trays in the final interpretation."
                    ),
                    "value": "Qualitative confirmed result",
                    "source": "Second Trial workflow notes / final outputs",
                }
            )

        elif bundle.name == "First Trial":
            rows.append(
                {
                    "trial": bundle.name,
                    "topic": "Workflow role",
                    "finding": (
                        "The First Trial functioned as the pilot stage for learning the tray-imaging workflow and identifying the need for repeatable image-processing steps."
                    ),
                    "value": "Pilot stage",
                    "source": "Internship workflow context",
                }
            )

        best_microbe = best_group_by_score(
            bundle.group_synthesis,
            "Microbe Status",
        )

        if best_microbe is not None:
            score_column = pick_column(
                bundle.group_synthesis,
                GROUP_SCORE_COLUMNS,
            )

            rows.append(
                {
                    "trial": bundle.name,
                    "topic": "Microbes vs No Microbes",
                    "finding": (
                        f"{clean_text(best_microbe.get('group', 'N/A'))} had the higher mean group-level score."
                    ),
                    "value": format_number(
                        best_microbe.get(score_column, math.nan),
                        2,
                    ),
                    "source": Path(bundle.group_synthesis_path).name if bundle.group_synthesis_path else "",
                }
            )

        best_treatment = best_group_by_score(
            bundle.group_synthesis,
            "Treatment Type",
        )

        if best_treatment is not None:
            score_column = pick_column(
                bundle.group_synthesis,
                GROUP_SCORE_COLUMNS,
            )

            rows.append(
                {
                    "trial": bundle.name,
                    "topic": "Treatment comparison",
                    "finding": (
                        f"{clean_text(best_treatment.get('group', 'N/A'))} had the highest available treatment-level score."
                    ),
                    "value": format_number(
                        best_treatment.get(score_column, math.nan),
                        2,
                    ),
                    "source": Path(bundle.group_synthesis_path).name if bundle.group_synthesis_path else "",
                }
            )

        if not bundle.inside_outside.empty:
            for row in bundle.inside_outside.itertuples(index=False):
                treatment = clean_text(
                    getattr(row, "treatment", "N/A")
                )

                interpretation = clean_text(
                    getattr(row, "interpretation", "N/A")
                )

                value = format_number(
                    getattr(row, "inside_minus_outside_combined_score", math.nan),
                    2,
                )

                if value == "N/A":
                    value = format_number(
                        getattr(row, "inside_minus_outside_day7_ndvi", math.nan),
                        3,
                    )

                rows.append(
                    {
                        "trial": bundle.name,
                        "topic": f"Inside vs Outside — {treatment}",
                        "finding": interpretation,
                        "value": value,
                        "source": Path(bundle.inside_outside_path).name if bundle.inside_outside_path else "",
                    }
                )

    rows.append(
        {
            "trial": "All Trials",
            "topic": "Methodological development",
            "finding": (
                "The workflow progressed from early trial learning to a stable 7 × 10 cell-based RGB and multispectral analysis pipeline with final reporting packages."
            ),
            "value": "Completed pipeline",
            "source": "Scripts 01–12",
        }
    )

    rows.append(
        {
            "trial": "All Trials",
            "topic": "Day 7 adjustment policy",
            "finding": (
                "Observed image values remain the primary evidence. Adjusted values are kept separately only for likely bug-eaten cells and should be described as scenario estimates."
            ),
            "value": "Observed and adjusted kept separate",
            "source": "Third Trial Scripts 05, 08, 09 and 11",
        }
    )

    return pd.DataFrame(rows)


# ============================================================
# 9) EXCEL OUTPUT
# ============================================================

def style_excel(path: Path) -> None:
    workbook = load_workbook(path)

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.row_dimensions[1].height = 34

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )
            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = column_cells[0].column_letter

            longest = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[letter].width = min(
                max(12, longest + 2),
                70,
            )

    workbook.save(path)


def write_tables(
    key_findings: pd.DataFrame,
    inventory: pd.DataFrame,
    bundles: list[TrialBundle],
) -> None:
    key_findings.to_csv(
        KEY_FINDINGS_CSV,
        index=False,
    )

    inventory.to_csv(
        INPUT_INVENTORY_CSV,
        index=False,
    )

    readme = pd.DataFrame(
        {
            "Notes": [
                "This workbook supports the all-trials final synthesis report.",
                "First Trial may appear as a pilot-stage summary if no final tray table was found.",
                "Second Trial and Third Trial are summarised using the best available final synthesis tables.",
                "Observed and adjusted Day 7 values must remain separate.",
                "Combined scores are descriptive, not formal statistical significance tests.",
            ]
        }
    )

    with pd.ExcelWriter(
        SYNTHESIS_WORKBOOK,
        engine="openpyxl",
    ) as writer:
        key_findings.to_excel(
            writer,
            sheet_name="Key Findings",
            index=False,
        )

        inventory.to_excel(
            writer,
            sheet_name="Input Inventory",
            index=False,
        )

        for bundle in bundles:
            if not bundle.combined_tray.empty:
                sheet_name = f"{bundle.name[:20]} Trays"

                bundle.combined_tray.head(200).to_excel(
                    writer,
                    sheet_name=sheet_name,
                    index=False,
                )

            if not bundle.group_synthesis.empty:
                sheet_name = f"{bundle.name[:18]} Groups"

                bundle.group_synthesis.head(200).to_excel(
                    writer,
                    sheet_name=sheet_name,
                    index=False,
                )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_excel(
        SYNTHESIS_WORKBOOK
    )


# ============================================================
# 10) WORD HELPERS
# ============================================================

def set_document_style(document: Document) -> None:
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(11)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name in document.styles:
            document.styles[style_name].font.name = "Times New Roman"


def add_table(
    document: Document,
    dataframe: pd.DataFrame,
    columns: list[str],
    max_rows: int = 12,
) -> None:
    if dataframe.empty:
        document.add_paragraph(
            "No table data were available for this section."
        )
        return

    valid_columns = [
        column
        for column in columns
        if column in dataframe.columns
    ]

    if not valid_columns:
        document.add_paragraph(
            "The expected columns were not available in this table."
        )
        return

    display = dataframe[valid_columns].head(max_rows).copy()

    table = document.add_table(
        rows=1,
        cols=len(valid_columns),
    )

    table.style = "Table Grid"

    for index, column in enumerate(valid_columns):
        table.rows[0].cells[index].text = column.replace("_", " ").title()

    for _index, row in display.iterrows():
        cells = table.add_row().cells

        for column_index, column in enumerate(valid_columns):
            value = row[column]

            if pd.isna(value):
                cells[column_index].text = ""

            elif isinstance(value, float):
                cells[column_index].text = f"{value:.3f}"

            else:
                cells[column_index].text = str(value)


def add_output_description(
    document: Document,
    filename: str,
    description: str,
) -> None:
    paragraph = document.add_paragraph()

    paragraph.add_run(filename).bold = True
    paragraph.add_run(f": {description}")


def add_trial_result_paragraphs(
    document: Document,
    bundle: TrialBundle,
) -> None:
    best_tray = best_row_by_score(
        bundle.combined_tray
    )

    if best_tray is not None:
        score_column = pick_column(
            bundle.combined_tray,
            TRAY_SCORE_COLUMNS,
        )

        tray = clean_text(
            best_tray.get("tray", "N/A")
        )

        treatment = clean_text(
            best_tray.get("treatment", "N/A")
        )

        microbe = clean_text(
            best_tray.get("microbe_status", "N/A")
        )

        environment = clean_text(
            best_tray.get("environment_group", best_tray.get("fixed_environment", "N/A"))
        )

        document.add_paragraph(
            f"The strongest tray-level result available for this trial was {tray}. "
            f"It was classified as {microbe} | {treatment} | {environment}, "
            f"with an available synthesis score of {format_number(best_tray.get(score_column), 2)}."
        )

    elif bundle.name == "First Trial":
        document.add_paragraph(
            "The First Trial is treated as the pilot stage of the workflow. Its main value was not only the biological result, but also the practical learning: how the tray images should be captured, how the cell structure should be handled, and why the later workflows needed a fixed 7 × 10 cell-based pipeline."
        )

    elif bundle.name == "Second Trial":
        document.add_paragraph(
            "The Second Trial became the first mature version of the workflow. It established the reusable crop, grid, visible-emergence, RGB growth, multispectral index and treatment-comparison pipeline. Its final interpretation identified the Microbes Inside trays as the strongest treatment pattern, with Tray 8 and Tray 5 standing out in the overall result."
        )

    else:
        document.add_paragraph(
            "No tray-level synthesis table was available for this trial."
        )

    microbe_best = best_group_by_score(
        bundle.group_synthesis,
        "Microbe Status",
    )

    if microbe_best is not None:
        score_column = pick_column(
            bundle.group_synthesis,
            GROUP_SCORE_COLUMNS,
        )

        document.add_paragraph(
            f"For the microbial comparison, {clean_text(microbe_best.get('group', 'N/A'))} had the stronger available group-level score "
            f"({format_number(microbe_best.get(score_column), 2)})."
        )

    treatment_best = best_group_by_score(
        bundle.group_synthesis,
        "Treatment Type",
    )

    if treatment_best is not None:
        score_column = pick_column(
            bundle.group_synthesis,
            GROUP_SCORE_COLUMNS,
        )

        document.add_paragraph(
            f"For the treatment comparison, {clean_text(treatment_best.get('group', 'N/A'))} had the strongest available treatment-level score "
            f"({format_number(treatment_best.get(score_column), 2)})."
        )

    if bundle.notes:
        for note in bundle.notes:
            document.add_paragraph(
                note,
                style="List Bullet",
            )


# ============================================================
# 11) WORD REPORT
# ============================================================

def create_word_report(
    bundles: list[TrialBundle],
    key_findings: pd.DataFrame,
    inventory: pd.DataFrame,
) -> Path:
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run:\n"
            '& "C:\\Users\\tshib\\anaconda3\\python.exe" -m pip install python-docx'
        )

    document = Document()

    set_document_style(document)

    title = document.add_heading(
        "Final Synthesis Report Across First, Second and Third Germination Trials",
        level=0,
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This report brings together the three germination-trial workflows completed during the internship. It is written as a final technical synthesis rather than another image-processing step. The aim is to explain what was tested, how the workflow developed, what the strongest patterns were, and how the outputs should be used for future reporting or handover."
    )

    document.add_paragraph(
        "Across the internship, the work moved from early trial learning to a more reliable tray-based computer-vision workflow. The final version combined RGB evidence, which captured visible emergence and green cover, with multispectral evidence, which captured relative image-derived NDVI and NDRE. These evidence streams were then brought together into final synthesis and reporting packages."
    )

    # ------------------------------------------------------------
    # Overall findings
    # ------------------------------------------------------------

    document.add_heading(
        "1. Overall key findings",
        level=1,
    )

    add_table(
        document,
        key_findings,
        [
            "trial",
            "topic",
            "finding",
            "value",
            "source",
        ],
        max_rows=20,
    )

    # ------------------------------------------------------------
    # Trial structure
    # ------------------------------------------------------------

    document.add_heading(
        "2. Trial structure and progression",
        level=1,
    )

    document.add_paragraph(
        "The First Trial acted as the pilot stage. It helped identify practical issues around image capture, tray alignment, cell detection and repeatable reporting. The Second Trial then developed the workflow into a more complete RGB and multispectral pipeline using the 7 × 10 tray structure. The Third Trial extended the same workflow to 12 trays, seven observation days, and additional Heat and Moisture treatment schedules."
    )

    document.add_paragraph(
        "This progression was important because the image-analysis workflow did not start as a fixed solution. It had to be built around real tray images, imperfect lighting, crop orientation differences, RGB and multispectral camera differences, and seedlings that sometimes appeared outside the inner circular cup region."
    )

    # ------------------------------------------------------------
    # First Trial
    # ------------------------------------------------------------

    document.add_heading(
        "3. First Trial synthesis",
        level=1,
    )

    first = bundles[0]

    add_trial_result_paragraphs(
        document,
        first,
    )

    if first.summary_report_path:
        document.add_paragraph(
            f"Available First Trial report/reference file: {first.summary_report_path}"
        )

    # ------------------------------------------------------------
    # Second Trial
    # ------------------------------------------------------------

    document.add_heading(
        "4. Second Trial synthesis",
        level=1,
    )

    second = bundles[1]

    document.add_paragraph(
        "The Second Trial was the main stage where the workflow became repeatable. The pipeline used cropped D/RGB and multispectral images, a fixed 7 × 10 cell grid, daily visible-emergence detection, RGB green-cover metrics, independent multispectral grid detection, NDVI/NDRE extraction, and treatment comparison."
    )

    add_trial_result_paragraphs(
        document,
        second,
    )

    if not second.combined_tray.empty:
        add_table(
            document,
            second.combined_tray,
            [
                "tray",
                "microbe_status",
                "treatment",
                "environment_group",
                "combined_rgb_ms_score",
                "overall_adjusted_ms_score",
                "overall_adjusted_rgb_score",
                "final_interpretation",
            ],
            max_rows=12,
        )

    # ------------------------------------------------------------
    # Third Trial
    # ------------------------------------------------------------

    document.add_heading(
        "5. Third Trial synthesis",
        level=1,
    )

    third = bundles[2]

    document.add_paragraph(
        "The Third Trial extended the workflow to a larger and more complex treatment design. It included Microbes and No Microbes trays, Ideal trays, Heat-treated trays, Moisture-treated trays, Inside greenhouse conditions and Outside open conditions. Day 1 was corrected to 29 June 2026, and the final image set was taken on Day 7, 7 July 2026."
    )

    document.add_paragraph(
        "The Third Trial also introduced an important Day 7 issue: some cells that had shown earlier seedlings appeared empty by the final image set, likely because plants had been eaten by bugs. The workflow therefore preserved the observed Day 7 values while also creating separate adjusted estimates for likely bug-eaten cells. This prevented the observed and adjusted results from being mixed silently."
    )

    add_trial_result_paragraphs(
        document,
        third,
    )

    if not third.combined_tray.empty:
        add_table(
            document,
            third.combined_tray,
            [
                "combined_rank",
                "tray",
                "microbe_status",
                "treatment",
                "environment_group",
                "combined_rgb_ms_score",
                "rgb_score_normalised",
                "ms_score_normalised",
                "evidence_agreement",
                "total_day7_adjusted_cells",
            ],
            max_rows=12,
        )

    # ------------------------------------------------------------
    # Cross-trial method development
    # ------------------------------------------------------------

    document.add_heading(
        "6. Cross-trial workflow development",
        level=1,
    )

    method_points = [
        "The workflow moved from manual and exploratory inspection toward a repeatable script-based pipeline.",
        "The fixed 7 × 10 cell structure became the central unit of analysis, allowing every tray to be compared at cell, tray and treatment level.",
        "RGB analysis was used for visible emergence and green-cover tracking.",
        "Multispectral analysis was kept independent because RGB and MS images have different camera geometry and crop margins.",
        "Square ownership zones were used instead of only inner cup circles, so seedlings growing slightly outside the cup could still be assigned to the correct cell.",
        "Observed daily evidence and cumulative emergence were kept separate.",
        "Day 7 observed and adjusted values were kept separate in the Third Trial to handle likely bug-eaten cells transparently.",
        "The final package scripts collected charts, CSVs, Excel workbooks and Word reports into organised handover folders.",
    ]

    for point in method_points:
        document.add_paragraph(
            point,
            style="List Bullet",
        )

    # ------------------------------------------------------------
    # Treatment interpretation
    # ------------------------------------------------------------

    document.add_heading(
        "7. Treatment-level interpretation across trials",
        level=1,
    )

    document.add_paragraph(
        "The Second Trial provided the clearest early treatment result, with the Microbes Inside condition performing best overall. This was important because it gave a strong comparison point before the more complex Third Trial design."
    )

    document.add_paragraph(
        "The Third Trial should be interpreted with more caution because it had more treatment factors and a final-day bug-damage adjustment scenario. Its final synthesis report and combined tray table are the main sources for identifying the strongest trays and treatment patterns in that trial."
    )

    document.add_paragraph(
        "Inside-versus-Outside comparisons should be made only where the environment was fixed. In the Third Trial, this means Ideal and Moisture trays. Heat trays moved between inside and outside, so they should be interpreted as a dynamic stress-and-recovery treatment rather than a fixed environment comparison."
    )

    # ------------------------------------------------------------
    # Observed vs adjusted
    # ------------------------------------------------------------

    document.add_heading(
        "8. Observed versus adjusted values",
        level=1,
    )

    document.add_paragraph(
        "The observed values are the direct image records and must remain the primary evidence. Adjusted values are useful only as a separate analytical scenario. In the Third Trial, adjusted Day 7 values were used when a cell appeared to have a plant earlier but was missing at the final image set, which was treated as likely bug damage."
    )

    document.add_paragraph(
        "This distinction matters for reporting. The adjusted values can make growth-rate comparison fairer, but they should not be described as what was directly visible in the Day 7 images."
    )

    # ------------------------------------------------------------
    # Limitations
    # ------------------------------------------------------------

    document.add_heading(
        "9. Limitations",
        level=1,
    )

    limitations = [
        "The RGB and multispectral outputs measure related but different evidence. RGB measures visible green cover and emergence, while MS measures relative image-derived vegetation indices.",
        "NDVI and NDRE should be described as relative image-derived indices unless calibrated reflectance data are confirmed.",
        "Soil background, shadows, moisture, lighting, band alignment and crop margins can affect image-derived metrics.",
        "The tray-level scores are descriptive integration scores, not formal statistical significance tests.",
        "The number of trays per treatment is limited, so treatment differences should be interpreted cautiously.",
        "The Third Trial Day 7 adjusted values are estimates for likely bug-eaten plants, not direct observations.",
        "Missing or differently named earlier-trial output files can limit automated cross-trial extraction, so the input inventory should be checked.",
    ]

    for limitation in limitations:
        document.add_paragraph(
            limitation,
            style="List Bullet",
        )

    # ------------------------------------------------------------
    # Output guide
    # ------------------------------------------------------------

    document.add_heading(
        "10. Output guide",
        level=1,
    )

    add_output_description(
        document,
        "all_trials_final_synthesis_report.docx",
        "This Word report. It summarises the First, Second and Third Trial workflow and final interpretation.",
    )

    add_output_description(
        document,
        "all_trials_key_findings.csv",
        "Compact key-findings table used in the report.",
    )

    add_output_description(
        document,
        "all_trials_input_inventory.csv",
        "Inventory showing which trial files were found and used.",
    )

    add_output_description(
        document,
        "all_trials_synthesis_workbook.xlsx",
        "Excel workbook containing key findings, input inventory, and available tray/group tables.",
    )

    add_output_description(
        document,
        "all_trials_synthesis_settings.json",
        "Settings file recording input/output paths and processing assumptions.",
    )

    # ------------------------------------------------------------
    # Inventory
    # ------------------------------------------------------------

    document.add_heading(
        "11. Input inventory",
        level=1,
    )

    add_table(
        document,
        inventory,
        [
            "trial",
            "item",
            "exists",
            "path",
            "description",
        ],
        max_rows=30,
    )

    document.save(
        WORD_REPORT
    )

    return WORD_REPORT


# ============================================================
# 12) SETTINGS
# ============================================================

def save_settings(
    bundles: list[TrialBundle],
    key_findings: pd.DataFrame,
    inventory: pd.DataFrame,
) -> Path:
    settings = {
        "purpose": "Final synthesis report across First, Second and Third Trial workflows",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "project_root": str(PROJECT_ROOT),
        "outputs_root": str(OUTPUTS_ROOT),
        "final_output_root": str(FINAL_OUTPUT_ROOT),
        "word_report": str(WORD_REPORT),
        "key_findings_csv": str(KEY_FINDINGS_CSV),
        "input_inventory_csv": str(INPUT_INVENTORY_CSV),
        "synthesis_workbook": str(SYNTHESIS_WORKBOOK),
        "settings_json": str(SETTINGS_JSON),
        "trial_inputs": {
            bundle.name: {
                "output_root": str(bundle.output_root) if bundle.output_root else "",
                "combined_tray_path": str(bundle.combined_tray_path) if bundle.combined_tray_path else "",
                "group_synthesis_path": str(bundle.group_synthesis_path) if bundle.group_synthesis_path else "",
                "inside_outside_path": str(bundle.inside_outside_path) if bundle.inside_outside_path else "",
                "summary_report_path": str(bundle.summary_report_path) if bundle.summary_report_path else "",
                "package_zip_path": str(bundle.package_zip_path) if bundle.package_zip_path else "",
                "notes": bundle.notes,
            }
            for bundle in bundles
        },
        "key_findings_count": int(len(key_findings)),
        "inventory_rows": int(len(inventory)),
        "observed_adjusted_policy": (
            "Observed values remain the primary direct image evidence. Adjusted values are kept separately and used only as scenario estimates for likely bug-eaten cells."
        ),
        "interpretation_warning": (
            "Combined RGB + MS scores are descriptive synthesis scores, not formal statistical significance tests."
        ),
    }

    SETTINGS_JSON.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    return SETTINGS_JSON


# ============================================================
# 13) MAIN
# ============================================================

def main() -> int:
    print(
        "\nSCRIPT 12 — FULL FIRST + SECOND + THIRD TRIAL FINAL SYNTHESIS REPORT"
    )

    print("=" * 88)

    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run this first:\n"
            '& "C:\\Users\\tshib\\anaconda3\\python.exe" -m pip install python-docx'
        )

    ensure_output_folders()

    first = load_first_trial_bundle()
    second = load_second_trial_bundle()
    third = load_third_trial_bundle()

    bundles = [
        first,
        second,
        third,
    ]

    inventory = create_input_inventory(
        bundles
    )

    key_findings = create_key_findings(
        bundles
    )

    write_tables(
        key_findings,
        inventory,
        bundles,
    )

    report_path = create_word_report(
        bundles,
        key_findings,
        inventory,
    )

    settings_path = save_settings(
        bundles,
        key_findings,
        inventory,
    )

    print("\nSCRIPT 12 FINISHED")
    print("=" * 88)

    for bundle in bundles:
        print(f"\n{bundle.name}:")
        print(f"  Output root: {bundle.output_root}")
        print(f"  Combined tray table: {bundle.combined_tray_path}")
        print(f"  Group table: {bundle.group_synthesis_path}")
        print(f"  Inside/Outside table: {bundle.inside_outside_path}")
        print(f"  Summary report: {bundle.summary_report_path}")
        print(f"  Package ZIP: {bundle.package_zip_path}")
        print(f"  Tray rows loaded: {len(bundle.combined_tray)}")
        print(f"  Group rows loaded: {len(bundle.group_synthesis)}")

        if bundle.notes:
            for note in bundle.notes:
                print(f"  Note: {note}")

    print(f"\nWord report:\n{report_path}")
    print(f"\nKey findings CSV:\n{KEY_FINDINGS_CSV}")
    print(f"\nInput inventory CSV:\n{INPUT_INVENTORY_CSV}")
    print(f"\nExcel workbook:\n{SYNTHESIS_WORKBOOK}")
    print(f"\nSettings:\n{settings_path}")

    print(
        "\nCheck the input inventory. If the First or Second Trial file names are different, "
        "the report will still be created, but you may want to manually confirm the source files."
    )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())