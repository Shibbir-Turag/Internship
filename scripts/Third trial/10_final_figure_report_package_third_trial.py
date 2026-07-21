from __future__ import annotations

"""
SCRIPT 10 — THIRD TRIAL FINAL FIGURE AND REPORT PACKAGE

Purpose
-------
Collect the important Third Trial outputs into one final organised package.

This script does not recalculate RGB or multispectral results. It only copies,
indexes, and summarises the completed workflow outputs from Scripts 04–09.

Inputs
------
outputs/Third trial/04_Visible_Emergence
outputs/Third trial/05_RGB_Growth_Rate_Treatment_Comparison
outputs/Third trial/06_MS_Cell_Grid_Detection
outputs/Third trial/07_MS_Vegetation_Indices
outputs/Third trial/08_MS_Treatment_Comparison
outputs/Third trial/09_Final_RGB_MS_Synthesis

Main output
-----------
outputs/Third trial/10_Final_Figure_Report_Package

The package contains:
- final Word reports
- final Excel workbooks
- key CSV data tables
- RGB charts
- MS charts
- final synthesis charts
- grid/evidence overlays
- settings/config files
- package manifest CSV
- package manifest Excel workbook
- Word handover report
- ZIP archive of the package folder

Important
---------
This package is for reporting and handover. It keeps observed and adjusted
Day 7 outputs separate by copying the original files exactly.
"""

import argparse
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

THIRD_OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
)

SCRIPT04_ROOT = THIRD_OUTPUT_ROOT / "04_Visible_Emergence"

SCRIPT05_ROOT = THIRD_OUTPUT_ROOT / "05_RGB_Growth_Rate_Treatment_Comparison"

SCRIPT06_ROOT = THIRD_OUTPUT_ROOT / "06_MS_Cell_Grid_Detection"

SCRIPT07_ROOT = THIRD_OUTPUT_ROOT / "07_MS_Vegetation_Indices"

SCRIPT08_ROOT = THIRD_OUTPUT_ROOT / "08_MS_Treatment_Comparison"

SCRIPT09_ROOT = THIRD_OUTPUT_ROOT / "09_Final_RGB_MS_Synthesis"

OUTPUT_ROOT = THIRD_OUTPUT_ROOT / "10_Final_Figure_Report_Package"

PACKAGE_ROOT = OUTPUT_ROOT / "Third_Trial_Final_Package"

REPORTS_DEST = PACKAGE_ROOT / "01_Final_Word_Reports"

EXCEL_DEST = PACKAGE_ROOT / "02_Final_Excel_Workbooks"

TABLES_DEST = PACKAGE_ROOT / "03_Key_CSV_Tables"

CHARTS_DEST = PACKAGE_ROOT / "04_Final_Charts"

EVIDENCE_DEST = PACKAGE_ROOT / "05_Quality_Evidence_Overlays"

CONFIG_DEST = PACKAGE_ROOT / "06_Settings_and_Manifests"

PACKAGE_REPORTS_DEST = PACKAGE_ROOT / "00_Package_Index"

ZIP_PATH = OUTPUT_ROOT / "Third_Trial_Final_Figure_Report_Package.zip"


# ============================================================
# 2) OPTIONAL WORD SUPPORT
# ============================================================

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    DOCX_AVAILABLE = True

except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 3) FILE COLLECTION RULES
# ============================================================

WORD_REPORT_PATTERNS = [
    SCRIPT05_ROOT / "_reports" / "*.docx",
    SCRIPT06_ROOT / "_reports" / "*.docx",
    SCRIPT07_ROOT / "_reports" / "*.docx",
    SCRIPT08_ROOT / "_reports" / "*.docx",
    SCRIPT09_ROOT / "_reports" / "*.docx",
]

EXCEL_REPORT_PATTERNS = [
    SCRIPT04_ROOT / "_reports" / "*.xlsx",
    SCRIPT05_ROOT / "_reports" / "*.xlsx",
    SCRIPT06_ROOT / "_reports" / "*.xlsx",
    SCRIPT07_ROOT / "_reports" / "*.xlsx",
    SCRIPT08_ROOT / "_reports" / "*.xlsx",
    SCRIPT09_ROOT / "_reports" / "*.xlsx",
]

CSV_TABLE_PATTERNS = [
    SCRIPT04_ROOT / "_reports" / "*.csv",
    SCRIPT05_ROOT / "_reports" / "*.csv",
    SCRIPT06_ROOT / "_reports" / "*.csv",
    SCRIPT07_ROOT / "_reports" / "*.csv",
    SCRIPT08_ROOT / "_reports" / "*.csv",
    SCRIPT09_ROOT / "_reports" / "*.csv",
]

CONFIG_PATTERNS = [
    SCRIPT04_ROOT / "_config" / "*.json",
    SCRIPT05_ROOT / "_config" / "*.json",
    SCRIPT06_ROOT / "_config" / "*.json",
    SCRIPT07_ROOT / "_config" / "*.json",
    SCRIPT08_ROOT / "_config" / "*.json",
    SCRIPT09_ROOT / "_config" / "*.json",
]

CHART_PATTERNS = [
    SCRIPT05_ROOT / "charts" / "*.png",
    SCRIPT08_ROOT / "charts" / "*.png",
    SCRIPT09_ROOT / "charts" / "*.png",
]

# These evidence overlays can be large. The script copies a selected sample
# from each relevant stage, not every overlay image.
EVIDENCE_SAMPLE_LIMIT_PER_FOLDER = 12

EVIDENCE_SEARCH_FOLDERS = [
    SCRIPT04_ROOT,
    SCRIPT06_ROOT,
    SCRIPT07_ROOT,
]

EVIDENCE_FILE_PATTERNS = [
    "*overlay*.png",
    "*overlay*.jpg",
    "*grid_overlay*.jpg",
    "*polygon_overlay*.png",
    "*mask*.png",
]


# ============================================================
# 4) HELPERS
# ============================================================

def clean_output_folder(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)

    path.mkdir(
        parents=True,
        exist_ok=True,
    )


def ensure_folders() -> None:
    for folder in [
        OUTPUT_ROOT,
        PACKAGE_ROOT,
        REPORTS_DEST,
        EXCEL_DEST,
        TABLES_DEST,
        CHARTS_DEST,
        EVIDENCE_DEST,
        CONFIG_DEST,
        PACKAGE_REPORTS_DEST,
    ]:
        folder.mkdir(
            parents=True,
            exist_ok=True,
        )


def safe_filename(path: Path, prefix: str = "") -> str:
    name = path.name

    if prefix:
        name = f"{prefix}__{name}"

    return name.replace(" ", "_")


def infer_source_script(path: Path) -> str:
    text = str(path).replace("\\", "/").casefold()

    if "04_visible_emergence" in text:
        return "Script 04 - RGB Visible Emergence"

    if "05_rgb_growth_rate_treatment_comparison" in text:
        return "Script 05 - RGB Growth Rate and Treatment Comparison"

    if "06_ms_cell_grid_detection" in text:
        return "Script 06 - MS Cell Grid Detection"

    if "07_ms_vegetation_indices" in text:
        return "Script 07 - MS Vegetation Indices"

    if "08_ms_treatment_comparison" in text:
        return "Script 08 - MS Treatment Comparison"

    if "09_final_rgb_ms_synthesis" in text:
        return "Script 09 - Final RGB + MS Synthesis"

    return "Unknown"


def infer_category(path: Path) -> str:
    suffix = path.suffix.casefold()

    if suffix == ".docx":
        return "Word report"

    if suffix == ".xlsx":
        return "Excel workbook"

    if suffix == ".csv":
        return "CSV table"

    if suffix in {".png", ".jpg", ".jpeg"}:
        return "Chart or evidence image"

    if suffix == ".json":
        return "Settings/config file"

    return "Other"


def file_size_kb(path: Path) -> float:
    try:
        return round(path.stat().st_size / 1024, 2)
    except Exception:
        return 0.0


def copy_file(
    source: Path,
    destination_folder: Path,
    copied_names: set[str],
) -> tuple[Path, str]:
    source_script = infer_source_script(source)

    prefix = (
        source_script
        .split(" - ")[0]
        .replace("Script ", "S")
        .replace(" ", "_")
    )

    destination_name = safe_filename(
        source,
        prefix=prefix,
    )

    destination = destination_folder / destination_name

    counter = 2

    while destination.name in copied_names or destination.exists():
        destination = destination_folder / f"{destination.stem}_{counter}{destination.suffix}"
        counter += 1

    shutil.copy2(
        source,
        destination,
    )

    copied_names.add(
        destination.name
    )

    return destination, source_script


def glob_existing(patterns: list[Path]) -> list[Path]:
    files = []

    for pattern in patterns:
        parent = pattern.parent

        if not parent.exists():
            continue

        files.extend(
            sorted(
                parent.glob(pattern.name),
                key=lambda path: str(path).casefold(),
            )
        )

    unique = []

    seen = set()

    for file in files:
        resolved = str(file.resolve()).casefold()

        if resolved not in seen and file.is_file():
            seen.add(resolved)
            unique.append(file)

    return unique


def sample_evidence_files() -> list[Path]:
    selected = []

    for folder in EVIDENCE_SEARCH_FOLDERS:
        if not folder.exists():
            continue

        folder_candidates = []

        for pattern in EVIDENCE_FILE_PATTERNS:
            folder_candidates.extend(
                sorted(
                    folder.rglob(pattern),
                    key=lambda path: str(path).casefold(),
                )
            )

        filtered = []

        seen = set()

        for file in folder_candidates:
            if not file.is_file():
                continue

            resolved = str(file.resolve()).casefold()

            if resolved in seen:
                continue

            seen.add(resolved)

            # Avoid copying too many very similar files. Prefer first few across folders.
            filtered.append(file)

        selected.extend(
            filtered[:EVIDENCE_SAMPLE_LIMIT_PER_FOLDER]
        )

    return selected


def describe_file(path: Path) -> str:
    name = path.name.casefold()

    if "final_rgb_ms_synthesis_report" in name:
        return "Final Third Trial RGB + MS synthesis Word report."

    if "ms_treatment_comparison_report" in name:
        return "Multispectral treatment-comparison Word report."

    if "rgb" in name and "report" in name and path.suffix.casefold() == ".docx":
        return "RGB treatment or growth-comparison Word report."

    if "work_process" in name:
        return "Work-process report explaining the processing method."

    if "final_combined_tray_synthesis" in name:
        return "Main final tray-level RGB + MS synthesis table."

    if "final_group_synthesis" in name:
        return "Final group-level RGB + MS comparison table."

    if "inside_outside" in name:
        return "Inside-versus-Outside comparison for Ideal and Moisture trays."

    if "tray_growth" in name:
        return "Tray-level growth metrics."

    if "group_growth" in name:
        return "Group-level growth metrics."

    if "day7" in name and ("bug" in name or "adjust" in name or "imputed" in name):
        return "Day 7 adjusted or possible bug-eaten cell record."

    if "manifest" in name:
        return "Processing manifest or file inventory."

    if "settings" in name or path.suffix.casefold() == ".json":
        return "Settings file documenting processing parameters."

    if path.suffix.casefold() in {".png", ".jpg", ".jpeg"}:
        return "Visual chart or quality-check evidence image."

    return "Copied workflow output."


# ============================================================
# 5) PACKAGE CREATION
# ============================================================

def collect_package_files() -> pd.DataFrame:
    manifest_rows = []

    copied_names = set()

    collection_specs = [
        (
            WORD_REPORT_PATTERNS,
            REPORTS_DEST,
            "Final Word reports",
        ),
        (
            EXCEL_REPORT_PATTERNS,
            EXCEL_DEST,
            "Excel workbooks",
        ),
        (
            CSV_TABLE_PATTERNS,
            TABLES_DEST,
            "Key CSV tables",
        ),
        (
            CHART_PATTERNS,
            CHARTS_DEST,
            "Charts",
        ),
        (
            CONFIG_PATTERNS,
            CONFIG_DEST,
            "Settings and manifests",
        ),
    ]

    for patterns, destination, package_section in collection_specs:
        files = glob_existing(
            patterns
        )

        for source in files:
            copied, source_script = copy_file(
                source,
                destination,
                copied_names,
            )

            manifest_rows.append(
                {
                    "package_section": package_section,
                    "source_script": source_script,
                    "file_category": infer_category(source),
                    "original_file_name": source.name,
                    "package_file_name": copied.name,
                    "original_path": str(source),
                    "package_path": str(copied),
                    "size_kb": file_size_kb(copied),
                    "description": describe_file(source),
                    "copied": "Yes",
                }
            )

    evidence_files = sample_evidence_files()

    for source in evidence_files:
        copied, source_script = copy_file(
            source,
            EVIDENCE_DEST,
            copied_names,
        )

        manifest_rows.append(
            {
                "package_section": "Quality evidence overlays",
                "source_script": source_script,
                "file_category": infer_category(source),
                "original_file_name": source.name,
                "package_file_name": copied.name,
                "original_path": str(source),
                "package_path": str(copied),
                "size_kb": file_size_kb(copied),
                "description": describe_file(source),
                "copied": "Yes",
            }
        )

    return pd.DataFrame(
        manifest_rows
    )


def create_summary_table(
    manifest: pd.DataFrame,
) -> pd.DataFrame:
    if manifest.empty:
        return pd.DataFrame()

    summary = (
        manifest.groupby(
            [
                "package_section",
                "file_category",
            ],
            as_index=False,
            dropna=False,
        )
        .agg(
            file_count=(
                "package_file_name",
                "count",
            ),
            total_size_kb=(
                "size_kb",
                "sum",
            ),
        )
        .sort_values(
            [
                "package_section",
                "file_category",
            ]
        )
    )

    summary["total_size_kb"] = summary[
        "total_size_kb"
    ].round(2)

    return summary


def style_excel(path: Path) -> None:
    workbook = load_workbook(
        path
    )

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.row_dimensions[1].height = 34

        for cell in worksheet[1]:
            cell.fill = header_fill

            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )

            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = column_cells[0].column_letter

            longest = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[letter].width = min(
                max(12, longest + 2),
                70,
            )

    workbook.save(
        path
    )


def write_manifest_files(
    manifest: pd.DataFrame,
    summary: pd.DataFrame,
) -> dict[str, Path]:
    PACKAGE_REPORTS_DEST.mkdir(
        parents=True,
        exist_ok=True,
    )

    manifest_csv = PACKAGE_REPORTS_DEST / "third_trial_package_manifest.csv"

    summary_csv = PACKAGE_REPORTS_DEST / "third_trial_package_summary.csv"

    excel_path = PACKAGE_REPORTS_DEST / "third_trial_package_index.xlsx"

    manifest.to_csv(
        manifest_csv,
        index=False,
    )

    summary.to_csv(
        summary_csv,
        index=False,
    )

    readme = pd.DataFrame(
        {
            "Notes": [
                "This package collects final Third Trial reporting outputs from Scripts 04–09.",
                "The package is intended for supervisor review, report writing, and future handover.",
                "Observed Day 7 and adjusted Day 7 outputs remain separated in the source CSV/Excel/Word files.",
                "Adjusted Day 7 values are scenario estimates for likely bug-eaten cells, not direct image observations.",
                "RGB outputs describe visible emergence and green cover.",
                "MS outputs describe relative image-derived NDVI and NDRE.",
                "The final synthesis combines RGB and MS evidence descriptively; it is not a formal statistical test.",
                "Quality evidence overlays are sampled so the package remains manageable.",
            ]
        }
    )

    with pd.ExcelWriter(
        excel_path,
        engine="openpyxl",
    ) as writer:
        manifest.to_excel(
            writer,
            sheet_name="Package Manifest",
            index=False,
        )

        summary.to_excel(
            writer,
            sheet_name="Package Summary",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_excel(
        excel_path
    )

    return {
        "manifest_csv": manifest_csv,
        "summary_csv": summary_csv,
        "excel_index": excel_path,
    }


# ============================================================
# 6) WORD HANDOVER REPORT
# ============================================================

def add_docx_table(
    document,
    dataframe: pd.DataFrame,
    columns: list[str],
    maximum_rows: int = 20,
) -> None:
    if dataframe.empty:
        document.add_paragraph(
            "No table data were available."
        )
        return

    columns = [
        column
        for column in columns
        if column in dataframe.columns
    ]

    if not columns:
        document.add_paragraph(
            "No valid table columns were available."
        )
        return

    frame = dataframe[columns].head(maximum_rows).copy()

    table = document.add_table(
        rows=1,
        cols=len(columns),
    )

    table.style = "Table Grid"

    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column.replace("_", " ").title()

    for _index, row in frame.iterrows():
        cells = table.add_row().cells

        for index, column in enumerate(columns):
            value = row[column]

            cells[index].text = "" if pd.isna(value) else str(value)


def create_word_handover_report(
    path: Path,
    manifest: pd.DataFrame,
    summary: pd.DataFrame,
    index_paths: dict[str, Path],
) -> Path | None:
    if not DOCX_AVAILABLE:
        print(
            "Word handover report skipped because python-docx is not installed."
        )
        print(
            "Install it with: pip install python-docx"
        )
        return None

    document = Document()

    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            document.styles[style_name].font.name = "Times New Roman"

    title = document.add_heading(
        "Third Trial Final Figure and Report Package Handover",
        level=0,
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This handover document summarises the final packaged outputs for the "
        "Third Trial workflow. The package brings together the main Word reports, "
        "Excel workbooks, CSV result tables, charts, settings files, and selected "
        "quality-check overlays from Scripts 04 to 09."
    )

    document.add_paragraph(
        "The package is designed for final review, supervisor discussion, report "
        "writing, and continuation of the internship workflow in another chat or "
        "analysis session."
    )

    document.add_heading(
        "1. Package location",
        level=1,
    )

    document.add_paragraph(
        f"Package folder: {PACKAGE_ROOT}"
    )

    document.add_paragraph(
        f"ZIP archive: {ZIP_PATH}"
    )

    document.add_heading(
        "2. Package summary",
        level=1,
    )

    document.add_paragraph(
        f"Total copied files: {len(manifest)}."
    )

    if not manifest.empty:
        document.add_paragraph(
            f"Total package size before ZIP: {manifest['size_kb'].sum():.2f} KB."
        )

    add_docx_table(
        document,
        summary,
        [
            "package_section",
            "file_category",
            "file_count",
            "total_size_kb",
        ],
        maximum_rows=30,
    )

    document.add_heading(
        "3. Workflow coverage",
        level=1,
    )

    workflow_notes = [
        (
            "Script 04: RGB visible emergence and green-cover evidence, including "
            "daily and cumulative visible emergence outputs."
        ),
        (
            "Script 05: RGB growth-rate and treatment comparison outputs, including "
            "Day 7 observed and adjusted scenarios for possible bug-eaten plants."
        ),
        (
            "Script 06: independent multispectral 70-cell grid detection evidence."
        ),
        (
            "Script 07: relative image-derived NDVI and NDRE extraction from original "
            "multispectral TIFF bands."
        ),
        (
            "Script 08: multispectral treatment and growth comparison outputs."
        ),
        (
            "Script 09: final RGB + MS synthesis outputs and final combined tray ranking."
        ),
    ]

    for note in workflow_notes:
        document.add_paragraph(
            note,
            style="List Bullet",
        )

    document.add_heading(
        "4. Important interpretation notes",
        level=1,
    )

    document.add_paragraph(
        "Observed Day 7 values are the direct image observations and remain the "
        "primary evidence. Adjusted Day 7 values are kept separately and should "
        "be described as scenario-based estimates for cells where plants were "
        "likely eaten by bugs after earlier growth."
    )

    document.add_paragraph(
        "RGB outputs measure visible emergence and green cover. MS outputs measure "
        "relative image-derived NDVI and NDRE. These evidence streams are related "
        "but not identical, so the final RGB + MS score should be treated as a "
        "descriptive synthesis score, not a formal statistical test."
    )

    document.add_paragraph(
        "Unless calibrated reflectance information is available, NDVI and NDRE "
        "should be reported as relative image-derived indices."
    )

    document.add_heading(
        "5. Main package index files",
        level=1,
    )

    for label, output_path in index_paths.items():
        paragraph = document.add_paragraph()
        paragraph.add_run(label).bold = True
        paragraph.add_run(f": {output_path}")

    document.add_heading(
        "6. Selected manifest preview",
        level=1,
    )

    add_docx_table(
        document,
        manifest,
        [
            "package_section",
            "source_script",
            "file_category",
            "package_file_name",
            "description",
        ],
        maximum_rows=25,
    )

    document.add_heading(
        "7. Recommended next use",
        level=1,
    )

    document.add_paragraph(
        "Use the final Word reports and the final synthesis workbook first. The "
        "CSV files are included for traceability and additional analysis. The "
        "charts folder contains the figure files most likely to be used in a "
        "final report or presentation."
    )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        path
    )

    return path


# ============================================================
# 7) ZIP PACKAGE
# ============================================================

def create_zip_archive(
    package_root: Path,
    zip_path: Path,
) -> Path:
    if zip_path.exists():
        zip_path.unlink()

    zip_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with zipfile.ZipFile(
        zip_path,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=6,
    ) as zip_file:
        for file in sorted(
            package_root.rglob("*"),
            key=lambda path: str(path).casefold(),
        ):
            if file.is_file():
                zip_file.write(
                    file,
                    file.relative_to(package_root.parent),
                )

    return zip_path


def save_settings(
    path: Path,
    manifest: pd.DataFrame,
) -> Path:
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    settings = {
        "purpose": "Third Trial final figure and report package",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "project_root": str(PROJECT_ROOT),
        "third_output_root": str(THIRD_OUTPUT_ROOT),
        "package_root": str(PACKAGE_ROOT),
        "zip_path": str(ZIP_PATH),
        "source_scripts": {
            "script04": str(SCRIPT04_ROOT),
            "script05": str(SCRIPT05_ROOT),
            "script06": str(SCRIPT06_ROOT),
            "script07": str(SCRIPT07_ROOT),
            "script08": str(SCRIPT08_ROOT),
            "script09": str(SCRIPT09_ROOT),
        },
        "file_count": int(len(manifest)),
        "total_size_kb": float(manifest["size_kb"].sum()) if not manifest.empty else 0.0,
        "evidence_sample_limit_per_folder": EVIDENCE_SAMPLE_LIMIT_PER_FOLDER,
        "day7_policy": (
            "Observed Day 7 and adjusted Day 7 outputs are copied as separate "
            "source files. This package does not mix or recalculate them."
        ),
    }

    path.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    return path


# ============================================================
# 8) MAIN
# ============================================================

def run_analysis(args) -> int:
    print(
        "\nSCRIPT 10 — THIRD TRIAL FINAL FIGURE AND REPORT PACKAGE"
    )

    print(
        "=" * 78
    )

    print(
        f"Third Trial output root:\n{THIRD_OUTPUT_ROOT}"
    )

    print(
        f"\nPackage output folder:\n{OUTPUT_ROOT}"
    )

    if args.clean:
        clean_output_folder(
            OUTPUT_ROOT
        )

    ensure_folders()

    manifest = collect_package_files()

    summary = create_summary_table(
        manifest
    )

    index_paths = write_manifest_files(
        manifest,
        summary,
    )

    settings_path = save_settings(
        CONFIG_DEST / "third_trial_package_settings.json",
        manifest,
    )

    index_paths["settings_json"] = settings_path

    word_path = PACKAGE_REPORTS_DEST / "third_trial_final_package_handover.docx"

    created_word = create_word_handover_report(
        word_path,
        manifest,
        summary,
        index_paths,
    )

    if created_word is not None:
        index_paths["word_handover_report"] = created_word

    zip_path = create_zip_archive(
        PACKAGE_ROOT,
        ZIP_PATH,
    )

    print(
        "\n"
        + "=" * 78
    )

    print(
        "SCRIPT 10 FINISHED"
    )

    print(
        "=" * 78
    )

    print(
        f"Files copied: {len(manifest)}"
    )

    if not manifest.empty:
        print(
            f"Package size before ZIP: {manifest['size_kb'].sum():.2f} KB"
        )

    print(
        "\nPackage folder:"
    )

    print(
        PACKAGE_ROOT
    )

    print(
        "\nZIP archive:"
    )

    print(
        zip_path
    )

    print(
        "\nPackage manifest:"
    )

    print(
        index_paths["manifest_csv"]
    )

    print(
        "\nPackage Excel index:"
    )

    print(
        index_paths["excel_index"]
    )

    if created_word is not None:
        print(
            "\nWord handover report:"
        )

        print(
            created_word
        )

    else:
        print(
            "\nWord handover report skipped because python-docx is not installed."
        )

    if manifest.empty:
        print(
            "\nWARNING: No files were copied. Check that Scripts 04–09 output folders exist."
        )

        return 1

    print(
        "\nPackage sections:"
    )

    if not summary.empty:
        print(
            summary.to_string(index=False)
        )

    return 0


# ============================================================
# 9) CLI
# ============================================================

def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Trial 3 Script 10: final figure and report package."
        )
    )

    parser.add_argument(
        "--clean",
        action="store_true",
        default=True,
        help=(
            "Clean and rebuild the package folder. Enabled by default."
        ),
    )

    args = parser.parse_args()

    return run_analysis(
        args
    )


if __name__ == "__main__":
    raise SystemExit(
        main()
    )