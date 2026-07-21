from __future__ import annotations

"""
SCRIPT 08 — THIRD TRIAL MULTISPECTRAL TREATMENT AND GROWTH COMPARISON

Purpose
-------
Use the cell-level relative NDVI and NDRE outputs from Script 07 to calculate
tray-level and treatment-level multispectral comparisons for Trial 3.

This script performs:
- Microbes vs No Microbes comparison
- Ideal vs Heat vs Moisture comparison
- Microbes × Treatment comparison
- Ideal Inside vs Outside comparison
- Moisture Inside vs Outside comparison
- Heat-treatment phase analysis
- Moisture-treatment phase analysis
- observed Day 7 NDVI/NDRE reporting
- separately flagged adjusted Day 7 NDVI/NDRE estimation for likely
  bug-eaten or missing seedlings
- tray-level multispectral growth-rate calculations
- descriptive tray ranking
- CSV, Excel, chart, and Word report generation

Important Day 7 rule
--------------------
Observed Day 7 values are never overwritten.

Cells previously identified as possibly bug-eaten are given separate adjusted
Day 7 NDVI and NDRE estimates. The adjusted estimate is derived from the
previous valid multispectral trend and is clearly flagged.

These adjusted results must not be described as directly observed data.

Input
-----
outputs/Third trial/07_MS_Vegetation_Indices/_reports/
    ms_index_cell_results.csv
    ms_index_tray_summary.csv

Optional bug-eaten-cell sources
-------------------------------
Preferred:
outputs/Third trial/05_RGB_Growth_Rate_Treatment_Comparison/_reports/
    possible_day7_bug_eaten_cells.csv

Fallback:
outputs/Third trial/04_Visible_Emergence/_reports/
    visible_emergence_cell_results.csv

Main output
-----------
outputs/Third trial/08_MS_Treatment_Comparison/

The Word report describes:
- the processing method
- the generated charts
- every main CSV file
- the Excel workbook
- observed versus adjusted Day 7 results
- interpretation limitations
"""

import argparse
import json
import math
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

SCRIPT07_REPORTS = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "07_MS_Vegetation_Indices"
    / "_reports"
)

SCRIPT07_CELL_RESULTS = (
    SCRIPT07_REPORTS
    / "ms_index_cell_results.csv"
)

SCRIPT07_TRAY_SUMMARY = (
    SCRIPT07_REPORTS
    / "ms_index_tray_summary.csv"
)

SCRIPT05_REPORTS = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "05_RGB_Growth_Rate_Treatment_Comparison"
    / "_reports"
)

SCRIPT05_BUG_CELLS = (
    SCRIPT05_REPORTS
    / "possible_day7_bug_eaten_cells.csv"
)

SCRIPT04_CELL_RESULTS = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "04_Visible_Emergence"
    / "_reports"
    / "visible_emergence_cell_results.csv"
)

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "08_MS_Treatment_Comparison"
)

CHARTS_ROOT = OUTPUT_ROOT / "charts"
REPORTS_ROOT = OUTPUT_ROOT / "_reports"
CONFIG_ROOT = OUTPUT_ROOT / "_config"


# ============================================================
# 2) TRIAL SETTINGS
# ============================================================

EXPECTED_CELLS_PER_TRAY = 70
EXPECTED_TRAYS = 12
EXPECTED_OBSERVATION_DAYS = 7
EXPECTED_TRAY_DAY_ROWS = 84
EXPECTED_CELL_DAY_ROWS = 5880

DATE_MAP = {
    1: "2026-06-29",
    2: "2026-06-30",
    3: "2026-07-01",
    4: "2026-07-02",
    5: "2026-07-03",
    6: "2026-07-04",
    7: "2026-07-07",
}

DAYS_SINCE_DAY1 = {
    1: 0,
    2: 1,
    3: 2,
    4: 3,
    5: 4,
    6: 5,
    7: 8,
}

DAYS_SINCE_PLANTING = {
    1: 2,
    2: 3,
    3: 4,
    4: 5,
    5: 6,
    6: 7,
    7: 10,
}

DAYS_SINCE_PREVIOUS_PHOTO = {
    1: 0,
    2: 1,
    3: 1,
    4: 1,
    5: 1,
    6: 1,
    7: 3,
}

MICROBE_ORDER = [
    "No Microbes",
    "Microbes",
]

TREATMENT_ORDER = [
    "Ideal",
    "Heat",
    "Moisture",
]

ENVIRONMENT_ORDER = [
    "Inside",
    "Outside",
]

FIXED_ENVIRONMENT_TREATMENTS = {
    "Ideal",
    "Moisture",
}

PASS_SCRIPT07_STATUSES = {
    "PASS",
    "CHECK",
}

MAX_INDEX_SLOPE_PER_DAY = 0.20
MIN_VALID_PRIOR_POINTS_FOR_LINEAR_TREND = 2


# ============================================================
# 3) CORRECTED TRIAL 3 TRAY DESIGN
# ============================================================

TRAY_METADATA = {
    1: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Inside",
    },
    2: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Outside",
    },
    3: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Outside",
    },
    4: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
    },
    5: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Inside",
    },
    6: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
    },
    7: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Outside",
    },
    8: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
    },
    9: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Inside",
    },
    10: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Outside",
    },
    11: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
    },
    12: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Inside",
    },
}


# ============================================================
# 4) TREATMENT-PHASE METADATA
# ============================================================

HEAT_PHASE_BY_DAY = {
    1: "Inside baseline",
    2: "Inside baseline",
    3: "Moved outside - heat exposure start",
    4: "Outside heat exposure",
    5: "Outside heat exposure",
    6: "Returned inside - recovery start",
    7: "Inside recovery",
}

HEAT_ENVIRONMENT_BY_DAY = {
    1: "Inside",
    2: "Inside",
    3: "Outside",
    4: "Outside",
    5: "Outside",
    6: "Inside",
    7: "Inside",
}

MOISTURE_PHASE_BY_DAY = {
    1: "Drying after Day 0 watering",
    2: "Drying after Day 0 watering",
    3: "Watered on Day 3",
    4: "Unwatered drying period",
    5: "Unwatered drying period",
    6: "Watered on Day 6",
    7: "Final unwatered interval",
}

MOISTURE_WATERED_TODAY = {
    1: "No",
    2: "No",
    3: "Yes",
    4: "No",
    5: "No",
    6: "Yes",
    7: "No",
}


# ============================================================
# 5) OPTIONAL WORD SUPPORT
# ============================================================

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True

except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 6) GENERAL HELPERS
# ============================================================

def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(
            r"(\d+)",
            str(value),
        )
    ]


def parse_yes(value: object) -> bool:
    return (
        str(value)
        .strip()
        .casefold()
        in {
            "yes",
            "y",
            "true",
            "1",
            "p",
        }
    )


def safe_float(
    value: object,
    default: float = math.nan,
) -> float:
    try:
        if pd.isna(value):
            return default

        return float(value)

    except Exception:
        return default


def format_number(
    value: object,
    decimals: int = 3,
) -> str:
    try:
        if pd.isna(value):
            return "N/A"

        return f"{float(value):.{decimals}f}"

    except Exception:
        return "N/A"


def require_columns(
    dataframe: pd.DataFrame,
    columns: list[str],
    source_name: str,
) -> None:
    missing = [
        column
        for column in columns
        if column not in dataframe.columns
    ]

    if missing:
        raise ValueError(
            f"{source_name} is missing required column(s): "
            + ", ".join(missing)
        )


def safe_numeric(
    dataframe: pd.DataFrame,
    columns: list[str],
) -> pd.DataFrame:
    output = dataframe.copy()

    for column in columns:
        if column in output.columns:
            output[column] = pd.to_numeric(
                output[column],
                errors="coerce",
            )

    return output


def safe_round_dataframe(
    dataframe: pd.DataFrame,
    decimals: int = 4,
) -> pd.DataFrame:
    output = dataframe.copy()

    numeric_columns = output.select_dtypes(
        include=["number"]
    ).columns

    output[numeric_columns] = output[
        numeric_columns
    ].round(decimals)

    return output


def mean_or_nan(
    dataframe: pd.DataFrame,
    column: str,
) -> float:
    if (
        dataframe.empty
        or column not in dataframe.columns
    ):
        return math.nan

    return float(
        pd.to_numeric(
            dataframe[column],
            errors="coerce",
        ).mean()
    )


def count_unique_or_zero(
    dataframe: pd.DataFrame,
    column: str,
) -> int:
    if (
        dataframe.empty
        or column not in dataframe.columns
    ):
        return 0

    return int(
        dataframe[column].nunique()
    )


def minmax_score(
    series: pd.Series,
) -> pd.Series:
    values = pd.to_numeric(
        series,
        errors="coerce",
    )

    valid = values.dropna()

    if valid.empty:
        return pd.Series(
            [math.nan] * len(values),
            index=values.index,
        )

    minimum = float(
        valid.min()
    )

    maximum = float(
        valid.max()
    )

    if maximum == minimum:
        return pd.Series(
            [
                50.0
                if pd.notna(value)
                else math.nan
                for value in values
            ],
            index=values.index,
        )

    return (
        values - minimum
    ) / (
        maximum - minimum
    ) * 100.0


def index_auc(
    dataframe: pd.DataFrame,
    value_column: str,
) -> float:
    frame = dataframe[
        [
            "days_since_day1",
            value_column,
        ]
    ].dropna().sort_values(
        "days_since_day1"
    )

    if len(frame) < 2:
        return math.nan

    x = frame[
        "days_since_day1"
    ].to_numpy(dtype=float)

    y = frame[
        value_column
    ].to_numpy(dtype=float)

    return float(
        np.trapz(
            y,
            x,
        )
    )


def linear_rate(
    start_value: object,
    end_value: object,
    start_elapsed: object,
    end_elapsed: object,
) -> float:
    start_value = safe_float(
        start_value
    )

    end_value = safe_float(
        end_value
    )

    start_elapsed = safe_float(
        start_elapsed
    )

    end_elapsed = safe_float(
        end_elapsed
    )

    if (
        not np.isfinite(start_value)
        or not np.isfinite(end_value)
        or not np.isfinite(start_elapsed)
        or not np.isfinite(end_elapsed)
        or end_elapsed <= start_elapsed
    ):
        return math.nan

    return (
        end_value - start_value
    ) / (
        end_elapsed - start_elapsed
    )


# ============================================================
# 7) METADATA
# ============================================================

def treatment_metadata(
    tray_no: int,
    day_order: int,
) -> dict:
    design = TRAY_METADATA.get(
        int(tray_no),
        {
            "microbe_status": "Unknown",
            "treatment": "Unknown",
            "fixed_environment": "Unknown",
        },
    )

    treatment = design[
        "treatment"
    ]

    fixed_environment = design[
        "fixed_environment"
    ]

    if treatment == "Heat":
        observed_environment = (
            HEAT_ENVIRONMENT_BY_DAY.get(
                day_order,
                "Unknown",
            )
        )

        environment_group = (
            "Dynamic Heat"
        )

        heat_phase = HEAT_PHASE_BY_DAY.get(
            day_order,
            ""
        )

    else:
        observed_environment = (
            fixed_environment
        )

        environment_group = (
            fixed_environment
        )

        heat_phase = "Not Heat"

    if treatment == "Moisture":
        moisture_phase = (
            MOISTURE_PHASE_BY_DAY.get(
                day_order,
                ""
            )
        )

        moisture_watered_today = (
            MOISTURE_WATERED_TODAY.get(
                day_order,
                ""
            )
        )

    else:
        moisture_phase = "Not Moisture"

        moisture_watered_today = (
            "Daily watering"
            if treatment in {
                "Ideal",
                "Heat",
            }
            else ""
        )

    return {
        "microbe_status": design[
            "microbe_status"
        ],
        "treatment": treatment,
        "fixed_environment": fixed_environment,
        "observed_environment": observed_environment,
        "environment_group": environment_group,
        "heat_phase": heat_phase,
        "moisture_phase": moisture_phase,
        "moisture_watered_today": moisture_watered_today,
        "microbe_treatment": (
            design["microbe_status"]
            + " | "
            + treatment
        ),
        "treatment_environment": (
            treatment
            + " | "
            + environment_group
        ),
        "microbe_environment": (
            design["microbe_status"]
            + " | "
            + environment_group
        ),
    }


def add_metadata(
    dataframe: pd.DataFrame,
) -> pd.DataFrame:
    output = dataframe.copy()

    output["day_order"] = pd.to_numeric(
        output["day_order"],
        errors="coerce",
    )

    output["tray_no"] = pd.to_numeric(
        output["tray_no"],
        errors="coerce",
    )

    output = output.dropna(
        subset=[
            "day_order",
            "tray_no",
        ]
    ).copy()

    output["day_order"] = output[
        "day_order"
    ].astype(int)

    output["tray_no"] = output[
        "tray_no"
    ].astype(int)

    output["calendar_date"] = output[
        "day_order"
    ].map(DATE_MAP)

    output["days_since_day1"] = output[
        "day_order"
    ].map(DAYS_SINCE_DAY1)

    output["days_since_planting"] = output[
        "day_order"
    ].map(DAYS_SINCE_PLANTING)

    output[
        "days_since_previous_photo"
    ] = output[
        "day_order"
    ].map(
        DAYS_SINCE_PREVIOUS_PHOTO
    )

    metadata_rows = output.apply(
        lambda row: pd.Series(
            treatment_metadata(
                int(row["tray_no"]),
                int(row["day_order"]),
            )
        ),
        axis=1,
    )

    for column in metadata_rows.columns:
        output[column] = metadata_rows[
            column
        ]

    return output


# ============================================================
# 8) LOAD SCRIPT 07 OUTPUT
# ============================================================

def load_script07_outputs() -> tuple[
    pd.DataFrame,
    pd.DataFrame,
]:
    if not SCRIPT07_CELL_RESULTS.exists():
        raise FileNotFoundError(
            f"Missing Script 07 cell results:\n"
            f"{SCRIPT07_CELL_RESULTS}"
        )

    if not SCRIPT07_TRAY_SUMMARY.exists():
        raise FileNotFoundError(
            f"Missing Script 07 tray summary:\n"
            f"{SCRIPT07_TRAY_SUMMARY}"
        )

    cell = pd.read_csv(
        SCRIPT07_CELL_RESULTS
    )

    tray = pd.read_csv(
        SCRIPT07_TRAY_SUMMARY
    )

    require_columns(
        cell,
        [
            "day_order",
            "day",
            "tray",
            "tray_no",
            "capture_id",
            "cell_id",
            "row",
            "column",
            "ndvi_mean",
            "ndre_mean",
        ],
        SCRIPT07_CELL_RESULTS.name,
    )

    require_columns(
        tray,
        [
            "day_order",
            "day",
            "tray",
            "tray_no",
            "capture_id",
            "status",
        ],
        SCRIPT07_TRAY_SUMMARY.name,
    )

    cell = safe_numeric(
        cell,
        [
            "day_order",
            "tray_no",
            "cell_id",
            "row",
            "column",
            "ndvi_mean",
            "ndvi_median",
            "ndvi_std",
            "ndvi_p10",
            "ndvi_p90",
            "ndvi_positive_fraction",
            "ndre_mean",
            "ndre_median",
            "ndre_std",
            "ndre_p10",
            "ndre_p90",
            "ndre_positive_fraction",
        ],
    )

    tray = safe_numeric(
        tray,
        [
            "day_order",
            "tray_no",
            "cells_processed",
            "cells_with_valid_ndvi",
            "cells_with_valid_ndre",
            "mean_cell_ndvi",
            "median_cell_ndvi",
            "mean_cell_ndre",
            "median_cell_ndre",
        ],
    )

    tray["status"] = (
        tray["status"]
        .astype(str)
        .str.upper()
    )

    tray = tray.loc[
        tray["status"].isin(
            PASS_SCRIPT07_STATUSES
        )
    ].copy()

    valid_keys = tray[
        [
            "day_order",
            "tray_no",
            "capture_id",
        ]
    ].drop_duplicates()

    cell = cell.merge(
        valid_keys,
        on=[
            "day_order",
            "tray_no",
            "capture_id",
        ],
        how="inner",
        validate="many_to_one",
    )

    cell = add_metadata(
        cell
    )

    tray = add_metadata(
        tray
    )

    cell["cell_id"] = pd.to_numeric(
        cell["cell_id"],
        errors="coerce",
    ).astype("Int64")

    cell = cell.dropna(
        subset=[
            "cell_id",
        ]
    ).copy()

    cell["cell_id"] = cell[
        "cell_id"
    ].astype(int)

    return tray, cell


# ============================================================
# 9) LOAD POSSIBLE BUG-EATEN CELL FLAGS
# ============================================================

def bug_flags_from_script05() -> set[
    tuple[int, int]
]:
    if not SCRIPT05_BUG_CELLS.exists():
        return set()

    try:
        frame = pd.read_csv(
            SCRIPT05_BUG_CELLS
        )

    except Exception:
        return set()

    if (
        "tray_no" not in frame.columns
        or "cell_id" not in frame.columns
    ):
        return set()

    frame["tray_no"] = pd.to_numeric(
        frame["tray_no"],
        errors="coerce",
    )

    frame["cell_id"] = pd.to_numeric(
        frame["cell_id"],
        errors="coerce",
    )

    if "day7_imputed" in frame.columns:
        frame = frame.loc[
            frame["day7_imputed"].apply(
                parse_yes
            )
        ].copy()

    frame = frame.dropna(
        subset=[
            "tray_no",
            "cell_id",
        ]
    )

    return {
        (
            int(row.tray_no),
            int(row.cell_id),
        )
        for row in frame.itertuples(
            index=False
        )
    }


def bug_flags_from_script04() -> set[
    tuple[int, int]
]:
    if not SCRIPT04_CELL_RESULTS.exists():
        return set()

    try:
        frame = pd.read_csv(
            SCRIPT04_CELL_RESULTS
        )

    except Exception:
        return set()

    required = {
        "day_order",
        "tray_no",
        "cell_id",
        "raw_current_green_evidence",
    }

    if not required.issubset(
        frame.columns
    ):
        return set()

    frame = safe_numeric(
        frame,
        [
            "day_order",
            "tray_no",
            "cell_id",
        ],
    )

    frame = frame.dropna(
        subset=[
            "day_order",
            "tray_no",
            "cell_id",
        ]
    ).copy()

    frame[
        [
            "day_order",
            "tray_no",
            "cell_id",
        ]
    ] = frame[
        [
            "day_order",
            "tray_no",
            "cell_id",
        ]
    ].astype(int)

    flags = set()

    for (
        tray_no,
        cell_id,
    ), group in frame.groupby(
        [
            "tray_no",
            "cell_id",
        ],
        sort=False,
    ):
        prior = group.loc[
            group["day_order"] < 7
        ]

        day7 = group.loc[
            group["day_order"] == 7
        ]

        had_prior_crop = prior[
            "raw_current_green_evidence"
        ].apply(parse_yes).any()

        day7_has_crop = (
            day7[
                "raw_current_green_evidence"
            ].apply(parse_yes).any()
            if not day7.empty
            else False
        )

        explicitly_flagged = (
            day7[
                "possible_day7_bug_eaten"
            ].apply(parse_yes).any()
            if (
                not day7.empty
                and "possible_day7_bug_eaten"
                in day7.columns
            )
            else False
        )

        if (
            explicitly_flagged
            or (
                had_prior_crop
                and not day7_has_crop
            )
        ):
            flags.add(
                (
                    int(tray_no),
                    int(cell_id),
                )
            )

    return flags


def load_bug_flags() -> tuple[
    set[tuple[int, int]],
    str,
]:
    script05_flags = (
        bug_flags_from_script05()
    )

    if script05_flags:
        return (
            script05_flags,
            "Script 05 possible_day7_bug_eaten_cells.csv",
        )

    script04_flags = (
        bug_flags_from_script04()
    )

    if script04_flags:
        return (
            script04_flags,
            "Derived from Script 04 daily RGB emergence evidence",
        )

    return (
        set(),
        "No bug-eaten-cell flag source was available",
    )


# ============================================================
# 10) DAY 7 INDEX ESTIMATION
# ============================================================

def estimate_index_from_prior_trend(
    group: pd.DataFrame,
    observed_column: str,
    target_elapsed: float,
) -> tuple[
    float,
    str,
    float,
]:
    prior = group.loc[
        (
            group["day_order"] < 7
        )
        & pd.to_numeric(
            group[observed_column],
            errors="coerce",
        ).notna()
    ].copy()

    prior = prior.sort_values(
        "days_since_day1"
    )

    if prior.empty:
        return (
            math.nan,
            "No valid earlier multispectral value",
            math.nan,
        )

    prior = prior.tail(3)

    if (
        len(prior)
        >= MIN_VALID_PRIOR_POINTS_FOR_LINEAR_TREND
        and prior[
            "days_since_day1"
        ].nunique() >= 2
    ):
        x = prior[
            "days_since_day1"
        ].to_numpy(dtype=float)

        y = prior[
            observed_column
        ].to_numpy(dtype=float)

        slope, intercept = np.polyfit(
            x,
            y,
            1,
        )

        slope = float(
            np.clip(
                slope,
                -MAX_INDEX_SLOPE_PER_DAY,
                MAX_INDEX_SLOPE_PER_DAY,
            )
        )

        latest_x = float(
            x[-1]
        )

        latest_y = float(
            y[-1]
        )

        estimate = (
            latest_y
            + slope
            * (
                target_elapsed
                - latest_x
            )
        )

        method = (
            "Linear projection from the latest valid "
            "multispectral trend"
        )

    else:
        latest = prior.iloc[-1]

        estimate = float(
            latest[observed_column]
        )

        slope = 0.0

        method = (
            "Carried forward latest valid "
            "multispectral value"
        )

    estimate = float(
        np.clip(
            estimate,
            -1.0,
            1.0,
        )
    )

    return (
        estimate,
        method,
        slope,
    )


def create_adjusted_cell_table(
    cell: pd.DataFrame,
    bug_flags: set[tuple[int, int]],
    bug_flag_source: str,
) -> pd.DataFrame:
    output = cell.copy()

    output["observed_ndvi_mean"] = (
        pd.to_numeric(
            output["ndvi_mean"],
            errors="coerce",
        )
    )

    output["observed_ndre_mean"] = (
        pd.to_numeric(
            output["ndre_mean"],
            errors="coerce",
        )
    )

    output["adjusted_ndvi_mean"] = (
        output["observed_ndvi_mean"]
    )

    output["adjusted_ndre_mean"] = (
        output["observed_ndre_mean"]
    )

    output["day7_imputed"] = "No"
    output["bug_flag_source"] = ""
    output["imputation_method"] = ""
    output["ndvi_prior_slope_per_day"] = math.nan
    output["ndre_prior_slope_per_day"] = math.nan
    output["ndvi_adjustment_difference"] = 0.0
    output["ndre_adjustment_difference"] = 0.0

    for (
        tray_no,
        cell_id,
    ), group in output.groupby(
        [
            "tray_no",
            "cell_id",
        ],
        sort=False,
    ):
        key = (
            int(tray_no),
            int(cell_id),
        )

        if key not in bug_flags:
            continue

        day7_index = group.index[
            group["day_order"].eq(7)
        ]

        if len(day7_index) == 0:
            continue

        target_index = (
            day7_index[0]
        )

        target_elapsed = float(
            output.loc[
                target_index,
                "days_since_day1",
            ]
        )

        ndvi_estimate, ndvi_method, ndvi_slope = (
            estimate_index_from_prior_trend(
                group,
                "observed_ndvi_mean",
                target_elapsed,
            )
        )

        ndre_estimate, ndre_method, ndre_slope = (
            estimate_index_from_prior_trend(
                group,
                "observed_ndre_mean",
                target_elapsed,
            )
        )

        adjusted_any = False

        if np.isfinite(
            ndvi_estimate
        ):
            output.loc[
                target_index,
                "adjusted_ndvi_mean",
            ] = ndvi_estimate

            adjusted_any = True

        if np.isfinite(
            ndre_estimate
        ):
            output.loc[
                target_index,
                "adjusted_ndre_mean",
            ] = ndre_estimate

            adjusted_any = True

        if adjusted_any:
            output.loc[
                target_index,
                "day7_imputed",
            ] = "Yes"

            output.loc[
                target_index,
                "bug_flag_source",
            ] = bug_flag_source

            output.loc[
                target_index,
                "imputation_method",
            ] = (
                f"NDVI: {ndvi_method}; "
                f"NDRE: {ndre_method}"
            )

            output.loc[
                target_index,
                "ndvi_prior_slope_per_day",
            ] = ndvi_slope

            output.loc[
                target_index,
                "ndre_prior_slope_per_day",
            ] = ndre_slope

    output[
        "ndvi_adjustment_difference"
    ] = (
        output["adjusted_ndvi_mean"]
        - output["observed_ndvi_mean"]
    )

    output[
        "ndre_adjustment_difference"
    ] = (
        output["adjusted_ndre_mean"]
        - output["observed_ndre_mean"]
    )

    return output.sort_values(
        [
            "day_order",
            "tray_no",
            "cell_id",
        ]
    ).reset_index(drop=True)


# ============================================================
# 11) TRAY-DAILY METRICS
# ============================================================

def create_tray_daily_metrics(
    adjusted_cell: pd.DataFrame,
) -> pd.DataFrame:
    group_columns = [
        "day_order",
        "day",
        "calendar_date",
        "days_since_day1",
        "days_since_planting",
        "days_since_previous_photo",
        "tray",
        "tray_no",
        "capture_id",
        "microbe_status",
        "treatment",
        "fixed_environment",
        "observed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
        "heat_phase",
        "moisture_phase",
        "moisture_watered_today",
    ]

    daily = (
        adjusted_cell.groupby(
            group_columns,
            as_index=False,
            dropna=False,
        )
        .agg(
            cells_processed=(
                "cell_id",
                "count",
            ),
            valid_observed_ndvi_cells=(
                "observed_ndvi_mean",
                lambda series: int(
                    series.notna().sum()
                ),
            ),
            valid_observed_ndre_cells=(
                "observed_ndre_mean",
                lambda series: int(
                    series.notna().sum()
                ),
            ),
            mean_observed_ndvi=(
                "observed_ndvi_mean",
                "mean",
            ),
            median_observed_ndvi=(
                "observed_ndvi_mean",
                "median",
            ),
            mean_adjusted_ndvi=(
                "adjusted_ndvi_mean",
                "mean",
            ),
            median_adjusted_ndvi=(
                "adjusted_ndvi_mean",
                "median",
            ),
            mean_observed_ndre=(
                "observed_ndre_mean",
                "mean",
            ),
            median_observed_ndre=(
                "observed_ndre_mean",
                "median",
            ),
            mean_adjusted_ndre=(
                "adjusted_ndre_mean",
                "mean",
            ),
            median_adjusted_ndre=(
                "adjusted_ndre_mean",
                "median",
            ),
            day7_imputed_cells=(
                "day7_imputed",
                lambda series: int(
                    series.eq("Yes").sum()
                ),
            ),
            mean_ndvi_adjustment=(
                "ndvi_adjustment_difference",
                "mean",
            ),
            mean_ndre_adjustment=(
                "ndre_adjustment_difference",
                "mean",
            ),
        )
    )

    daily["status"] = np.where(
        daily[
            "cells_processed"
        ].eq(
            EXPECTED_CELLS_PER_TRAY
        ),
        "PASS",
        "CHECK",
    )

    return daily.sort_values(
        [
            "day_order",
            "tray_no",
        ]
    ).reset_index(drop=True)


# ============================================================
# 12) TRAY GROWTH METRICS
# ============================================================

def value_for_day(
    rows_by_day: dict[int, pd.Series],
    day_order: int,
    column: str,
):
    row = rows_by_day.get(
        day_order
    )

    if row is None:
        return math.nan

    return row.get(
        column,
        math.nan,
    )


def create_tray_growth_metrics(
    tray_daily: pd.DataFrame,
) -> pd.DataFrame:
    records = []

    for tray_no, group in tray_daily.groupby(
        "tray_no",
        sort=True,
    ):
        group = group.sort_values(
            "day_order"
        )

        rows_by_day = {
            int(row["day_order"]): row
            for _index, row
            in group.iterrows()
        }

        reference = group.iloc[0]

        day1_elapsed = value_for_day(
            rows_by_day,
            1,
            "days_since_day1",
        )

        day6_elapsed = value_for_day(
            rows_by_day,
            6,
            "days_since_day1",
        )

        day7_elapsed = value_for_day(
            rows_by_day,
            7,
            "days_since_day1",
        )

        record = {
            "tray_no": int(
                tray_no
            ),
            "tray": str(
                reference["tray"]
            ),
            "microbe_status": str(
                reference[
                    "microbe_status"
                ]
            ),
            "treatment": str(
                reference["treatment"]
            ),
            "fixed_environment": str(
                reference[
                    "fixed_environment"
                ]
            ),
            "environment_group": str(
                reference[
                    "environment_group"
                ]
            ),
            "microbe_treatment": str(
                reference[
                    "microbe_treatment"
                ]
            ),
            "treatment_environment": str(
                reference[
                    "treatment_environment"
                ]
            ),
            "microbe_environment": str(
                reference[
                    "microbe_environment"
                ]
            ),
            "available_day_count": int(
                group[
                    "day_order"
                ].nunique()
            ),

            "day1_observed_ndvi": value_for_day(
                rows_by_day,
                1,
                "mean_observed_ndvi",
            ),
            "day6_observed_ndvi": value_for_day(
                rows_by_day,
                6,
                "mean_observed_ndvi",
            ),
            "day7_observed_ndvi": value_for_day(
                rows_by_day,
                7,
                "mean_observed_ndvi",
            ),
            "day7_adjusted_ndvi": value_for_day(
                rows_by_day,
                7,
                "mean_adjusted_ndvi",
            ),

            "day1_observed_ndre": value_for_day(
                rows_by_day,
                1,
                "mean_observed_ndre",
            ),
            "day6_observed_ndre": value_for_day(
                rows_by_day,
                6,
                "mean_observed_ndre",
            ),
            "day7_observed_ndre": value_for_day(
                rows_by_day,
                7,
                "mean_observed_ndre",
            ),
            "day7_adjusted_ndre": value_for_day(
                rows_by_day,
                7,
                "mean_adjusted_ndre",
            ),

            "day7_imputed_cells": value_for_day(
                rows_by_day,
                7,
                "day7_imputed_cells",
            ),

            "observed_ndvi_rate_day1_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    1,
                    "mean_observed_ndvi",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_observed_ndvi",
                ),
                day1_elapsed,
                day7_elapsed,
            ),

            "adjusted_ndvi_rate_day1_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    1,
                    "mean_adjusted_ndvi",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_adjusted_ndvi",
                ),
                day1_elapsed,
                day7_elapsed,
            ),

            "observed_ndre_rate_day1_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    1,
                    "mean_observed_ndre",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_observed_ndre",
                ),
                day1_elapsed,
                day7_elapsed,
            ),

            "adjusted_ndre_rate_day1_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    1,
                    "mean_adjusted_ndre",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_adjusted_ndre",
                ),
                day1_elapsed,
                day7_elapsed,
            ),

            "observed_ndvi_rate_day6_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    6,
                    "mean_observed_ndvi",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_observed_ndvi",
                ),
                day6_elapsed,
                day7_elapsed,
            ),

            "adjusted_ndvi_rate_day6_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    6,
                    "mean_adjusted_ndvi",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_adjusted_ndvi",
                ),
                day6_elapsed,
                day7_elapsed,
            ),

            "observed_ndre_rate_day6_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    6,
                    "mean_observed_ndre",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_observed_ndre",
                ),
                day6_elapsed,
                day7_elapsed,
            ),

            "adjusted_ndre_rate_day6_to_day7_per_day": linear_rate(
                value_for_day(
                    rows_by_day,
                    6,
                    "mean_adjusted_ndre",
                ),
                value_for_day(
                    rows_by_day,
                    7,
                    "mean_adjusted_ndre",
                ),
                day6_elapsed,
                day7_elapsed,
            ),

            "observed_ndvi_auc": index_auc(
                group,
                "mean_observed_ndvi",
            ),
            "adjusted_ndvi_auc": index_auc(
                group,
                "mean_adjusted_ndvi",
            ),
            "observed_ndre_auc": index_auc(
                group,
                "mean_observed_ndre",
            ),
            "adjusted_ndre_auc": index_auc(
                group,
                "mean_adjusted_ndre",
            ),
        }

        record[
            "day7_ndvi_adjustment"
        ] = (
            safe_float(
                record[
                    "day7_adjusted_ndvi"
                ]
            )
            - safe_float(
                record[
                    "day7_observed_ndvi"
                ]
            )
        )

        record[
            "day7_ndre_adjustment"
        ] = (
            safe_float(
                record[
                    "day7_adjusted_ndre"
                ]
            )
            - safe_float(
                record[
                    "day7_observed_ndre"
                ]
            )
        )

        records.append(
            record
        )

    metrics = pd.DataFrame(
        records
    )

    score_components = [
        "day7_adjusted_ndvi",
        "day7_adjusted_ndre",
        "adjusted_ndvi_rate_day1_to_day7_per_day",
        "adjusted_ndre_rate_day1_to_day7_per_day",
        "adjusted_ndvi_auc",
        "adjusted_ndre_auc",
    ]

    score_columns = []

    for component in score_components:
        score_column = (
            f"{component}_score"
        )

        metrics[score_column] = (
            minmax_score(
                metrics[component]
            )
        )

        score_columns.append(
            score_column
        )

    metrics[
        "overall_adjusted_ms_score"
    ] = metrics[
        score_columns
    ].mean(
        axis=1,
        skipna=True,
    )

    metrics[
        "overall_adjusted_ms_rank"
    ] = metrics[
        "overall_adjusted_ms_score"
    ].rank(
        ascending=False,
        method="min",
    ).astype("Int64")

    return metrics.sort_values(
        [
            "overall_adjusted_ms_rank",
            "tray_no",
        ],
        na_position="last",
    ).reset_index(drop=True)


# ============================================================
# 13) GROUP COMPARISON TABLES
# ============================================================

def create_group_daily(
    tray_daily: pd.DataFrame,
    group_column: str,
    group_type: str,
) -> pd.DataFrame:
    result = (
        tray_daily.groupby(
            [
                group_column,
                "day_order",
                "day",
                "calendar_date",
                "days_since_day1",
            ],
            as_index=False,
            dropna=False,
        )
        .agg(
            tray_count=(
                "tray_no",
                "nunique",
            ),
            mean_observed_ndvi=(
                "mean_observed_ndvi",
                "mean",
            ),
            sd_observed_ndvi=(
                "mean_observed_ndvi",
                "std",
            ),
            mean_adjusted_ndvi=(
                "mean_adjusted_ndvi",
                "mean",
            ),
            sd_adjusted_ndvi=(
                "mean_adjusted_ndvi",
                "std",
            ),
            mean_observed_ndre=(
                "mean_observed_ndre",
                "mean",
            ),
            sd_observed_ndre=(
                "mean_observed_ndre",
                "std",
            ),
            mean_adjusted_ndre=(
                "mean_adjusted_ndre",
                "mean",
            ),
            sd_adjusted_ndre=(
                "mean_adjusted_ndre",
                "std",
            ),
            mean_day7_imputed_cells=(
                "day7_imputed_cells",
                "mean",
            ),
        )
        .rename(
            columns={
                group_column: "group"
            }
        )
    )

    result["group_type"] = (
        group_type
    )

    for column in result.columns:
        if column.startswith(
            "sd_"
        ):
            result[column] = (
                result[column]
                .fillna(0.0)
            )

    return result


def create_group_growth(
    tray_growth: pd.DataFrame,
    group_column: str,
    group_type: str,
) -> pd.DataFrame:
    metrics = [
        "day7_observed_ndvi",
        "day7_adjusted_ndvi",
        "day7_observed_ndre",
        "day7_adjusted_ndre",
        "day7_imputed_cells",
        "observed_ndvi_rate_day1_to_day7_per_day",
        "adjusted_ndvi_rate_day1_to_day7_per_day",
        "observed_ndre_rate_day1_to_day7_per_day",
        "adjusted_ndre_rate_day1_to_day7_per_day",
        "observed_ndvi_rate_day6_to_day7_per_day",
        "adjusted_ndvi_rate_day6_to_day7_per_day",
        "observed_ndre_rate_day6_to_day7_per_day",
        "adjusted_ndre_rate_day6_to_day7_per_day",
        "observed_ndvi_auc",
        "adjusted_ndvi_auc",
        "observed_ndre_auc",
        "adjusted_ndre_auc",
        "overall_adjusted_ms_score",
    ]

    aggregations = {
        "tray_count": (
            "tray_no",
            "nunique",
        )
    }

    for metric in metrics:
        aggregations[
            f"mean_{metric}"
        ] = (
            metric,
            "mean",
        )

        aggregations[
            f"sd_{metric}"
        ] = (
            metric,
            "std",
        )

    result = (
        tray_growth.groupby(
            group_column,
            as_index=False,
            dropna=False,
        )
        .agg(
            **aggregations
        )
        .rename(
            columns={
                group_column: "group"
            }
        )
    )

    result["group_type"] = (
        group_type
    )

    for column in result.columns:
        if column.startswith(
            "sd_"
        ):
            result[column] = (
                result[column]
                .fillna(0.0)
            )

    return result


def create_all_group_tables(
    tray_daily: pd.DataFrame,
    tray_growth: pd.DataFrame,
) -> tuple[
    pd.DataFrame,
    pd.DataFrame,
]:
    specifications = [
        (
            "microbe_status",
            "Microbe Status",
        ),
        (
            "treatment",
            "Treatment Type",
        ),
        (
            "microbe_treatment",
            "Microbe x Treatment",
        ),
        (
            "treatment_environment",
            "Treatment x Environment",
        ),
        (
            "microbe_environment",
            "Microbe x Environment",
        ),
    ]

    daily_parts = []

    growth_parts = []

    for (
        group_column,
        group_type,
    ) in specifications:
        daily_parts.append(
            create_group_daily(
                tray_daily,
                group_column,
                group_type,
            )
        )

        growth_parts.append(
            create_group_growth(
                tray_growth,
                group_column,
                group_type,
            )
        )

    return (
        pd.concat(
            daily_parts,
            ignore_index=True,
        ),
        pd.concat(
            growth_parts,
            ignore_index=True,
        ),
    )


# ============================================================
# 14) IDEAL/MOISTURE INSIDE VS OUTSIDE
# ============================================================

def create_inside_outside_comparison(
    tray_growth: pd.DataFrame,
) -> pd.DataFrame:
    rows = []

    for treatment in [
        "Ideal",
        "Moisture",
    ]:
        subset = tray_growth.loc[
            tray_growth[
                "treatment"
            ].eq(
                treatment
            )
        ].copy()

        inside = subset.loc[
            subset[
                "fixed_environment"
            ].eq(
                "Inside"
            )
        ]

        outside = subset.loc[
            subset[
                "fixed_environment"
            ].eq(
                "Outside"
            )
        ]

        row = {
            "treatment": treatment,

            "inside_tray_count": count_unique_or_zero(
                inside,
                "tray_no",
            ),
            "outside_tray_count": count_unique_or_zero(
                outside,
                "tray_no",
            ),

            "inside_trays": ", ".join(
                inside[
                    "tray"
                ].astype(str)
            ),
            "outside_trays": ", ".join(
                outside[
                    "tray"
                ].astype(str)
            ),

            "inside_day7_adjusted_ndvi": mean_or_nan(
                inside,
                "day7_adjusted_ndvi",
            ),
            "outside_day7_adjusted_ndvi": mean_or_nan(
                outside,
                "day7_adjusted_ndvi",
            ),

            "inside_day7_adjusted_ndre": mean_or_nan(
                inside,
                "day7_adjusted_ndre",
            ),
            "outside_day7_adjusted_ndre": mean_or_nan(
                outside,
                "day7_adjusted_ndre",
            ),

            "inside_adjusted_ndvi_rate_per_day": mean_or_nan(
                inside,
                "adjusted_ndvi_rate_day1_to_day7_per_day",
            ),
            "outside_adjusted_ndvi_rate_per_day": mean_or_nan(
                outside,
                "adjusted_ndvi_rate_day1_to_day7_per_day",
            ),

            "inside_adjusted_ndre_rate_per_day": mean_or_nan(
                inside,
                "adjusted_ndre_rate_day1_to_day7_per_day",
            ),
            "outside_adjusted_ndre_rate_per_day": mean_or_nan(
                outside,
                "adjusted_ndre_rate_day1_to_day7_per_day",
            ),

            "inside_overall_ms_score": mean_or_nan(
                inside,
                "overall_adjusted_ms_score",
            ),
            "outside_overall_ms_score": mean_or_nan(
                outside,
                "overall_adjusted_ms_score",
            ),
        }

        row[
            "inside_minus_outside_day7_ndvi"
        ] = (
            row[
                "inside_day7_adjusted_ndvi"
            ]
            - row[
                "outside_day7_adjusted_ndvi"
            ]
        )

        row[
            "inside_minus_outside_day7_ndre"
        ] = (
            row[
                "inside_day7_adjusted_ndre"
            ]
            - row[
                "outside_day7_adjusted_ndre"
            ]
        )

        row[
            "inside_minus_outside_ndvi_rate"
        ] = (
            row[
                "inside_adjusted_ndvi_rate_per_day"
            ]
            - row[
                "outside_adjusted_ndvi_rate_per_day"
            ]
        )

        row[
            "inside_minus_outside_ndre_rate"
        ] = (
            row[
                "inside_adjusted_ndre_rate_per_day"
            ]
            - row[
                "outside_adjusted_ndre_rate_per_day"
            ]
        )

        if (
            row["inside_tray_count"] == 0
            or row["outside_tray_count"] == 0
        ):
            interpretation = (
                "Insufficient Inside or Outside tray data."
            )

        elif (
            row[
                "inside_minus_outside_day7_ndvi"
            ] > 0
        ):
            interpretation = (
                "Inside trays had higher adjusted final NDVI."
            )

        elif (
            row[
                "inside_minus_outside_day7_ndvi"
            ] < 0
        ):
            interpretation = (
                "Outside trays had higher adjusted final NDVI."
            )

        else:
            interpretation = (
                "Inside and Outside adjusted final NDVI were equal."
            )

        row["interpretation"] = (
            interpretation
        )

        rows.append(
            row
        )

    return pd.DataFrame(
        rows
    )


# ============================================================
# 15) HEAT AND MOISTURE PHASE TABLES
# ============================================================

def create_phase_tables(
    tray_daily: pd.DataFrame,
) -> tuple[
    pd.DataFrame,
    pd.DataFrame,
]:
    heat = tray_daily.loc[
        tray_daily[
            "treatment"
        ].eq(
            "Heat"
        )
    ].copy()

    moisture = tray_daily.loc[
        tray_daily[
            "treatment"
        ].eq(
            "Moisture"
        )
    ].copy()

    heat_phase = (
        heat.groupby(
            [
                "heat_phase",
                "observed_environment",
                "microbe_status",
                "day_order",
                "day",
                "days_since_day1",
            ],
            as_index=False,
            dropna=False,
        )
        .agg(
            tray_count=(
                "tray_no",
                "nunique",
            ),
            mean_observed_ndvi=(
                "mean_observed_ndvi",
                "mean",
            ),
            mean_adjusted_ndvi=(
                "mean_adjusted_ndvi",
                "mean",
            ),
            mean_observed_ndre=(
                "mean_observed_ndre",
                "mean",
            ),
            mean_adjusted_ndre=(
                "mean_adjusted_ndre",
                "mean",
            ),
            mean_day7_imputed_cells=(
                "day7_imputed_cells",
                "mean",
            ),
        )
        .sort_values(
            [
                "day_order",
                "microbe_status",
            ]
        )
        .reset_index(drop=True)
    )

    moisture_phase = (
        moisture.groupby(
            [
                "moisture_phase",
                "moisture_watered_today",
                "fixed_environment",
                "microbe_status",
                "day_order",
                "day",
                "days_since_day1",
            ],
            as_index=False,
            dropna=False,
        )
        .agg(
            tray_count=(
                "tray_no",
                "nunique",
            ),
            mean_observed_ndvi=(
                "mean_observed_ndvi",
                "mean",
            ),
            mean_adjusted_ndvi=(
                "mean_adjusted_ndvi",
                "mean",
            ),
            mean_observed_ndre=(
                "mean_observed_ndre",
                "mean",
            ),
            mean_adjusted_ndre=(
                "mean_adjusted_ndre",
                "mean",
            ),
            mean_day7_imputed_cells=(
                "day7_imputed_cells",
                "mean",
            ),
        )
        .sort_values(
            [
                "day_order",
                "fixed_environment",
                "microbe_status",
            ]
        )
        .reset_index(drop=True)
    )

    return (
        heat_phase,
        moisture_phase,
    )


# ============================================================
# 16) CHART HELPERS
# ============================================================

def no_data_chart(
    title: str,
    output_path: Path,
    message: str = "No valid data available",
) -> None:
    figure, axis = plt.subplots(
        figsize=(9, 5.5)
    )

    axis.text(
        0.5,
        0.5,
        message,
        horizontalalignment="center",
        verticalalignment="center",
        fontsize=12,
    )

    axis.set_title(
        title
    )

    axis.axis(
        "off"
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_group_trend_chart(
    group_daily: pd.DataFrame,
    group_type: str,
    group_order: list[str],
    value_column: str,
    sd_column: str,
    title: str,
    y_label: str,
    output_path: Path,
) -> None:
    frame = group_daily.loc[
        group_daily[
            "group_type"
        ].eq(
            group_type
        )
    ].copy()

    if (
        frame.empty
        or value_column not in frame.columns
    ):
        no_data_chart(
            title,
            output_path,
        )

        return

    figure, axis = plt.subplots(
        figsize=(11.5, 6.5)
    )

    plotted = False

    for group in group_order:
        subset = frame.loc[
            frame["group"].eq(
                group
            )
        ].sort_values(
            "days_since_day1"
        )

        if subset.empty:
            continue

        values = pd.to_numeric(
            subset[value_column],
            errors="coerce",
        )

        if values.notna().sum() == 0:
            continue

        errors = (
            pd.to_numeric(
                subset[sd_column],
                errors="coerce",
            )
            if sd_column in subset.columns
            else None
        )

        axis.errorbar(
            subset[
                "days_since_day1"
            ],
            values,
            yerr=errors,
            marker="o",
            linewidth=2,
            capsize=4,
            label=(
                f"{group} "
                f"(n={int(subset['tray_count'].max())})"
            ),
        )

        plotted = True

    if not plotted:
        plt.close(
            figure
        )

        no_data_chart(
            title,
            output_path,
        )

        return

    ticks = (
        frame[
            [
                "day",
                "days_since_day1",
            ]
        ]
        .drop_duplicates()
        .sort_values(
            "days_since_day1"
        )
    )

    axis.set_title(
        title
    )

    axis.set_xlabel(
        "Observation day"
    )

    axis.set_ylabel(
        y_label
    )

    axis.set_xticks(
        ticks[
            "days_since_day1"
        ]
    )

    axis.set_xticklabels(
        ticks["day"]
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    axis.legend(
        loc="best"
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_observed_adjusted_chart(
    tray_growth: pd.DataFrame,
    observed_column: str,
    adjusted_column: str,
    title: str,
    y_label: str,
    output_path: Path,
) -> None:
    frame = tray_growth.sort_values(
        "tray_no"
    ).copy()

    if frame.empty:
        no_data_chart(
            title,
            output_path,
        )

        return

    x = np.arange(
        len(frame)
    )

    width = 0.38

    figure, axis = plt.subplots(
        figsize=(12, 6.5)
    )

    axis.bar(
        x - width / 2,
        frame[
            observed_column
        ],
        width,
        label="Observed Day 7",
    )

    axis.bar(
        x + width / 2,
        frame[
            adjusted_column
        ],
        width,
        label="Adjusted Day 7",
    )

    axis.set_title(
        title
    )

    axis.set_xlabel(
        "Tray"
    )

    axis.set_ylabel(
        y_label
    )

    axis.set_xticks(
        x
    )

    axis.set_xticklabels(
        frame["tray"],
        rotation=30,
        horizontalalignment="right",
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    axis.legend(
        loc="best"
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_tray_ranking_chart(
    tray_growth: pd.DataFrame,
    output_path: Path,
) -> None:
    frame = tray_growth.dropna(
        subset=[
            "overall_adjusted_ms_score",
        ]
    ).sort_values(
        "overall_adjusted_ms_score"
    ).copy()

    if frame.empty:
        no_data_chart(
            "Trial 3 adjusted multispectral tray ranking",
            output_path,
        )

        return

    frame["chart_label"] = (
        frame["tray"]
        + " - "
        + frame["microbe_status"]
        + " | "
        + frame["treatment"]
        + " | "
        + frame["environment_group"]
    )

    figure, axis = plt.subplots(
        figsize=(12.5, 7.2)
    )

    axis.barh(
        frame["chart_label"],
        frame[
            "overall_adjusted_ms_score"
        ],
    )

    axis.set_title(
        "Trial 3 adjusted multispectral tray ranking"
    )

    axis.set_xlabel(
        "Adjusted multispectral performance score"
    )

    axis.set_ylabel(
        "Tray"
    )

    axis.grid(
        True,
        axis="x",
        alpha=0.30,
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_inside_outside_chart(
    comparison: pd.DataFrame,
    inside_column: str,
    outside_column: str,
    title: str,
    y_label: str,
    output_path: Path,
) -> None:
    if comparison.empty:
        no_data_chart(
            title,
            output_path,
        )

        return

    inside_values = pd.to_numeric(
        comparison[
            inside_column
        ],
        errors="coerce",
    )

    outside_values = pd.to_numeric(
        comparison[
            outside_column
        ],
        errors="coerce",
    )

    if (
        inside_values.notna().sum() == 0
        and outside_values.notna().sum() == 0
    ):
        no_data_chart(
            title,
            output_path,
            "Inside/Outside values could not be calculated",
        )

        return

    x = np.arange(
        len(comparison)
    )

    width = 0.36

    figure, axis = plt.subplots(
        figsize=(9.5, 5.8)
    )

    axis.bar(
        x - width / 2,
        inside_values,
        width,
        label="Inside",
    )

    axis.bar(
        x + width / 2,
        outside_values,
        width,
        label="Outside",
    )

    axis.set_title(
        title
    )

    axis.set_xlabel(
        "Treatment"
    )

    axis.set_ylabel(
        y_label
    )

    axis.set_xticks(
        x
    )

    axis.set_xticklabels(
        comparison[
            "treatment"
        ]
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    axis.legend(
        loc="best"
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_imputed_cells_chart(
    tray_growth: pd.DataFrame,
    output_path: Path,
) -> None:
    frame = tray_growth.sort_values(
        "tray_no"
    )

    if frame.empty:
        no_data_chart(
            "Possible Day 7 bug-eaten cells used for MS adjustment",
            output_path,
        )

        return

    figure, axis = plt.subplots(
        figsize=(11, 6)
    )

    axis.bar(
        frame["tray"],
        pd.to_numeric(
            frame[
                "day7_imputed_cells"
            ],
            errors="coerce",
        ).fillna(0),
    )

    axis.set_title(
        "Possible Day 7 bug-eaten cells used for MS adjustment"
    )

    axis.set_xlabel(
        "Tray"
    )

    axis.set_ylabel(
        "Cells adjusted"
    )

    axis.tick_params(
        axis="x",
        rotation=30,
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_phase_trend_chart(
    phase_frame: pd.DataFrame,
    grouping_column: str,
    value_column: str,
    title: str,
    y_label: str,
    output_path: Path,
) -> None:
    if (
        phase_frame.empty
        or grouping_column
        not in phase_frame.columns
        or value_column
        not in phase_frame.columns
    ):
        no_data_chart(
            title,
            output_path,
        )

        return

    figure, axis = plt.subplots(
        figsize=(11, 6.5)
    )

    plotted = False

    for group_name, group in phase_frame.groupby(
        grouping_column,
        sort=False,
    ):
        group = group.sort_values(
            "days_since_day1"
        )

        values = pd.to_numeric(
            group[value_column],
            errors="coerce",
        )

        if values.notna().sum() == 0:
            continue

        axis.plot(
            group[
                "days_since_day1"
            ],
            values,
            marker="o",
            linewidth=2,
            label=str(
                group_name
            ),
        )

        plotted = True

    if not plotted:
        plt.close(
            figure
        )

        no_data_chart(
            title,
            output_path,
        )

        return

    ticks = (
        phase_frame[
            [
                "day",
                "days_since_day1",
            ]
        ]
        .drop_duplicates()
        .sort_values(
            "days_since_day1"
        )
    )

    axis.set_xticks(
        ticks[
            "days_since_day1"
        ]
    )

    axis.set_xticklabels(
        ticks["day"]
    )

    axis.set_title(
        title
    )

    axis.set_xlabel(
        "Observation day"
    )

    axis.set_ylabel(
        y_label
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    axis.legend(
        loc="best"
    )

    figure.tight_layout()

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=220,
    )

    plt.close(
        figure
    )


def create_charts(
    group_daily: pd.DataFrame,
    tray_growth: pd.DataFrame,
    inside_outside: pd.DataFrame,
    heat_phase: pd.DataFrame,
    moisture_phase: pd.DataFrame,
) -> dict[str, Path]:
    CHARTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    charts = {}

    path = (
        CHARTS_ROOT
        / "01_microbes_adjusted_ndvi_trend.png"
    )

    save_group_trend_chart(
        group_daily,
        "Microbe Status",
        MICROBE_ORDER,
        "mean_adjusted_ndvi",
        "sd_adjusted_ndvi",
        "Adjusted relative NDVI: Microbes vs No Microbes",
        "Mean relative NDVI",
        path,
    )

    charts[
        "microbes_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "02_microbes_adjusted_ndre_trend.png"
    )

    save_group_trend_chart(
        group_daily,
        "Microbe Status",
        MICROBE_ORDER,
        "mean_adjusted_ndre",
        "sd_adjusted_ndre",
        "Adjusted relative NDRE: Microbes vs No Microbes",
        "Mean relative NDRE",
        path,
    )

    charts[
        "microbes_ndre"
    ] = path

    path = (
        CHARTS_ROOT
        / "03_treatment_adjusted_ndvi_trend.png"
    )

    save_group_trend_chart(
        group_daily,
        "Treatment Type",
        TREATMENT_ORDER,
        "mean_adjusted_ndvi",
        "sd_adjusted_ndvi",
        "Adjusted relative NDVI by treatment",
        "Mean relative NDVI",
        path,
    )

    charts[
        "treatment_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "04_treatment_adjusted_ndre_trend.png"
    )

    save_group_trend_chart(
        group_daily,
        "Treatment Type",
        TREATMENT_ORDER,
        "mean_adjusted_ndre",
        "sd_adjusted_ndre",
        "Adjusted relative NDRE by treatment",
        "Mean relative NDRE",
        path,
    )

    charts[
        "treatment_ndre"
    ] = path

    path = (
        CHARTS_ROOT
        / "05_day7_observed_vs_adjusted_ndvi_by_tray.png"
    )

    save_observed_adjusted_chart(
        tray_growth,
        "day7_observed_ndvi",
        "day7_adjusted_ndvi",
        "Day 7 relative NDVI: observed vs adjusted",
        "Mean relative NDVI",
        path,
    )

    charts[
        "observed_adjusted_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "06_day7_observed_vs_adjusted_ndre_by_tray.png"
    )

    save_observed_adjusted_chart(
        tray_growth,
        "day7_observed_ndre",
        "day7_adjusted_ndre",
        "Day 7 relative NDRE: observed vs adjusted",
        "Mean relative NDRE",
        path,
    )

    charts[
        "observed_adjusted_ndre"
    ] = path

    path = (
        CHARTS_ROOT
        / "07_adjusted_ms_tray_ranking.png"
    )

    save_tray_ranking_chart(
        tray_growth,
        path,
    )

    charts[
        "tray_ranking"
    ] = path

    path = (
        CHARTS_ROOT
        / "08_inside_outside_adjusted_ndvi_ideal_moisture.png"
    )

    save_inside_outside_chart(
        inside_outside,
        "inside_day7_adjusted_ndvi",
        "outside_day7_adjusted_ndvi",
        "Ideal and Moisture: Inside vs Outside adjusted Day 7 NDVI",
        "Adjusted relative NDVI",
        path,
    )

    charts[
        "inside_outside_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "09_inside_outside_adjusted_ndre_ideal_moisture.png"
    )

    save_inside_outside_chart(
        inside_outside,
        "inside_day7_adjusted_ndre",
        "outside_day7_adjusted_ndre",
        "Ideal and Moisture: Inside vs Outside adjusted Day 7 NDRE",
        "Adjusted relative NDRE",
        path,
    )

    charts[
        "inside_outside_ndre"
    ] = path

    path = (
        CHARTS_ROOT
        / "10_heat_phase_adjusted_ndvi.png"
    )

    save_phase_trend_chart(
        heat_phase,
        "microbe_status",
        "mean_adjusted_ndvi",
        "Heat-treatment phase response: adjusted relative NDVI",
        "Mean relative NDVI",
        path,
    )

    charts[
        "heat_phase_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "11_moisture_environment_adjusted_ndvi.png"
    )

    save_phase_trend_chart(
        moisture_phase,
        "fixed_environment",
        "mean_adjusted_ndvi",
        "Moisture treatment: Inside vs Outside adjusted relative NDVI",
        "Mean relative NDVI",
        path,
    )

    charts[
        "moisture_ndvi"
    ] = path

    path = (
        CHARTS_ROOT
        / "12_day7_imputed_cells_by_tray.png"
    )

    save_imputed_cells_chart(
        tray_growth,
        path,
    )

    charts[
        "imputed_cells"
    ] = path

    return charts


# ============================================================
# 17) EXCEL OUTPUT
# ============================================================

def style_workbook(
    path: Path,
) -> None:
    workbook = load_workbook(
        path
    )

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = (
            "A2"
        )

        worksheet.auto_filter.ref = (
            worksheet.dimensions
        )

        worksheet.row_dimensions[
            1
        ].height = 34

        for cell in worksheet[1]:
            cell.fill = (
                header_fill
            )

            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )

            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = (
                column_cells[0]
                .column_letter
            )

            longest = max(
                len(
                    str(
                        cell.value
                    )
                )
                if cell.value is not None
                else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[
                letter
            ].width = min(
                max(
                    12,
                    longest + 2,
                ),
                58,
            )

    workbook.save(
        path
    )


def save_tables(
    adjusted_cell: pd.DataFrame,
    tray_daily: pd.DataFrame,
    tray_growth: pd.DataFrame,
    group_daily: pd.DataFrame,
    group_growth: pd.DataFrame,
    inside_outside: pd.DataFrame,
    heat_phase: pd.DataFrame,
    moisture_phase: pd.DataFrame,
) -> dict[str, Path]:
    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    paths = {
        "adjusted_cell": (
            REPORTS_ROOT
            / "ms_cell_indices_with_day7_adjustment.csv"
        ),
        "tray_daily": (
            REPORTS_ROOT
            / "ms_tray_daily_metrics.csv"
        ),
        "tray_growth": (
            REPORTS_ROOT
            / "ms_tray_growth_metrics.csv"
        ),
        "group_daily": (
            REPORTS_ROOT
            / "ms_group_daily_metrics.csv"
        ),
        "group_growth": (
            REPORTS_ROOT
            / "ms_group_growth_metrics.csv"
        ),
        "inside_outside": (
            REPORTS_ROOT
            / "inside_outside_comparison_ideal_moisture.csv"
        ),
        "heat_phase": (
            REPORTS_ROOT
            / "heat_phase_response.csv"
        ),
        "moisture_phase": (
            REPORTS_ROOT
            / "moisture_phase_response.csv"
        ),
        "imputed_cells": (
            REPORTS_ROOT
            / "possible_day7_bug_eaten_ms_cells.csv"
        ),
        "excel": (
            REPORTS_ROOT
            / "ms_treatment_comparison_report.xlsx"
        ),
    }

    adjusted_cell.to_csv(
        paths["adjusted_cell"],
        index=False,
    )

    tray_daily.to_csv(
        paths["tray_daily"],
        index=False,
    )

    tray_growth.to_csv(
        paths["tray_growth"],
        index=False,
    )

    group_daily.to_csv(
        paths["group_daily"],
        index=False,
    )

    group_growth.to_csv(
        paths["group_growth"],
        index=False,
    )

    inside_outside.to_csv(
        paths["inside_outside"],
        index=False,
    )

    heat_phase.to_csv(
        paths["heat_phase"],
        index=False,
    )

    moisture_phase.to_csv(
        paths["moisture_phase"],
        index=False,
    )

    imputed_cells = adjusted_cell.loc[
        adjusted_cell[
            "day7_imputed"
        ].eq(
            "Yes"
        )
    ].copy()

    imputed_cells.to_csv(
        paths["imputed_cells"],
        index=False,
    )

    readme = pd.DataFrame(
        {
            "Notes": [
                (
                    "This workbook summarises Trial 3 "
                    "relative NDVI and NDRE treatment comparisons."
                ),
                (
                    "Observed Day 7 NDVI/NDRE values are preserved."
                ),
                (
                    "Adjusted Day 7 values are stored separately "
                    "and only applied to flagged possible bug-eaten cells."
                ),
                (
                    "Adjusted values are estimated from previous valid "
                    "multispectral trends and are not direct observations."
                ),
                (
                    "Day 1 was 29/06/2026 and Day 7 was 07/07/2026."
                ),
                (
                    "Day 7 is eight elapsed days after Day 1."
                ),
                (
                    "Inside vs Outside comparisons are restricted to "
                    "Ideal and Moisture treatments."
                ),
                (
                    "Heat trays are analysed by movement phase because "
                    "their environment changed during the trial."
                ),
                (
                    "NDVI and NDRE are relative image-derived indices "
                    "unless calibrated reflectance data are available."
                ),
                (
                    "Tray performance scores are descriptive rankings, "
                    "not formal statistical significance tests."
                ),
            ]
        }
    )

    with pd.ExcelWriter(
        paths["excel"],
        engine="openpyxl",
    ) as writer:
        safe_round_dataframe(
            tray_growth
        ).to_excel(
            writer,
            sheet_name="Tray Growth Metrics",
            index=False,
        )

        safe_round_dataframe(
            tray_daily
        ).to_excel(
            writer,
            sheet_name="Tray Daily Metrics",
            index=False,
        )

        safe_round_dataframe(
            group_daily
        ).to_excel(
            writer,
            sheet_name="Group Daily Metrics",
            index=False,
        )

        safe_round_dataframe(
            group_growth
        ).to_excel(
            writer,
            sheet_name="Group Growth Metrics",
            index=False,
        )

        safe_round_dataframe(
            inside_outside
        ).to_excel(
            writer,
            sheet_name="Inside Outside Compare",
            index=False,
        )

        safe_round_dataframe(
            heat_phase
        ).to_excel(
            writer,
            sheet_name="Heat Phase Response",
            index=False,
        )

        safe_round_dataframe(
            moisture_phase
        ).to_excel(
            writer,
            sheet_name="Moisture Phase Response",
            index=False,
        )

        safe_round_dataframe(
            imputed_cells
        ).to_excel(
            writer,
            sheet_name="Day7 Adjusted Cells",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_workbook(
        paths["excel"]
    )

    return paths


# ============================================================
# 18) WORD REPORT
# ============================================================

def add_docx_table(
    document,
    dataframe: pd.DataFrame,
    columns: list[str],
    maximum_rows: int = 12,
) -> None:
    if dataframe.empty:
        document.add_paragraph(
            "No table data were available."
        )

        return

    columns = [
        column
        for column in columns
        if column in dataframe.columns
    ]

    if not columns:
        document.add_paragraph(
            "No valid columns were available."
        )

        return

    frame = safe_round_dataframe(
        dataframe[
            columns
        ].head(
            maximum_rows
        ),
        decimals=3,
    )

    table = document.add_table(
        rows=1,
        cols=len(columns),
    )

    table.style = "Table Grid"

    for index, column in enumerate(
        columns
    ):
        table.rows[0].cells[
            index
        ].text = (
            column
            .replace("_", " ")
            .title()
        )

    for _index, row in frame.iterrows():
        cells = (
            table.add_row().cells
        )

        for index, column in enumerate(
            columns
        ):
            value = row[column]

            cells[index].text = (
                ""
                if pd.isna(value)
                else str(value)
            )


def add_picture_if_exists(
    document,
    path: Path | None,
    width_inches: float = 6.3,
) -> None:
    if (
        path is None
        or not Path(path).exists()
    ):
        return

    document.add_picture(
        str(path),
        width=Inches(
            width_inches
        ),
    )


def describe_output_file(
    document,
    filename: str,
    description: str,
) -> None:
    paragraph = (
        document.add_paragraph()
    )

    paragraph.add_run(
        filename
    ).bold = True

    paragraph.add_run(
        f": {description}"
    )


def best_group_sentence(
    group_growth: pd.DataFrame,
    group_type: str,
    metric: str,
    metric_label: str,
) -> str:
    subset = group_growth.loc[
        group_growth[
            "group_type"
        ].eq(
            group_type
        )
    ].copy()

    if (
        subset.empty
        or metric not in subset.columns
    ):
        return (
            f"No valid {group_type.lower()} result was "
            f"available for {metric_label}."
        )

    subset = subset.dropna(
        subset=[
            metric,
        ]
    )

    if subset.empty:
        return (
            f"No valid {group_type.lower()} result was "
            f"available for {metric_label}."
        )

    best = subset.loc[
        subset[
            metric
        ].idxmax()
    ]

    return (
        f"The highest mean {metric_label} in the "
        f"{group_type.lower()} comparison was recorded by "
        f"{best['group']} ({format_number(best[metric])})."
    )


def create_word_report(
    path: Path,
    adjusted_cell: pd.DataFrame,
    tray_daily: pd.DataFrame,
    tray_growth: pd.DataFrame,
    group_growth: pd.DataFrame,
    inside_outside: pd.DataFrame,
    charts: dict[str, Path],
    bug_flag_source: str,
) -> Path | None:
    if not DOCX_AVAILABLE:
        print(
            "Word report skipped because python-docx "
            "is not installed."
        )

        print(
            "Install it with: pip install python-docx"
        )

        return None

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Times New Roman"

    document.styles[
        "Normal"
    ].font.size = Pt(11)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name in document.styles:
            document.styles[
                style_name
            ].font.name = "Times New Roman"

    title = document.add_heading(
        (
            "Trial 3 Multispectral Treatment and "
            "Growth Comparison Report"
        ),
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_paragraph(
        "This report summarises the Trial 3 relative NDVI and NDRE "
        "treatment-comparison workflow. Script 08 used the cell-level "
        "multispectral results from Script 07 to calculate tray-level "
        "growth rates, treatment-group comparisons, environmental "
        "comparisons, and observed-versus-adjusted Day 7 results."
    )

    document.add_paragraph(
        "The vegetation indices should be described as relative "
        "image-derived NDVI and NDRE unless calibrated reflectance "
        "data are available. The analysis is descriptive and should "
        "not be interpreted as a formal statistical significance test."
    )

    document.add_heading(
        "1. Data processed",
        level=1,
    )

    document.add_paragraph(
        f"The analysis contained {len(adjusted_cell)} cell-day rows, "
        f"{len(tray_daily)} tray-day rows and "
        f"{tray_growth['tray_no'].nunique()} unique trays."
    )

    document.add_paragraph(
        "The observation period was Day 1 on 29 June 2026 through "
        "Day 7 on 7 July 2026. Because images were not taken on "
        "5 and 6 July, Day 7 occurred eight elapsed days after the "
        "Day 1 image."
    )

    document.add_heading(
        "2. Day 7 observed and adjusted values",
        level=1,
    )

    imputed_count = int(
        adjusted_cell[
            "day7_imputed"
        ].eq(
            "Yes"
        ).sum()
    )

    document.add_paragraph(
        f"A total of {imputed_count} Day 7 cell records were given "
        f"separate adjusted NDVI/NDRE estimates. Bug-eaten-cell flags "
        f"were obtained from: {bug_flag_source}."
    )

    document.add_paragraph(
        "Observed Day 7 values were preserved without alteration. "
        "For a flagged cell, the adjusted value was estimated from "
        "the latest valid prior multispectral trend. Where at least "
        "two earlier valid observations were available, a linear "
        "trend was projected to Day 7. Otherwise, the latest valid "
        "value was carried forward. All estimated values were limited "
        "to the valid index range from -1 to 1."
    )

    add_picture_if_exists(
        document,
        charts.get(
            "observed_adjusted_ndvi"
        ),
    )

    add_picture_if_exists(
        document,
        charts.get(
            "observed_adjusted_ndre"
        ),
    )

    add_picture_if_exists(
        document,
        charts.get(
            "imputed_cells"
        ),
    )

    document.add_heading(
        "3. Tray-level multispectral results",
        level=1,
    )

    if not tray_growth.empty:
        best = tray_growth.sort_values(
            "overall_adjusted_ms_rank",
            na_position="last",
        ).iloc[0]

        document.add_paragraph(
            f"The highest descriptive adjusted multispectral score "
            f"was recorded by {best['tray']} "
            f"({best['microbe_status']} | {best['treatment']} | "
            f"{best['environment_group']}). Its score was "
            f"{format_number(best['overall_adjusted_ms_score'], 2)}."
        )

    add_docx_table(
        document,
        tray_growth.sort_values(
            "overall_adjusted_ms_rank",
            na_position="last",
        ),
        [
            "tray",
            "microbe_status",
            "treatment",
            "environment_group",
            "day7_adjusted_ndvi",
            "day7_adjusted_ndre",
            "day7_imputed_cells",
            "overall_adjusted_ms_score",
            "overall_adjusted_ms_rank",
        ],
        maximum_rows=12,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "tray_ranking"
        ),
    )

    document.add_heading(
        "4. Microbes vs No Microbes",
        level=1,
    )

    document.add_paragraph(
        best_group_sentence(
            group_growth,
            "Microbe Status",
            "mean_day7_adjusted_ndvi",
            "adjusted Day 7 NDVI",
        )
    )

    document.add_paragraph(
        best_group_sentence(
            group_growth,
            "Microbe Status",
            "mean_day7_adjusted_ndre",
            "adjusted Day 7 NDRE",
        )
    )

    microbe_table = group_growth.loc[
        group_growth[
            "group_type"
        ].eq(
            "Microbe Status"
        )
    ]

    add_docx_table(
        document,
        microbe_table,
        [
            "group",
            "tray_count",
            "mean_day7_adjusted_ndvi",
            "mean_day7_adjusted_ndre",
            "mean_adjusted_ndvi_rate_day1_to_day7_per_day",
            "mean_adjusted_ndre_rate_day1_to_day7_per_day",
            "mean_overall_adjusted_ms_score",
        ],
        maximum_rows=5,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "microbes_ndvi"
        ),
    )

    add_picture_if_exists(
        document,
        charts.get(
            "microbes_ndre"
        ),
    )

    document.add_heading(
        "5. Ideal vs Heat vs Moisture",
        level=1,
    )

    document.add_paragraph(
        best_group_sentence(
            group_growth,
            "Treatment Type",
            "mean_day7_adjusted_ndvi",
            "adjusted Day 7 NDVI",
        )
    )

    document.add_paragraph(
        best_group_sentence(
            group_growth,
            "Treatment Type",
            "mean_day7_adjusted_ndre",
            "adjusted Day 7 NDRE",
        )
    )

    treatment_table = group_growth.loc[
        group_growth[
            "group_type"
        ].eq(
            "Treatment Type"
        )
    ]

    add_docx_table(
        document,
        treatment_table,
        [
            "group",
            "tray_count",
            "mean_day7_adjusted_ndvi",
            "mean_day7_adjusted_ndre",
            "mean_adjusted_ndvi_rate_day1_to_day7_per_day",
            "mean_adjusted_ndre_rate_day1_to_day7_per_day",
            "mean_overall_adjusted_ms_score",
        ],
        maximum_rows=10,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "treatment_ndvi"
        ),
    )

    add_picture_if_exists(
        document,
        charts.get(
            "treatment_ndre"
        ),
    )

    document.add_heading(
        "6. Ideal and Moisture: Inside vs Outside",
        level=1,
    )

    document.add_paragraph(
        "The direct Inside-versus-Outside comparison was restricted "
        "to Ideal and Moisture trays because these treatments had "
        "fixed environmental placement. Heat trays were excluded "
        "from this direct comparison because they were moved between "
        "Inside and Outside during the experiment."
    )

    add_docx_table(
        document,
        inside_outside,
        [
            "treatment",
            "inside_tray_count",
            "outside_tray_count",
            "inside_trays",
            "outside_trays",
            "inside_day7_adjusted_ndvi",
            "outside_day7_adjusted_ndvi",
            "inside_minus_outside_day7_ndvi",
            "inside_day7_adjusted_ndre",
            "outside_day7_adjusted_ndre",
            "inside_minus_outside_day7_ndre",
            "interpretation",
        ],
        maximum_rows=5,
    )

    for row in inside_outside.itertuples(
        index=False
    ):
        document.add_paragraph(
            f"For {row.treatment}, the Inside-minus-Outside "
            f"adjusted Day 7 NDVI difference was "
            f"{format_number(row.inside_minus_outside_day7_ndvi)} "
            f"and the NDRE difference was "
            f"{format_number(row.inside_minus_outside_day7_ndre)}. "
            f"{row.interpretation}"
        )

    add_picture_if_exists(
        document,
        charts.get(
            "inside_outside_ndvi"
        ),
    )

    add_picture_if_exists(
        document,
        charts.get(
            "inside_outside_ndre"
        ),
    )

    document.add_heading(
        "7. Heat-treatment phase response",
        level=1,
    )

    document.add_paragraph(
        "Heat trays were Inside during the baseline period, moved "
        "Outside on Day 3, remained Outside through Days 4 and 5, "
        "and were returned Inside on Day 6. The heat-phase chart "
        "shows the adjusted relative NDVI response over this movement "
        "schedule, separated by microbial treatment."
    )

    add_picture_if_exists(
        document,
        charts.get(
            "heat_phase_ndvi"
        ),
    )

    document.add_heading(
        "8. Moisture-treatment response",
        level=1,
    )

    document.add_paragraph(
        "Moisture trays were watered on Day 0, Day 3 and Day 6. "
        "The remaining observation days represent drying intervals. "
        "The moisture chart compares the adjusted relative NDVI "
        "response of fixed Inside and fixed Outside moisture trays."
    )

    add_picture_if_exists(
        document,
        charts.get(
            "moisture_ndvi"
        ),
    )

    document.add_heading(
        "9. Description of generated CSV files",
        level=1,
    )

    output_descriptions = [
        (
            "ms_cell_indices_with_day7_adjustment.csv",
            (
                "Full cell-level multispectral table. It preserves "
                "observed NDVI/NDRE values and adds separate adjusted "
                "Day 7 values, imputation flags, slope estimates and "
                "the source of each bug-eaten-cell flag."
            ),
        ),
        (
            "ms_tray_daily_metrics.csv",
            (
                "Tray-by-day summary containing observed and adjusted "
                "mean NDVI/NDRE, valid-cell totals, treatment metadata "
                "and Day 7 adjusted-cell counts."
            ),
        ),
        (
            "ms_tray_growth_metrics.csv",
            (
                "One row per tray containing Day 1, Day 6 and Day 7 "
                "indices, observed and adjusted growth rates, index "
                "area-under-curve values, adjustment differences and "
                "the descriptive multispectral ranking."
            ),
        ),
        (
            "ms_group_daily_metrics.csv",
            (
                "Daily mean and standard-deviation results for microbial, "
                "treatment and environment-based groups. These values are "
                "used to create the trend charts."
            ),
        ),
        (
            "ms_group_growth_metrics.csv",
            (
                "Final group-level growth table for Microbes, treatments, "
                "Microbe × Treatment, Treatment × Environment and "
                "Microbe × Environment comparisons."
            ),
        ),
        (
            "inside_outside_comparison_ideal_moisture.csv",
            (
                "Dedicated fixed-environment comparison for Ideal and "
                "Moisture trays, including NDVI, NDRE and growth-rate "
                "differences."
            ),
        ),
        (
            "heat_phase_response.csv",
            (
                "Heat-tray NDVI/NDRE results organised by Inside baseline, "
                "Outside exposure and Inside recovery phases."
            ),
        ),
        (
            "moisture_phase_response.csv",
            (
                "Moisture-tray NDVI/NDRE results organised by watering "
                "events, drying intervals, environment and microbial status."
            ),
        ),
        (
            "possible_day7_bug_eaten_ms_cells.csv",
            (
                "Filtered list of Day 7 cell records that received separate "
                "adjusted multispectral estimates."
            ),
        ),
    ]

    for (
        filename,
        description,
    ) in output_descriptions:
        describe_output_file(
            document,
            filename,
            description,
        )

    document.add_heading(
        "10. Description of the Excel workbook",
        level=1,
    )

    document.add_paragraph(
        "The file ms_treatment_comparison_report.xlsx combines the "
        "main CSV outputs into one workbook. Its sheets are Tray Growth "
        "Metrics, Tray Daily Metrics, Group Daily Metrics, Group Growth "
        "Metrics, Inside Outside Compare, Heat Phase Response, Moisture "
        "Phase Response, Day7 Adjusted Cells and Read Me."
    )

    document.add_heading(
        "11. Description of generated charts",
        level=1,
    )

    chart_descriptions = [
        (
            "01_microbes_adjusted_ndvi_trend.png",
            "Adjusted relative NDVI trend for Microbes and No Microbes.",
        ),
        (
            "02_microbes_adjusted_ndre_trend.png",
            "Adjusted relative NDRE trend for Microbes and No Microbes.",
        ),
        (
            "03_treatment_adjusted_ndvi_trend.png",
            "Adjusted relative NDVI trend for Ideal, Heat and Moisture.",
        ),
        (
            "04_treatment_adjusted_ndre_trend.png",
            "Adjusted relative NDRE trend for Ideal, Heat and Moisture.",
        ),
        (
            "05_day7_observed_vs_adjusted_ndvi_by_tray.png",
            "Observed and adjusted final NDVI for every tray.",
        ),
        (
            "06_day7_observed_vs_adjusted_ndre_by_tray.png",
            "Observed and adjusted final NDRE for every tray.",
        ),
        (
            "07_adjusted_ms_tray_ranking.png",
            "Descriptive tray ranking based on adjusted NDVI/NDRE outcomes.",
        ),
        (
            "08_inside_outside_adjusted_ndvi_ideal_moisture.png",
            "Inside and Outside adjusted Day 7 NDVI for Ideal and Moisture.",
        ),
        (
            "09_inside_outside_adjusted_ndre_ideal_moisture.png",
            "Inside and Outside adjusted Day 7 NDRE for Ideal and Moisture.",
        ),
        (
            "10_heat_phase_adjusted_ndvi.png",
            "Heat-treatment NDVI response across movement phases.",
        ),
        (
            "11_moisture_environment_adjusted_ndvi.png",
            "Inside and Outside moisture-treatment NDVI response.",
        ),
        (
            "12_day7_imputed_cells_by_tray.png",
            "Number of Day 7 cells receiving adjusted MS estimates by tray.",
        ),
    ]

    for (
        filename,
        description,
    ) in chart_descriptions:
        describe_output_file(
            document,
            filename,
            description,
        )

    document.add_heading(
        "12. Interpretation limitations",
        level=1,
    )

    document.add_paragraph(
        "The NDVI and NDRE values are relative image-derived indices. "
        "They may be affected by band calibration, exposure, soil moisture, "
        "shadows, band registration, polygon placement and the proportion "
        "of soil within each cell zone."
    )

    document.add_paragraph(
        "The Day 7 adjusted results are an analytical scenario based on "
        "the assumption that previously visible seedlings disappeared "
        "because of bug damage. The observed Day 7 values must remain "
        "the primary record of what was visible in the final images."
    )

    document.add_paragraph(
        "The overall multispectral score is a descriptive normalised "
        "ranking. It is useful for comparison and visual reporting but "
        "does not constitute a formal biological or statistical test."
    )

    document.save(
        path
    )

    return path


# ============================================================
# 19) SETTINGS
# ============================================================

def save_settings(
    path: Path,
    bug_flag_source: str,
) -> Path:
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    settings = {
        "purpose": (
            "Third Trial multispectral treatment "
            "and growth comparison"
        ),
        "script07_cell_results": str(
            SCRIPT07_CELL_RESULTS
        ),
        "script07_tray_summary": str(
            SCRIPT07_TRAY_SUMMARY
        ),
        "output_root": str(
            OUTPUT_ROOT
        ),
        "expected_cells_per_tray": (
            EXPECTED_CELLS_PER_TRAY
        ),
        "expected_trays": (
            EXPECTED_TRAYS
        ),
        "expected_tray_day_rows": (
            EXPECTED_TRAY_DAY_ROWS
        ),
        "expected_cell_day_rows": (
            EXPECTED_CELL_DAY_ROWS
        ),
        "date_map": DATE_MAP,
        "days_since_day1": (
            DAYS_SINCE_DAY1
        ),
        "bug_flag_source": (
            bug_flag_source
        ),
        "day7_observed_policy": (
            "Observed NDVI and NDRE values are preserved."
        ),
        "day7_adjusted_policy": (
            "Only flagged possible bug-eaten cells receive "
            "separate adjusted Day 7 values."
        ),
        "imputation_method": (
            "Linear projection from up to the latest three valid "
            "prior observations. If fewer than two valid observations "
            "exist, the latest valid value is carried forward."
        ),
        "maximum_index_slope_per_day": (
            MAX_INDEX_SLOPE_PER_DAY
        ),
        "index_bounds": [
            -1.0,
            1.0,
        ],
        "inside_outside_policy": (
            "Direct comparison is limited to Ideal and Moisture "
            "trays because Heat trays changed environment."
        ),
        "word_report": (
            "_reports/ms_treatment_comparison_report.docx"
        ),
        "interpretation_warning": (
            "NDVI/NDRE are relative image-derived indices unless "
            "calibrated reflectance is available."
        ),
    }

    path.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    return path


# ============================================================
# 20) MAIN WORKFLOW
# ============================================================

def run_analysis(
    _args,
) -> int:
    print(
        "\nSCRIPT 08 — THIRD TRIAL "
        "MS TREATMENT COMPARISON"
    )

    print(
        "=" * 78
    )

    print(
        f"Script 07 reports:\n"
        f"{SCRIPT07_REPORTS}"
    )

    print(
        f"\nOutput folder:\n"
        f"{OUTPUT_ROOT}"
    )

    tray07, cell07 = (
        load_script07_outputs()
    )

    print(
        f"\nAccepted Script 07 tray rows: "
        f"{len(tray07)}"
    )

    print(
        f"Accepted Script 07 cell rows: "
        f"{len(cell07)}"
    )

    bug_flags, bug_flag_source = (
        load_bug_flags()
    )

    print(
        f"\nPossible bug-eaten cell flags: "
        f"{len(bug_flags)}"
    )

    print(
        f"Bug flag source:\n"
        f"{bug_flag_source}"
    )

    adjusted_cell = (
        create_adjusted_cell_table(
            cell07,
            bug_flags,
            bug_flag_source,
        )
    )

    tray_daily = (
        create_tray_daily_metrics(
            adjusted_cell
        )
    )

    tray_growth = (
        create_tray_growth_metrics(
            tray_daily
        )
    )

    (
        group_daily,
        group_growth,
    ) = create_all_group_tables(
        tray_daily,
        tray_growth,
    )

    inside_outside = (
        create_inside_outside_comparison(
            tray_growth
        )
    )

    (
        heat_phase,
        moisture_phase,
    ) = create_phase_tables(
        tray_daily
    )

    print(
        "\nInside vs Outside comparison:"
    )

    print(
        inside_outside[
            [
                "treatment",
                "inside_trays",
                "outside_trays",
                "inside_day7_adjusted_ndvi",
                "outside_day7_adjusted_ndvi",
                "inside_day7_adjusted_ndre",
                "outside_day7_adjusted_ndre",
            ]
        ].to_string(
            index=False
        )
    )

    charts = create_charts(
        group_daily,
        tray_growth,
        inside_outside,
        heat_phase,
        moisture_phase,
    )

    table_paths = save_tables(
        adjusted_cell,
        tray_daily,
        tray_growth,
        group_daily,
        group_growth,
        inside_outside,
        heat_phase,
        moisture_phase,
    )

    word_path = (
        REPORTS_ROOT
        / "ms_treatment_comparison_report.docx"
    )

    create_word_report(
        word_path,
        adjusted_cell,
        tray_daily,
        tray_growth,
        group_growth,
        inside_outside,
        charts,
        bug_flag_source,
    )

    settings_path = save_settings(
        CONFIG_ROOT
        / "ms_treatment_comparison_settings.json",
        bug_flag_source,
    )

    adjusted_count = int(
        adjusted_cell[
            "day7_imputed"
        ].eq(
            "Yes"
        ).sum()
    )

    print(
        "\n"
        + "=" * 78
    )

    print(
        "SCRIPT 08 FINISHED"
    )

    print(
        "=" * 78
    )

    print(
        f"Tray-day rows: "
        f"{len(tray_daily)} "
        f"(expected {EXPECTED_TRAY_DAY_ROWS})"
    )

    print(
        f"Cell-day rows: "
        f"{len(adjusted_cell)} "
        f"(expected {EXPECTED_CELL_DAY_ROWS})"
    )

    print(
        f"Day 7 adjusted cell records: "
        f"{adjusted_count}"
    )

    if not tray_growth.empty:
        best = tray_growth.sort_values(
            "overall_adjusted_ms_rank",
            na_position="last",
        ).iloc[0]

        print(
            "\nTop adjusted multispectral tray:"
        )

        print(
            f"  {best['tray']} — "
            f"{best['microbe_status']} | "
            f"{best['treatment']} | "
            f"{best['environment_group']}"
        )

        print(
            f"  Score: "
            f"{best['overall_adjusted_ms_score']:.2f}"
        )

    print(
        "\nMain output files:"
    )

    print(
        f"\nAdjusted cell table:\n"
        f"{table_paths['adjusted_cell']}"
    )

    print(
        f"\nTray daily metrics:\n"
        f"{table_paths['tray_daily']}"
    )

    print(
        f"\nTray growth metrics:\n"
        f"{table_paths['tray_growth']}"
    )

    print(
        f"\nGroup daily metrics:\n"
        f"{table_paths['group_daily']}"
    )

    print(
        f"\nGroup growth metrics:\n"
        f"{table_paths['group_growth']}"
    )

    print(
        f"\nInside vs Outside comparison:\n"
        f"{table_paths['inside_outside']}"
    )

    print(
        f"\nHeat phase response:\n"
        f"{table_paths['heat_phase']}"
    )

    print(
        f"\nMoisture phase response:\n"
        f"{table_paths['moisture_phase']}"
    )

    print(
        f"\nAdjusted Day 7 cells:\n"
        f"{table_paths['imputed_cells']}"
    )

    print(
        f"\nExcel report:\n"
        f"{table_paths['excel']}"
    )

    if DOCX_AVAILABLE:
        print(
            f"\nWord comparison report:\n"
            f"{word_path}"
        )

    else:
        print(
            "\nWord report was skipped because "
            "python-docx is not installed."
        )

    print(
        f"\nSettings:\n"
        f"{settings_path}"
    )

    print(
        f"\nCharts folder:\n"
        f"{CHARTS_ROOT}"
    )

    return 0


# ============================================================
# 21) CLI
# ============================================================

def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Trial 3 Script 08: multispectral "
            "treatment and growth comparison."
        )
    )

    args = parser.parse_args()

    return run_analysis(
        args
    )


if __name__ == "__main__":
    raise SystemExit(
        main()
    )