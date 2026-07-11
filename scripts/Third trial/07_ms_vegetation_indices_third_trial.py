from __future__ import annotations

"""
SCRIPT 07 — THIRD TRIAL MULTISPECTRAL VEGETATION INDICES

Calculates relative image-derived NDVI and NDRE from the original cropped
multispectral TIFF bands using the accepted 70-cell polygon grid from Script 06.
Observed Day 7 values are preserved; no bug-eaten/missing-plant imputation is
performed in this script.
"""

import argparse
import json
import math
import re
from pathlib import Path

import cv2
import numpy as np
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image, ImageOps

PROJECT_ROOT = Path(r"C:\Users\tshib\OneDrive\Desktop\Internship")

CROP_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "01_Crop_Dual_Reference"
)

GRID_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "06_MS_Cell_Grid_Detection"
)

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "07_MS_Vegetation_Indices"
)

GRID_MANIFEST_CSV = (
    GRID_ROOT
    / "_reports"
    / "ms_grid_manifest.csv"
)

GRID_COORDINATES_CSV = (
    GRID_ROOT
    / "_reports"
    / "ms_cell_coordinates.csv"
)

REPORTS_ROOT = OUTPUT_ROOT / "_reports"
CONFIG_ROOT = OUTPUT_ROOT / "_config"

EXPECTED_CELLS = 70
EXPECTED_IMAGE_SETS = 84

ACCEPTED_GRID_STATUSES = {
    "PASS_AUTO",
    "PASS_MANUAL",
}

MIN_VALID_PIXELS_PER_CELL = 20
MIN_CHECK_VALID_CELLS = 65
EPSILON = 1e-6

DATE_MAP = {
    1: "2026-06-29",
    2: "2026-06-30",
    3: "2026-07-01",
    4: "2026-07-02",
    5: "2026-07-03",
    6: "2026-07-04",
    7: "2026-07-07",
}

DAYS_SINCE_DAY1 = {
    1: 0,
    2: 1,
    3: 2,
    4: 3,
    5: 4,
    6: 5,
    7: 8,
}

DAYS_SINCE_PLANTING = {
    1: 2,
    2: 3,
    3: 4,
    4: 5,
    5: 6,
    6: 7,
    7: 10,
}

DAYS_SINCE_PREVIOUS = {
    1: 0,
    2: 1,
    3: 1,
    4: 1,
    5: 1,
    6: 1,
    7: 3,
}

BAND_COLUMNS = {
    "MS_G": "ms_g_path",
    "MS_R": "ms_r_path",
    "MS_RE": "ms_re_path",
    "MS_NIR": "ms_nir_path",
}


try:
    import tifffile

    TIFFFILE_AVAILABLE = True
except Exception:
    TIFFFILE_AVAILABLE = False


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r"(\d+)", str(value))
    ]


def parse_filter(value: str | None):
    if not value:
        return None

    return {
        item.strip().casefold()
        for item in value.split(",")
        if item.strip()
    }


def safe_name(value: object) -> str:
    return re.sub(
        r"[^A-Za-z0-9_.-]+",
        "_",
        str(value),
    )


def relative_path(path: Path) -> str:
    try:
        return str(path.relative_to(OUTPUT_ROOT))
    except ValueError:
        return str(path)


def resolve_path(value: object) -> Path | None:
    text = str(value).strip()

    if not text or text.casefold() == "nan":
        return None

    path = Path(
        text.replace("\\", "/")
    )

    if path.is_absolute():
        return path

    candidates = (
        CROP_ROOT / path,
        PROJECT_ROOT / path,
        GRID_ROOT / path,
    )

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return CROP_ROOT / path


def metadata(day_order: int) -> dict:
    return {
        "calendar_date": DATE_MAP.get(
            day_order,
            "",
        ),
        "days_since_day1": DAYS_SINCE_DAY1.get(
            day_order,
            math.nan,
        ),
        "days_since_planting": DAYS_SINCE_PLANTING.get(
            day_order,
            math.nan,
        ),
        "days_since_previous_photo": DAYS_SINCE_PREVIOUS.get(
            day_order,
            math.nan,
        ),
    }


def require_columns(
    frame: pd.DataFrame,
    columns: list[str],
    source: Path,
) -> None:
    missing = [
        column
        for column in columns
        if column not in frame.columns
    ]

    if missing:
        raise ValueError(
            f"{source.name} is missing: "
            + ", ".join(missing)
        )


def load_inputs() -> tuple[pd.DataFrame, pd.DataFrame]:
    if (
        not GRID_MANIFEST_CSV.exists()
        or not GRID_COORDINATES_CSV.exists()
    ):
        raise FileNotFoundError(
            "Script 06 manifest or coordinate CSV was not found."
        )

    manifest = pd.read_csv(
        GRID_MANIFEST_CSV
    )

    coordinates = pd.read_csv(
        GRID_COORDINATES_CSV
    )

    require_columns(
        manifest,
        [
            "day_order",
            "day",
            "tray",
            "tray_no",
            "capture_id",
            "status",
        ],
        GRID_MANIFEST_CSV,
    )

    require_columns(
        coordinates,
        [
            "day_order",
            "day",
            "tray",
            "tray_no",
            "capture_id",
            "cell_id",
            "row",
            "column",
            "x",
            "y",
            "grid_status",
            *BAND_COLUMNS.values(),
        ],
        GRID_COORDINATES_CSV,
    )

    polygon_columns = [
        "poly_tl_x",
        "poly_tl_y",
        "poly_tr_x",
        "poly_tr_y",
        "poly_br_x",
        "poly_br_y",
        "poly_bl_x",
        "poly_bl_y",
    ]

    if not all(
        column in coordinates.columns
        for column in polygon_columns
    ):
        require_columns(
            coordinates,
            [
                "square_x0",
                "square_y0",
                "square_x1",
                "square_y1",
            ],
            GRID_COORDINATES_CSV,
        )

        coordinates["poly_tl_x"] = coordinates["square_x0"]
        coordinates["poly_tl_y"] = coordinates["square_y0"]

        coordinates["poly_tr_x"] = coordinates["square_x1"]
        coordinates["poly_tr_y"] = coordinates["square_y0"]

        coordinates["poly_br_x"] = coordinates["square_x1"]
        coordinates["poly_br_y"] = coordinates["square_y1"]

        coordinates["poly_bl_x"] = coordinates["square_x0"]
        coordinates["poly_bl_y"] = coordinates["square_y1"]

    for frame in (
        manifest,
        coordinates,
    ):
        frame["day_order"] = pd.to_numeric(
            frame["day_order"],
            errors="coerce",
        )

        frame["tray_no"] = pd.to_numeric(
            frame["tray_no"],
            errors="coerce",
        )

    numeric_coordinate_columns = [
        "cell_id",
        "row",
        "column",
        "x",
        "y",
        *polygon_columns,
    ]

    for column in numeric_coordinate_columns:
        coordinates[column] = pd.to_numeric(
            coordinates[column],
            errors="coerce",
        )

    manifest = manifest.dropna(
        subset=[
            "day_order",
            "tray_no",
        ]
    ).copy()

    coordinates = coordinates.dropna(
        subset=[
            "day_order",
            "tray_no",
            "cell_id",
            "row",
            "column",
            "x",
            "y",
            *polygon_columns,
        ]
    ).copy()

    manifest[
        [
            "day_order",
            "tray_no",
        ]
    ] = manifest[
        [
            "day_order",
            "tray_no",
        ]
    ].astype(int)

    coordinates[
        [
            "day_order",
            "tray_no",
            "cell_id",
            "row",
            "column",
        ]
    ] = coordinates[
        [
            "day_order",
            "tray_no",
            "cell_id",
            "row",
            "column",
        ]
    ].astype(int)

    manifest["status"] = (
        manifest["status"]
        .astype(str)
        .str.upper()
    )

    coordinates["grid_status"] = (
        coordinates["grid_status"]
        .astype(str)
        .str.upper()
    )

    manifest = manifest.loc[
        manifest["status"].isin(
            ACCEPTED_GRID_STATUSES
        )
    ].copy()

    coordinates = coordinates.loc[
        coordinates["grid_status"].isin(
            ACCEPTED_GRID_STATUSES
        )
    ].copy()

    for frame in (
        manifest,
        coordinates,
    ):
        frame["grid_key"] = (
            frame[
                [
                    "day",
                    "tray",
                    "capture_id",
                ]
            ]
            .astype(str)
            .agg("|".join, axis=1)
        )

    accepted_keys = set(
        manifest["grid_key"]
    )

    coordinates = coordinates.loc[
        coordinates["grid_key"].isin(
            accepted_keys
        )
    ].copy()

    return manifest, coordinates


def read_band(path: Path) -> np.ndarray:
    if not path.exists():
        raise FileNotFoundError(
            f"Band not found: {path}"
        )

    if TIFFFILE_AVAILABLE:
        array = tifffile.imread(
            str(path)
        )
    else:
        with Image.open(path) as image:
            array = np.asarray(
                ImageOps.exif_transpose(image)
            )

    if array.ndim == 2:
        return array.astype(
            np.float32
        )

    if array.ndim == 3:
        return (
            array[:, :, :3]
            .astype(np.float32)
            .mean(axis=2)
        )

    raise ValueError(
        f"Unsupported band dimensions: {array.shape}"
    )


def calculate_indices(
    red: np.ndarray,
    red_edge: np.ndarray,
    nir: np.ndarray,
):
    if (
        red.shape != nir.shape
        or red_edge.shape != nir.shape
    ):
        raise ValueError(
            "Band sizes differ: "
            f"R={red.shape}, "
            f"RE={red_edge.shape}, "
            f"NIR={nir.shape}"
        )

    ndvi_denominator = nir + red
    ndre_denominator = nir + red_edge

    ndvi_valid = (
        np.isfinite(nir)
        & np.isfinite(red)
        & (
            ndvi_denominator
            > EPSILON
        )
    )

    ndre_valid = (
        np.isfinite(nir)
        & np.isfinite(red_edge)
        & (
            ndre_denominator
            > EPSILON
        )
    )

    ndvi = np.full(
        nir.shape,
        np.nan,
        dtype=np.float32,
    )

    ndre = np.full(
        nir.shape,
        np.nan,
        dtype=np.float32,
    )

    ndvi[ndvi_valid] = (
        nir[ndvi_valid]
        - red[ndvi_valid]
    ) / ndvi_denominator[ndvi_valid]

    ndre[ndre_valid] = (
        nir[ndre_valid]
        - red_edge[ndre_valid]
    ) / ndre_denominator[ndre_valid]

    return (
        ndvi,
        ndre,
        ndvi_valid,
        ndre_valid,
    )


def polygon_from_row(
    row: object,
) -> np.ndarray:
    return np.array(
        [
            [
                row.poly_tl_x,
                row.poly_tl_y,
            ],
            [
                row.poly_tr_x,
                row.poly_tr_y,
            ],
            [
                row.poly_br_x,
                row.poly_br_y,
            ],
            [
                row.poly_bl_x,
                row.poly_bl_y,
            ],
        ],
        dtype=np.float32,
    )


def polygon_values(
    index_array: np.ndarray,
    polygon: np.ndarray,
):
    height, width = index_array.shape

    polygon[:, 0] = np.clip(
        polygon[:, 0],
        0,
        width - 1,
    )

    polygon[:, 1] = np.clip(
        polygon[:, 1],
        0,
        height - 1,
    )

    x0 = max(
        0,
        int(
            np.floor(
                polygon[:, 0].min()
            )
        ),
    )

    y0 = max(
        0,
        int(
            np.floor(
                polygon[:, 1].min()
            )
        ),
    )

    x1 = min(
        width,
        int(
            np.ceil(
                polygon[:, 0].max()
            )
        ) + 1,
    )

    y1 = min(
        height,
        int(
            np.ceil(
                polygon[:, 1].max()
            )
        ) + 1,
    )

    if (
        x1 <= x0
        or y1 <= y0
    ):
        return (
            np.array(
                [],
                dtype=np.float32,
            ),
            0,
        )

    local_polygon = polygon.copy()

    local_polygon[:, 0] -= x0
    local_polygon[:, 1] -= y0

    mask = np.zeros(
        (
            y1 - y0,
            x1 - x0,
        ),
        dtype=np.uint8,
    )

    cv2.fillPoly(
        mask,
        [
            np.round(
                local_polygon
            ).astype(np.int32)
        ],
        1,
    )

    region = index_array[
        y0:y1,
        x0:x1,
    ]

    valid = (
        (mask > 0)
        & np.isfinite(region)
    )

    values = region[
        valid
    ].astype(np.float32)

    return (
        values,
        int(mask.sum()),
    )


def calculate_statistics(
    values: np.ndarray,
    zone_pixels: int,
) -> dict:
    values = values[
        np.isfinite(values)
    ]

    valid_pixels = len(values)

    base = {
        "zone_pixels": zone_pixels,
        "valid_pixels": valid_pixels,
        "valid_fraction": (
            valid_pixels / zone_pixels
            if zone_pixels
            else math.nan
        ),
    }

    if (
        valid_pixels
        < MIN_VALID_PIXELS_PER_CELL
    ):
        return {
            **base,
            "mean": math.nan,
            "median": math.nan,
            "std": math.nan,
            "p10": math.nan,
            "p90": math.nan,
            "positive_fraction": math.nan,
        }

    return {
        **base,
        "mean": float(
            values.mean()
        ),
        "median": float(
            np.median(values)
        ),
        "std": float(
            values.std()
        ),
        "p10": float(
            np.percentile(
                values,
                10,
            )
        ),
        "p90": float(
            np.percentile(
                values,
                90,
            )
        ),
        "positive_fraction": float(
            np.mean(
                values > 0
            )
        ),
    }


def evaluate_cells(
    ndvi: np.ndarray,
    ndre: np.ndarray,
    grid: pd.DataFrame,
) -> list[dict]:
    output = []

    for row in (
        grid.sort_values("cell_id")
        .itertuples(index=False)
    ):
        polygon = polygon_from_row(
            row
        )

        ndvi_values, ndvi_zone = polygon_values(
            ndvi,
            polygon.copy(),
        )

        ndre_values, ndre_zone = polygon_values(
            ndre,
            polygon.copy(),
        )

        ndvi_statistics = calculate_statistics(
            ndvi_values,
            ndvi_zone,
        )

        ndre_statistics = calculate_statistics(
            ndre_values,
            ndre_zone,
        )

        output.append(
            {
                "cell_id": int(
                    row.cell_id
                ),
                "row": int(
                    row.row
                ),
                "column": int(
                    row.column
                ),
                "x": float(
                    row.x
                ),
                "y": float(
                    row.y
                ),
                "poly_tl_x": float(
                    row.poly_tl_x
                ),
                "poly_tl_y": float(
                    row.poly_tl_y
                ),
                "poly_tr_x": float(
                    row.poly_tr_x
                ),
                "poly_tr_y": float(
                    row.poly_tr_y
                ),
                "poly_br_x": float(
                    row.poly_br_x
                ),
                "poly_br_y": float(
                    row.poly_br_y
                ),
                "poly_bl_x": float(
                    row.poly_bl_x
                ),
                "poly_bl_y": float(
                    row.poly_bl_y
                ),
                "coordinate_source": str(
                    getattr(
                        row,
                        "coordinate_source",
                        "",
                    )
                ),
                "grid_status": str(
                    getattr(
                        row,
                        "grid_status",
                        "",
                    )
                ),
                "needs_review": str(
                    getattr(
                        row,
                        "needs_review",
                        "",
                    )
                ),
                **{
                    f"ndvi_{key}": value
                    for key, value
                    in ndvi_statistics.items()
                },
                **{
                    f"ndre_{key}": value
                    for key, value
                    in ndre_statistics.items()
                },
            }
        )

    return output


def colour_index(
    array: np.ndarray,
) -> np.ndarray:
    valid = np.isfinite(
        array
    )

    scaled = np.zeros(
        array.shape,
        dtype=np.uint8,
    )

    clipped = np.clip(
        array,
        -1,
        1,
    )

    scaled[valid] = np.round(
        (
            clipped[valid]
            + 1
        )
        * 127.5
    ).astype(np.uint8)

    image = cv2.applyColorMap(
        scaled,
        cv2.COLORMAP_TURBO,
    )

    image[~valid] = 0

    return image


def add_header(
    image: np.ndarray,
    text: str,
) -> np.ndarray:
    output = np.full(
        (
            image.shape[0] + 48,
            image.shape[1],
            3,
        ),
        255,
        dtype=np.uint8,
    )

    output[48:] = image

    cv2.putText(
        output,
        text,
        (12, 31),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.62,
        (0, 0, 0),
        2,
        cv2.LINE_AA,
    )

    return output


def save_preview(
    array: np.ndarray,
    path: Path,
    title: str,
) -> None:
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    cv2.imwrite(
        str(path),
        add_header(
            colour_index(array),
            title,
        ),
        [
            cv2.IMWRITE_PNG_COMPRESSION,
            3,
        ],
    )


def save_overlay(
    array: np.ndarray,
    cells: list[dict],
    path: Path,
    title: str,
) -> None:
    image = colour_index(
        array
    )

    for cell in cells:
        polygon = np.array(
            [
                [
                    cell["poly_tl_x"],
                    cell["poly_tl_y"],
                ],
                [
                    cell["poly_tr_x"],
                    cell["poly_tr_y"],
                ],
                [
                    cell["poly_br_x"],
                    cell["poly_br_y"],
                ],
                [
                    cell["poly_bl_x"],
                    cell["poly_bl_y"],
                ],
            ],
            dtype=np.int32,
        )

        colour = (
            (0, 180, 255)
            if str(
                cell["needs_review"]
            ).casefold()
            == "yes"
            else (0, 180, 0)
        )

        cv2.polylines(
            image,
            [polygon],
            True,
            colour,
            2,
            cv2.LINE_AA,
        )

        cv2.putText(
            image,
            str(
                cell["cell_id"]
            ),
            (
                max(
                    0,
                    int(cell["x"]) - 7,
                ),
                max(
                    12,
                    int(cell["y"]) + 5,
                ),
            ),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.4,
            colour,
            1,
            cv2.LINE_AA,
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    cv2.imwrite(
        str(path),
        add_header(
            image,
            title,
        ),
        [
            cv2.IMWRITE_PNG_COMPRESSION,
            3,
        ],
    )


def style_excel(
    path: Path,
) -> None:
    workbook = load_workbook(
        path
    )

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions

        for cell in worksheet[1]:
            cell.fill = header_fill

            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )

            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column in worksheet.columns:
            letter = (
                column[0]
                .column_letter
            )

            longest = max(
                len(
                    str(cell.value)
                )
                if cell.value is not None
                else 0
                for cell in column
            )

            worksheet.column_dimensions[
                letter
            ].width = min(
                max(
                    12,
                    longest + 2,
                ),
                58,
            )

    workbook.save(
        path
    )


def create_excel(
    path: Path,
    tray_rows: list[dict],
    cell_rows: list[dict],
) -> None:
    tray_frame = pd.DataFrame(
        tray_rows
    )

    cell_frame = pd.DataFrame(
        cell_rows
    )

    if not tray_frame.empty:
        status_summary = (
            tray_frame["status"]
            .value_counts()
            .rename_axis("status")
            .reset_index(name="count")
        )
    else:
        status_summary = pd.DataFrame()

    readme = pd.DataFrame(
        {
            "Notes": [
                "NDVI = (MS_NIR - MS_R) / (MS_NIR + MS_R).",
                "NDRE = (MS_NIR - MS_RE) / (MS_NIR + MS_RE).",
                "Original TIFF values are used; normalisation is only for previews.",
                "Script 06 polygon zones are used for cell extraction.",
                "Day 1 = 29/06/2026.",
                "Day 7 = 07/07/2026 and is eight elapsed days after Day 1.",
                "Observed Day 7 values are preserved.",
                "No bug-eaten plant imputation occurs in Script 07.",
                "Report as relative image-derived indices unless calibrated reflectance is available.",
            ]
        }
    )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with pd.ExcelWriter(
        path,
        engine="openpyxl",
    ) as writer:
        tray_frame.to_excel(
            writer,
            sheet_name="Tray Summary",
            index=False,
        )

        cell_frame.to_excel(
            writer,
            sheet_name="Cell Results",
            index=False,
        )

        status_summary.to_excel(
            writer,
            sheet_name="Status Summary",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_excel(
        path
    )


def create_word(
    path: Path,
    tray_rows: list[dict],
    cell_rows: list[dict],
) -> None:
    if not DOCX_AVAILABLE:
        print(
            "Word report skipped. Install python-docx: "
            "pip install python-docx"
        )
        return

    tray_frame = pd.DataFrame(
        tray_rows
    )

    cell_frame = pd.DataFrame(
        cell_rows
    )

    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Times New Roman"

    document.styles[
        "Normal"
    ].font.size = Pt(11)

    title = document.add_heading(
        "Trial 3 Script 07 Work Process Report — "
        "Multispectral Vegetation Indices",
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_paragraph(
        "Script 07 calculated relative image-derived NDVI and "
        "NDRE from original cropped multispectral TIFF values "
        "and the accepted independent polygon grid from Script 06."
    )

    document.add_heading(
        "1. Inputs and formulas",
        level=1,
    )

    document.add_paragraph(
        f"Grid manifest: {GRID_MANIFEST_CSV}"
    )

    document.add_paragraph(
        f"Cell coordinates: {GRID_COORDINATES_CSV}"
    )

    document.add_paragraph(
        "NDVI = (MS_NIR - MS_R) / (MS_NIR + MS_R). "
        "NDRE = (MS_NIR - MS_RE) / (MS_NIR + MS_RE). "
        "Invalid or near-zero denominators were excluded."
    )

    document.add_heading(
        "2. Processing method",
        level=1,
    )

    document.add_paragraph(
        "The script verified all four multispectral bands, "
        "calculated full-image NDVI and NDRE, and extracted "
        "valid pixels from each of the 70 Script 06 polygons. "
        "Cell outputs include mean, median, standard deviation, "
        "percentiles, valid-pixel fraction and positive-index fraction."
    )

    document.add_heading(
        "3. Timing and Day 7",
        level=1,
    )

    document.add_paragraph(
        "Day 1 was 29/06/2026, Day 6 was 04/07/2026 and "
        "Day 7 was 07/07/2026. Day 7 is eight elapsed days "
        "after Day 1. Observed Day 7 values are preserved. "
        "Possible bug-eaten plants will be handled later as a "
        "separate flagged adjusted scenario."
    )

    document.add_heading(
        "4. Status summary",
        level=1,
    )

    if not tray_frame.empty:
        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.style = "Table Grid"

        table.rows[0].cells[0].text = "Status"
        table.rows[0].cells[1].text = "Count"

        for status, count in (
            tray_frame["status"]
            .value_counts()
            .items()
        ):
            row = table.add_row().cells

            row[0].text = str(status)
            row[1].text = str(count)

    document.add_paragraph(
        f"Tray/image rows: {len(tray_frame)}. "
        f"Cell rows: {len(cell_frame)}."
    )

    document.add_heading(
        "5. CSV, Excel and visual outputs",
        level=1,
    )

    output_descriptions = [
        (
            "ms_index_tray_summary.csv",
            "One row per day, tray and capture with NDVI/NDRE "
            "summaries, valid-cell counts, timing metadata and output paths.",
        ),
        (
            "ms_index_cell_results.csv",
            "One row per cell with polygon coordinates and "
            "detailed NDVI/NDRE statistics.",
        ),
        (
            "ms_vegetation_index_report.xlsx",
            "Tray Summary, Cell Results, Status Summary and Read Me sheets.",
        ),
        (
            "relative_indices.npz",
            "Full NDVI, NDRE and validity arrays.",
        ),
        (
            "NDVI/NDRE previews and polygon overlays",
            "Fixed-scale visual checks for the image and all 70 cells.",
        ),
    ]

    for filename, description in output_descriptions:
        paragraph = document.add_paragraph()

        paragraph.add_run(
            filename
        ).bold = True

        paragraph.add_run(
            f": {description}"
        )

    document.add_heading(
        "6. Interpretation and limitations",
        level=1,
    )

    document.add_paragraph(
        "Unless the source bands are calibrated reflectance, "
        "report the results as relative image-derived indices. "
        "Soil, shadows, moisture, exposure, band alignment and "
        "polygon placement can affect the values. Review overlays "
        "and valid-pixel counts before treatment comparison."
    )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        path
    )


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Trial 3 Script 07: relative NDVI/NDRE "
            "using Script 06 polygons."
        )
    )

    parser.add_argument(
        "--days",
        help='Example: --days "Day 1,Day 7"',
    )

    parser.add_argument(
        "--trays",
        help='Example: --trays "Tray 1,Tray 12"',
    )

    parser.add_argument(
        "--overwrite",
        action="store_true",
    )

    parser.add_argument(
        "--dry-run",
        action="store_true",
    )

    args = parser.parse_args()

    manifest, coordinates = load_inputs()

    day_filter = parse_filter(
        args.days
    )

    tray_filter = parse_filter(
        args.trays
    )

    if day_filter:
        manifest = manifest.loc[
            manifest["day"]
            .astype(str)
            .str.casefold()
            .isin(day_filter)
        ].copy()

    if tray_filter:
        manifest = manifest.loc[
            manifest["tray"]
            .astype(str)
            .str.casefold()
            .isin(tray_filter)
        ].copy()

    if manifest.empty:
        print(
            "No accepted Script 06 records matched the filters."
        )
        return 1

    coordinate_groups = {
        key: group.copy()
        for key, group
        in coordinates.groupby(
            "grid_key"
        )
    }

    print(
        "\nSCRIPT 07 — THIRD TRIAL MS VEGETATION INDICES"
    )

    print(
        "=" * 72
    )

    for job in manifest.sort_values(
        [
            "day_order",
            "tray_no",
        ]
    ).itertuples(index=False):
        cell_count = len(
            coordinate_groups.get(
                job.grid_key,
                [],
            )
        )

        state = (
            "READY"
            if cell_count == EXPECTED_CELLS
            else "INVALID"
        )

        print(
            f"{state}: "
            f"{job.day} > "
            f"{job.tray} > "
            f"{job.capture_id}"
        )

    if args.dry_run:
        return 0

    tray_rows = []
    cell_rows = []

    jobs = manifest.sort_values(
        [
            "day_order",
            "tray_no",
            "capture_id",
        ]
    )

    for job in jobs.itertuples(
        index=False
    ):
        day_order = int(
            job.day_order
        )

        day = str(
            job.day
        )

        tray = str(
            job.tray
        )

        tray_no = int(
            job.tray_no
        )

        capture_id = str(
            job.capture_id
        )

        base = {
            "day_order": day_order,
            "day": day,
            **metadata(day_order),
            "tray": tray,
            "tray_no": tray_no,
            "capture_id": capture_id,
            "grid_status": str(
                job.status
            ),
        }

        grid = coordinate_groups.get(
            job.grid_key
        )

        if (
            grid is None
            or len(grid)
            != EXPECTED_CELLS
        ):
            tray_rows.append(
                {
                    **base,
                    "cells_processed": 0,
                    "cells_with_valid_ndvi": 0,
                    "cells_with_valid_ndre": 0,
                    "status": "FAIL",
                    "notes": (
                        "Grid was not 70 cells."
                    ),
                }
            )

            continue

        output_folder = (
            OUTPUT_ROOT
            / day
            / tray
        )

        safe_capture = safe_name(
            capture_id
        )

        output_paths = {
            "npz": (
                output_folder
                / f"{safe_capture}_relative_indices.npz"
            ),
            "ndvi_preview": (
                output_folder
                / f"{safe_capture}_ndvi_preview.png"
            ),
            "ndre_preview": (
                output_folder
                / f"{safe_capture}_ndre_preview.png"
            ),
            "ndvi_overlay": (
                output_folder
                / f"{safe_capture}_ndvi_polygon_overlay.png"
            ),
            "ndre_overlay": (
                output_folder
                / f"{safe_capture}_ndre_polygon_overlay.png"
            ),
        }

        try:
            reuse_existing = (
                not args.overwrite
                and all(
                    path.exists()
                    for path
                    in output_paths.values()
                )
            )

            if reuse_existing:
                saved = np.load(
                    output_paths["npz"]
                )

                ndvi = saved["ndvi"]
                ndre = saved["ndre"]

                print(
                    f"REUSED: "
                    f"{day} > "
                    f"{tray} > "
                    f"{capture_id}"
                )

            else:
                first_cell = grid.iloc[0]

                band_paths = {
                    band: resolve_path(
                        first_cell[column]
                    )
                    for band, column
                    in BAND_COLUMNS.items()
                }

                missing_bands = [
                    band
                    for band, path
                    in band_paths.items()
                    if (
                        path is None
                        or not path.exists()
                    )
                ]

                if missing_bands:
                    raise FileNotFoundError(
                        "Missing bands: "
                        + ", ".join(
                            missing_bands
                        )
                    )

                bands = {
                    band: read_band(path)
                    for band, path
                    in band_paths.items()
                }

                dimensions = {
                    array.shape
                    for array
                    in bands.values()
                }

                if len(dimensions) != 1:
                    raise ValueError(
                        "Band dimensions do not match."
                    )

                (
                    ndvi,
                    ndre,
                    ndvi_valid,
                    ndre_valid,
                ) = calculate_indices(
                    bands["MS_R"],
                    bands["MS_RE"],
                    bands["MS_NIR"],
                )

                output_folder.mkdir(
                    parents=True,
                    exist_ok=True,
                )

                np.savez_compressed(
                    output_paths["npz"],
                    ndvi=ndvi,
                    ndre=ndre,
                    ndvi_valid_mask=ndvi_valid,
                    ndre_valid_mask=ndre_valid,
                )

            cells = evaluate_cells(
                ndvi,
                ndre,
                grid,
            )

            cell_frame = pd.DataFrame(
                cells
            )

            if not reuse_existing:
                title = (
                    f"{day} | "
                    f"{tray} | "
                    f"{capture_id}"
                )

                save_preview(
                    ndvi,
                    output_paths["ndvi_preview"],
                    (
                        f"{title} | "
                        "Relative image-derived NDVI"
                    ),
                )

                save_preview(
                    ndre,
                    output_paths["ndre_preview"],
                    (
                        f"{title} | "
                        "Relative image-derived NDRE"
                    ),
                )

                save_overlay(
                    ndvi,
                    cells,
                    output_paths["ndvi_overlay"],
                    (
                        f"{title} | "
                        "Relative NDVI | "
                        "70 polygons"
                    ),
                )

                save_overlay(
                    ndre,
                    cells,
                    output_paths["ndre_overlay"],
                    (
                        f"{title} | "
                        "Relative NDRE | "
                        "70 polygons"
                    ),
                )

            valid_ndvi_cells = int(
                cell_frame[
                    "ndvi_mean"
                ]
                .notna()
                .sum()
            )

            valid_ndre_cells = int(
                cell_frame[
                    "ndre_mean"
                ]
                .notna()
                .sum()
            )

            if (
                valid_ndvi_cells
                == EXPECTED_CELLS
                and valid_ndre_cells
                == EXPECTED_CELLS
            ):
                status = "PASS"

            elif (
                valid_ndvi_cells
                >= MIN_CHECK_VALID_CELLS
                and valid_ndre_cells
                >= MIN_CHECK_VALID_CELLS
            ):
                status = "CHECK"

            else:
                status = "FAIL"

            tray_rows.append(
                {
                    **base,
                    "cells_processed": len(cells),
                    "cells_with_valid_ndvi": valid_ndvi_cells,
                    "cells_with_valid_ndre": valid_ndre_cells,
                    "mean_cell_ndvi": float(
                        cell_frame[
                            "ndvi_mean"
                        ].mean()
                    ),
                    "median_cell_ndvi": float(
                        cell_frame[
                            "ndvi_median"
                        ].median()
                    ),
                    "mean_cell_ndre": float(
                        cell_frame[
                            "ndre_mean"
                        ].mean()
                    ),
                    "median_cell_ndre": float(
                        cell_frame[
                            "ndre_median"
                        ].median()
                    ),
                    "mean_ndvi_positive_fraction": float(
                        cell_frame[
                            "ndvi_positive_fraction"
                        ].mean()
                    ),
                    "mean_ndre_positive_fraction": float(
                        cell_frame[
                            "ndre_positive_fraction"
                        ].mean()
                    ),
                    "npz_path": relative_path(
                        output_paths["npz"]
                    ),
                    "ndvi_preview_path": relative_path(
                        output_paths[
                            "ndvi_preview"
                        ]
                    ),
                    "ndre_preview_path": relative_path(
                        output_paths[
                            "ndre_preview"
                        ]
                    ),
                    "ndvi_overlay_path": relative_path(
                        output_paths[
                            "ndvi_overlay"
                        ]
                    ),
                    "ndre_overlay_path": relative_path(
                        output_paths[
                            "ndre_overlay"
                        ]
                    ),
                    "status": status,
                    "notes": (
                        "All 70 polygons processed."
                        if status == "PASS"
                        else (
                            "Inspect cells with insufficient valid pixels."
                        )
                    ),
                }
            )

            cell_rows.extend(
                [
                    {
                        **base,
                        **cell,
                    }
                    for cell in cells
                ]
            )

            print(
                f"{status}: "
                f"{day} > "
                f"{tray} | "
                f"NDVI={valid_ndvi_cells}/70 | "
                f"NDRE={valid_ndre_cells}/70"
            )

        except Exception as error:
            tray_rows.append(
                {
                    **base,
                    "cells_processed": 0,
                    "cells_with_valid_ndvi": 0,
                    "cells_with_valid_ndre": 0,
                    "status": "FAIL",
                    "notes": str(error),
                }
            )

            print(
                f"FAIL: "
                f"{day} > "
                f"{tray} | "
                f"{error}"
            )

    tray_rows.sort(
        key=lambda row: (
            row["day_order"],
            row["tray_no"],
            natural_key(
                row["capture_id"]
            ),
        )
    )

    cell_rows.sort(
        key=lambda row: (
            row["day_order"],
            row["tray_no"],
            natural_key(
                row["capture_id"]
            ),
            row["cell_id"],
        )
    )

    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    CONFIG_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    tray_summary_csv = (
        REPORTS_ROOT
        / "ms_index_tray_summary.csv"
    )

    cell_results_csv = (
        REPORTS_ROOT
        / "ms_index_cell_results.csv"
    )

    excel_report = (
        REPORTS_ROOT
        / "ms_vegetation_index_report.xlsx"
    )

    word_report = (
        REPORTS_ROOT
        / "ms_vegetation_indices_work_process.docx"
    )

    settings_path = (
        CONFIG_ROOT
        / "index_settings.json"
    )

    pd.DataFrame(
        tray_rows
    ).to_csv(
        tray_summary_csv,
        index=False,
    )

    pd.DataFrame(
        cell_rows
    ).to_csv(
        cell_results_csv,
        index=False,
    )

    create_excel(
        excel_report,
        tray_rows,
        cell_rows,
    )

    create_word(
        word_report,
        tray_rows,
        cell_rows,
    )

    settings = {
        "ndvi_formula": (
            "(MS_NIR - MS_R) / "
            "(MS_NIR + MS_R)"
        ),
        "ndre_formula": (
            "(MS_NIR - MS_RE) / "
            "(MS_NIR + MS_RE)"
        ),
        "accepted_grid_statuses": sorted(
            ACCEPTED_GRID_STATUSES
        ),
        "minimum_valid_pixels_per_cell": (
            MIN_VALID_PIXELS_PER_CELL
        ),
        "date_map": DATE_MAP,
        "days_since_day1": DAYS_SINCE_DAY1,
        "day7_policy": (
            "Observed only; no imputation in Script 07."
        ),
        "interpretation": (
            "Relative image-derived indices unless "
            "calibrated reflectance is available."
        ),
    }

    settings_path.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    tray_frame = pd.DataFrame(
        tray_rows
    )

    status_counts = (
        tray_frame["status"]
        .value_counts()
        .to_dict()
        if not tray_frame.empty
        else {}
    )

    print(
        "\n"
        + "=" * 72
    )

    print(
        "SCRIPT 07 FINISHED"
    )

    print(
        "=" * 72
    )

    for status in (
        "PASS",
        "CHECK",
        "FAIL",
    ):
        print(
            f"{status}: "
            f"{status_counts.get(status, 0)}"
        )

    print(
        f"Expected full Trial 3 sets: "
        f"{EXPECTED_IMAGE_SETS}"
    )

    print(
        f"Processed rows: "
        f"{len(tray_rows)}"
    )

    print(
        f"Cell rows: "
        f"{len(cell_rows)}"
    )

    print(
        f"\nTray summary:\n"
        f"{tray_summary_csv}"
    )

    print(
        f"\nCell results:\n"
        f"{cell_results_csv}"
    )

    print(
        f"\nExcel report:\n"
        f"{excel_report}"
    )

    if DOCX_AVAILABLE:
        print(
            f"\nWord work-process report:\n"
            f"{word_report}"
        )

    print(
        f"\nSettings:\n"
        f"{settings_path}"
    )

    return (
        0
        if status_counts.get(
            "FAIL",
            0,
        ) == 0
        else 1
    )


if __name__ == "__main__":
    raise SystemExit(
        main()
    )