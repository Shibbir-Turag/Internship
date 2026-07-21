from __future__ import annotations

"""
SCRIPT 09 — THIRD TRIAL FINAL RGB + MS SYNTHESIS

Purpose
-------
Combine the completed Trial 3 RGB and multispectral comparison outputs into one
final synthesis.

This script uses:

RGB side:
outputs/Third trial/05_RGB_Growth_Rate_Treatment_Comparison/_reports/

MS side:
outputs/Third trial/08_MS_Treatment_Comparison/_reports/

Main tasks:
- combine RGB visible emergence / green-cover metrics with MS NDVI / NDRE metrics
- preserve observed Day 7 and adjusted Day 7 values separately
- calculate final combined RGB + MS tray scores
- rank all 12 trays
- compare Microbes vs No Microbes
- compare Ideal vs Heat vs Moisture
- compare Inside vs Outside for Ideal and Moisture
- summarise Heat and Moisture treatment evidence
- generate final CSV tables
- generate a final Excel workbook
- generate final charts
- generate a detailed Word synthesis report

Important
---------
This script does not modify any previous script outputs.
It only reads existing reports and writes a new synthesis folder.

Output folder
-------------
outputs/Third trial/09_Final_RGB_MS_Synthesis
"""

import argparse
import json
import math
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

RGB_REPORTS_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "05_RGB_Growth_Rate_Treatment_Comparison"
    / "_reports"
)

MS_REPORTS_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "08_MS_Treatment_Comparison"
    / "_reports"
)

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "09_Final_RGB_MS_Synthesis"
)

REPORTS_ROOT = OUTPUT_ROOT / "_reports"
CHARTS_ROOT = OUTPUT_ROOT / "charts"
CONFIG_ROOT = OUTPUT_ROOT / "_config"


# ============================================================
# 2) EXPECTED INPUT FILES
# ============================================================

RGB_TRAY_GROWTH_CANDIDATES = [
    RGB_REPORTS_ROOT / "rgb_tray_growth_metrics.csv",
    RGB_REPORTS_ROOT / "tray_growth_metrics.csv",
    RGB_REPORTS_ROOT / "rgb_growth_metrics.csv",
    RGB_REPORTS_ROOT / "growth_metrics_by_tray.csv",
]

RGB_GROUP_GROWTH_CANDIDATES = [
    RGB_REPORTS_ROOT / "rgb_group_growth_metrics.csv",
    RGB_REPORTS_ROOT / "group_growth_metrics.csv",
]

RGB_INSIDE_OUTSIDE_CANDIDATES = [
    RGB_REPORTS_ROOT / "inside_outside_comparison_ideal_moisture.csv",
    RGB_REPORTS_ROOT / "rgb_inside_outside_comparison_ideal_moisture.csv",
]

RGB_BUG_CELLS_CANDIDATES = [
    RGB_REPORTS_ROOT / "possible_day7_bug_eaten_cells.csv",
    RGB_REPORTS_ROOT / "possible_day7_bug_eaten_rgb_cells.csv",
]

MS_TRAY_GROWTH = (
    MS_REPORTS_ROOT
    / "ms_tray_growth_metrics.csv"
)

MS_GROUP_GROWTH = (
    MS_REPORTS_ROOT
    / "ms_group_growth_metrics.csv"
)

MS_INSIDE_OUTSIDE = (
    MS_REPORTS_ROOT
    / "inside_outside_comparison_ideal_moisture.csv"
)

MS_ADJUSTED_CELLS = (
    MS_REPORTS_ROOT
    / "possible_day7_bug_eaten_ms_cells.csv"
)


# ============================================================
# 3) TRIAL METADATA
# ============================================================

EXPECTED_TRAYS = 12
EXPECTED_TRAY_DAY_ROWS = 84
EXPECTED_CELL_DAY_ROWS = 5880

DATE_MAP = {
    1: "2026-06-29",
    2: "2026-06-30",
    3: "2026-07-01",
    4: "2026-07-02",
    5: "2026-07-03",
    6: "2026-07-04",
    7: "2026-07-07",
}

TRAY_METADATA = {
    1: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Inside",
        "environment_group": "Inside",
    },
    2: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Outside",
        "environment_group": "Outside",
    },
    3: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Outside",
        "environment_group": "Outside",
    },
    4: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
        "environment_group": "Dynamic Heat",
    },
    5: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Inside",
        "environment_group": "Inside",
    },
    6: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
        "environment_group": "Dynamic Heat",
    },
    7: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Outside",
        "environment_group": "Outside",
    },
    8: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
        "environment_group": "Dynamic Heat",
    },
    9: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "fixed_environment": "Inside",
        "environment_group": "Inside",
    },
    10: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Outside",
        "environment_group": "Outside",
    },
    11: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "fixed_environment": "Dynamic Heat",
        "environment_group": "Dynamic Heat",
    },
    12: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "fixed_environment": "Inside",
        "environment_group": "Inside",
    },
}

MICROBE_ORDER = [
    "No Microbes",
    "Microbes",
]

TREATMENT_ORDER = [
    "Ideal",
    "Heat",
    "Moisture",
]

ENVIRONMENT_ORDER = [
    "Inside",
    "Outside",
    "Dynamic Heat",
]


# ============================================================
# 4) OPTIONAL WORD SUPPORT
# ============================================================

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True

except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 5) GENERAL HELPERS
# ============================================================

def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(
            r"(\d+)",
            str(value),
        )
    ]


def find_first_existing(
    candidates: list[Path],
) -> Path | None:
    for path in candidates:
        if path.exists():
            return path

    return None


def safe_float(
    value: object,
    default: float = math.nan,
) -> float:
    try:
        if pd.isna(value):
            return default

        return float(value)

    except Exception:
        return default


def format_number(
    value: object,
    decimals: int = 3,
) -> str:
    try:
        if pd.isna(value):
            return "N/A"

        return f"{float(value):.{decimals}f}"

    except Exception:
        return "N/A"


def require_file(
    path: Path,
    description: str,
) -> None:
    if not path.exists():
        raise FileNotFoundError(
            f"Missing {description}:\n{path}"
        )


def pick_column(
    dataframe: pd.DataFrame,
    candidates: list[str],
) -> str | None:
    for column in candidates:
        if column in dataframe.columns:
            return column

    return None


def get_numeric_series(
    dataframe: pd.DataFrame,
    column: str | None,
) -> pd.Series:
    if column is None or column not in dataframe.columns:
        return pd.Series(
            [math.nan] * len(dataframe),
            index=dataframe.index,
            dtype=float,
        )

    return pd.to_numeric(
        dataframe[column],
        errors="coerce",
    )


def minmax_score(
    series: pd.Series,
) -> pd.Series:
    values = pd.to_numeric(
        series,
        errors="coerce",
    )

    valid = values.dropna()

    if valid.empty:
        return pd.Series(
            [math.nan] * len(values),
            index=values.index,
            dtype=float,
        )

    minimum = float(
        valid.min()
    )

    maximum = float(
        valid.max()
    )

    if maximum == minimum:
        return pd.Series(
            [
                50.0
                if pd.notna(value)
                else math.nan
                for value in values
            ],
            index=values.index,
            dtype=float,
        )

    return (
        values
        - minimum
    ) / (
        maximum
        - minimum
    ) * 100.0


def score_to_0_100(
    series: pd.Series,
) -> pd.Series:
    values = pd.to_numeric(
        series,
        errors="coerce",
    )

    valid = values.dropna()

    if valid.empty:
        return pd.Series(
            [math.nan] * len(values),
            index=values.index,
            dtype=float,
        )

    if (
        valid.min() >= 0
        and valid.max() <= 100
    ):
        return values

    return minmax_score(
        values
    )


def safe_round_dataframe(
    dataframe: pd.DataFrame,
    decimals: int = 4,
) -> pd.DataFrame:
    output = dataframe.copy()

    numeric_columns = output.select_dtypes(
        include=["number"]
    ).columns

    output[numeric_columns] = output[
        numeric_columns
    ].round(
        decimals
    )

    return output


def mean_or_nan(
    dataframe: pd.DataFrame,
    column: str,
) -> float:
    if (
        dataframe.empty
        or column not in dataframe.columns
    ):
        return math.nan

    return float(
        pd.to_numeric(
            dataframe[column],
            errors="coerce",
        ).mean()
    )


def unique_trays_text(
    dataframe: pd.DataFrame,
) -> str:
    if dataframe.empty or "tray" not in dataframe.columns:
        return ""

    return ", ".join(
        dataframe.sort_values(
            "tray_no"
        )["tray"].astype(str).tolist()
    )


# ============================================================
# 6) METADATA HELPERS
# ============================================================

def add_trial_metadata(
    dataframe: pd.DataFrame,
) -> pd.DataFrame:
    output = dataframe.copy()

    if "tray_no" not in output.columns:
        if "tray" in output.columns:
            output["tray_no"] = (
                output["tray"]
                .astype(str)
                .str.extract(r"(\d+)")
                .iloc[:, 0]
            )
        else:
            raise ValueError(
                "Input table does not contain tray_no or tray."
            )

    output["tray_no"] = pd.to_numeric(
        output["tray_no"],
        errors="coerce",
    )

    output = output.dropna(
        subset=[
            "tray_no",
        ]
    ).copy()

    output["tray_no"] = output[
        "tray_no"
    ].astype(int)

    if "tray" not in output.columns:
        output["tray"] = output[
            "tray_no"
        ].apply(
            lambda value: f"Tray {value}"
        )

    for column in [
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
    ]:
        if column not in output.columns:
            output[column] = ""

    for index, row in output.iterrows():
        tray_no = int(
            row["tray_no"]
        )

        metadata = TRAY_METADATA.get(
            tray_no,
            {},
        )

        for column in [
            "microbe_status",
            "treatment",
            "fixed_environment",
            "environment_group",
        ]:
            existing = str(
                output.at[
                    index,
                    column,
                ]
            ).strip()

            if (
                not existing
                or existing.casefold()
                in {
                    "nan",
                    "none",
                    "unknown",
                }
            ):
                output.at[
                    index,
                    column,
                ] = metadata.get(
                    column,
                    "Unknown",
                )

    output["microbe_treatment"] = (
        output["microbe_status"].astype(str)
        + " | "
        + output["treatment"].astype(str)
    )

    output["treatment_environment"] = (
        output["treatment"].astype(str)
        + " | "
        + output["environment_group"].astype(str)
    )

    output["microbe_environment"] = (
        output["microbe_status"].astype(str)
        + " | "
        + output["environment_group"].astype(str)
    )

    return output


# ============================================================
# 7) LOAD RGB OUTPUTS
# ============================================================

def load_rgb_tray_growth() -> tuple[pd.DataFrame, Path | None]:
    path = find_first_existing(
        RGB_TRAY_GROWTH_CANDIDATES
    )

    if path is None:
        raise FileNotFoundError(
            "Could not find the RGB tray growth metrics file in:\n"
            f"{RGB_REPORTS_ROOT}"
        )

    frame = pd.read_csv(
        path
    )

    frame = add_trial_metadata(
        frame
    )

    return frame, path


def compute_rgb_score(
    frame: pd.DataFrame,
) -> pd.Series:
    direct_score_column = pick_column(
        frame,
        [
            "overall_adjusted_rgb_score",
            "overall_rgb_score",
            "overall_adjusted_growth_score",
            "overall_adjusted_score",
            "adjusted_rgb_score",
            "rgb_score",
        ],
    )

    if direct_score_column is not None:
        return score_to_0_100(
            get_numeric_series(
                frame,
                direct_score_column,
            )
        )

    candidate_columns = [
        "day7_adjusted_green_cover_percent",
        "day7_adjusted_green_cover",
        "adjusted_day7_green_cover",
        "final_adjusted_green_cover",
        "final_green_cover_percent",
        "final_green_cover",
        "final_tracked_visible_emergence_percent",
        "final_visible_emergence_percent",
        "day7_tracked_visible_emergence_percent",
        "day7_visible_emergence_percent",
        "final_emergence_percent",
        "adjusted_green_cover_rate_day1_to_day7_per_day",
        "green_cover_rate_day1_to_day7_per_day",
        "adjusted_emergence_rate_day1_to_day7_per_day",
        "emergence_rate_day1_to_day7_per_day",
        "green_cover_auc",
        "adjusted_green_cover_auc",
    ]

    available = [
        column
        for column in candidate_columns
        if column in frame.columns
    ]

    if not available:
        return pd.Series(
            [math.nan] * len(frame),
            index=frame.index,
            dtype=float,
        )

    score_parts = []

    for column in available:
        score_parts.append(
            minmax_score(
                get_numeric_series(
                    frame,
                    column,
                )
            )
        )

    return pd.concat(
        score_parts,
        axis=1,
    ).mean(
        axis=1,
        skipna=True,
    )


def standardise_rgb_tray(
    frame: pd.DataFrame,
) -> pd.DataFrame:
    output = pd.DataFrame()

    output["tray_no"] = frame[
        "tray_no"
    ].astype(int)

    output["tray"] = frame[
        "tray"
    ].astype(str)

    for column in [
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
    ]:
        output[column] = frame[column].astype(str)

    output["rgb_score"] = compute_rgb_score(
        frame
    )

    rgb_final_emergence_column = pick_column(
        frame,
        [
            "final_tracked_visible_emergence_percent",
            "final_visible_emergence_percent",
            "day7_tracked_visible_emergence_percent",
            "day7_visible_emergence_percent",
            "final_emergence_percent",
            "tracked_emergence_percent",
            "emergence_percent",
        ],
    )

    rgb_day7_observed_green_column = pick_column(
        frame,
        [
            "day7_observed_green_cover_percent",
            "day7_observed_green_cover",
            "observed_day7_green_cover",
            "day7_green_cover_percent",
            "day7_green_cover",
            "final_observed_green_cover",
            "final_green_cover_percent",
            "final_green_cover",
        ],
    )

    rgb_day7_adjusted_green_column = pick_column(
        frame,
        [
            "day7_adjusted_green_cover_percent",
            "day7_adjusted_green_cover",
            "adjusted_day7_green_cover",
            "final_adjusted_green_cover",
            "adjusted_green_cover_day7",
        ],
    )

    rgb_growth_rate_column = pick_column(
        frame,
        [
            "adjusted_green_cover_rate_day1_to_day7_per_day",
            "green_cover_rate_day1_to_day7_per_day",
            "adjusted_rgb_growth_rate_day1_to_day7_per_day",
            "rgb_growth_rate_day1_to_day7_per_day",
            "growth_rate_day1_to_day7_per_day",
        ],
    )

    rgb_imputed_cells_column = pick_column(
        frame,
        [
            "day7_imputed_cells",
            "day7_adjusted_cells",
            "possible_day7_bug_eaten_cells",
            "bug_eaten_cells",
        ],
    )

    output["rgb_final_visible_emergence_percent"] = get_numeric_series(
        frame,
        rgb_final_emergence_column,
    )

    output["rgb_day7_observed_green_cover"] = get_numeric_series(
        frame,
        rgb_day7_observed_green_column,
    )

    output["rgb_day7_adjusted_green_cover"] = get_numeric_series(
        frame,
        rgb_day7_adjusted_green_column,
    )

    output["rgb_growth_rate_day1_to_day7_per_day"] = get_numeric_series(
        frame,
        rgb_growth_rate_column,
    )

    output["rgb_day7_imputed_cells"] = get_numeric_series(
        frame,
        rgb_imputed_cells_column,
    ).fillna(0)

    output["rgb_source_score_note"] = (
        "Direct RGB score column used"
        if pick_column(
            frame,
            [
                "overall_adjusted_rgb_score",
                "overall_rgb_score",
                "overall_adjusted_growth_score",
                "overall_adjusted_score",
                "adjusted_rgb_score",
                "rgb_score",
            ],
        )
        else "RGB score computed from available emergence/green-cover metrics"
    )

    return output.sort_values(
        "tray_no"
    ).reset_index(drop=True)


# ============================================================
# 8) LOAD MS OUTPUTS
# ============================================================

def load_ms_tray_growth() -> pd.DataFrame:
    require_file(
        MS_TRAY_GROWTH,
        "Script 08 MS tray growth metrics",
    )

    frame = pd.read_csv(
        MS_TRAY_GROWTH
    )

    frame = add_trial_metadata(
        frame
    )

    return frame


def standardise_ms_tray(
    frame: pd.DataFrame,
) -> pd.DataFrame:
    output = pd.DataFrame()

    output["tray_no"] = frame[
        "tray_no"
    ].astype(int)

    output["tray"] = frame[
        "tray"
    ].astype(str)

    for column in [
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
    ]:
        output[column] = frame[column].astype(str)

    ms_score_column = pick_column(
        frame,
        [
            "overall_adjusted_ms_score",
            "overall_ms_score",
            "adjusted_ms_score",
        ],
    )

    output["ms_score"] = score_to_0_100(
        get_numeric_series(
            frame,
            ms_score_column,
        )
    )

    output["ms_day7_observed_ndvi"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "day7_observed_ndvi",
            ],
        ),
    )

    output["ms_day7_adjusted_ndvi"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "day7_adjusted_ndvi",
            ],
        ),
    )

    output["ms_day7_observed_ndre"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "day7_observed_ndre",
            ],
        ),
    )

    output["ms_day7_adjusted_ndre"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "day7_adjusted_ndre",
            ],
        ),
    )

    output["ms_adjusted_ndvi_rate_day1_to_day7_per_day"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "adjusted_ndvi_rate_day1_to_day7_per_day",
            ],
        ),
    )

    output["ms_adjusted_ndre_rate_day1_to_day7_per_day"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "adjusted_ndre_rate_day1_to_day7_per_day",
            ],
        ),
    )

    output["ms_day7_imputed_cells"] = get_numeric_series(
        frame,
        pick_column(
            frame,
            [
                "day7_imputed_cells",
            ],
        ),
    ).fillna(0)

    return output.sort_values(
        "tray_no"
    ).reset_index(drop=True)


# ============================================================
# 9) COMBINE RGB AND MS TRAY TABLES
# ============================================================

def combine_rgb_ms_trays(
    rgb: pd.DataFrame,
    ms: pd.DataFrame,
) -> pd.DataFrame:
    combined = rgb.merge(
        ms,
        on="tray_no",
        how="outer",
        suffixes=(
            "_rgb",
            "_ms",
        ),
    )

    for column in [
        "tray",
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",
    ]:
        rgb_column = f"{column}_rgb"
        ms_column = f"{column}_ms"

        if rgb_column in combined.columns and ms_column in combined.columns:
            combined[column] = combined[
                rgb_column
            ].combine_first(
                combined[
                    ms_column
                ]
            )

        elif rgb_column in combined.columns:
            combined[column] = combined[
                rgb_column
            ]

        elif ms_column in combined.columns:
            combined[column] = combined[
                ms_column
            ]

        else:
            combined[column] = ""

    combined["tray"] = combined.apply(
        lambda row: (
            row["tray"]
            if str(row["tray"]).strip()
            and str(row["tray"]).strip().casefold() != "nan"
            else f"Tray {int(row['tray_no'])}"
        ),
        axis=1,
    )

    combined = add_trial_metadata(
        combined
    )

    combined["rgb_score_normalised"] = score_to_0_100(
        combined.get(
            "rgb_score",
            pd.Series(
                [math.nan] * len(combined),
                index=combined.index,
            ),
        )
    )

    combined["ms_score_normalised"] = score_to_0_100(
        combined.get(
            "ms_score",
            pd.Series(
                [math.nan] * len(combined),
                index=combined.index,
            ),
        )
    )

    combined["combined_rgb_ms_score"] = combined[
        [
            "rgb_score_normalised",
            "ms_score_normalised",
        ]
    ].mean(
        axis=1,
        skipna=True,
    )

    combined["combined_rank"] = combined[
        "combined_rgb_ms_score"
    ].rank(
        ascending=False,
        method="min",
    ).astype("Int64")

    combined["rgb_ms_score_difference"] = (
        combined["rgb_score_normalised"]
        - combined["ms_score_normalised"]
    )

    combined["evidence_agreement"] = np.where(
        combined[
            "rgb_score_normalised"
        ].isna()
        | combined[
            "ms_score_normalised"
        ].isna(),
        "Partial evidence only",
        np.where(
            combined[
                "rgb_ms_score_difference"
            ].abs()
            <= 15,
            "RGB and MS broadly agree",
            np.where(
                combined[
                    "rgb_ms_score_difference"
                ]
                > 15,
                "RGB stronger than MS",
                "MS stronger than RGB",
            ),
        ),
    )

    combined["total_day7_adjusted_cells"] = (
        pd.to_numeric(
            combined.get(
                "rgb_day7_imputed_cells",
                0,
            ),
            errors="coerce",
        ).fillna(0)
        + pd.to_numeric(
            combined.get(
                "ms_day7_imputed_cells",
                0,
            ),
            errors="coerce",
        ).fillna(0)
    )

    combined["final_interpretation"] = combined.apply(
        interpret_tray_row,
        axis=1,
    )

    keep_columns = [
        "combined_rank",
        "tray_no",
        "tray",
        "microbe_status",
        "treatment",
        "fixed_environment",
        "environment_group",
        "microbe_treatment",
        "treatment_environment",
        "microbe_environment",

        "combined_rgb_ms_score",
        "rgb_score_normalised",
        "ms_score_normalised",
        "rgb_ms_score_difference",
        "evidence_agreement",

        "rgb_final_visible_emergence_percent",
        "rgb_day7_observed_green_cover",
        "rgb_day7_adjusted_green_cover",
        "rgb_growth_rate_day1_to_day7_per_day",
        "rgb_day7_imputed_cells",

        "ms_day7_observed_ndvi",
        "ms_day7_adjusted_ndvi",
        "ms_day7_observed_ndre",
        "ms_day7_adjusted_ndre",
        "ms_adjusted_ndvi_rate_day1_to_day7_per_day",
        "ms_adjusted_ndre_rate_day1_to_day7_per_day",
        "ms_day7_imputed_cells",

        "total_day7_adjusted_cells",
        "rgb_source_score_note",
        "final_interpretation",
    ]

    keep_columns = [
        column
        for column in keep_columns
        if column in combined.columns
    ]

    return combined[
        keep_columns
    ].sort_values(
        [
            "combined_rank",
            "tray_no",
        ],
        na_position="last",
    ).reset_index(drop=True)


def interpret_tray_row(
    row: pd.Series,
) -> str:
    score = safe_float(
        row.get(
            "combined_rgb_ms_score",
            math.nan,
        )
    )

    agreement = str(
        row.get(
            "evidence_agreement",
            "",
        )
    )

    if not np.isfinite(score):
        return (
            "Combined interpretation could not be calculated because "
            "the required RGB/MS score values were incomplete."
        )

    if score >= 75:
        strength = "strong overall performance"
    elif score >= 55:
        strength = "moderate overall performance"
    elif score >= 35:
        strength = "weak-to-moderate overall performance"
    else:
        strength = "weak overall performance"

    return (
        f"{row.get('tray', 'This tray')} showed {strength}. "
        f"The evidence pattern was classified as: {agreement}."
    )


# ============================================================
# 10) GROUP SYNTHESIS
# ============================================================

def create_group_synthesis(
    combined: pd.DataFrame,
) -> pd.DataFrame:
    group_specs = [
        (
            "microbe_status",
            "Microbe Status",
        ),
        (
            "treatment",
            "Treatment Type",
        ),
        (
            "environment_group",
            "Environment Group",
        ),
        (
            "microbe_treatment",
            "Microbe x Treatment",
        ),
        (
            "treatment_environment",
            "Treatment x Environment",
        ),
        (
            "microbe_environment",
            "Microbe x Environment",
        ),
    ]

    rows = []

    for group_column, group_type in group_specs:
        if group_column not in combined.columns:
            continue

        for group_name, group in combined.groupby(
            group_column,
            dropna=False,
            sort=False,
        ):
            group = group.copy()

            best_tray = ""

            if (
                "combined_rgb_ms_score"
                in group.columns
                and group[
                    "combined_rgb_ms_score"
                ].notna().any()
            ):
                best_row = group.loc[
                    group[
                        "combined_rgb_ms_score"
                    ].idxmax()
                ]

                best_tray = str(
                    best_row[
                        "tray"
                    ]
                )

            rows.append(
                {
                    "group_type": group_type,
                    "group": str(
                        group_name
                    ),
                    "tray_count": int(
                        group[
                            "tray_no"
                        ].nunique()
                    ),
                    "trays": unique_trays_text(
                        group
                    ),
                    "mean_combined_rgb_ms_score": mean_or_nan(
                        group,
                        "combined_rgb_ms_score",
                    ),
                    "mean_rgb_score": mean_or_nan(
                        group,
                        "rgb_score_normalised",
                    ),
                    "mean_ms_score": mean_or_nan(
                        group,
                        "ms_score_normalised",
                    ),
                    "mean_rgb_final_visible_emergence_percent": mean_or_nan(
                        group,
                        "rgb_final_visible_emergence_percent",
                    ),
                    "mean_rgb_day7_adjusted_green_cover": mean_or_nan(
                        group,
                        "rgb_day7_adjusted_green_cover",
                    ),
                    "mean_ms_day7_adjusted_ndvi": mean_or_nan(
                        group,
                        "ms_day7_adjusted_ndvi",
                    ),
                    "mean_ms_day7_adjusted_ndre": mean_or_nan(
                        group,
                        "ms_day7_adjusted_ndre",
                    ),
                    "mean_total_day7_adjusted_cells": mean_or_nan(
                        group,
                        "total_day7_adjusted_cells",
                    ),
                    "best_tray_in_group": best_tray,
                }
            )

    result = pd.DataFrame(
        rows
    )

    if not result.empty:
        result["group_rank_within_type"] = (
            result.groupby(
                "group_type"
            )[
                "mean_combined_rgb_ms_score"
            ]
            .rank(
                ascending=False,
                method="min",
            )
            .astype("Int64")
        )

        result = result.sort_values(
            [
                "group_type",
                "group_rank_within_type",
                "group",
            ],
            na_position="last",
        ).reset_index(drop=True)

    return result


# ============================================================
# 11) INSIDE VS OUTSIDE SYNTHESIS
# ============================================================

def create_inside_outside_synthesis(
    combined: pd.DataFrame,
) -> pd.DataFrame:
    rows = []

    for treatment in [
        "Ideal",
        "Moisture",
    ]:
        subset = combined.loc[
            combined[
                "treatment"
            ].eq(
                treatment
            )
        ].copy()

        inside = subset.loc[
            subset[
                "fixed_environment"
            ].eq(
                "Inside"
            )
        ]

        outside = subset.loc[
            subset[
                "fixed_environment"
            ].eq(
                "Outside"
            )
        ]

        row = {
            "treatment": treatment,
            "inside_tray_count": int(
                inside[
                    "tray_no"
                ].nunique()
            ),
            "outside_tray_count": int(
                outside[
                    "tray_no"
                ].nunique()
            ),
            "inside_trays": unique_trays_text(
                inside
            ),
            "outside_trays": unique_trays_text(
                outside
            ),
            "inside_mean_combined_score": mean_or_nan(
                inside,
                "combined_rgb_ms_score",
            ),
            "outside_mean_combined_score": mean_or_nan(
                outside,
                "combined_rgb_ms_score",
            ),
            "inside_mean_rgb_score": mean_or_nan(
                inside,
                "rgb_score_normalised",
            ),
            "outside_mean_rgb_score": mean_or_nan(
                outside,
                "rgb_score_normalised",
            ),
            "inside_mean_ms_score": mean_or_nan(
                inside,
                "ms_score_normalised",
            ),
            "outside_mean_ms_score": mean_or_nan(
                outside,
                "ms_score_normalised",
            ),
            "inside_mean_rgb_green_cover": mean_or_nan(
                inside,
                "rgb_day7_adjusted_green_cover",
            ),
            "outside_mean_rgb_green_cover": mean_or_nan(
                outside,
                "rgb_day7_adjusted_green_cover",
            ),
            "inside_mean_ms_ndvi": mean_or_nan(
                inside,
                "ms_day7_adjusted_ndvi",
            ),
            "outside_mean_ms_ndvi": mean_or_nan(
                outside,
                "ms_day7_adjusted_ndvi",
            ),
            "inside_mean_ms_ndre": mean_or_nan(
                inside,
                "ms_day7_adjusted_ndre",
            ),
            "outside_mean_ms_ndre": mean_or_nan(
                outside,
                "ms_day7_adjusted_ndre",
            ),
        }

        for metric in [
            "combined_score",
            "rgb_score",
            "ms_score",
            "rgb_green_cover",
            "ms_ndvi",
            "ms_ndre",
        ]:
            inside_column = f"inside_mean_{metric}"
            outside_column = f"outside_mean_{metric}"

            row[
                f"inside_minus_outside_{metric}"
            ] = (
                row[
                    inside_column
                ]
                - row[
                    outside_column
                ]
            )

        difference = safe_float(
            row[
                "inside_minus_outside_combined_score"
            ]
        )

        if not np.isfinite(
            difference
        ):
            interpretation = (
                "Inside/Outside comparison could not be calculated."
            )

        elif difference > 0:
            interpretation = (
                "Inside trays performed better on the combined RGB + MS score."
            )

        elif difference < 0:
            interpretation = (
                "Outside trays performed better on the combined RGB + MS score."
            )

        else:
            interpretation = (
                "Inside and Outside trays had equal combined RGB + MS scores."
            )

        row[
            "interpretation"
        ] = interpretation

        rows.append(
            row
        )

    return pd.DataFrame(
        rows
    )


# ============================================================
# 12) INPUT INVENTORY
# ============================================================

def create_input_inventory(
    rgb_tray_path: Path,
) -> pd.DataFrame:
    rgb_group_path = find_first_existing(
        RGB_GROUP_GROWTH_CANDIDATES
    )

    rgb_inside_outside_path = find_first_existing(
        RGB_INSIDE_OUTSIDE_CANDIDATES
    )

    rgb_bug_cells_path = find_first_existing(
        RGB_BUG_CELLS_CANDIDATES
    )

    rows = [
        {
            "source": "RGB tray growth metrics",
            "path": str(
                rgb_tray_path
            ),
            "exists": rgb_tray_path.exists(),
            "used": "Yes",
            "notes": "Main RGB tray-level input.",
        },
        {
            "source": "RGB group growth metrics",
            "path": str(
                rgb_group_path
            )
            if rgb_group_path
            else "",
            "exists": bool(
                rgb_group_path
                and rgb_group_path.exists()
            ),
            "used": "Reference only",
            "notes": "Group metrics are recomputed in this synthesis for consistency.",
        },
        {
            "source": "RGB Inside vs Outside comparison",
            "path": str(
                rgb_inside_outside_path
            )
            if rgb_inside_outside_path
            else "",
            "exists": bool(
                rgb_inside_outside_path
                and rgb_inside_outside_path.exists()
            ),
            "used": "Reference only",
            "notes": "Inside/Outside comparison is recomputed from tray-level data.",
        },
        {
            "source": "RGB possible bug-eaten cells",
            "path": str(
                rgb_bug_cells_path
            )
            if rgb_bug_cells_path
            else "",
            "exists": bool(
                rgb_bug_cells_path
                and rgb_bug_cells_path.exists()
            ),
            "used": "Indirect",
            "notes": "Already incorporated into Script 05 RGB adjusted values.",
        },
        {
            "source": "MS tray growth metrics",
            "path": str(
                MS_TRAY_GROWTH
            ),
            "exists": MS_TRAY_GROWTH.exists(),
            "used": "Yes",
            "notes": "Main MS tray-level input.",
        },
        {
            "source": "MS group growth metrics",
            "path": str(
                MS_GROUP_GROWTH
            ),
            "exists": MS_GROUP_GROWTH.exists(),
            "used": "Reference only",
            "notes": "Group metrics are recomputed in this synthesis for consistency.",
        },
        {
            "source": "MS Inside vs Outside comparison",
            "path": str(
                MS_INSIDE_OUTSIDE
            ),
            "exists": MS_INSIDE_OUTSIDE.exists(),
            "used": "Reference only",
            "notes": "Inside/Outside comparison is recomputed from tray-level data.",
        },
        {
            "source": "MS possible bug-eaten adjusted cells",
            "path": str(
                MS_ADJUSTED_CELLS
            ),
            "exists": MS_ADJUSTED_CELLS.exists(),
            "used": "Indirect",
            "notes": "Already incorporated into Script 08 MS adjusted values.",
        },
    ]

    return pd.DataFrame(
        rows
    )


# ============================================================
# 13) CHART HELPERS
# ============================================================

def no_data_chart(
    title: str,
    path: Path,
    message: str = "No valid data available",
) -> None:
    figure, axis = plt.subplots(
        figsize=(
            9,
            5,
        )
    )

    axis.text(
        0.5,
        0.5,
        message,
        horizontalalignment="center",
        verticalalignment="center",
        fontsize=12,
    )

    axis.set_title(
        title
    )

    axis.axis(
        "off"
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_combined_ranking_chart(
    combined: pd.DataFrame,
    path: Path,
) -> None:
    frame = combined.dropna(
        subset=[
            "combined_rgb_ms_score",
        ]
    ).sort_values(
        "combined_rgb_ms_score"
    ).copy()

    if frame.empty:
        no_data_chart(
            "Trial 3 final combined RGB + MS tray ranking",
            path,
        )

        return

    frame[
        "chart_label"
    ] = (
        frame[
            "tray"
        ].astype(str)
        + " — "
        + frame[
            "microbe_status"
        ].astype(str)
        + " | "
        + frame[
            "treatment"
        ].astype(str)
        + " | "
        + frame[
            "environment_group"
        ].astype(str)
    )

    figure, axis = plt.subplots(
        figsize=(
            12.5,
            7.2,
        )
    )

    axis.barh(
        frame[
            "chart_label"
        ],
        frame[
            "combined_rgb_ms_score"
        ],
    )

    axis.set_title(
        "Trial 3 final combined RGB + MS tray ranking"
    )

    axis.set_xlabel(
        "Combined RGB + MS score"
    )

    axis.set_ylabel(
        "Tray"
    )

    axis.grid(
        True,
        axis="x",
        alpha=0.30,
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_rgb_vs_ms_scatter(
    combined: pd.DataFrame,
    path: Path,
) -> None:
    frame = combined.dropna(
        subset=[
            "rgb_score_normalised",
            "ms_score_normalised",
        ]
    ).copy()

    if frame.empty:
        no_data_chart(
            "RGB score vs MS score",
            path,
        )

        return

    figure, axis = plt.subplots(
        figsize=(
            8,
            7,
        )
    )

    axis.scatter(
        frame[
            "rgb_score_normalised"
        ],
        frame[
            "ms_score_normalised"
        ],
        s=80,
    )

    for row in frame.itertuples(
        index=False
    ):
        axis.text(
            row.rgb_score_normalised,
            row.ms_score_normalised,
            str(
                row.tray
            ),
            fontsize=8,
            ha="left",
            va="bottom",
        )

    axis.plot(
        [
            0,
            100,
        ],
        [
            0,
            100,
        ],
        linestyle="--",
        linewidth=1,
    )

    axis.set_xlim(
        0,
        105,
    )

    axis.set_ylim(
        0,
        105,
    )

    axis.set_title(
        "Trial 3 RGB score versus MS score"
    )

    axis.set_xlabel(
        "RGB score"
    )

    axis.set_ylabel(
        "MS score"
    )

    axis.grid(
        True,
        alpha=0.30,
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_group_bar_chart(
    group_synthesis: pd.DataFrame,
    group_type: str,
    path: Path,
) -> None:
    frame = group_synthesis.loc[
        group_synthesis[
            "group_type"
        ].eq(
            group_type
        )
    ].copy()

    frame = frame.dropna(
        subset=[
            "mean_combined_rgb_ms_score",
        ]
    ).sort_values(
        "mean_combined_rgb_ms_score",
        ascending=False,
    )

    if frame.empty:
        no_data_chart(
            f"{group_type}: combined RGB + MS score",
            path,
        )

        return

    figure, axis = plt.subplots(
        figsize=(
            10.5,
            6,
        )
    )

    axis.bar(
        frame[
            "group"
        ],
        frame[
            "mean_combined_rgb_ms_score"
        ],
    )

    axis.set_title(
        f"{group_type}: mean combined RGB + MS score"
    )

    axis.set_xlabel(
        group_type
    )

    axis.set_ylabel(
        "Mean combined score"
    )

    axis.tick_params(
        axis="x",
        rotation=25,
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_inside_outside_chart(
    inside_outside: pd.DataFrame,
    path: Path,
) -> None:
    if inside_outside.empty:
        no_data_chart(
            "Ideal and Moisture: Inside vs Outside combined score",
            path,
        )

        return

    x = np.arange(
        len(
            inside_outside
        )
    )

    width = 0.36

    figure, axis = plt.subplots(
        figsize=(
            9.5,
            5.8,
        )
    )

    axis.bar(
        x - width / 2,
        pd.to_numeric(
            inside_outside[
                "inside_mean_combined_score"
            ],
            errors="coerce",
        ),
        width,
        label="Inside",
    )

    axis.bar(
        x + width / 2,
        pd.to_numeric(
            inside_outside[
                "outside_mean_combined_score"
            ],
            errors="coerce",
        ),
        width,
        label="Outside",
    )

    axis.set_title(
        "Ideal and Moisture: Inside vs Outside combined RGB + MS score"
    )

    axis.set_xlabel(
        "Treatment"
    )

    axis.set_ylabel(
        "Mean combined score"
    )

    axis.set_xticks(
        x
    )

    axis.set_xticklabels(
        inside_outside[
            "treatment"
        ]
    )

    axis.legend(
        loc="best"
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def save_adjusted_cells_chart(
    combined: pd.DataFrame,
    path: Path,
) -> None:
    frame = combined.sort_values(
        "tray_no"
    ).copy()

    if frame.empty:
        no_data_chart(
            "Day 7 adjusted cells by tray",
            path,
        )

        return

    x = np.arange(
        len(
            frame
        )
    )

    width = 0.36

    figure, axis = plt.subplots(
        figsize=(
            11,
            6,
        )
    )

    axis.bar(
        x - width / 2,
        pd.to_numeric(
            frame.get(
                "rgb_day7_imputed_cells",
                0,
            ),
            errors="coerce",
        ).fillna(0),
        width,
        label="RGB adjusted cells",
    )

    axis.bar(
        x + width / 2,
        pd.to_numeric(
            frame.get(
                "ms_day7_imputed_cells",
                0,
            ),
            errors="coerce",
        ).fillna(0),
        width,
        label="MS adjusted cells",
    )

    axis.set_title(
        "Day 7 adjusted cell counts by tray"
    )

    axis.set_xlabel(
        "Tray"
    )

    axis.set_ylabel(
        "Adjusted cell records"
    )

    axis.set_xticks(
        x
    )

    axis.set_xticklabels(
        frame[
            "tray"
        ],
        rotation=30,
        horizontalalignment="right",
    )

    axis.legend(
        loc="best"
    )

    axis.grid(
        True,
        axis="y",
        alpha=0.30,
    )

    figure.tight_layout()

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        path,
        dpi=220,
    )

    plt.close(
        figure
    )


def create_charts(
    combined: pd.DataFrame,
    group_synthesis: pd.DataFrame,
    inside_outside: pd.DataFrame,
) -> dict[str, Path]:
    CHARTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    charts = {}

    charts[
        "combined_tray_ranking"
    ] = CHARTS_ROOT / "01_final_combined_rgb_ms_tray_ranking.png"

    save_combined_ranking_chart(
        combined,
        charts[
            "combined_tray_ranking"
        ],
    )

    charts[
        "rgb_vs_ms_scatter"
    ] = CHARTS_ROOT / "02_rgb_vs_ms_score_scatter.png"

    save_rgb_vs_ms_scatter(
        combined,
        charts[
            "rgb_vs_ms_scatter"
        ],
    )

    charts[
        "microbe_group_score"
    ] = CHARTS_ROOT / "03_combined_score_by_microbe_status.png"

    save_group_bar_chart(
        group_synthesis,
        "Microbe Status",
        charts[
            "microbe_group_score"
        ],
    )

    charts[
        "treatment_group_score"
    ] = CHARTS_ROOT / "04_combined_score_by_treatment.png"

    save_group_bar_chart(
        group_synthesis,
        "Treatment Type",
        charts[
            "treatment_group_score"
        ],
    )

    charts[
        "inside_outside_score"
    ] = CHARTS_ROOT / "05_inside_outside_combined_score_ideal_moisture.png"

    save_inside_outside_chart(
        inside_outside,
        charts[
            "inside_outside_score"
        ],
    )

    charts[
        "adjusted_cells"
    ] = CHARTS_ROOT / "06_day7_adjusted_cells_rgb_ms.png"

    save_adjusted_cells_chart(
        combined,
        charts[
            "adjusted_cells"
        ],
    )

    return charts


# ============================================================
# 14) EXCEL OUTPUT
# ============================================================

def style_workbook(
    path: Path,
) -> None:
    workbook = load_workbook(
        path
    )

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.row_dimensions[1].height = 34

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )
            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = column_cells[0].column_letter

            longest = max(
                len(
                    str(
                        cell.value
                    )
                )
                if cell.value is not None
                else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[
                letter
            ].width = min(
                max(
                    12,
                    longest + 2,
                ),
                58,
            )

    workbook.save(
        path
    )


def save_tables(
    combined: pd.DataFrame,
    group_synthesis: pd.DataFrame,
    inside_outside: pd.DataFrame,
    input_inventory: pd.DataFrame,
) -> dict[str, Path]:
    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    paths = {
        "combined_tray_synthesis": REPORTS_ROOT / "final_combined_tray_synthesis.csv",
        "group_synthesis": REPORTS_ROOT / "final_group_synthesis.csv",
        "inside_outside_synthesis": REPORTS_ROOT / "final_inside_outside_synthesis.csv",
        "input_inventory": REPORTS_ROOT / "input_file_inventory.csv",
        "excel": REPORTS_ROOT / "final_rgb_ms_synthesis_report.xlsx",
    }

    combined.to_csv(
        paths[
            "combined_tray_synthesis"
        ],
        index=False,
    )

    group_synthesis.to_csv(
        paths[
            "group_synthesis"
        ],
        index=False,
    )

    inside_outside.to_csv(
        paths[
            "inside_outside_synthesis"
        ],
        index=False,
    )

    input_inventory.to_csv(
        paths[
            "input_inventory"
        ],
        index=False,
    )

    readme = pd.DataFrame(
        {
            "Notes": [
                "This workbook combines Trial 3 RGB and MS comparison outputs.",
                "RGB evidence includes visible emergence and green-cover metrics.",
                "MS evidence includes relative image-derived NDVI and NDRE metrics.",
                "Observed Day 7 values and adjusted Day 7 estimates remain separate.",
                "Adjusted Day 7 values are scenario estimates for possible bug-eaten cells, not direct observations.",
                "Combined RGB + MS score is a descriptive synthesis score, not a formal statistical test.",
                "Inside vs Outside comparison is restricted to Ideal and Moisture treatments.",
                "Heat trays are not directly included in fixed Inside/Outside comparison because they were moved during the trial.",
                "Report NDVI/NDRE as relative image-derived indices unless calibrated reflectance data are available.",
            ]
        }
    )

    with pd.ExcelWriter(
        paths[
            "excel"
        ],
        engine="openpyxl",
    ) as writer:
        safe_round_dataframe(
            combined
        ).to_excel(
            writer,
            sheet_name="Combined Tray Synthesis",
            index=False,
        )

        safe_round_dataframe(
            group_synthesis
        ).to_excel(
            writer,
            sheet_name="Group Synthesis",
            index=False,
        )

        safe_round_dataframe(
            inside_outside
        ).to_excel(
            writer,
            sheet_name="Inside Outside Synthesis",
            index=False,
        )

        input_inventory.to_excel(
            writer,
            sheet_name="Input File Inventory",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_workbook(
        paths[
            "excel"
        ]
    )

    return paths


# ============================================================
# 15) WORD REPORT
# ============================================================

def add_docx_table(
    document,
    dataframe: pd.DataFrame,
    columns: list[str],
    maximum_rows: int = 12,
) -> None:
    if dataframe.empty:
        document.add_paragraph(
            "No table data were available."
        )

        return

    columns = [
        column
        for column in columns
        if column in dataframe.columns
    ]

    if not columns:
        document.add_paragraph(
            "No valid columns were available for this table."
        )

        return

    frame = safe_round_dataframe(
        dataframe[
            columns
        ].head(
            maximum_rows
        ),
        decimals=3,
    )

    table = document.add_table(
        rows=1,
        cols=len(
            columns
        ),
    )

    table.style = "Table Grid"

    for index, column in enumerate(
        columns
    ):
        table.rows[0].cells[
            index
        ].text = (
            column.replace(
                "_",
                " ",
            ).title()
        )

    for _index, row in frame.iterrows():
        cells = table.add_row().cells

        for index, column in enumerate(
            columns
        ):
            value = row[
                column
            ]

            cells[
                index
            ].text = (
                ""
                if pd.isna(
                    value
                )
                else str(
                    value
                )
            )


def add_picture_if_exists(
    document,
    path: Path | None,
    width_inches: float = 6.3,
) -> None:
    if (
        path is None
        or not Path(path).exists()
    ):
        return

    document.add_picture(
        str(path),
        width=Inches(
            width_inches
        ),
    )


def describe_output_file(
    document,
    filename: str,
    description: str,
) -> None:
    paragraph = document.add_paragraph()

    paragraph.add_run(
        filename
    ).bold = True

    paragraph.add_run(
        f": {description}"
    )


def best_group_sentence(
    group_synthesis: pd.DataFrame,
    group_type: str,
) -> str:
    subset = group_synthesis.loc[
        group_synthesis[
            "group_type"
        ].eq(
            group_type
        )
    ].dropna(
        subset=[
            "mean_combined_rgb_ms_score",
        ]
    )

    if subset.empty:
        return (
            f"No valid {group_type.lower()} synthesis result was available."
        )

    best = subset.loc[
        subset[
            "mean_combined_rgb_ms_score"
        ].idxmax()
    ]

    return (
        f"For {group_type.lower()}, the highest mean combined RGB + MS score "
        f"was recorded by {best['group']} "
        f"({format_number(best['mean_combined_rgb_ms_score'], 2)})."
    )


def create_word_report(
    path: Path,
    combined: pd.DataFrame,
    group_synthesis: pd.DataFrame,
    inside_outside: pd.DataFrame,
    input_inventory: pd.DataFrame,
    charts: dict[str, Path],
) -> Path | None:
    if not DOCX_AVAILABLE:
        print(
            "Word report skipped because python-docx is not installed."
        )

        print(
            "Install it with: pip install python-docx"
        )

        return None

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Times New Roman"

    document.styles[
        "Normal"
    ].font.size = Pt(11)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name in document.styles:
            document.styles[
                style_name
            ].font.name = "Times New Roman"

    title = document.add_heading(
        "Trial 3 Final RGB and Multispectral Synthesis Report",
        level=0,
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This report combines the RGB and multispectral outputs from the "
        "Third Trial workflow. The RGB results describe visible emergence "
        "and green-cover development, while the multispectral results describe "
        "relative image-derived NDVI and NDRE responses. The aim is to provide "
        "one final tray-level and treatment-level interpretation for Trial 3."
    )

    document.add_paragraph(
        "Observed Day 7 values are preserved as the primary record. Adjusted "
        "Day 7 values remain separate and represent a flagged analytical "
        "scenario for cells where earlier seedlings may have disappeared due "
        "to bug damage."
    )

    document.add_heading(
        "1. Input files checked",
        level=1,
    )

    document.add_paragraph(
        "The synthesis used the main tray-level outputs from Script 05 and "
        "Script 08. Group and Inside/Outside tables were also checked as "
        "reference outputs, but the final group summaries were recomputed from "
        "tray-level values so that RGB and MS were combined consistently."
    )

    add_docx_table(
        document,
        input_inventory,
        [
            "source",
            "exists",
            "used",
            "notes",
        ],
        maximum_rows=12,
    )

    document.add_heading(
        "2. Synthesis method",
        level=1,
    )

    document.add_paragraph(
        "For each tray, the script standardised the available RGB score and "
        "the available MS score to a 0–100 scale. The final combined score was "
        "calculated as the mean of the standardised RGB and MS scores. If only "
        "one evidence stream was available for a tray, the combined score used "
        "the available evidence and flagged the case as partial evidence."
    )

    document.add_paragraph(
        "The combined score is descriptive. It is useful for ranking trays and "
        "summarising the pattern across treatments, but it should not be treated "
        "as a formal statistical significance result."
    )

    document.add_heading(
        "3. Final tray ranking",
        level=1,
    )

    if not combined.empty:
        best = combined.sort_values(
            "combined_rank",
            na_position="last",
        ).iloc[0]

        document.add_paragraph(
            f"The highest final combined RGB + MS score was recorded by "
            f"{best['tray']} ({best['microbe_status']} | {best['treatment']} | "
            f"{best['environment_group']}). Its combined score was "
            f"{format_number(best['combined_rgb_ms_score'], 2)}."
        )

    add_docx_table(
        document,
        combined.sort_values(
            "combined_rank",
            na_position="last",
        ),
        [
            "combined_rank",
            "tray",
            "microbe_status",
            "treatment",
            "environment_group",
            "combined_rgb_ms_score",
            "rgb_score_normalised",
            "ms_score_normalised",
            "evidence_agreement",
            "total_day7_adjusted_cells",
        ],
        maximum_rows=12,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "combined_tray_ranking"
        ),
    )

    document.add_heading(
        "4. RGB and MS agreement",
        level=1,
    )

    agreement_counts = (
        combined[
            "evidence_agreement"
        ]
        .value_counts()
        .to_dict()
        if "evidence_agreement" in combined.columns
        else {}
    )

    document.add_paragraph(
        "The RGB-versus-MS scatter chart checks whether visible growth evidence "
        "and multispectral vegetation-index evidence point in the same general "
        "direction. Trays close to the diagonal have broadly similar RGB and MS "
        "scores. Trays far from the diagonal may have strong visible emergence "
        "but weaker vegetation-index response, or the reverse."
    )

    document.add_paragraph(
        "Evidence agreement summary: "
        + "; ".join(
            [
                f"{key}: {value}"
                for key, value in agreement_counts.items()
            ]
        )
        + "."
    )

    add_picture_if_exists(
        document,
        charts.get(
            "rgb_vs_ms_scatter"
        ),
    )

    document.add_heading(
        "5. Microbes vs No Microbes",
        level=1,
    )

    document.add_paragraph(
        best_group_sentence(
            group_synthesis,
            "Microbe Status",
        )
    )

    add_docx_table(
        document,
        group_synthesis.loc[
            group_synthesis[
                "group_type"
            ].eq(
                "Microbe Status"
            )
        ],
        [
            "group",
            "tray_count",
            "trays",
            "mean_combined_rgb_ms_score",
            "mean_rgb_score",
            "mean_ms_score",
            "best_tray_in_group",
        ],
        maximum_rows=5,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "microbe_group_score"
        ),
    )

    document.add_heading(
        "6. Ideal vs Heat vs Moisture",
        level=1,
    )

    document.add_paragraph(
        best_group_sentence(
            group_synthesis,
            "Treatment Type",
        )
    )

    add_docx_table(
        document,
        group_synthesis.loc[
            group_synthesis[
                "group_type"
            ].eq(
                "Treatment Type"
            )
        ],
        [
            "group",
            "tray_count",
            "trays",
            "mean_combined_rgb_ms_score",
            "mean_rgb_score",
            "mean_ms_score",
            "best_tray_in_group",
        ],
        maximum_rows=10,
    )

    add_picture_if_exists(
        document,
        charts.get(
            "treatment_group_score"
        ),
    )

    document.add_heading(
        "7. Ideal and Moisture: Inside vs Outside",
        level=1,
    )

    document.add_paragraph(
        "Inside-versus-Outside synthesis was calculated only for Ideal and "
        "Moisture treatments because those trays had fixed environmental "
        "placement. Heat trays were excluded from this direct comparison "
        "because they were moved between inside and outside during the trial."
    )

    add_docx_table(
        document,
        inside_outside,
        [
            "treatment",
            "inside_trays",
            "outside_trays",
            "inside_mean_combined_score",
            "outside_mean_combined_score",
            "inside_minus_outside_combined_score",
            "inside_mean_rgb_score",
            "outside_mean_rgb_score",
            "inside_mean_ms_score",
            "outside_mean_ms_score",
            "interpretation",
        ],
        maximum_rows=5,
    )

    for row in inside_outside.itertuples(
        index=False
    ):
        document.add_paragraph(
            f"For {row.treatment}, the Inside-minus-Outside combined score "
            f"difference was {format_number(row.inside_minus_outside_combined_score, 2)}. "
            f"{row.interpretation}"
        )

    add_picture_if_exists(
        document,
        charts.get(
            "inside_outside_score"
        ),
    )

    document.add_heading(
        "8. Day 7 observed and adjusted evidence",
        level=1,
    )

    adjusted_total = safe_float(
        combined[
            "total_day7_adjusted_cells"
        ].sum()
        if "total_day7_adjusted_cells" in combined.columns
        else 0
    )

    document.add_paragraph(
        f"Across all trays, the combined RGB and MS workflow recorded "
        f"{format_number(adjusted_total, 0)} adjusted Day 7 cell records. "
        "These records are not direct observations. They are kept as a separate "
        "adjusted scenario to avoid mixing visible final-image evidence with "
        "estimated values."
    )

    add_picture_if_exists(
        document,
        charts.get(
            "adjusted_cells"
        ),
    )

    document.add_heading(
        "9. Description of generated CSV files",
        level=1,
    )

    output_descriptions = [
        (
            "final_combined_tray_synthesis.csv",
            "Main tray-level output. It combines RGB and MS evidence, includes final scores, ranks, observed/adjusted Day 7 indicators, and interpretation notes.",
        ),
        (
            "final_group_synthesis.csv",
            "Group-level summary for Microbe Status, Treatment Type, Environment Group, Microbe × Treatment, Treatment × Environment, and Microbe × Environment.",
        ),
        (
            "final_inside_outside_synthesis.csv",
            "Dedicated Inside-versus-Outside comparison for Ideal and Moisture treatments using combined RGB + MS scores.",
        ),
        (
            "input_file_inventory.csv",
            "List of the input files checked by Script 09, including whether each file existed and how it was used.",
        ),
    ]

    for filename, description in output_descriptions:
        describe_output_file(
            document,
            filename,
            description,
        )

    document.add_heading(
        "10. Description of the Excel workbook",
        level=1,
    )

    document.add_paragraph(
        "The workbook final_rgb_ms_synthesis_report.xlsx contains the Combined "
        "Tray Synthesis, Group Synthesis, Inside Outside Synthesis, Input File "
        "Inventory, and Read Me sheets. It is the easiest single file for checking "
        "final Trial 3 synthesis values."
    )

    document.add_heading(
        "11. Description of generated charts",
        level=1,
    )

    chart_descriptions = [
        (
            "01_final_combined_rgb_ms_tray_ranking.png",
            "Final tray ranking using the combined RGB + MS score.",
        ),
        (
            "02_rgb_vs_ms_score_scatter.png",
            "Agreement check between RGB and MS evidence streams.",
        ),
        (
            "03_combined_score_by_microbe_status.png",
            "Mean combined score for Microbes and No Microbes.",
        ),
        (
            "04_combined_score_by_treatment.png",
            "Mean combined score for Ideal, Heat and Moisture treatments.",
        ),
        (
            "05_inside_outside_combined_score_ideal_moisture.png",
            "Inside vs Outside combined score for Ideal and Moisture treatments.",
        ),
        (
            "06_day7_adjusted_cells_rgb_ms.png",
            "Adjusted Day 7 cell counts from RGB and MS workflows by tray.",
        ),
    ]

    for filename, description in chart_descriptions:
        describe_output_file(
            document,
            filename,
            description,
        )

    document.add_heading(
        "12. Interpretation limitations",
        level=1,
    )

    document.add_paragraph(
        "The synthesis combines different kinds of evidence. RGB metrics are "
        "based on visible emergence and green cover, whereas MS metrics are "
        "based on relative NDVI and NDRE. These do not measure the exact same "
        "biological property, so perfect agreement is not expected."
    )

    document.add_paragraph(
        "The multispectral indices should be described as relative image-derived "
        "indices unless reflectance calibration is confirmed. Results can be "
        "affected by shadows, soil moisture, exposure, band alignment, cell-zone "
        "placement and the amount of soil visible within each cell."
    )

    document.add_paragraph(
        "The final combined score is a descriptive integration score. It is "
        "appropriate for ranking, reporting and visual comparison, but it should "
        "not be used as a formal statistical significance test."
    )

    document.save(
        path
    )

    return path


# ============================================================
# 16) SETTINGS OUTPUT
# ============================================================

def save_settings(
    path: Path,
    rgb_tray_path: Path,
) -> Path:
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    settings = {
        "purpose": "Third Trial final RGB + MS synthesis",
        "rgb_reports_root": str(
            RGB_REPORTS_ROOT
        ),
        "ms_reports_root": str(
            MS_REPORTS_ROOT
        ),
        "rgb_tray_growth_file": str(
            rgb_tray_path
        ),
        "ms_tray_growth_file": str(
            MS_TRAY_GROWTH
        ),
        "output_root": str(
            OUTPUT_ROOT
        ),
        "expected_trays": EXPECTED_TRAYS,
        "expected_tray_day_rows": EXPECTED_TRAY_DAY_ROWS,
        "expected_cell_day_rows": EXPECTED_CELL_DAY_ROWS,
        "date_map": DATE_MAP,
        "combined_score_method": (
            "Mean of normalised RGB score and normalised MS score."
        ),
        "rgb_evidence": (
            "Visible emergence and green-cover based indicators."
        ),
        "ms_evidence": (
            "Relative image-derived NDVI and NDRE indicators."
        ),
        "day7_policy": (
            "Observed Day 7 and adjusted Day 7 values remain separate."
        ),
        "inside_outside_policy": (
            "Only Ideal and Moisture treatments are used for fixed Inside/Outside comparison."
        ),
        "interpretation_warning": (
            "Combined score is descriptive, not a formal statistical test."
        ),
    }

    path.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    return path


# ============================================================
# 17) MAIN WORKFLOW
# ============================================================

def run_analysis(
    _args,
) -> int:
    print(
        "\nSCRIPT 09 — THIRD TRIAL FINAL RGB + MS SYNTHESIS"
    )

    print(
        "=" * 78
    )

    print(
        f"RGB reports folder:\n{RGB_REPORTS_ROOT}"
    )

    print(
        f"\nMS reports folder:\n{MS_REPORTS_ROOT}"
    )

    print(
        f"\nOutput folder:\n{OUTPUT_ROOT}"
    )

    rgb_raw, rgb_tray_path = load_rgb_tray_growth()

    ms_raw = load_ms_tray_growth()

    print(
        f"\nRGB tray table loaded:\n{rgb_tray_path}"
    )

    print(
        f"RGB tray rows: {len(rgb_raw)}"
    )

    print(
        f"\nMS tray table loaded:\n{MS_TRAY_GROWTH}"
    )

    print(
        f"MS tray rows: {len(ms_raw)}"
    )

    rgb = standardise_rgb_tray(
        rgb_raw
    )

    ms = standardise_ms_tray(
        ms_raw
    )

    combined = combine_rgb_ms_trays(
        rgb,
        ms,
    )

    group_synthesis = create_group_synthesis(
        combined
    )

    inside_outside = create_inside_outside_synthesis(
        combined
    )

    input_inventory = create_input_inventory(
        rgb_tray_path
    )

    charts = create_charts(
        combined,
        group_synthesis,
        inside_outside,
    )

    table_paths = save_tables(
        combined,
        group_synthesis,
        inside_outside,
        input_inventory,
    )

    word_path = (
        REPORTS_ROOT
        / "final_rgb_ms_synthesis_report.docx"
    )

    create_word_report(
        word_path,
        combined,
        group_synthesis,
        inside_outside,
        input_inventory,
        charts,
    )

    settings_path = save_settings(
        CONFIG_ROOT
        / "final_rgb_ms_synthesis_settings.json",
        rgb_tray_path,
    )

    print(
        "\n"
        + "=" * 78
    )

    print(
        "SCRIPT 09 FINISHED"
    )

    print(
        "=" * 78
    )

    print(
        f"Combined tray rows: {len(combined)} "
        f"(expected {EXPECTED_TRAYS})"
    )

    if not combined.empty:
        best = combined.sort_values(
            "combined_rank",
            na_position="last",
        ).iloc[0]

        print(
            "\nTop combined RGB + MS tray:"
        )

        print(
            f"  {best['tray']} — "
            f"{best['microbe_status']} | "
            f"{best['treatment']} | "
            f"{best['environment_group']}"
        )

        print(
            f"  Combined score: "
            f"{best['combined_rgb_ms_score']:.2f}"
        )

    print(
        "\nMain output files:"
    )

    print(
        f"\nCombined tray synthesis:\n"
        f"{table_paths['combined_tray_synthesis']}"
    )

    print(
        f"\nGroup synthesis:\n"
        f"{table_paths['group_synthesis']}"
    )

    print(
        f"\nInside vs Outside synthesis:\n"
        f"{table_paths['inside_outside_synthesis']}"
    )

    print(
        f"\nInput inventory:\n"
        f"{table_paths['input_inventory']}"
    )

    print(
        f"\nExcel report:\n"
        f"{table_paths['excel']}"
    )

    if DOCX_AVAILABLE:
        print(
            f"\nWord synthesis report:\n"
            f"{word_path}"
        )

    else:
        print(
            "\nWord synthesis report skipped because python-docx is not installed."
        )

    print(
        f"\nSettings:\n"
        f"{settings_path}"
    )

    print(
        f"\nCharts folder:\n"
        f"{CHARTS_ROOT}"
    )

    if len(
        combined
    ) != EXPECTED_TRAYS:
        print(
            "\nWARNING: Combined tray count is not 12. "
            "Check input files for missing tray rows."
        )

        return 1

    return 0


# ============================================================
# 18) CLI
# ============================================================

def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Trial 3 Script 09: final RGB + MS synthesis."
        )
    )

    args = parser.parse_args()

    return run_analysis(
        args
    )


if __name__ == "__main__":
    raise SystemExit(
        main()
    )