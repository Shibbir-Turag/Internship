from __future__ import annotations

"""
SCRIPT 06 — THIRD TRIAL INDEPENDENT MULTISPECTRAL 70-CELL GRID DETECTION

Purpose
-------
Detect the complete 7 × 10 = 70-cell tray lattice directly from each cropped
Trial 3 multispectral image set. MS_NIR is used as the reference image.

Why this version is different
-----------------------------
Earlier Trial 3 versions either:
- accepted an incorrect NIR-only grid too easily, or
- transferred RGB geometry into MS image space, which left a true row/column
  uncovered because RGB and MS crops do not have identical margins.

This version does not use RGB coordinates. It detects the repeated 10-column
and 7-row structure directly from the MS_NIR image using:
1. regular darkness projection profiles across the full image;
2. local cup-centre refinement around the projected lattice;
3. robust smooth-surface fitting to preserve a regular tray geometry;
4. optional separately evaluated Hough-circle candidates if the projection
   candidate is uncertain;
5. strict coverage checks requiring the lattice to reach the real outer rows
   and columns.

Important rule
--------------
Hough configurations are evaluated separately. Their circles are never merged
across threshold/configuration outputs.

Main outputs
------------
outputs/Third trial/06_MS_Cell_Grid_Detection/_reports/
    ms_grid_manifest.csv
    ms_cell_coordinates.csv
    ms_cell_grid_report.xlsx
    ms_independent_grid_work_process.docx

Accepted statuses for Script 07
-------------------------------
    PASS_AUTO
    PASS_MANUAL

CHECK_AUTO and FAIL must be inspected before NDVI/NDRE extraction.
"""

import argparse
import csv
import json
import math
import re
from collections import defaultdict
from pathlib import Path

import cv2
import numpy as np
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image, ImageOps


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

CROP_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "01_Crop_Dual_Reference"
)

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
    / "06_MS_Cell_Grid_Detection"
)

REPORTS_ROOT = OUTPUT_ROOT / "_reports"
CONFIG_ROOT = OUTPUT_ROOT / "_config"
MANUAL_POINTS_JSON = CONFIG_ROOT / "manual_ms_grid_points.json"


# ============================================================
# 2) GRID AND DETECTION SETTINGS
# ============================================================

ROWS = 7
COLS = 10
EXPECTED_CELLS = ROWS * COLS

DAY_NAME_TO_ORDER = {
    "day 1": 1,
    "day 2": 2,
    "day 3": 3,
    "day 4": 4,
    "day 5": 5,
    "day 6": 6,
    "day 7": 7,
}

REQUIRED_MS_BANDS = [
    "G",
    "R",
    "RE",
    "NIR",
]

REFERENCE_BAND = "NIR"

# Ownership-zone width relative to halfway-to-neighbour boundaries.
# 0.90 gives a small gap and reduces neighbouring-cell contamination.
ZONE_SHRINK_DEFAULT = 0.90

# Projection-lattice search ranges.
X_SPACING_DIVISOR_MIN = 11.5
X_SPACING_DIVISOR_MAX = 9.2

Y_SPACING_DIVISOR_MIN = 8.5
Y_SPACING_DIVISOR_MAX = 6.2

# Local cup-centre refinement.
LOCAL_SEARCH_RADIUS_RATIO = 0.32
LOCAL_PRIOR_SIGMA_RATIO = 0.70

# Robust surface fitting.
MAX_SURFACE_FIT_ITERATIONS = 4
OUTLIER_MAD_MULTIPLIER = 3.0

# Strict outer coverage.
# These stop a mathematically complete but physically compressed grid
# from receiving PASS_AUTO.
MAX_FIRST_COLUMN_X_RATIO = 0.14
MIN_LAST_COLUMN_X_RATIO = 0.86

MAX_FIRST_ROW_Y_RATIO = 0.18
MIN_LAST_ROW_Y_RATIO = 0.82

MIN_HORIZONTAL_COVERAGE_RATIO = 0.76
MIN_VERTICAL_COVERAGE_RATIO = 0.64

# Automatic quality thresholds.
PASS_MAX_SPACING_CV = 0.10
CHECK_MAX_SPACING_CV = 0.18

PASS_MAX_MEDIAN_FIT_RESIDUAL_RATIO = 0.10
CHECK_MAX_MEDIAN_FIT_RESIDUAL_RATIO = 0.18

PASS_MIN_MEAN_CUP_CONTRAST = 15.0
CHECK_MIN_MEAN_CUP_CONTRAST = 6.0

PASS_MIN_CONTRAST_P05 = 2.0
CHECK_MIN_CONTRAST_P05 = -8.0

PASS_MAX_EDGE_CLIPPED_CELLS = 12
CHECK_MAX_EDGE_CLIPPED_CELLS = 24

# If the profile candidate passes these checks, Hough evaluation is skipped
# unless --always-evaluate-hough is supplied.
PROFILE_FAST_ACCEPT_SCORE = 0.62

# Hough is a fallback/secondary method.
# Every configuration is evaluated independently.
HOUGH_CONFIGS = [
    {
        "dp": 1.2,
        "param1": 45,
        "param2": 12,
    },
    {
        "dp": 1.2,
        "param1": 55,
        "param2": 15,
    },
    {
        "dp": 1.2,
        "param1": 65,
        "param2": 18,
    },
    {
        "dp": 1.5,
        "param1": 55,
        "param2": 14,
    },
    {
        "dp": 1.5,
        "param1": 70,
        "param2": 18,
    },
    {
        "dp": 1.8,
        "param1": 75,
        "param2": 20,
    },
]

ACCEPTED_STATUSES = {
    "PASS_AUTO",
    "PASS_MANUAL",
}


# ============================================================
# 3) OPTIONAL IMPORTS
# ============================================================

try:
    import tifffile

    TIFFFILE_AVAILABLE = True
except Exception:
    TIFFFILE_AVAILABLE = False


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 4) GENERAL HELPERS
# ============================================================

def natural_key(value: object):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(
            r"(\d+)",
            str(value),
        )
    ]


def yes_no(value: bool) -> str:
    return "Yes" if value else "No"


def parse_filter_list(value: str | None):
    if not value:
        return None

    return {
        item.strip().casefold()
        for item in value.split(",")
        if item.strip()
    }


def tray_number_from_name(name: str):
    match = re.search(
        r"(\d+)",
        name,
    )

    return int(match.group(1)) if match else ""


def day_sort_key(folder: Path):
    return (
        DAY_NAME_TO_ORDER.get(
            folder.name.casefold(),
            999,
        ),
        natural_key(folder.name),
    )


def record_key(
    day: str,
    tray: str,
    capture_id: str,
):
    return f"{day}|{tray}|{capture_id}"


def relative_path(
    path: Path | None,
    root: Path,
):
    if path is None:
        return ""

    try:
        return str(path.relative_to(root))
    except ValueError:
        return str(path)


def safe_float(
    value,
    default=math.nan,
):
    try:
        if pd.isna(value):
            return default

        return float(value)

    except Exception:
        return default


def safe_int(
    value,
    default="",
):
    try:
        if pd.isna(value):
            return default

        return int(float(value))

    except Exception:
        return default


# ============================================================
# 5) FIND MULTISPECTRAL IMAGE SETS
# ============================================================

def parse_ms_file(path: Path):
    if path.suffix.casefold() not in {
        ".tif",
        ".tiff",
    }:
        return None

    match = re.match(
        r"^(?P<capture>.+)_MS_(?P<band>G|R|RE|NIR)$",
        path.stem.upper(),
    )

    if not match:
        return None

    return {
        "capture_id": match.group("capture"),
        "band": match.group("band"),
        "path": path,
    }


def find_ms_sets(tray_folder: Path):
    grouped = defaultdict(dict)

    for file in tray_folder.rglob("*"):
        if not file.is_file():
            continue

        parsed = parse_ms_file(file)

        if parsed is None:
            continue

        grouped[
            parsed["capture_id"]
        ][
            parsed["band"]
        ] = parsed["path"]

    rows = []

    for capture_id, bands in grouped.items():
        missing_bands = [
            band
            for band in REQUIRED_MS_BANDS
            if band not in bands
        ]

        rows.append(
            {
                "capture_id": capture_id,
                "bands": bands,
                "missing_bands": missing_bands,
                "complete": len(missing_bands) == 0,
            }
        )

    return sorted(
        rows,
        key=lambda item: natural_key(
            item["capture_id"]
        ),
    )


def collect_jobs(
    days_filter=None,
    trays_filter=None,
):
    if not CROP_ROOT.exists():
        raise FileNotFoundError(
            f"Input crop folder not found:\n{CROP_ROOT}"
        )

    jobs = []

    day_folders = sorted(
        [
            folder
            for folder in CROP_ROOT.iterdir()
            if folder.is_dir()
            and folder.name.casefold() in DAY_NAME_TO_ORDER
        ],
        key=day_sort_key,
    )

    for day_folder in day_folders:
        if (
            days_filter
            and day_folder.name.casefold() not in days_filter
        ):
            continue

        tray_folders = sorted(
            [
                folder
                for folder in day_folder.iterdir()
                if folder.is_dir()
                and folder.name.casefold().startswith("tray")
            ],
            key=lambda folder: natural_key(folder.name),
        )

        for tray_folder in tray_folders:
            if (
                trays_filter
                and tray_folder.name.casefold() not in trays_filter
            ):
                continue

            ms_sets = find_ms_sets(
                tray_folder
            )

            for sequence_index, ms_set in enumerate(ms_sets):
                jobs.append(
                    {
                        "day": day_folder.name,
                        "day_order": DAY_NAME_TO_ORDER[
                            day_folder.name.casefold()
                        ],
                        "tray": tray_folder.name,
                        "tray_no": tray_number_from_name(
                            tray_folder.name
                        ),
                        "capture_id": ms_set["capture_id"],
                        "capture_sequence_index": sequence_index,
                        "bands": ms_set["bands"],
                        "missing_bands": ms_set["missing_bands"],
                        "complete": ms_set["complete"],
                    }
                )

    return sorted(
        jobs,
        key=lambda item: (
            item["day_order"],
            int(item["tray_no"]),
            item["capture_sequence_index"],
            natural_key(item["capture_id"]),
        ),
    )


# ============================================================
# 6) IMAGE READING AND PREPROCESSING
# ============================================================

def read_ms_image(path: Path) -> np.ndarray:
    if TIFFFILE_AVAILABLE:
        array = tifffile.imread(
            str(path)
        )
    else:
        with Image.open(path) as image:
            image = ImageOps.exif_transpose(
                image
            )
            array = np.asarray(image)

    if array.ndim == 3:
        array = array[:, :, 0]

    return array.astype(
        np.float32
    )


def normalise_to_uint8(
    array: np.ndarray,
) -> np.ndarray:
    valid = array[
        np.isfinite(array)
    ]

    if valid.size == 0:
        return np.zeros_like(
            array,
            dtype=np.uint8,
        )

    low, high = np.percentile(
        valid,
        [1, 99],
    )

    if high <= low:
        low = float(valid.min())
        high = float(valid.max())

    if high <= low:
        return np.zeros_like(
            array,
            dtype=np.uint8,
        )

    scaled = (
        (array - low)
        / (high - low)
        * 255.0
    )

    return np.clip(
        scaled,
        0,
        255,
    ).astype(np.uint8)


def build_darkness_map(
    gray: np.ndarray,
) -> np.ndarray:
    """
    Create a stable dark-cup response while reducing slow background drift.
    """

    gray_float = gray.astype(
        np.float32
    )

    global_darkness = (
        255.0
        - gray_float
    )

    typical_scale = max(
        9.0,
        min(
            gray.shape[1] / COLS,
            gray.shape[0] / ROWS,
        ),
    )

    background = cv2.GaussianBlur(
        gray_float,
        (0, 0),
        sigmaX=typical_scale * 0.48,
        sigmaY=typical_scale * 0.48,
    )

    local_darkness = np.maximum(
        background - gray_float,
        0.0,
    )

    local_high = float(
        np.percentile(
            local_darkness,
            99,
        )
    )

    if local_high > 0:
        local_darkness = (
            local_darkness
            / local_high
            * 255.0
        )

    darkness = (
        0.72 * global_darkness
        + 0.28 * local_darkness
    )

    darkness = cv2.GaussianBlur(
        darkness.astype(np.float32),
        (0, 0),
        sigmaX=max(
            4.0,
            typical_scale * 0.075,
        ),
        sigmaY=max(
            4.0,
            typical_scale * 0.075,
        ),
    )

    return darkness


def smooth_profile(
    profile: np.ndarray,
    sigma: float,
) -> np.ndarray:
    kernel_size = max(
        3,
        int(round(sigma * 6)),
    )

    if kernel_size % 2 == 0:
        kernel_size += 1

    smoothed = cv2.GaussianBlur(
        profile.astype(
            np.float32
        ).reshape(1, -1),
        (kernel_size, 1),
        sigmaX=sigma,
    )

    return smoothed.ravel()


def robust_normalise_profile(
    profile: np.ndarray,
) -> np.ndarray:
    low = float(
        np.percentile(
            profile,
            5,
        )
    )

    high = float(
        np.percentile(
            profile,
            95,
        )
    )

    if high <= low:
        return np.zeros_like(
            profile,
            dtype=np.float32,
        )

    return (
        (profile - low)
        / (high - low)
    ).astype(np.float32)


# ============================================================
# 7) REGULAR PROJECTION-LATTICE SEARCH
# ============================================================

def best_regular_lattice_on_profile(
    profile: np.ndarray,
    count: int,
    spacing_min: float,
    spacing_max: float,
    first_min: float,
    first_max: float,
    local_peak_radius: int,
):
    normalised = robust_normalise_profile(
        profile
    )

    length = len(normalised)

    best = None

    spacing_values = np.linspace(
        spacing_min,
        spacing_max,
        max(
            12,
            int(
                round(
                    spacing_max - spacing_min
                )
            ) + 1,
        ),
    )

    for spacing in spacing_values:
        allowed_first_max = min(
            first_max,
            length
            - 1
            - (count - 1) * spacing,
        )

        if allowed_first_max < first_min:
            continue

        first_values = np.linspace(
            first_min,
            allowed_first_max,
            max(
                16,
                int(
                    round(
                        allowed_first_max
                        - first_min
                    )
                ) + 1,
            ),
        )

        for first in first_values:
            positions = (
                first
                + spacing
                * np.arange(
                    count,
                    dtype=float,
                )
            )

            peak_values = []
            peak_offsets = []

            for position in positions:
                left = max(
                    0,
                    int(
                        round(
                            position
                            - local_peak_radius
                        )
                    ),
                )

                right = min(
                    length,
                    int(
                        round(
                            position
                            + local_peak_radius
                            + 1
                        )
                    ),
                )

                if right <= left:
                    peak_values.append(0.0)
                    peak_offsets.append(0.0)
                    continue

                local_index = int(
                    np.argmax(
                        normalised[left:right]
                    )
                )

                peak_index = (
                    left
                    + local_index
                )

                peak_values.append(
                    float(
                        normalised[
                            peak_index
                        ]
                    )
                )

                peak_offsets.append(
                    float(
                        peak_index
                        - position
                    )
                )

            mean_peak = float(
                np.mean(
                    peak_values
                )
            )

            minimum_peak = float(
                np.min(
                    peak_values
                )
            )

            offset_sd = float(
                np.std(
                    peak_offsets
                )
            )

            outer_margin = min(
                float(positions[0]),
                float(
                    length
                    - 1
                    - positions[-1]
                ),
            )

            score = (
                mean_peak
                + 0.18 * minimum_peak
                - 0.018 * offset_sd
                + 0.00035 * outer_margin
            )

            result = {
                "score": score,
                "first": float(first),
                "spacing": float(spacing),
                "positions": positions,
                "mean_peak": mean_peak,
                "minimum_peak": minimum_peak,
                "offset_sd": offset_sd,
            }

            if (
                best is None
                or result["score"] > best["score"]
            ):
                best = result

    return best


def create_projection_candidate(
    gray: np.ndarray,
    darkness: np.ndarray,
):
    height, width = gray.shape[:2]

    x_profile = darkness.mean(
        axis=0
    )

    y_profile = darkness.mean(
        axis=1
    )

    typical_spacing = min(
        width / COLS,
        height / ROWS,
    )

    x_profile = smooth_profile(
        x_profile,
        sigma=max(
            4.0,
            typical_spacing * 0.045,
        ),
    )

    y_profile = smooth_profile(
        y_profile,
        sigma=max(
            4.0,
            typical_spacing * 0.045,
        ),
    )

    x_result = best_regular_lattice_on_profile(
        x_profile,
        count=COLS,
        spacing_min=(
            width
            / X_SPACING_DIVISOR_MIN
        ),
        spacing_max=(
            width
            / X_SPACING_DIVISOR_MAX
        ),
        first_min=width * 0.015,
        first_max=width * 0.13,
        local_peak_radius=max(
            5,
            int(
                round(
                    typical_spacing
                    * 0.09
                )
            ),
        ),
    )

    y_result = best_regular_lattice_on_profile(
        y_profile,
        count=ROWS,
        spacing_min=(
            height
            / Y_SPACING_DIVISOR_MIN
        ),
        spacing_max=(
            height
            / Y_SPACING_DIVISOR_MAX
        ),
        first_min=height * 0.015,
        first_max=height * 0.17,
        local_peak_radius=max(
            5,
            int(
                round(
                    typical_spacing
                    * 0.09
                )
            ),
        ),
    )

    if (
        x_result is None
        or y_result is None
    ):
        return None

    centres = {}

    for row_index, y in enumerate(
        y_result["positions"]
    ):
        for col_index, x in enumerate(
            x_result["positions"]
        ):
            centres[
                (
                    row_index,
                    col_index,
                )
            ] = np.asarray(
                [
                    float(x),
                    float(y),
                ],
                dtype=float,
            )

    return {
        "method": "projection_regular_lattice",
        "initial_centres": centres,
        "source_points": [],
        "x_profile_score": x_result["score"],
        "y_profile_score": y_result["score"],
        "x_profile_mean_peak": x_result["mean_peak"],
        "y_profile_mean_peak": y_result["mean_peak"],
        "x_profile_minimum_peak": x_result["minimum_peak"],
        "y_profile_minimum_peak": y_result["minimum_peak"],
        "x_initial_spacing": x_result["spacing"],
        "y_initial_spacing": y_result["spacing"],
    }


# ============================================================
# 8) SEPARATELY EVALUATED HOUGH CANDIDATES
# ============================================================

def non_max_suppression_points(
    points: list[
        tuple[
            float,
            float,
            float,
        ]
    ],
    minimum_distance: float,
):
    points = sorted(
        points,
        key=lambda item: item[2],
        reverse=True,
    )

    kept = []

    for point in points:
        x, y, _radius = point

        if all(
            math.hypot(
                x - existing[0],
                y - existing[1],
            )
            >= minimum_distance
            for existing in kept
        ):
            kept.append(point)

    return kept


def kmeans_1d(
    values: np.ndarray,
    cluster_count: int,
    iterations: int = 50,
):
    values = np.asarray(
        values,
        dtype=float,
    )

    if values.size < cluster_count:
        return None

    centres = np.percentile(
        values,
        np.linspace(
            0,
            100,
            cluster_count,
        ),
    )

    for _ in range(iterations):
        distances = np.abs(
            values[:, None]
            - centres[None, :]
        )

        labels = distances.argmin(
            axis=1
        )

        new_centres = centres.copy()

        for index in range(
            cluster_count
        ):
            members = values[
                labels == index
            ]

            if members.size > 0:
                new_centres[
                    index
                ] = members.mean()

        if np.allclose(
            new_centres,
            centres,
            atol=1e-3,
        ):
            break

        centres = new_centres

    return np.sort(
        centres
    )


def create_hough_candidates(
    gray: np.ndarray,
):
    height, width = gray.shape[:2]

    typical_spacing = min(
        width / COLS,
        height / ROWS,
    )

    enhanced = cv2.createCLAHE(
        clipLimit=2.0,
        tileGridSize=(8, 8),
    ).apply(gray)

    variants = {
        "enhanced": enhanced,
        "inverted": cv2.bitwise_not(
            enhanced
        ),
        "blurred": cv2.GaussianBlur(
            enhanced,
            (5, 5),
            0,
        ),
        "inverted_blurred": cv2.GaussianBlur(
            cv2.bitwise_not(
                enhanced
            ),
            (5, 5),
            0,
        ),
    }

    min_distance = max(
        10,
        int(
            round(
                typical_spacing
                * 0.55
            )
        ),
    )

    min_radius = max(
        5,
        int(
            round(
                typical_spacing
                * 0.19
            )
        ),
    )

    max_radius = max(
        min_radius + 3,
        int(
            round(
                typical_spacing
                * 0.48
            )
        ),
    )

    candidates = []

    for variant_name, image in variants.items():
        for config_index, config in enumerate(
            HOUGH_CONFIGS,
            start=1,
        ):
            circles = cv2.HoughCircles(
                image,
                cv2.HOUGH_GRADIENT,
                dp=config["dp"],
                minDist=min_distance,
                param1=config["param1"],
                param2=config["param2"],
                minRadius=min_radius,
                maxRadius=max_radius,
            )

            if circles is None:
                continue

            points = [
                (
                    float(x),
                    float(y),
                    float(radius),
                )
                for x, y, radius
                in circles[0]
            ]

            points = non_max_suppression_points(
                points,
                minimum_distance=(
                    typical_spacing
                    * 0.38
                ),
            )

            if len(points) < 35:
                continue

            point_array = np.asarray(
                [
                    [
                        point[0],
                        point[1],
                    ]
                    for point in points
                ],
                dtype=float,
            )

            x_centres = kmeans_1d(
                point_array[:, 0],
                COLS,
            )

            y_centres = kmeans_1d(
                point_array[:, 1],
                ROWS,
            )

            if (
                x_centres is None
                or y_centres is None
            ):
                continue

            centres = {}

            for row_index, y in enumerate(
                y_centres
            ):
                for col_index, x in enumerate(
                    x_centres
                ):
                    centres[
                        (
                            row_index,
                            col_index,
                        )
                    ] = np.asarray(
                        [
                            float(x),
                            float(y),
                        ],
                        dtype=float,
                    )

            candidates.append(
                {
                    "method": (
                        f"hough_{variant_name}"
                        f"_config_{config_index}"
                    ),
                    "initial_centres": centres,
                    "source_points": points,
                    "x_profile_score": math.nan,
                    "y_profile_score": math.nan,
                    "x_profile_mean_peak": math.nan,
                    "y_profile_mean_peak": math.nan,
                    "x_profile_minimum_peak": math.nan,
                    "y_profile_minimum_peak": math.nan,
                    "x_initial_spacing": float(
                        np.median(
                            np.diff(
                                x_centres
                            )
                        )
                    ),
                    "y_initial_spacing": float(
                        np.median(
                            np.diff(
                                y_centres
                            )
                        )
                    ),
                }
            )

    return candidates


# ============================================================
# 9) LOCAL CENTRE REFINEMENT
# ============================================================

def typical_spacing_from_centres(
    centres,
):
    horizontal = []
    vertical = []

    for row_index in range(ROWS):
        for col_index in range(
            COLS - 1
        ):
            horizontal.append(
                float(
                    np.linalg.norm(
                        centres[
                            (
                                row_index,
                                col_index + 1,
                            )
                        ]
                        - centres[
                            (
                                row_index,
                                col_index,
                            )
                        ]
                    )
                )
            )

    for row_index in range(
        ROWS - 1
    ):
        for col_index in range(COLS):
            vertical.append(
                float(
                    np.linalg.norm(
                        centres[
                            (
                                row_index + 1,
                                col_index,
                            )
                        ]
                        - centres[
                            (
                                row_index,
                                col_index,
                            )
                        ]
                    )
                )
            )

    all_values = (
        horizontal
        + vertical
    )

    return max(
        1.0,
        float(
            np.median(
                all_values
            )
        ),
    )


def refine_centres_locally(
    darkness: np.ndarray,
    initial_centres,
):
    height, width = darkness.shape[:2]

    spacing = typical_spacing_from_centres(
        initial_centres
    )

    radius = max(
        8,
        int(
            round(
                spacing
                * LOCAL_SEARCH_RADIUS_RATIO
            )
        ),
    )

    sigma = max(
        3.0,
        radius
        * LOCAL_PRIOR_SIGMA_RATIO,
    )

    refined = {}
    movement_values = []
    support_values = []

    for key, centre in initial_centres.items():
        x = float(centre[0])
        y = float(centre[1])

        x0 = max(
            0,
            int(
                round(
                    x - radius
                )
            ),
        )

        x1 = min(
            width,
            int(
                round(
                    x + radius + 1
                )
            ),
        )

        y0 = max(
            0,
            int(
                round(
                    y - radius
                )
            ),
        )

        y1 = min(
            height,
            int(
                round(
                    y + radius + 1
                )
            ),
        )

        patch = darkness[
            y0:y1,
            x0:x1,
        ]

        if patch.size < 20:
            refined[key] = centre.copy()
            movement_values.append(0.0)
            support_values.append(0.0)
            continue

        yy, xx = np.mgrid[
            y0:y1,
            x0:x1,
        ]

        threshold = float(
            np.percentile(
                patch,
                55,
            )
        )

        weights = np.maximum(
            patch.astype(np.float32)
            - threshold,
            0.0,
        )

        distance_squared = (
            (xx - x) ** 2
            + (yy - y) ** 2
        )

        prior = np.exp(
            -distance_squared
            / (
                2.0
                * sigma
                * sigma
            )
        )

        combined = (
            weights
            * prior
        )

        if combined.sum() <= 1e-6:
            refined_point = centre.copy()
            support = 0.0
        else:
            refined_point = np.asarray(
                [
                    float(
                        (
                            combined
                            * xx
                        ).sum()
                        / combined.sum()
                    ),
                    float(
                        (
                            combined
                            * yy
                        ).sum()
                        / combined.sum()
                    ),
                ],
                dtype=float,
            )

            support = float(
                np.percentile(
                    patch,
                    85,
                )
                - np.percentile(
                    patch,
                    35,
                )
            )

        maximum_movement = (
            spacing
            * 0.20
        )

        movement = float(
            np.linalg.norm(
                refined_point
                - centre
            )
        )

        if movement > maximum_movement:
            refined_point = centre.copy()
            movement = 0.0

        refined[key] = refined_point
        movement_values.append(
            movement
        )
        support_values.append(
            support
        )

    return refined, {
        "typical_spacing_initial": spacing,
        "local_search_radius_px": radius,
        "mean_local_movement_px": float(
            np.mean(
                movement_values
            )
        ),
        "max_local_movement_px": float(
            np.max(
                movement_values
            )
        ),
        "mean_local_support": float(
            np.mean(
                support_values
            )
        ),
    }


# ============================================================
# 10) ROBUST SMOOTH-SURFACE FITTING
# ============================================================

def polynomial_terms(
    row_position: float,
    col_position: float,
):
    return np.asarray(
        [
            1.0,
            float(col_position),
            float(row_position),
            float(
                row_position
                * col_position
            ),
            float(
                col_position
                * col_position
            ),
            float(
                row_position
                * row_position
            ),
        ],
        dtype=float,
    )


def fit_robust_surface(
    refined_centres,
):
    keys = list(
        refined_centres.keys()
    )

    matrix = np.asarray(
        [
            polynomial_terms(
                row_index,
                col_index,
            )
            for row_index, col_index
            in keys
        ],
        dtype=float,
    )

    x_values = np.asarray(
        [
            refined_centres[key][0]
            for key in keys
        ],
        dtype=float,
    )

    y_values = np.asarray(
        [
            refined_centres[key][1]
            for key in keys
        ],
        dtype=float,
    )

    inlier_mask = np.ones(
        len(keys),
        dtype=bool,
    )

    for _ in range(
        MAX_SURFACE_FIT_ITERATIONS
    ):
        x_coefficients, *_ = np.linalg.lstsq(
            matrix[inlier_mask],
            x_values[inlier_mask],
            rcond=None,
        )

        y_coefficients, *_ = np.linalg.lstsq(
            matrix[inlier_mask],
            y_values[inlier_mask],
            rcond=None,
        )

        predicted_x = (
            matrix
            @ x_coefficients
        )

        predicted_y = (
            matrix
            @ y_coefficients
        )

        residuals = np.sqrt(
            (
                predicted_x
                - x_values
            )
            ** 2
            + (
                predicted_y
                - y_values
            )
            ** 2
        )

        active_residuals = residuals[
            inlier_mask
        ]

        median = float(
            np.median(
                active_residuals
            )
        )

        mad = float(
            np.median(
                np.abs(
                    active_residuals
                    - median
                )
            )
        )

        threshold = (
            median
            + OUTLIER_MAD_MULTIPLIER
            * 1.4826
            * mad
            + 2.0
        )

        new_mask = (
            residuals
            <= threshold
        )

        if new_mask.sum() < 20:
            break

        if np.array_equal(
            new_mask,
            inlier_mask,
        ):
            break

        inlier_mask = new_mask

    return {
        "x_coefficients": x_coefficients,
        "y_coefficients": y_coefficients,
        "residuals": residuals,
        "inlier_mask": inlier_mask,
    }


def evaluate_surface(
    surface,
    row_position: float,
    col_position: float,
):
    terms = polynomial_terms(
        row_position,
        col_position,
    )

    return np.asarray(
        [
            float(
                np.dot(
                    terms,
                    surface[
                        "x_coefficients"
                    ],
                )
            ),
            float(
                np.dot(
                    terms,
                    surface[
                        "y_coefficients"
                    ],
                )
            ),
        ],
        dtype=float,
    )


def clip_polygon(
    polygon,
    image_width: int,
    image_height: int,
):
    original = polygon.copy()
    clipped = polygon.copy()

    clipped[:, 0] = np.clip(
        clipped[:, 0],
        0,
        image_width - 1,
    )

    clipped[:, 1] = np.clip(
        clipped[:, 1],
        0,
        image_height - 1,
    )

    was_clipped = not np.allclose(
        original,
        clipped,
        atol=1e-6,
    )

    return clipped, bool(
        was_clipped
    )


def polygon_area(points):
    x = points[:, 0]
    y = points[:, 1]

    return float(
        0.5
        * abs(
            np.dot(
                x,
                np.roll(
                    y,
                    -1,
                ),
            )
            - np.dot(
                y,
                np.roll(
                    x,
                    -1,
                ),
            )
        )
    )


def build_grid_rows_from_surface(
    surface,
    image_width: int,
    image_height: int,
    zone_shrink: float,
):
    rows = []
    edge_clipped_count = 0
    centre_outside_count = 0
    cell_id = 1

    half_step = (
        0.5
        * zone_shrink
    )

    for row_index in range(ROWS):
        for col_index in range(COLS):
            centre = evaluate_surface(
                surface,
                row_index,
                col_index,
            )

            polygon = np.vstack(
                [
                    evaluate_surface(
                        surface,
                        row_index - half_step,
                        col_index - half_step,
                    ),
                    evaluate_surface(
                        surface,
                        row_index - half_step,
                        col_index + half_step,
                    ),
                    evaluate_surface(
                        surface,
                        row_index + half_step,
                        col_index + half_step,
                    ),
                    evaluate_surface(
                        surface,
                        row_index + half_step,
                        col_index - half_step,
                    ),
                ]
            )

            polygon, was_clipped = clip_polygon(
                polygon,
                image_width,
                image_height,
            )

            centre_inside = (
                0
                <= centre[0]
                < image_width
                and 0
                <= centre[1]
                < image_height
            )

            if was_clipped:
                edge_clipped_count += 1

            if not centre_inside:
                centre_outside_count += 1

            x_values = polygon[:, 0]
            y_values = polygon[:, 1]

            rows.append(
                {
                    "cell_id": cell_id,
                    "row": row_index + 1,
                    "column": col_index + 1,
                    "x": float(
                        centre[0]
                    ),
                    "y": float(
                        centre[1]
                    ),
                    "poly_tl_x": float(
                        polygon[0, 0]
                    ),
                    "poly_tl_y": float(
                        polygon[0, 1]
                    ),
                    "poly_tr_x": float(
                        polygon[1, 0]
                    ),
                    "poly_tr_y": float(
                        polygon[1, 1]
                    ),
                    "poly_br_x": float(
                        polygon[2, 0]
                    ),
                    "poly_br_y": float(
                        polygon[2, 1]
                    ),
                    "poly_bl_x": float(
                        polygon[3, 0]
                    ),
                    "poly_bl_y": float(
                        polygon[3, 1]
                    ),
                    "square_x0": float(
                        np.min(
                            x_values
                        )
                    ),
                    "square_y0": float(
                        np.min(
                            y_values
                        )
                    ),
                    "square_x1": float(
                        np.max(
                            x_values
                        )
                    ),
                    "square_y1": float(
                        np.max(
                            y_values
                        )
                    ),
                    "zone_area_px": polygon_area(
                        polygon
                    ),
                    "centre_inside_image": yes_no(
                        centre_inside
                    ),
                    "polygon_clipped_at_image_edge": yes_no(
                        was_clipped
                    ),
                }
            )

            cell_id += 1

    return rows, {
        "edge_clipped_cell_count": edge_clipped_count,
        "centre_outside_count": centre_outside_count,
    }


# ============================================================
# 11) QUALITY METRICS
# ============================================================

def final_centres_from_rows(
    grid_rows,
):
    return {
        (
            int(row["row"]) - 1,
            int(row["column"]) - 1,
        ): np.asarray(
            [
                float(row["x"]),
                float(row["y"]),
            ],
            dtype=float,
        )
        for row in grid_rows
    }


def calculate_spacing_metrics(
    centres,
):
    horizontal = []
    vertical = []

    for row_index in range(ROWS):
        for col_index in range(
            COLS - 1
        ):
            horizontal.append(
                float(
                    np.linalg.norm(
                        centres[
                            (
                                row_index,
                                col_index + 1,
                            )
                        ]
                        - centres[
                            (
                                row_index,
                                col_index,
                            )
                        ]
                    )
                )
            )

    for row_index in range(
        ROWS - 1
    ):
        for col_index in range(COLS):
            vertical.append(
                float(
                    np.linalg.norm(
                        centres[
                            (
                                row_index + 1,
                                col_index,
                            )
                        ]
                        - centres[
                            (
                                row_index,
                                col_index,
                            )
                        ]
                    )
                )
            )

    horizontal = np.asarray(
        horizontal,
        dtype=float,
    )

    vertical = np.asarray(
        vertical,
        dtype=float,
    )

    horizontal_mean = float(
        np.mean(
            horizontal
        )
    )

    vertical_mean = float(
        np.mean(
            vertical
        )
    )

    horizontal_cv = (
        float(
            np.std(horizontal)
            / horizontal_mean
        )
        if horizontal_mean > 0
        else 999.0
    )

    vertical_cv = (
        float(
            np.std(vertical)
            / vertical_mean
        )
        if vertical_mean > 0
        else 999.0
    )

    typical_spacing = max(
        1.0,
        float(
            np.median(
                np.concatenate(
                    [
                        horizontal,
                        vertical,
                    ]
                )
            )
        ),
    )

    return {
        "horizontal_spacing_mean": horizontal_mean,
        "vertical_spacing_mean": vertical_mean,
        "horizontal_spacing_cv": horizontal_cv,
        "vertical_spacing_cv": vertical_cv,
        "max_spacing_cv": max(
            horizontal_cv,
            vertical_cv,
        ),
        "typical_spacing": typical_spacing,
    }


def calculate_cup_contrast(
    gray: np.ndarray,
    centres,
    typical_spacing: float,
):
    height, width = gray.shape[:2]

    contrasts = []

    outer_radius = max(
        8,
        int(
            round(
                typical_spacing
                * 0.49
            )
        ),
    )

    for centre in centres.values():
        x = float(centre[0])
        y = float(centre[1])

        x0 = max(
            0,
            int(
                round(
                    x - outer_radius
                )
            ),
        )

        x1 = min(
            width,
            int(
                round(
                    x + outer_radius + 1
                )
            ),
        )

        y0 = max(
            0,
            int(
                round(
                    y - outer_radius
                )
            ),
        )

        y1 = min(
            height,
            int(
                round(
                    y + outer_radius + 1
                )
            ),
        )

        patch = gray[
            y0:y1,
            x0:x1,
        ].astype(
            np.float32
        )

        if patch.size < 20:
            contrasts.append(
                -255.0
            )
            continue

        yy, xx = np.mgrid[
            y0:y1,
            x0:x1,
        ]

        radius_ratio = (
            np.sqrt(
                (xx - x) ** 2
                + (yy - y) ** 2
            )
            / typical_spacing
        )

        inner = patch[
            radius_ratio <= 0.25
        ]

        ring = patch[
            (
                radius_ratio >= 0.34
            )
            & (
                radius_ratio <= 0.48
            )
        ]

        if (
            inner.size == 0
            or ring.size == 0
        ):
            contrasts.append(
                -255.0
            )
            continue

        contrasts.append(
            float(
                np.mean(ring)
                - np.mean(inner)
            )
        )

    contrast_array = np.asarray(
        contrasts,
        dtype=float,
    )

    return {
        "mean_cup_contrast": float(
            np.mean(
                contrast_array
            )
        ),
        "minimum_cup_contrast": float(
            np.min(
                contrast_array
            )
        ),
        "cup_contrast_p05": float(
            np.percentile(
                contrast_array,
                5,
            )
        ),
        "cup_contrast_median": float(
            np.median(
                contrast_array
            )
        ),
    }


def calculate_coverage_metrics(
    centres,
    width: int,
    height: int,
):
    first_column_x = float(
        np.mean(
            [
                centres[
                    (
                        row_index,
                        0,
                    )
                ][0]
                for row_index
                in range(ROWS)
            ]
        )
    )

    last_column_x = float(
        np.mean(
            [
                centres[
                    (
                        row_index,
                        COLS - 1,
                    )
                ][0]
                for row_index
                in range(ROWS)
            ]
        )
    )

    first_row_y = float(
        np.mean(
            [
                centres[
                    (
                        0,
                        col_index,
                    )
                ][1]
                for col_index
                in range(COLS)
            ]
        )
    )

    last_row_y = float(
        np.mean(
            [
                centres[
                    (
                        ROWS - 1,
                        col_index,
                    )
                ][1]
                for col_index
                in range(COLS)
            ]
        )
    )

    horizontal_coverage = (
        last_column_x
        - first_column_x
    ) / max(
        width,
        1,
    )

    vertical_coverage = (
        last_row_y
        - first_row_y
    ) / max(
        height,
        1,
    )

    outer_coverage_pass = (
        first_column_x / width
        <= MAX_FIRST_COLUMN_X_RATIO
        and last_column_x / width
        >= MIN_LAST_COLUMN_X_RATIO
        and first_row_y / height
        <= MAX_FIRST_ROW_Y_RATIO
        and last_row_y / height
        >= MIN_LAST_ROW_Y_RATIO
        and horizontal_coverage
        >= MIN_HORIZONTAL_COVERAGE_RATIO
        and vertical_coverage
        >= MIN_VERTICAL_COVERAGE_RATIO
    )

    return {
        "first_column_x_ratio": (
            first_column_x
            / width
        ),
        "last_column_x_ratio": (
            last_column_x
            / width
        ),
        "first_row_y_ratio": (
            first_row_y
            / height
        ),
        "last_row_y_ratio": (
            last_row_y
            / height
        ),
        "horizontal_coverage_ratio": horizontal_coverage,
        "vertical_coverage_ratio": vertical_coverage,
        "outer_coverage_pass": yes_no(
            outer_coverage_pass
        ),
    }


def calculate_hough_support(
    source_points,
    centres,
    typical_spacing,
):
    if not source_points:
        return {
            "source_point_count": 0,
            "source_supported_cell_count": 0,
            "source_mean_distance_ratio": math.nan,
        }

    source_array = np.asarray(
        [
            [
                point[0],
                point[1],
            ]
            for point in source_points
        ],
        dtype=float,
    )

    grid_array = np.asarray(
        list(
            centres.values()
        ),
        dtype=float,
    )

    distances = np.sqrt(
        (
            (
                source_array[:, None, :]
                - grid_array[None, :, :]
            )
            ** 2
        ).sum(
            axis=2
        )
    )

    nearest_grid = distances.argmin(
        axis=1
    )

    nearest_distance = distances.min(
        axis=1
    )

    supported = (
        nearest_distance
        <= typical_spacing * 0.42
    )

    supported_cells = len(
        set(
            nearest_grid[
                supported
            ].tolist()
        )
    )

    mean_distance_ratio = (
        float(
            np.mean(
                nearest_distance[
                    supported
                ]
            )
            / typical_spacing
        )
        if supported.any()
        else math.nan
    )

    return {
        "source_point_count": len(
            source_points
        ),
        "source_supported_cell_count": supported_cells,
        "source_mean_distance_ratio": mean_distance_ratio,
    }


# ============================================================
# 12) STATUS AND CANDIDATE SCORING
# ============================================================

def assign_status(metrics):
    coverage_pass = (
        metrics[
            "outer_coverage_pass"
        ]
        == "Yes"
    )

    pass_conditions = (
        coverage_pass
        and metrics[
            "centre_outside_count"
        ] == 0
        and metrics[
            "edge_clipped_cell_count"
        ] <= PASS_MAX_EDGE_CLIPPED_CELLS
        and metrics[
            "max_spacing_cv"
        ] <= PASS_MAX_SPACING_CV
        and metrics[
            "median_fit_residual_ratio"
        ] <= PASS_MAX_MEDIAN_FIT_RESIDUAL_RATIO
        and metrics[
            "mean_cup_contrast"
        ] >= PASS_MIN_MEAN_CUP_CONTRAST
        and metrics[
            "cup_contrast_p05"
        ] >= PASS_MIN_CONTRAST_P05
    )

    if pass_conditions:
        return "PASS_AUTO", ""

    check_conditions = (
        coverage_pass
        and metrics[
            "centre_outside_count"
        ] == 0
        and metrics[
            "edge_clipped_cell_count"
        ] <= CHECK_MAX_EDGE_CLIPPED_CELLS
        and metrics[
            "max_spacing_cv"
        ] <= CHECK_MAX_SPACING_CV
        and metrics[
            "median_fit_residual_ratio"
        ] <= CHECK_MAX_MEDIAN_FIT_RESIDUAL_RATIO
        and metrics[
            "mean_cup_contrast"
        ] >= CHECK_MIN_MEAN_CUP_CONTRAST
        and metrics[
            "cup_contrast_p05"
        ] >= CHECK_MIN_CONTRAST_P05
    )

    if check_conditions:
        reasons = []

        if (
            metrics["max_spacing_cv"]
            > PASS_MAX_SPACING_CV
        ):
            reasons.append(
                "spacing variation above PASS threshold"
            )

        if (
            metrics[
                "median_fit_residual_ratio"
            ]
            > PASS_MAX_MEDIAN_FIT_RESIDUAL_RATIO
        ):
            reasons.append(
                "surface-fit residual above PASS threshold"
            )

        if (
            metrics[
                "mean_cup_contrast"
            ]
            < PASS_MIN_MEAN_CUP_CONTRAST
        ):
            reasons.append(
                "mean cup contrast below PASS threshold"
            )

        if (
            metrics[
                "cup_contrast_p05"
            ]
            < PASS_MIN_CONTRAST_P05
        ):
            reasons.append(
                "some cells have weak cup contrast"
            )

        if (
            metrics[
                "edge_clipped_cell_count"
            ]
            > PASS_MAX_EDGE_CLIPPED_CELLS
        ):
            reasons.append(
                "many outer polygons touch the image edge"
            )

        return (
            "CHECK_AUTO",
            "; ".join(reasons),
        )

    failure_reasons = []

    if not coverage_pass:
        failure_reasons.append(
            "grid does not cover the true outer rows/columns"
        )

    if (
        metrics[
            "centre_outside_count"
        ]
        > 0
    ):
        failure_reasons.append(
            "one or more centres are outside the image"
        )

    if (
        metrics[
            "max_spacing_cv"
        ]
        > CHECK_MAX_SPACING_CV
    ):
        failure_reasons.append(
            "irregular cell spacing"
        )

    if (
        metrics[
            "median_fit_residual_ratio"
        ]
        > CHECK_MAX_MEDIAN_FIT_RESIDUAL_RATIO
    ):
        failure_reasons.append(
            "unstable surface fit"
        )

    if (
        metrics[
            "mean_cup_contrast"
        ]
        < CHECK_MIN_MEAN_CUP_CONTRAST
    ):
        failure_reasons.append(
            "weak cup-to-background contrast"
        )

    return (
        "FAIL",
        "; ".join(
            failure_reasons
        ),
    )


def candidate_score(
    metrics,
    method: str,
):
    coverage_bonus = (
        240.0
        if metrics[
            "outer_coverage_pass"
        ] == "Yes"
        else -500.0
    )

    profile_values = [
        safe_float(
            metrics.get(
                "x_profile_score"
            )
        ),
        safe_float(
            metrics.get(
                "y_profile_score"
            )
        ),
    ]

    finite_profiles = [
        value
        for value in profile_values
        if np.isfinite(value)
    ]

    profile_score = (
        float(
            np.mean(
                finite_profiles
            )
        )
        if finite_profiles
        else 0.0
    )

    source_support = safe_float(
        metrics.get(
            "source_supported_cell_count"
        ),
        0.0,
    )

    score = (
        coverage_bonus
        + metrics[
            "mean_cup_contrast"
        ] * 2.2
        + metrics[
            "cup_contrast_p05"
        ] * 0.8
        + profile_score * 120.0
        + source_support * 1.5
        - metrics[
            "max_spacing_cv"
        ] * 500.0
        - metrics[
            "median_fit_residual_ratio"
        ] * 450.0
        - metrics[
            "edge_clipped_cell_count"
        ] * 2.0
        - metrics[
            "centre_outside_count"
        ] * 100.0
    )

    if method.startswith(
        "projection"
    ):
        score += 50.0

    return float(score)


# ============================================================
# 13) EVALUATE ONE CANDIDATE
# ============================================================

def evaluate_candidate(
    gray: np.ndarray,
    darkness: np.ndarray,
    candidate: dict,
    zone_shrink: float,
):
    image_height, image_width = gray.shape[:2]

    refined_centres, refinement_metrics = refine_centres_locally(
        darkness,
        candidate[
            "initial_centres"
        ],
    )

    surface = fit_robust_surface(
        refined_centres
    )

    grid_rows, boundary_metrics = build_grid_rows_from_surface(
        surface,
        image_width=image_width,
        image_height=image_height,
        zone_shrink=zone_shrink,
    )

    centres = final_centres_from_rows(
        grid_rows
    )

    spacing_metrics = calculate_spacing_metrics(
        centres
    )

    contrast_metrics = calculate_cup_contrast(
        gray,
        centres,
        spacing_metrics[
            "typical_spacing"
        ],
    )

    coverage_metrics = calculate_coverage_metrics(
        centres,
        image_width,
        image_height,
    )

    hough_metrics = calculate_hough_support(
        candidate.get(
            "source_points",
            [],
        ),
        centres,
        spacing_metrics[
            "typical_spacing"
        ],
    )

    residuals = surface[
        "residuals"
    ]

    median_fit_residual = float(
        np.median(
            residuals
        )
    )

    maximum_fit_residual = float(
        np.max(
            residuals
        )
    )

    metrics = {
        "method": candidate["method"],
        "grid_rows": grid_rows,
        "cell_count": len(grid_rows),
        "x_profile_score": candidate.get(
            "x_profile_score",
            math.nan,
        ),
        "y_profile_score": candidate.get(
            "y_profile_score",
            math.nan,
        ),
        "x_profile_mean_peak": candidate.get(
            "x_profile_mean_peak",
            math.nan,
        ),
        "y_profile_mean_peak": candidate.get(
            "y_profile_mean_peak",
            math.nan,
        ),
        "x_profile_minimum_peak": candidate.get(
            "x_profile_minimum_peak",
            math.nan,
        ),
        "y_profile_minimum_peak": candidate.get(
            "y_profile_minimum_peak",
            math.nan,
        ),
        "x_initial_spacing": candidate.get(
            "x_initial_spacing",
            math.nan,
        ),
        "y_initial_spacing": candidate.get(
            "y_initial_spacing",
            math.nan,
        ),
        "fit_inlier_count": int(
            surface[
                "inlier_mask"
            ].sum()
        ),
        "median_fit_residual_px": median_fit_residual,
        "maximum_fit_residual_px": maximum_fit_residual,
        "median_fit_residual_ratio": (
            median_fit_residual
            / spacing_metrics[
                "typical_spacing"
            ]
        ),
        "maximum_fit_residual_ratio": (
            maximum_fit_residual
            / spacing_metrics[
                "typical_spacing"
            ]
        ),
    }

    metrics.update(
        refinement_metrics
    )

    metrics.update(
        boundary_metrics
    )

    metrics.update(
        spacing_metrics
    )

    metrics.update(
        contrast_metrics
    )

    metrics.update(
        coverage_metrics
    )

    metrics.update(
        hough_metrics
    )

    status, notes = assign_status(
        metrics
    )

    metrics["status"] = status
    metrics["notes"] = notes

    metrics["candidate_score"] = candidate_score(
        metrics,
        candidate["method"],
    )

    return metrics


def auto_detect_grid(
    gray: np.ndarray,
    args,
):
    darkness = build_darkness_map(
        gray
    )

    evaluated = []

    projection_candidate = create_projection_candidate(
        gray,
        darkness,
    )

    if projection_candidate is not None:
        projection_result = evaluate_candidate(
            gray,
            darkness,
            projection_candidate,
            zone_shrink=args.zone_shrink,
        )

        evaluated.append(
            projection_result
        )

        profile_values = [
            projection_result[
                "x_profile_score"
            ],
            projection_result[
                "y_profile_score"
            ],
        ]

        finite_values = [
            value
            for value in profile_values
            if np.isfinite(value)
        ]

        projection_profile_score = (
            float(
                np.mean(
                    finite_values
                )
            )
            if finite_values
            else 0.0
        )

        if (
            projection_result["status"]
            == "PASS_AUTO"
            and projection_profile_score
            >= PROFILE_FAST_ACCEPT_SCORE
            and not args.always_evaluate_hough
        ):
            projection_result[
                "candidate_count"
            ] = 1

            projection_result[
                "candidate_methods_evaluated"
            ] = projection_result[
                "method"
            ]

            return projection_result

    if not args.profile_only:
        hough_candidates = create_hough_candidates(
            gray
        )

        for hough_candidate in hough_candidates:
            result = evaluate_candidate(
                gray,
                darkness,
                hough_candidate,
                zone_shrink=args.zone_shrink,
            )

            evaluated.append(
                result
            )

    if not evaluated:
        return None

    evaluated.sort(
        key=lambda item: (
            2
            if item["status"]
            == "PASS_AUTO"
            else 1
            if item["status"]
            == "CHECK_AUTO"
            else 0,
            item[
                "candidate_score"
            ],
        ),
        reverse=True,
    )

    best = evaluated[0]

    best[
        "candidate_count"
    ] = len(evaluated)

    best[
        "candidate_methods_evaluated"
    ] = ", ".join(
        result["method"]
        for result in evaluated
    )

    return best


# ============================================================
# 14) OPTIONAL TARGETED MANUAL CORRECTION
# ============================================================

def load_manual_points():
    if not MANUAL_POINTS_JSON.exists():
        return {}

    try:
        with MANUAL_POINTS_JSON.open(
            "r",
            encoding="utf-8",
        ) as file:
            return json.load(file)

    except Exception:
        return {}


def save_manual_points(
    data: dict,
):
    MANUAL_POINTS_JSON.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with MANUAL_POINTS_JSON.open(
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            data,
            file,
            indent=2,
        )


def request_manual_points(
    gray: np.ndarray,
    title: str,
):
    import matplotlib.pyplot as plt

    print(
        "\nTargeted manual correction required."
    )
    print(
        "Click four CELL CENTRES in this order:"
    )
    print(
        "1) Cell 1   — top-left"
    )
    print(
        "2) Cell 10  — top-right"
    )
    print(
        "3) Cell 61  — bottom-left"
    )
    print(
        "4) Cell 70  — bottom-right\n"
    )

    figure, axis = plt.subplots(
        figsize=(12, 8),
    )

    axis.imshow(
        gray,
        cmap="gray",
    )

    axis.set_title(
        title
        + "\nClick Cell 1, Cell 10, Cell 61, Cell 70",
        fontsize=11,
    )

    axis.axis("off")

    points = plt.ginput(
        4,
        timeout=0,
    )

    plt.close(
        figure
    )

    if len(points) != 4:
        raise RuntimeError(
            "Manual correction incomplete. Four points are required."
        )

    return [
        [
            float(x),
            float(y),
        ]
        for x, y in points
    ]


def create_manual_candidate(
    manual_points,
):
    top_left = np.asarray(
        manual_points[0],
        dtype=float,
    )

    top_right = np.asarray(
        manual_points[1],
        dtype=float,
    )

    bottom_left = np.asarray(
        manual_points[2],
        dtype=float,
    )

    bottom_right = np.asarray(
        manual_points[3],
        dtype=float,
    )

    centres = {}

    for row_index in range(ROWS):
        row_fraction = (
            row_index
            / max(
                ROWS - 1,
                1,
            )
        )

        left = (
            top_left
            * (
                1.0
                - row_fraction
            )
            + bottom_left
            * row_fraction
        )

        right = (
            top_right
            * (
                1.0
                - row_fraction
            )
            + bottom_right
            * row_fraction
        )

        for col_index in range(COLS):
            col_fraction = (
                col_index
                / max(
                    COLS - 1,
                    1,
                )
            )

            centres[
                (
                    row_index,
                    col_index,
                )
            ] = (
                left
                * (
                    1.0
                    - col_fraction
                )
                + right
                * col_fraction
            )

    return {
        "method": "manual_four_outer_cell_centres",
        "initial_centres": centres,
        "source_points": [],
        "x_profile_score": math.nan,
        "y_profile_score": math.nan,
        "x_profile_mean_peak": math.nan,
        "y_profile_mean_peak": math.nan,
        "x_profile_minimum_peak": math.nan,
        "y_profile_minimum_peak": math.nan,
        "x_initial_spacing": math.nan,
        "y_initial_spacing": math.nan,
    }


def manual_result(
    gray,
    manual_points,
    zone_shrink: float,
):
    darkness = build_darkness_map(
        gray
    )

    result = evaluate_candidate(
        gray,
        darkness,
        create_manual_candidate(
            manual_points
        ),
        zone_shrink=zone_shrink,
    )

    result[
        "status"
    ] = "PASS_MANUAL"

    result[
        "notes"
    ] = "Targeted manual four-corner correction."

    result[
        "candidate_count"
    ] = 1

    result[
        "candidate_methods_evaluated"
    ] = result["method"]

    return result


# ============================================================
# 15) OVERLAY OUTPUT
# ============================================================

def polygon_from_row(
    row: dict,
):
    return np.asarray(
        [
            [
                row["poly_tl_x"],
                row["poly_tl_y"],
            ],
            [
                row["poly_tr_x"],
                row["poly_tr_y"],
            ],
            [
                row["poly_br_x"],
                row["poly_br_y"],
            ],
            [
                row["poly_bl_x"],
                row["poly_bl_y"],
            ],
        ],
        dtype=np.int32,
    )


def draw_grid_overlay(
    gray: np.ndarray,
    grid_rows,
    output_path: Path,
    title: str,
    status: str,
):
    image_height, image_width = gray.shape[:2]

    image_bgr = cv2.cvtColor(
        gray,
        cv2.COLOR_GRAY2BGR,
    )

    if status == "PASS_AUTO":
        colour = (
            0,
            190,
            0,
        )
    elif status == "PASS_MANUAL":
        colour = (
            255,
            150,
            0,
        )
    elif status == "CHECK_AUTO":
        colour = (
            0,
            220,
            255,
        )
    else:
        colour = (
            0,
            0,
            220,
        )

    median_side = (
        float(
            np.median(
                [
                    math.sqrt(
                        max(
                            row[
                                "zone_area_px"
                            ],
                            1.0,
                        )
                    )
                    for row
                    in grid_rows
                ]
            )
        )
        if grid_rows
        else 35.0
    )

    line_width = max(
        1,
        int(
            round(
                median_side
                / 34
            )
        ),
    )

    font_scale = max(
        0.30,
        min(
            0.72,
            median_side
            / 92,
        ),
    )

    font_thickness = max(
        1,
        int(
            round(
                font_scale
                * 2
            )
        ),
    )

    header_height = max(
        52,
        int(
            round(
                median_side
                * 0.65
            )
        ),
    )

    canvas = np.full(
        (
            image_height
            + header_height,
            image_width,
            3,
        ),
        255,
        dtype=np.uint8,
    )

    canvas[
        header_height:,
        :,
        :,
    ] = image_bgr

    y_offset = header_height

    for row in grid_rows:
        polygon = polygon_from_row(
            row
        )

        polygon[:, 1] += (
            y_offset
        )

        cv2.polylines(
            canvas,
            [polygon],
            isClosed=True,
            color=colour,
            thickness=line_width,
            lineType=cv2.LINE_AA,
        )

        x = int(
            round(
                row["x"]
            )
        )

        y = (
            int(
                round(
                    row["y"]
                )
            )
            + y_offset
        )

        cv2.circle(
            canvas,
            (
                x,
                y,
            ),
            max(
                2,
                line_width,
            ),
            colour,
            thickness=-1,
        )

        cv2.putText(
            canvas,
            str(
                row[
                    "cell_id"
                ]
            ),
            (
                max(
                    0,
                    x - 8,
                ),
                min(
                    canvas.shape[0]
                    - 2,
                    y + 5,
                ),
            ),
            cv2.FONT_HERSHEY_SIMPLEX,
            font_scale,
            colour,
            font_thickness,
            cv2.LINE_AA,
        )

    cv2.putText(
        canvas,
        (
            f"{title} | {status} | "
            f"Cells: {len(grid_rows)}"
        ),
        (
            12,
            int(
                header_height
                * 0.68
            ),
        ),
        cv2.FONT_HERSHEY_SIMPLEX,
        max(
            0.45,
            min(
                0.85,
                font_scale,
            ),
        ),
        (
            0,
            0,
            0,
        ),
        2,
        cv2.LINE_AA,
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    cv2.imwrite(
        str(
            output_path
        ),
        canvas,
        [
            cv2.IMWRITE_JPEG_QUALITY,
            95,
        ],
    )


# ============================================================
# 16) CSV / EXCEL / WORD REPORTING
# ============================================================

def write_csv(
    path: Path,
    rows: list[dict],
):
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    if not rows:
        path.write_text(
            "",
            encoding="utf-8",
        )
        return

    fieldnames = list(
        rows[0].keys()
    )

    with path.open(
        "w",
        newline="",
        encoding="utf-8",
    ) as file:
        writer = csv.DictWriter(
            file,
            fieldnames=fieldnames,
        )

        writer.writeheader()
        writer.writerows(
            rows
        )


def style_excel(
    path: Path,
):
    workbook = load_workbook(
        path
    )

    header_fill = PatternFill(
        "solid",
        fgColor="1F4E78",
    )

    for worksheet in workbook.worksheets:
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.row_dimensions[1].height = 34

        for cell in worksheet[1]:
            cell.fill = header_fill

            cell.font = Font(
                color="FFFFFF",
                bold=True,
            )

            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )

        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(
                    vertical="top",
                    wrap_text=True,
                )

        for column_cells in worksheet.columns:
            letter = (
                column_cells[0]
                .column_letter
            )

            longest = max(
                len(
                    str(
                        cell.value
                    )
                )
                if cell.value is not None
                else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[
                letter
            ].width = min(
                max(
                    12,
                    longest + 2,
                ),
                58,
            )

    workbook.save(
        path
    )


def create_word_work_process_report(
    output_path: Path,
    manifest_frame: pd.DataFrame,
    settings: dict,
):
    if not DOCX_AVAILABLE:
        print(
            "WARNING: python-docx is not installed. "
            "Word report was skipped."
        )
        print(
            "Install it with: pip install python-docx"
        )
        return None

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Times New Roman"

    document.styles[
        "Normal"
    ].font.size = Pt(11)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name in document.styles:
            document.styles[
                style_name
            ].font.name = "Times New Roman"

    title = document.add_heading(
        (
            "Trial 3 Script 06 Work Process Report — "
            "Independent MS Grid Detection"
        ),
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_paragraph(
        "This report documents the workflow used to detect the complete "
        "7 × 10 multispectral tray grid for Trial 3. The method operates "
        "directly on the cropped MS_NIR image and does not transfer RGB/D "
        "coordinates."
    )

    document.add_heading(
        "1. Reason for replacing the earlier method",
        level=1,
    )

    document.add_paragraph(
        "The earlier RGB-guided method produced a mathematically complete "
        "70-cell grid but did not cover the true rightmost column and bottom "
        "row in some MS images. This happened because RGB and multispectral "
        "crops have different margins and camera geometry. A correct MS grid "
        "must therefore be detected independently within each MS image."
    )

    document.add_heading(
        "2. Input data",
        level=1,
    )

    document.add_paragraph(
        (
            "Cropped image folder: "
            f"{settings.get('crop_root', '')}"
        )
    )

    document.add_paragraph(
        "Each complete multispectral set contains MS_G, MS_R, MS_RE, and "
        "MS_NIR. MS_NIR is used as the grid-reference band because the tray "
        "cups and soil normally produce a repeated dark circular pattern "
        "against the lighter tray."
    )

    document.add_heading(
        "3. Automatic detection process",
        level=1,
    )

    process_steps = [
        (
            "The MS_NIR image is normalised to 8-bit only for grid detection "
            "and overlay creation; original band values are not modified."
        ),
        (
            "A darkness response map is created while reducing slow "
            "illumination variation."
        ),
        (
            "Horizontal and vertical darkness profiles are calculated across "
            "the full image."
        ),
        (
            "The script searches for the strongest regular ten-column and "
            "seven-row lattice, including the true outer rows and columns."
        ),
        (
            "Each predicted cell centre is locally refined within a limited "
            "search radius so it cannot drift to a neighbouring cell or "
            "handwritten label."
        ),
        (
            "A robust smooth polynomial surface is fitted to all 70 refined "
            "centres. Outliers are controlled through iterative residual "
            "filtering."
        ),
        (
            "Polygon ownership zones are generated around every final centre "
            "for later NDVI and NDRE extraction."
        ),
        (
            "Strict outer-coverage checks prevent a compressed grid from "
            "being accepted."
        ),
        (
            "If the profile result is uncertain, Hough-circle configurations "
            "are evaluated separately. Their detections are never merged "
            "across configurations."
        ),
    ]

    for step in process_steps:
        document.add_paragraph(
            step,
            style="List Number",
        )

    document.add_heading(
        "4. Validation status summary",
        level=1,
    )

    if (
        not manifest_frame.empty
        and "status" in manifest_frame.columns
    ):
        summary = (
            manifest_frame[
                "status"
            ]
            .value_counts()
            .reset_index()
        )

        summary.columns = [
            "Status",
            "Count",
        ]

        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.style = "Table Grid"

        table.rows[0].cells[0].text = "Status"
        table.rows[0].cells[1].text = "Count"

        for row in summary.itertuples(
            index=False
        ):
            cells = table.add_row().cells

            cells[0].text = str(
                row.Status
            )

            cells[1].text = str(
                row.Count
            )

    else:
        document.add_paragraph(
            "No status data was available."
        )

    document.add_paragraph(
        "PASS_AUTO and PASS_MANUAL are acceptable inputs for Script 07 "
        "after visual overlay checking. CHECK_AUTO and FAIL rows should "
        "not be used for vegetation-index extraction until corrected."
    )

    document.add_heading(
        "5. Generated files",
        level=1,
    )

    output_descriptions = [
        (
            "ms_grid_manifest.csv",
            (
                "One row per MS image set, including method, status, "
                "coverage, spacing, contrast, fit residuals, and overlay path."
            ),
        ),
        (
            "ms_cell_coordinates.csv",
            (
                "One row per cell, containing the final MS-specific centre, "
                "polygon corners, bounding box, image paths, and validation "
                "status."
            ),
        ),
        (
            "ms_cell_grid_report.xlsx",
            (
                "Excel workbook containing the manifest, coordinate table, "
                "status summary, and read-me guidance."
            ),
        ),
        (
            "ms_independent_grid_work_process.docx",
            (
                "This work-process report describing the method and "
                "validation requirements."
            ),
        ),
        (
            "MS_NIR grid overlay images",
            (
                "Visual evidence showing whether all seven rows and ten "
                "columns are correctly covered."
            ),
        ),
    ]

    for filename, description in output_descriptions:
        paragraph = document.add_paragraph()

        paragraph.add_run(
            filename
        ).bold = True

        paragraph.add_run(
            f": {description}"
        )

    document.add_heading(
        "6. Visual checking requirements",
        level=1,
    )

    document.add_paragraph(
        "Inspect overlays from different days and trays. Cell 1 must "
        "correspond to the real top-left cup, Cell 10 to the real top-right "
        "cup, Cell 61 to the real bottom-left cup, and Cell 70 to the real "
        "bottom-right cup. The handwritten label must not be interpreted "
        "as an additional column."
    )

    document.add_heading(
        "7. Use in Script 07",
        level=1,
    )

    document.add_paragraph(
        "Script 07 should use the polygon columns in "
        "ms_cell_coordinates.csv. Observed multispectral values must be "
        "extracted from the original TIFF bands, not from the normalised "
        "overlay image."
    )

    document.add_heading(
        "8. Limitations",
        level=1,
    )

    document.add_paragraph(
        "Strong image cropping errors, severe rotation, incomplete trays, "
        "or unusual lighting can still produce CHECK_AUTO or FAIL. Only "
        "those exceptional image sets should use targeted manual "
        "four-corner correction."
    )

    document.save(
        output_path
    )

    return output_path


def write_reports(
    manifest_rows,
    coordinate_rows,
    settings,
):
    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    CONFIG_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    manifest_csv = (
        REPORTS_ROOT
        / "ms_grid_manifest.csv"
    )

    coordinates_csv = (
        REPORTS_ROOT
        / "ms_cell_coordinates.csv"
    )

    excel_path = (
        REPORTS_ROOT
        / "ms_cell_grid_report.xlsx"
    )

    settings_path = (
        CONFIG_ROOT
        / "ms_grid_detection_settings.json"
    )

    word_path = (
        REPORTS_ROOT
        / "ms_independent_grid_work_process.docx"
    )

    write_csv(
        manifest_csv,
        manifest_rows,
    )

    write_csv(
        coordinates_csv,
        coordinate_rows,
    )

    manifest_frame = pd.DataFrame(
        manifest_rows
    )

    coordinates_frame = pd.DataFrame(
        coordinate_rows
    )

    if (
        not manifest_frame.empty
        and "status" in manifest_frame.columns
    ):
        status_summary = (
            manifest_frame[
                "status"
            ]
            .value_counts()
            .reset_index()
        )

        status_summary.columns = [
            "status",
            "count",
        ]

    else:
        status_summary = pd.DataFrame()

    readme = pd.DataFrame(
        {
            "Notes": [
                (
                    "This workbook contains independent MS_NIR "
                    "70-cell grid detection results."
                ),
                (
                    "The method does not use RGB/D coordinates."
                ),
                (
                    "The primary detector uses regular darkness "
                    "projections across the complete image."
                ),
                (
                    "Hough configurations are evaluated separately "
                    "and are never merged."
                ),
                (
                    "PASS_AUTO and PASS_MANUAL may be used by "
                    "Script 07 after overlay inspection."
                ),
                (
                    "CHECK_AUTO and FAIL require review or targeted "
                    "correction."
                ),
                (
                    "The final coordinate table includes polygon zones "
                    "for NDVI/NDRE extraction."
                ),
                (
                    "Cell 1/10/61/70 must align with the real four "
                    "outer tray cells."
                ),
            ]
        }
    )

    with pd.ExcelWriter(
        excel_path,
        engine="openpyxl",
    ) as writer:
        manifest_frame.to_excel(
            writer,
            sheet_name="MS Grid Manifest",
            index=False,
        )

        coordinates_frame.to_excel(
            writer,
            sheet_name="MS Cell Coordinates",
            index=False,
        )

        status_summary.to_excel(
            writer,
            sheet_name="Status Summary",
            index=False,
        )

        readme.to_excel(
            writer,
            sheet_name="Read Me",
            index=False,
        )

    style_excel(
        excel_path
    )

    with settings_path.open(
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            settings,
            file,
            indent=2,
        )

    create_word_work_process_report(
        word_path,
        manifest_frame,
        settings,
    )

    return {
        "manifest_csv": manifest_csv,
        "coordinates_csv": coordinates_csv,
        "excel_path": excel_path,
        "settings_path": settings_path,
        "word_path": word_path,
    }


# ============================================================
# 17) PROCESS ONE IMAGE SET
# ============================================================

def incomplete_manifest(
    job,
    missing,
):
    return {
        "day_order": job["day_order"],
        "day": job["day"],
        "tray": job["tray"],
        "tray_no": job["tray_no"],
        "capture_id": job["capture_id"],
        "capture_sequence_index": job.get(
            "capture_sequence_index",
            "",
        ),
        "status": "SKIPPED_INCOMPLETE_MS_SET",
        "method": "",
        "reference_band": REFERENCE_BAND,
        "cell_count": 0,
        "image_width": "",
        "image_height": "",
        "candidate_count": 0,
        "candidate_methods_evaluated": "",
        "candidate_score": "",
        "outer_coverage_pass": "No",
        "first_column_x_ratio": "",
        "last_column_x_ratio": "",
        "first_row_y_ratio": "",
        "last_row_y_ratio": "",
        "horizontal_coverage_ratio": "",
        "vertical_coverage_ratio": "",
        "horizontal_spacing_mean": "",
        "vertical_spacing_mean": "",
        "max_spacing_cv": "",
        "median_fit_residual_px": "",
        "median_fit_residual_ratio": "",
        "maximum_fit_residual_px": "",
        "maximum_fit_residual_ratio": "",
        "fit_inlier_count": "",
        "mean_cup_contrast": "",
        "cup_contrast_p05": "",
        "minimum_cup_contrast": "",
        "edge_clipped_cell_count": "",
        "centre_outside_count": "",
        "source_point_count": "",
        "source_supported_cell_count": "",
        "source_mean_distance_ratio": "",
        "x_profile_score": "",
        "y_profile_score": "",
        "overlay_path": "",
        "reference_image_path": "",
        "missing_bands": missing,
        "notes": (
            f"Incomplete MS set. Missing bands: {missing}"
        ),
    }


def process_job(
    job,
    args,
    manual_points_store,
):
    day = job["day"]
    tray = job["tray"]
    capture_id = job["capture_id"]

    key = record_key(
        day,
        tray,
        capture_id,
    )

    base_output = (
        OUTPUT_ROOT
        / day
        / tray
    )

    overlay_path = (
        base_output
        / "overlays"
        / f"{capture_id}_MS_NIR_grid_overlay.jpg"
    )

    if not job["complete"]:
        missing = ", ".join(
            job["missing_bands"]
        )

        return (
            incomplete_manifest(
                job,
                missing,
            ),
            [],
        )

    reference_path = job["bands"][
        REFERENCE_BAND
    ]

    raw = read_ms_image(
        reference_path
    )

    gray = normalise_to_uint8(
        raw
    )

    image_height, image_width = gray.shape[:2]

    if (
        args.force_manual
        or key in manual_points_store
    ):
        if (
            key in manual_points_store
            and not args.reclick
        ):
            manual_points = manual_points_store[
                key
            ]
        else:
            manual_points = request_manual_points(
                gray,
                (
                    f"{day} | {tray} | "
                    f"{capture_id} | MS_NIR"
                ),
            )

            manual_points_store[
                key
            ] = manual_points

            save_manual_points(
                manual_points_store
            )

        result = manual_result(
            gray,
            manual_points,
            zone_shrink=args.zone_shrink,
        )

    else:
        result = auto_detect_grid(
            gray,
            args,
        )

        if result is None:
            result = {
                "status": "FAIL",
                "method": "no_valid_candidate",
                "grid_rows": [],
                "cell_count": 0,
                "candidate_count": 0,
                "candidate_methods_evaluated": "",
                "candidate_score": -99999.0,
                "outer_coverage_pass": "No",
                "first_column_x_ratio": math.nan,
                "last_column_x_ratio": math.nan,
                "first_row_y_ratio": math.nan,
                "last_row_y_ratio": math.nan,
                "horizontal_coverage_ratio": math.nan,
                "vertical_coverage_ratio": math.nan,
                "horizontal_spacing_mean": math.nan,
                "vertical_spacing_mean": math.nan,
                "max_spacing_cv": math.nan,
                "median_fit_residual_px": math.nan,
                "median_fit_residual_ratio": math.nan,
                "maximum_fit_residual_px": math.nan,
                "maximum_fit_residual_ratio": math.nan,
                "fit_inlier_count": 0,
                "mean_cup_contrast": math.nan,
                "cup_contrast_p05": math.nan,
                "minimum_cup_contrast": math.nan,
                "edge_clipped_cell_count": 0,
                "centre_outside_count": 0,
                "source_point_count": 0,
                "source_supported_cell_count": 0,
                "source_mean_distance_ratio": math.nan,
                "x_profile_score": math.nan,
                "y_profile_score": math.nan,
                "notes": (
                    "No valid automatic grid candidate was found."
                ),
            }

        if (
            args.manual_fallback
            and result["status"]
            in {
                "CHECK_AUTO",
                "FAIL",
            }
        ):
            manual_points = request_manual_points(
                gray,
                (
                    f"{day} | {tray} | "
                    f"{capture_id} | MS_NIR"
                ),
            )

            manual_points_store[
                key
            ] = manual_points

            save_manual_points(
                manual_points_store
            )

            result = manual_result(
                gray,
                manual_points,
                zone_shrink=args.zone_shrink,
            )

    grid_rows = result.get(
        "grid_rows",
        [],
    )

    if (
        grid_rows
        and (
            args.overwrite
            or not overlay_path.exists()
        )
    ):
        draw_grid_overlay(
            gray,
            grid_rows,
            overlay_path,
            title=(
                f"{day} | {tray} | "
                f"{capture_id} | MS_NIR"
            ),
            status=result["status"],
        )

    needs_review = (
        result["status"]
        not in ACCEPTED_STATUSES
    )

    coordinate_rows = []

    for row in grid_rows:
        coordinate_rows.append(
            {
                "day_order": job["day_order"],
                "day": day,
                "tray": tray,
                "tray_no": job["tray_no"],
                "capture_id": capture_id,
                "capture_sequence_index": job.get(
                    "capture_sequence_index",
                    "",
                ),
                "reference_band": REFERENCE_BAND,
                "cell_id": row["cell_id"],
                "row": row["row"],
                "column": row["column"],
                "x": row["x"],
                "y": row["y"],
                "poly_tl_x": row["poly_tl_x"],
                "poly_tl_y": row["poly_tl_y"],
                "poly_tr_x": row["poly_tr_x"],
                "poly_tr_y": row["poly_tr_y"],
                "poly_br_x": row["poly_br_x"],
                "poly_br_y": row["poly_br_y"],
                "poly_bl_x": row["poly_bl_x"],
                "poly_bl_y": row["poly_bl_y"],
                "square_x0": row["square_x0"],
                "square_y0": row["square_y0"],
                "square_x1": row["square_x1"],
                "square_y1": row["square_y1"],
                "zone_area_px": row["zone_area_px"],
                "centre_inside_image": row[
                    "centre_inside_image"
                ],
                "polygon_clipped_at_image_edge": row[
                    "polygon_clipped_at_image_edge"
                ],
                "coordinate_source": result["method"],
                "grid_status": result["status"],
                "needs_review": yes_no(
                    needs_review
                ),
                "image_width": image_width,
                "image_height": image_height,
                "ms_g_path": str(
                    job["bands"].get(
                        "G",
                        "",
                    )
                ),
                "ms_r_path": str(
                    job["bands"].get(
                        "R",
                        "",
                    )
                ),
                "ms_re_path": str(
                    job["bands"].get(
                        "RE",
                        "",
                    )
                ),
                "ms_nir_path": str(
                    job["bands"].get(
                        "NIR",
                        "",
                    )
                ),
            }
        )

    manifest = {
        "day_order": job["day_order"],
        "day": day,
        "tray": tray,
        "tray_no": job["tray_no"],
        "capture_id": capture_id,
        "capture_sequence_index": job.get(
            "capture_sequence_index",
            "",
        ),
        "status": result["status"],
        "method": result["method"],
        "reference_band": REFERENCE_BAND,
        "cell_count": len(
            coordinate_rows
        ),
        "image_width": image_width,
        "image_height": image_height,
        "candidate_count": result.get(
            "candidate_count",
            "",
        ),
        "candidate_methods_evaluated": result.get(
            "candidate_methods_evaluated",
            "",
        ),
        "candidate_score": result.get(
            "candidate_score",
            "",
        ),
        "outer_coverage_pass": result.get(
            "outer_coverage_pass",
            "",
        ),
        "first_column_x_ratio": result.get(
            "first_column_x_ratio",
            "",
        ),
        "last_column_x_ratio": result.get(
            "last_column_x_ratio",
            "",
        ),
        "first_row_y_ratio": result.get(
            "first_row_y_ratio",
            "",
        ),
        "last_row_y_ratio": result.get(
            "last_row_y_ratio",
            "",
        ),
        "horizontal_coverage_ratio": result.get(
            "horizontal_coverage_ratio",
            "",
        ),
        "vertical_coverage_ratio": result.get(
            "vertical_coverage_ratio",
            "",
        ),
        "horizontal_spacing_mean": result.get(
            "horizontal_spacing_mean",
            "",
        ),
        "vertical_spacing_mean": result.get(
            "vertical_spacing_mean",
            "",
        ),
        "max_spacing_cv": result.get(
            "max_spacing_cv",
            "",
        ),
        "median_fit_residual_px": result.get(
            "median_fit_residual_px",
            "",
        ),
        "median_fit_residual_ratio": result.get(
            "median_fit_residual_ratio",
            "",
        ),
        "maximum_fit_residual_px": result.get(
            "maximum_fit_residual_px",
            "",
        ),
        "maximum_fit_residual_ratio": result.get(
            "maximum_fit_residual_ratio",
            "",
        ),
        "fit_inlier_count": result.get(
            "fit_inlier_count",
            "",
        ),
        "mean_cup_contrast": result.get(
            "mean_cup_contrast",
            "",
        ),
        "cup_contrast_p05": result.get(
            "cup_contrast_p05",
            "",
        ),
        "minimum_cup_contrast": result.get(
            "minimum_cup_contrast",
            "",
        ),
        "edge_clipped_cell_count": result.get(
            "edge_clipped_cell_count",
            "",
        ),
        "centre_outside_count": result.get(
            "centre_outside_count",
            "",
        ),
        "source_point_count": result.get(
            "source_point_count",
            "",
        ),
        "source_supported_cell_count": result.get(
            "source_supported_cell_count",
            "",
        ),
        "source_mean_distance_ratio": result.get(
            "source_mean_distance_ratio",
            "",
        ),
        "x_profile_score": result.get(
            "x_profile_score",
            "",
        ),
        "y_profile_score": result.get(
            "y_profile_score",
            "",
        ),
        "overlay_path": (
            relative_path(
                overlay_path,
                OUTPUT_ROOT,
            )
            if grid_rows
            else ""
        ),
        "reference_image_path": str(
            reference_path
        ),
        "missing_bands": "",
        "notes": result.get(
            "notes",
            "",
        ),
    }

    return (
        manifest,
        coordinate_rows,
    )


# ============================================================
# 18) MAIN WORKFLOW
# ============================================================

def run_analysis(args):
    days_filter = parse_filter_list(
        args.days
    )

    trays_filter = parse_filter_list(
        args.trays
    )

    jobs = collect_jobs(
        days_filter=days_filter,
        trays_filter=trays_filter,
    )

    print(
        "\nSCRIPT 06 — THIRD TRIAL INDEPENDENT MS GRID DETECTION"
    )
    print(
        "=" * 78
    )
    print(
        f"Crop folder:\n{CROP_ROOT}"
    )
    print(
        f"\nOutput folder:\n{OUTPUT_ROOT}"
    )
    print(
        f"\nJobs found: {len(jobs)}\n"
    )

    for job in jobs:
        status = (
            "complete"
            if job["complete"]
            else "incomplete"
        )

        missing = (
            ""
            if job["complete"]
            else (
                " | missing: "
                + ", ".join(
                    job[
                        "missing_bands"
                    ]
                )
            )
        )

        print(
            f"{job['day']} > "
            f"{job['tray']} > "
            f"{job['capture_id']} "
            f"({status}){missing}"
        )

    if args.dry_run:
        print(
            "\nDry run complete. No outputs created."
        )
        return 0

    if not jobs:
        print(
            "\nNo multispectral image sets were found."
        )
        return 1

    manual_points_store = load_manual_points()

    manifest_rows = []
    coordinate_rows = []

    for job in jobs:
        print(
            "\nProcessing: "
            f"{job['day']} > "
            f"{job['tray']} > "
            f"{job['capture_id']}"
        )

        manifest, coordinates = process_job(
            job,
            args,
            manual_points_store,
        )

        manifest_rows.append(
            manifest
        )

        coordinate_rows.extend(
            coordinates
        )

        print(
            f"{manifest['status']} | "
            f"cells={manifest['cell_count']} | "
            f"method={manifest['method']} | "
            f"coverage={manifest['outer_coverage_pass']} | "
            f"contrast={manifest['mean_cup_contrast']}"
        )

        if manifest.get(
            "notes"
        ):
            print(
                f"  Notes: {manifest['notes']}"
            )

    settings = {
        "purpose": (
            "Third Trial independent multispectral "
            "70-cell grid detection"
        ),
        "crop_root": str(
            CROP_ROOT
        ),
        "output_root": str(
            OUTPUT_ROOT
        ),
        "rows": ROWS,
        "columns": COLS,
        "expected_cells": EXPECTED_CELLS,
        "required_ms_bands": REQUIRED_MS_BANDS,
        "reference_band": REFERENCE_BAND,
        "zone_shrink": args.zone_shrink,
        "primary_method": (
            "regular darkness projection lattice on MS_NIR"
        ),
        "secondary_method": (
            "separately evaluated Hough-circle configurations; "
            "no cross-config merging"
        ),
        "outer_coverage_checks": {
            "max_first_column_x_ratio": MAX_FIRST_COLUMN_X_RATIO,
            "min_last_column_x_ratio": MIN_LAST_COLUMN_X_RATIO,
            "max_first_row_y_ratio": MAX_FIRST_ROW_Y_RATIO,
            "min_last_row_y_ratio": MIN_LAST_ROW_Y_RATIO,
            "min_horizontal_coverage_ratio": MIN_HORIZONTAL_COVERAGE_RATIO,
            "min_vertical_coverage_ratio": MIN_VERTICAL_COVERAGE_RATIO,
        },
        "accepted_statuses_for_script07": sorted(
            ACCEPTED_STATUSES
        ),
        "profile_only": bool(
            args.profile_only
        ),
        "always_evaluate_hough": bool(
            args.always_evaluate_hough
        ),
        "manual_fallback": bool(
            args.manual_fallback
        ),
        "manual_points_json": str(
            MANUAL_POINTS_JSON
        ),
        "word_report": (
            "_reports/ms_independent_grid_work_process.docx"
        ),
    }

    output_paths = write_reports(
        manifest_rows,
        coordinate_rows,
        settings,
    )

    status_counts = (
        pd.DataFrame(
            manifest_rows
        )[
            "status"
        ]
        .value_counts()
        .to_dict()
        if manifest_rows
        else {}
    )

    print(
        "\n"
        + "=" * 78
    )
    print(
        "SCRIPT 06 FINISHED"
    )
    print(
        "=" * 78
    )

    for status, count in sorted(
        status_counts.items()
    ):
        print(
            f"{status}: {count}"
        )

    print(
        f"\nManifest:\n"
        f"{output_paths['manifest_csv']}"
    )

    print(
        f"\nCoordinates:\n"
        f"{output_paths['coordinates_csv']}"
    )

    print(
        f"\nExcel report:\n"
        f"{output_paths['excel_path']}"
    )

    print(
        f"\nSettings:\n"
        f"{output_paths['settings_path']}"
    )

    if DOCX_AVAILABLE:
        print(
            f"\nWord work-process report:\n"
            f"{output_paths['word_path']}"
        )
    else:
        print(
            "\nWord report skipped because python-docx "
            "is not installed."
        )

    accepted_count = sum(
        1
        for row in manifest_rows
        if row["status"]
        in ACCEPTED_STATUSES
    )

    print(
        f"\nAccepted MS grids: "
        f"{accepted_count}/{len(manifest_rows)} "
        "(full Trial 3 target: 84)"
    )

    problematic = [
        row
        for row in manifest_rows
        if row["status"]
        not in ACCEPTED_STATUSES
    ]

    if problematic:
        print(
            "\nSome image sets need review before Script 07."
        )
        print(
            "Inspect CHECK_AUTO/FAIL rows in "
            "ms_cell_grid_report.xlsx."
        )
        print(
            "Use targeted --force-manual with --days and "
            "--trays only for the remaining problematic image sets."
        )
        return 1

    print(
        "\nAll image sets are accepted, subject to visual "
        "overlay checking."
    )

    return 0


# ============================================================
# 19) CLI
# ============================================================

def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Trial 3 Script 06: independent MS_NIR "
            "70-cell grid detection."
        )
    )

    parser.add_argument(
        "--days",
        help=(
            'Process selected days only. '
            'Example: --days "Day 1,Day 7"'
        ),
    )

    parser.add_argument(
        "--trays",
        help=(
            'Process selected trays only. '
            'Example: --trays "Tray 1,Tray 12"'
        ),
    )

    parser.add_argument(
        "--overwrite",
        action="store_true",
        help=(
            "Overwrite existing overlay images."
        ),
    )

    parser.add_argument(
        "--dry-run",
        action="store_true",
        help=(
            "List image sets without processing."
        ),
    )

    parser.add_argument(
        "--profile-only",
        action="store_true",
        help=(
            "Use only the fast regular projection-lattice "
            "method. Do not evaluate Hough fallback candidates."
        ),
    )

    parser.add_argument(
        "--always-evaluate-hough",
        action="store_true",
        help=(
            "Evaluate all separate Hough configurations even "
            "when the projection candidate already passes."
        ),
    )

    parser.add_argument(
        "--manual-fallback",
        action="store_true",
        help=(
            "Open a four-corner click window only when an image "
            "produces CHECK_AUTO or FAIL."
        ),
    )

    parser.add_argument(
        "--force-manual",
        action="store_true",
        help=(
            "Force targeted manual four-corner correction. "
            "Use with --days and --trays for a small number of "
            "problem images."
        ),
    )

    parser.add_argument(
        "--reclick",
        action="store_true",
        help=(
            "Ignore saved manual points and click them again."
        ),
    )

    parser.add_argument(
        "--zone-shrink",
        type=float,
        default=ZONE_SHRINK_DEFAULT,
        help=(
            "Cell polygon shrink factor. Default: "
            f"{ZONE_SHRINK_DEFAULT}."
        ),
    )

    args = parser.parse_args()

    if (
        args.zone_shrink <= 0
        or args.zone_shrink > 1
    ):
        raise ValueError(
            "--zone-shrink must be greater than 0 "
            "and less than or equal to 1."
        )

    return run_analysis(
        args
    )


if __name__ == "__main__":
    raise SystemExit(
        main()
    )