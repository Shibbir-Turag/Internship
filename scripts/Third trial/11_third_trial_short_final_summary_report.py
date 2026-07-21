from __future__ import annotations

"""
SCRIPT 11 — THIRD TRIAL SHORT FINAL SUMMARY REPORT

Purpose
-------
Create a short but complete Word summary report for the Third Trial.

This script does not recalculate image-processing results. It reads the final
RGB + MS synthesis outputs from Script 09 and the final package index from
Script 10, then writes a human-readable Word report.

Main input
----------
outputs/Third trial/09_Final_RGB_MS_Synthesis/_reports/
    final_combined_tray_synthesis.csv
    final_group_synthesis.csv
    final_inside_outside_synthesis.csv
    input_file_inventory.csv

Optional input
--------------
outputs/Third trial/10_Final_Figure_Report_Package/
    Third_Trial_Final_Package/00_Package_Index/third_trial_package_manifest.csv

Main output
-----------
outputs/Third trial/11_Third_Trial_Short_Final_Summary_Report/_reports/
    third_trial_short_final_summary_report.docx
    third_trial_summary_key_findings.csv
    third_trial_summary_report_settings.json

Report sections
---------------
1. Trial overview
2. Treatment structure
3. Workflow summary
4. RGB evidence summary
5. Multispectral evidence summary
6. Final combined RGB + MS interpretation
7. Microbes vs No Microbes
8. Ideal vs Heat vs Moisture
9. Inside vs Outside for Ideal and Moisture
10. Heat and Moisture response
11. Day 7 observed vs adjusted handling
12. Limitations
13. Output package guide
"""

import json
import math
import re
import shutil
from datetime import datetime
from pathlib import Path

import pandas as pd


# ============================================================
# 1) PATHS
# ============================================================

PROJECT_ROOT = Path(
    r"C:\Users\tshib\OneDrive\Desktop\Internship"
)

THIRD_OUTPUT_ROOT = (
    PROJECT_ROOT
    / "outputs"
    / "Third trial"
)

SCRIPT09_ROOT = (
    THIRD_OUTPUT_ROOT
    / "09_Final_RGB_MS_Synthesis"
)

SCRIPT09_REPORTS = (
    SCRIPT09_ROOT
    / "_reports"
)

SCRIPT09_CHARTS = (
    SCRIPT09_ROOT
    / "charts"
)

SCRIPT10_PACKAGE_ROOT = (
    THIRD_OUTPUT_ROOT
    / "10_Final_Figure_Report_Package"
    / "Third_Trial_Final_Package"
)

SCRIPT10_PACKAGE_INDEX = (
    SCRIPT10_PACKAGE_ROOT
    / "00_Package_Index"
)

OUTPUT_ROOT = (
    THIRD_OUTPUT_ROOT
    / "11_Third_Trial_Short_Final_Summary_Report"
)

REPORTS_ROOT = (
    OUTPUT_ROOT
    / "_reports"
)

CHARTS_USED_ROOT = (
    OUTPUT_ROOT
    / "charts_used"
)

CONFIG_ROOT = (
    OUTPUT_ROOT
    / "_config"
)


# ============================================================
# 2) INPUT FILES
# ============================================================

FINAL_COMBINED_TRAY_SYNTHESIS = (
    SCRIPT09_REPORTS
    / "final_combined_tray_synthesis.csv"
)

FINAL_GROUP_SYNTHESIS = (
    SCRIPT09_REPORTS
    / "final_group_synthesis.csv"
)

FINAL_INSIDE_OUTSIDE_SYNTHESIS = (
    SCRIPT09_REPORTS
    / "final_inside_outside_synthesis.csv"
)

INPUT_FILE_INVENTORY = (
    SCRIPT09_REPORTS
    / "input_file_inventory.csv"
)

PACKAGE_MANIFEST = (
    SCRIPT10_PACKAGE_INDEX
    / "third_trial_package_manifest.csv"
)

PACKAGE_SUMMARY = (
    SCRIPT10_PACKAGE_INDEX
    / "third_trial_package_summary.csv"
)

PACKAGE_ZIP = (
    THIRD_OUTPUT_ROOT
    / "10_Final_Figure_Report_Package"
    / "Third_Trial_Final_Figure_Report_Package.zip"
)


# ============================================================
# 3) OUTPUT FILES
# ============================================================

WORD_REPORT = (
    REPORTS_ROOT
    / "third_trial_short_final_summary_report.docx"
)

KEY_FINDINGS_CSV = (
    REPORTS_ROOT
    / "third_trial_summary_key_findings.csv"
)

SETTINGS_JSON = (
    CONFIG_ROOT
    / "third_trial_summary_report_settings.json"
)


# ============================================================
# 4) TRIAL METADATA
# ============================================================

TRAY_METADATA = {
    1: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "environment": "Inside",
    },
    2: {
        "microbe_status": "No Microbes",
        "treatment": "Ideal",
        "environment": "Outside",
    },
    3: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "environment": "Outside",
    },
    4: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "environment": "Dynamic Heat",
    },
    5: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "environment": "Inside",
    },
    6: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "environment": "Dynamic Heat",
    },
    7: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "environment": "Outside",
    },
    8: {
        "microbe_status": "Microbes",
        "treatment": "Heat",
        "environment": "Dynamic Heat",
    },
    9: {
        "microbe_status": "Microbes",
        "treatment": "Ideal",
        "environment": "Inside",
    },
    10: {
        "microbe_status": "Microbes",
        "treatment": "Moisture",
        "environment": "Outside",
    },
    11: {
        "microbe_status": "No Microbes",
        "treatment": "Heat",
        "environment": "Dynamic Heat",
    },
    12: {
        "microbe_status": "No Microbes",
        "treatment": "Moisture",
        "environment": "Inside",
    },
}

OBSERVATION_DATES = {
    "Day 0": "27/06/2026 — seeds planted",
    "Day 1": "29/06/2026 — first image set",
    "Day 2": "30/06/2026",
    "Day 3": "01/07/2026",
    "Day 4": "02/07/2026",
    "Day 5": "03/07/2026",
    "Day 6": "04/07/2026",
    "Day 7": "07/07/2026 — final image set",
}


# ============================================================
# 5) OPTIONAL WORD SUPPORT
# ============================================================

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True

except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 6) HELPERS
# ============================================================

def require_file(path: Path, description: str) -> None:
    if not path.exists():
        raise FileNotFoundError(
            f"Missing {description}:\n{path}"
        )


def read_csv_required(path: Path, description: str) -> pd.DataFrame:
    require_file(path, description)

    return pd.read_csv(path)


def read_csv_optional(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()

    try:
        return pd.read_csv(path)

    except Exception:
        return pd.DataFrame()


def safe_float(value: object, default: float = math.nan) -> float:
    try:
        if pd.isna(value):
            return default

        return float(value)

    except Exception:
        return default


def format_number(value: object, decimals: int = 2) -> str:
    value = safe_float(value)

    if not math.isfinite(value):
        return "N/A"

    return f"{value:.{decimals}f}"


def format_count(value: object) -> str:
    value = safe_float(value)

    if not math.isfinite(value):
        return "0"

    return str(int(round(value)))


def clean_text(value: object) -> str:
    text = str(value).strip()

    if text.casefold() in {
        "nan",
        "none",
        "",
    }:
        return "N/A"

    return text


def first_valid_text(row: pd.Series, columns: list[str], default: str = "N/A") -> str:
    for column in columns:
        if column in row.index:
            text = clean_text(row[column])

            if text != "N/A":
                return text

    return default


def first_valid_number(row: pd.Series, columns: list[str], default: float = math.nan) -> float:
    for column in columns:
        if column in row.index:
            value = safe_float(row[column])

            if math.isfinite(value):
                return value

    return default


def sort_by_rank(frame: pd.DataFrame) -> pd.DataFrame:
    if frame.empty:
        return frame

    if "combined_rank" in frame.columns:
        temp = frame.copy()

        temp["combined_rank_numeric"] = pd.to_numeric(
            temp["combined_rank"],
            errors="coerce",
        )

        return (
            temp.sort_values(
                [
                    "combined_rank_numeric",
                    "tray_no",
                ],
                na_position="last",
            )
            .drop(columns=["combined_rank_numeric"])
            .reset_index(drop=True)
        )

    if "combined_rgb_ms_score" in frame.columns:
        temp = frame.copy()

        temp["combined_rgb_ms_score_numeric"] = pd.to_numeric(
            temp["combined_rgb_ms_score"],
            errors="coerce",
        )

        return (
            temp.sort_values(
                "combined_rgb_ms_score_numeric",
                ascending=False,
                na_position="last",
            )
            .drop(columns=["combined_rgb_ms_score_numeric"])
            .reset_index(drop=True)
        )

    return frame


def get_best_tray(combined: pd.DataFrame) -> pd.Series | None:
    if combined.empty:
        return None

    sorted_frame = sort_by_rank(combined)

    if sorted_frame.empty:
        return None

    return sorted_frame.iloc[0]


def filter_group(
    group_synthesis: pd.DataFrame,
    group_type: str,
) -> pd.DataFrame:
    if group_synthesis.empty or "group_type" not in group_synthesis.columns:
        return pd.DataFrame()

    return group_synthesis.loc[
        group_synthesis["group_type"].astype(str).eq(group_type)
    ].copy()


def best_group(
    group_synthesis: pd.DataFrame,
    group_type: str,
) -> pd.Series | None:
    subset = filter_group(
        group_synthesis,
        group_type,
    )

    if subset.empty or "mean_combined_rgb_ms_score" not in subset.columns:
        return None

    subset["score_numeric"] = pd.to_numeric(
        subset["mean_combined_rgb_ms_score"],
        errors="coerce",
    )

    subset = subset.dropna(
        subset=["score_numeric"]
    )

    if subset.empty:
        return None

    return subset.loc[
        subset["score_numeric"].idxmax()
    ]


def copy_chart(source: Path) -> Path | None:
    if not source.exists():
        return None

    CHARTS_USED_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    destination = CHARTS_USED_ROOT / source.name

    shutil.copy2(
        source,
        destination,
    )

    return destination


def expected_chart_paths() -> dict[str, Path]:
    return {
        "combined_ranking": SCRIPT09_CHARTS / "01_final_combined_rgb_ms_tray_ranking.png",
        "rgb_vs_ms": SCRIPT09_CHARTS / "02_rgb_vs_ms_score_scatter.png",
        "microbe_score": SCRIPT09_CHARTS / "03_combined_score_by_microbe_status.png",
        "treatment_score": SCRIPT09_CHARTS / "04_combined_score_by_treatment.png",
        "inside_outside": SCRIPT09_CHARTS / "05_inside_outside_combined_score_ideal_moisture.png",
        "adjusted_cells": SCRIPT09_CHARTS / "06_day7_adjusted_cells_rgb_ms.png",
    }


# ============================================================
# 7) FINDINGS TABLE
# ============================================================

def create_key_findings(
    combined: pd.DataFrame,
    group_synthesis: pd.DataFrame,
    inside_outside: pd.DataFrame,
    package_manifest: pd.DataFrame,
) -> pd.DataFrame:
    rows = []

    best = get_best_tray(
        combined
    )

    if best is not None:
        rows.append(
            {
                "topic": "Top tray",
                "finding": (
                    f"{clean_text(best.get('tray', 'N/A'))} had the highest final combined RGB + MS score."
                ),
                "value": format_number(
                    best.get(
                        "combined_rgb_ms_score",
                        math.nan,
                    ),
                    2,
                ),
                "source_table": "final_combined_tray_synthesis.csv",
            }
        )

    microbe_best = best_group(
        group_synthesis,
        "Microbe Status",
    )

    if microbe_best is not None:
        rows.append(
            {
                "topic": "Microbes vs No Microbes",
                "finding": (
                    f"{clean_text(microbe_best.get('group', 'N/A'))} had the higher mean combined RGB + MS score."
                ),
                "value": format_number(
                    microbe_best.get(
                        "mean_combined_rgb_ms_score",
                        math.nan,
                    ),
                    2,
                ),
                "source_table": "final_group_synthesis.csv",
            }
        )

    treatment_best = best_group(
        group_synthesis,
        "Treatment Type",
    )

    if treatment_best is not None:
        rows.append(
            {
                "topic": "Treatment comparison",
                "finding": (
                    f"{clean_text(treatment_best.get('group', 'N/A'))} had the highest mean combined RGB + MS score among the treatment types."
                ),
                "value": format_number(
                    treatment_best.get(
                        "mean_combined_rgb_ms_score",
                        math.nan,
                    ),
                    2,
                ),
                "source_table": "final_group_synthesis.csv",
            }
        )

    for row in inside_outside.itertuples(index=False):
        treatment = clean_text(
            getattr(row, "treatment", "N/A")
        )

        difference = format_number(
            getattr(
                row,
                "inside_minus_outside_combined_score",
                math.nan,
            ),
            2,
        )

        interpretation = clean_text(
            getattr(
                row,
                "interpretation",
                "N/A",
            )
        )

        rows.append(
            {
                "topic": f"Inside vs Outside — {treatment}",
                "finding": interpretation,
                "value": difference,
                "source_table": "final_inside_outside_synthesis.csv",
            }
        )

    adjusted_cells = (
        pd.to_numeric(
            combined.get(
                "total_day7_adjusted_cells",
                pd.Series(dtype=float),
            ),
            errors="coerce",
        )
        .fillna(0)
        .sum()
        if not combined.empty
        else 0
    )

    rows.append(
        {
            "topic": "Day 7 adjusted scenario",
            "finding": (
                "Observed Day 7 values were preserved, while likely bug-eaten cells were kept as a separate adjusted scenario."
            ),
            "value": format_count(adjusted_cells),
            "source_table": "final_combined_tray_synthesis.csv",
        }
    )

    if not package_manifest.empty:
        rows.append(
            {
                "topic": "Output package",
                "finding": (
                    "The final package collected Word reports, Excel workbooks, CSV tables, charts and evidence files from Scripts 04–09."
                ),
                "value": str(len(package_manifest)),
                "source_table": "third_trial_package_manifest.csv",
            }
        )

    return pd.DataFrame(rows)


# ============================================================
# 8) WORD DOCUMENT HELPERS
# ============================================================

def set_document_style(document: Document) -> None:
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(11)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        if style_name in document.styles:
            document.styles[style_name].font.name = "Times New Roman"


def add_table(
    document: Document,
    dataframe: pd.DataFrame,
    columns: list[str],
    max_rows: int = 12,
) -> None:
    if dataframe.empty:
        document.add_paragraph(
            "No table data were available for this section."
        )
        return

    valid_columns = [
        column
        for column in columns
        if column in dataframe.columns
    ]

    if not valid_columns:
        document.add_paragraph(
            "The expected columns were not available in this table."
        )
        return

    display = dataframe[valid_columns].head(max_rows).copy()

    table = document.add_table(
        rows=1,
        cols=len(valid_columns),
    )

    table.style = "Table Grid"

    for index, column in enumerate(valid_columns):
        table.rows[0].cells[index].text = (
            column.replace("_", " ").title()
        )

    for _index, row in display.iterrows():
        cells = table.add_row().cells

        for column_index, column in enumerate(valid_columns):
            value = row[column]

            if pd.isna(value):
                cells[column_index].text = ""

            elif isinstance(value, float):
                cells[column_index].text = f"{value:.3f}"

            else:
                cells[column_index].text = str(value)


def add_picture(
    document: Document,
    path: Path | None,
    width_inches: float = 6.2,
) -> None:
    if path is None or not Path(path).exists():
        return

    document.add_picture(
        str(path),
        width=Inches(width_inches),
    )


def add_output_file_description(
    document: Document,
    filename: str,
    description: str,
) -> None:
    paragraph = document.add_paragraph()

    paragraph.add_run(filename).bold = True
    paragraph.add_run(f": {description}")


# ============================================================
# 9) WORD REPORT CREATION
# ============================================================

def create_word_report(
    combined: pd.DataFrame,
    group_synthesis: pd.DataFrame,
    inside_outside: pd.DataFrame,
    input_inventory: pd.DataFrame,
    package_manifest: pd.DataFrame,
    key_findings: pd.DataFrame,
) -> Path:
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run:\n"
            '& "C:\\Users\\tshib\\anaconda3\\python.exe" -m pip install python-docx'
        )

    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    set_document_style(
        document
    )

    title = document.add_heading(
        "Third Trial Short Final Summary Report",
        level=0,
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This short report summarises the final results of the Third Trial. "
        "It brings together the RGB evidence from visible emergence and green-cover analysis "
        "with the multispectral evidence from relative NDVI and NDRE. The aim is to give a clear, "
        "readable summary of what the trial showed and how the final outputs should be used."
    )

    document.add_paragraph(
        "The report is based on the completed workflow outputs from Scripts 04 to 10. "
        "It does not recalculate the image-processing results. Instead, it reads the final synthesis "
        "tables and uses them to prepare a concise written interpretation."
    )

    # ------------------------------------------------------------
    # 1. Trial overview
    # ------------------------------------------------------------

    document.add_heading(
        "1. Trial overview",
        level=1,
    )

    document.add_paragraph(
        "The Third Trial was designed to compare seedling emergence and growth under different "
        "microbial, environmental, heat and moisture conditions. Each tray contained a 7 × 10 grid, "
        "giving 70 cells per tray. Twelve trays were included, so a complete observation day contained "
        "840 cell-level positions across all trays."
    )

    document.add_paragraph(
        "Seeds were planted on Day 0, 27 June 2026. Image collection started on Day 1, "
        "29 June 2026, and continued through Day 7, 7 July 2026. Images were not taken on "
        "5 and 6 July because most seeds were already visibly germinated, and the final image set "
        "was taken later to support growth-rate comparison."
    )

    date_table = pd.DataFrame(
        [
            {
                "trial_day": key,
                "date_or_note": value,
            }
            for key, value in OBSERVATION_DATES.items()
        ]
    )

    add_table(
        document,
        date_table,
        [
            "trial_day",
            "date_or_note",
        ],
        max_rows=10,
    )

    # ------------------------------------------------------------
    # 2. Treatment structure
    # ------------------------------------------------------------

    document.add_heading(
        "2. Treatment structure",
        level=1,
    )

    document.add_paragraph(
        "The tray labels used M for Microbes, N/M for No Microbes, In for Inside greenhouse, "
        "and Out for Outside in the open. Ideal trays stayed fully inside or fully outside. "
        "Heat trays were moved during the trial, while Moisture trays followed a watering schedule."
    )

    treatment_table = pd.DataFrame(
        [
            {
                "tray": f"Tray {tray_no}",
                "microbe_status": values["microbe_status"],
                "treatment": values["treatment"],
                "environment": values["environment"],
            }
            for tray_no, values in TRAY_METADATA.items()
        ]
    )

    add_table(
        document,
        treatment_table,
        [
            "tray",
            "microbe_status",
            "treatment",
            "environment",
        ],
        max_rows=12,
    )

    document.add_paragraph(
        "For the Heat treatment, the heat-labelled trays were inside at planting, moved outside on "
        "Day 3, stayed outside for Days 4 and 5, and were returned inside on Day 6. For the Moisture "
        "treatment, trays were watered on Day 0, Day 3 and Day 6, with drying intervals in between."
    )

    # ------------------------------------------------------------
    # 3. Workflow summary
    # ------------------------------------------------------------

    document.add_heading(
        "3. Workflow summary",
        level=1,
    )

    workflow_points = [
        "Script 04 detected visible emergence and RGB green-cover evidence from the cropped RGB images.",
        "Script 05 compared RGB emergence, green-cover and growth-rate outcomes across treatments.",
        "Script 06 produced independent multispectral 70-cell grids directly from the MS_NIR images.",
        "Script 07 calculated relative image-derived NDVI and NDRE from the original multispectral TIFF bands.",
        "Script 08 compared multispectral treatment outcomes and created observed and adjusted Day 7 summaries.",
        "Script 09 combined RGB and multispectral evidence into a final tray-level synthesis.",
        "Script 10 collected the key figures, reports, tables and evidence files into one final package.",
    ]

    for point in workflow_points:
        document.add_paragraph(
            point,
            style="List Bullet",
        )

    # ------------------------------------------------------------
    # 4. Key findings
    # ------------------------------------------------------------

    document.add_heading(
        "4. Key findings",
        level=1,
    )

    add_table(
        document,
        key_findings,
        [
            "topic",
            "finding",
            "value",
            "source_table",
        ],
        max_rows=12,
    )

    # ------------------------------------------------------------
    # 5. RGB evidence
    # ------------------------------------------------------------

    document.add_heading(
        "5. RGB visible emergence and green-cover evidence",
        level=1,
    )

    document.add_paragraph(
        "The RGB workflow measured what was directly visible in the cropped colour images. "
        "This included visible emergence, green-cover development and adjusted Day 7 green-cover "
        "handling for cells where seedlings had appeared earlier but were missing later."
    )

    top_rgb = combined.copy()

    if "rgb_score_normalised" in top_rgb.columns:
        top_rgb["rgb_score_numeric"] = pd.to_numeric(
            top_rgb["rgb_score_normalised"],
            errors="coerce",
        )

        top_rgb = top_rgb.dropna(
            subset=["rgb_score_numeric"]
        ).sort_values(
            "rgb_score_numeric",
            ascending=False,
        )

        if not top_rgb.empty:
            row = top_rgb.iloc[0]

            document.add_paragraph(
                f"The strongest RGB evidence was recorded by {clean_text(row.get('tray', 'N/A'))}, "
                f"with a normalised RGB score of {format_number(row.get('rgb_score_normalised'), 2)}. "
                f"This score summarises the available visible emergence and green-cover indicators."
            )

    # ------------------------------------------------------------
    # 6. MS evidence
    # ------------------------------------------------------------

    document.add_heading(
        "6. Multispectral NDVI and NDRE evidence",
        level=1,
    )

    document.add_paragraph(
        "The multispectral workflow used the original MS_G, MS_R, MS_RE and MS_NIR TIFF bands. "
        "The main vegetation indices were relative NDVI and relative NDRE. These should be described "
        "as image-derived relative indices unless calibrated reflectance data are available."
    )

    top_ms = combined.copy()

    if "ms_score_normalised" in top_ms.columns:
        top_ms["ms_score_numeric"] = pd.to_numeric(
            top_ms["ms_score_normalised"],
            errors="coerce",
        )

        top_ms = top_ms.dropna(
            subset=["ms_score_numeric"]
        ).sort_values(
            "ms_score_numeric",
            ascending=False,
        )

        if not top_ms.empty:
            row = top_ms.iloc[0]

            document.add_paragraph(
                f"The strongest multispectral evidence was recorded by {clean_text(row.get('tray', 'N/A'))}, "
                f"with a normalised MS score of {format_number(row.get('ms_score_normalised'), 2)}. "
                f"The MS score summarises the available NDVI and NDRE-based indicators."
            )

    # ------------------------------------------------------------
    # 7. Final combined interpretation
    # ------------------------------------------------------------

    document.add_heading(
        "7. Final combined RGB + MS interpretation",
        level=1,
    )

    best = get_best_tray(
        combined
    )

    if best is not None:
        document.add_paragraph(
            f"The highest final combined RGB + MS score was recorded by "
            f"{clean_text(best.get('tray', 'N/A'))} "
            f"({clean_text(best.get('microbe_status', 'N/A'))} | "
            f"{clean_text(best.get('treatment', 'N/A'))} | "
            f"{clean_text(best.get('environment_group', 'N/A'))}). "
            f"Its combined score was {format_number(best.get('combined_rgb_ms_score'), 2)}."
        )

        interpretation = clean_text(
            best.get(
                "final_interpretation",
                "",
            )
        )

        if interpretation != "N/A":
            document.add_paragraph(
                interpretation
            )

    document.add_paragraph(
        "The combined score should be treated as a descriptive synthesis. It helps rank trays and "
        "summarise the overall pattern, but it is not a formal statistical test."
    )

    add_table(
        document,
        sort_by_rank(combined),
        [
            "combined_rank",
            "tray",
            "microbe_status",
            "treatment",
            "environment_group",
            "combined_rgb_ms_score",
            "rgb_score_normalised",
            "ms_score_normalised",
            "evidence_agreement",
            "total_day7_adjusted_cells",
        ],
        max_rows=12,
    )

    charts = {
        key: copy_chart(path)
        for key, path in expected_chart_paths().items()
    }

    add_picture(
        document,
        charts.get("combined_ranking"),
    )

    add_picture(
        document,
        charts.get("rgb_vs_ms"),
    )

    # ------------------------------------------------------------
    # 8. Microbes vs No Microbes
    # ------------------------------------------------------------

    document.add_heading(
        "8. Microbes vs No Microbes",
        level=1,
    )

    microbe_table = filter_group(
        group_synthesis,
        "Microbe Status",
    )

    microbe_best = best_group(
        group_synthesis,
        "Microbe Status",
    )

    if microbe_best is not None:
        document.add_paragraph(
            f"In the Microbes vs No Microbes comparison, "
            f"{clean_text(microbe_best.get('group', 'N/A'))} had the higher mean combined score "
            f"({format_number(microbe_best.get('mean_combined_rgb_ms_score'), 2)})."
        )

    add_table(
        document,
        microbe_table,
        [
            "group",
            "tray_count",
            "trays",
            "mean_combined_rgb_ms_score",
            "mean_rgb_score",
            "mean_ms_score",
            "best_tray_in_group",
        ],
        max_rows=5,
    )

    add_picture(
        document,
        charts.get("microbe_score"),
    )

    # ------------------------------------------------------------
    # 9. Treatment comparison
    # ------------------------------------------------------------

    document.add_heading(
        "9. Ideal vs Heat vs Moisture",
        level=1,
    )

    treatment_table_result = filter_group(
        group_synthesis,
        "Treatment Type",
    )

    treatment_best = best_group(
        group_synthesis,
        "Treatment Type",
    )

    if treatment_best is not None:
        document.add_paragraph(
            f"Across the three treatment types, "
            f"{clean_text(treatment_best.get('group', 'N/A'))} had the highest mean combined score "
            f"({format_number(treatment_best.get('mean_combined_rgb_ms_score'), 2)})."
        )

    add_table(
        document,
        treatment_table_result,
        [
            "group",
            "tray_count",
            "trays",
            "mean_combined_rgb_ms_score",
            "mean_rgb_score",
            "mean_ms_score",
            "best_tray_in_group",
        ],
        max_rows=10,
    )

    add_picture(
        document,
        charts.get("treatment_score"),
    )

    # ------------------------------------------------------------
    # 10. Inside vs Outside
    # ------------------------------------------------------------

    document.add_heading(
        "10. Inside vs Outside for Ideal and Moisture trays",
        level=1,
    )

    document.add_paragraph(
        "The Inside vs Outside comparison was limited to Ideal and Moisture trays, because those trays "
        "had fixed environmental placement. Heat trays were not included in the direct fixed-environment "
        "comparison because they were moved between inside and outside during the trial."
    )

    add_table(
        document,
        inside_outside,
        [
            "treatment",
            "inside_trays",
            "outside_trays",
            "inside_mean_combined_score",
            "outside_mean_combined_score",
            "inside_minus_outside_combined_score",
            "inside_mean_rgb_score",
            "outside_mean_rgb_score",
            "inside_mean_ms_score",
            "outside_mean_ms_score",
            "interpretation",
        ],
        max_rows=5,
    )

    for row in inside_outside.itertuples(index=False):
        treatment = clean_text(
            getattr(row, "treatment", "N/A")
        )

        difference = format_number(
            getattr(
                row,
                "inside_minus_outside_combined_score",
                math.nan,
            ),
            2,
        )

        interpretation = clean_text(
            getattr(
                row,
                "interpretation",
                "N/A",
            )
        )

        document.add_paragraph(
            f"For {treatment}, the Inside-minus-Outside combined score difference was {difference}. "
            f"{interpretation}"
        )

    add_picture(
        document,
        charts.get("inside_outside"),
    )

    # ------------------------------------------------------------
    # 11. Heat and Moisture response
    # ------------------------------------------------------------

    document.add_heading(
        "11. Heat and Moisture treatment response",
        level=1,
    )

    document.add_paragraph(
        "Heat treatment was interpreted by movement phase rather than fixed environment. The trays were "
        "inside at baseline, moved outside during the heat-exposure period, and returned inside before the "
        "final image set. This means the Heat treatment should be read as a dynamic stress-and-recovery "
        "condition, not a simple Inside or Outside comparison."
    )

    document.add_paragraph(
        "Moisture treatment was interpreted around the watering schedule. Moisture trays were watered on "
        "Day 0, Day 3 and Day 6, with drying intervals between watering events. The fixed Inside/Outside "
        "comparison for Moisture trays is therefore useful, but it should still be interpreted alongside "
        "the watering timing."
    )

    # ------------------------------------------------------------
    # 12. Day 7 observed vs adjusted
    # ------------------------------------------------------------

    document.add_heading(
        "12. Day 7 observed vs adjusted handling",
        level=1,
    )

    adjusted_total = (
        pd.to_numeric(
            combined.get(
                "total_day7_adjusted_cells",
                pd.Series(dtype=float),
            ),
            errors="coerce",
        )
        .fillna(0)
        .sum()
        if not combined.empty
        else 0
    )

    document.add_paragraph(
        f"The workflow recorded {format_count(adjusted_total)} total adjusted Day 7 cell records across "
        "the RGB and MS outputs. These adjusted values are not direct observations. They represent a "
        "separate scenario for cells where seedlings had appeared earlier but were missing on Day 7, "
        "which was treated as possible bug damage."
    )

    document.add_paragraph(
        "This separation is important. The observed Day 7 image values remain the main record of what was "
        "visible in the final images. The adjusted values are included only to support a fairer growth-rate "
        "comparison where plants may have been eaten after earlier growth."
    )

    add_picture(
        document,
        charts.get("adjusted_cells"),
    )

    # ------------------------------------------------------------
    # 13. Limitations
    # ------------------------------------------------------------

    document.add_heading(
        "13. Limitations",
        level=1,
    )

    limitations = [
        "RGB and multispectral results measure related but different evidence. RGB measures visible green cover and emergence, while MS measures image-derived vegetation indices.",
        "NDVI and NDRE should be reported as relative image-derived indices unless reflectance calibration is confirmed.",
        "Soil, shadows, moisture, exposure, band alignment and polygon placement can affect MS values.",
        "The Day 7 adjusted values are estimated scenario values, not direct final-image observations.",
        "The final combined RGB + MS score is descriptive and should not be treated as a formal statistical significance test.",
        "The number of biological replicates per treatment group is limited, so group-level differences should be interpreted cautiously.",
    ]

    for limitation in limitations:
        document.add_paragraph(
            limitation,
            style="List Bullet",
        )

    # ------------------------------------------------------------
    # 14. Output package guide
    # ------------------------------------------------------------

    document.add_heading(
        "14. Output package guide",
        level=1,
    )

    document.add_paragraph(
        "The main completed evidence package is the Script 10 final package. It contains the Word reports, "
        "Excel workbooks, CSV tables, charts and selected overlays from the Trial 3 workflow."
    )

    if PACKAGE_ZIP.exists():
        document.add_paragraph(
            f"Final package ZIP: {PACKAGE_ZIP}"
        )

    if not package_manifest.empty:
        document.add_paragraph(
            f"The package manifest lists {len(package_manifest)} copied output files."
        )

    add_table(
        document,
        input_inventory,
        [
            "source",
            "exists",
            "used",
            "notes",
        ],
        max_rows=12,
    )

    document.add_heading(
        "15. Main files created by this script",
        level=1,
    )

    add_output_file_description(
        document,
        "third_trial_short_final_summary_report.docx",
        "The short final Word report for the Third Trial.",
    )

    add_output_file_description(
        document,
        "third_trial_summary_key_findings.csv",
        "A compact table of the main findings used in this report.",
    )

    add_output_file_description(
        document,
        "third_trial_summary_report_settings.json",
        "A settings file recording the input and output locations for this summary step.",
    )

    document.save(
        WORD_REPORT
    )

    return WORD_REPORT


# ============================================================
# 10) SETTINGS
# ============================================================

def save_settings(
    key_findings: pd.DataFrame,
) -> Path:
    CONFIG_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    settings = {
        "purpose": "Third Trial short final summary report",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "project_root": str(PROJECT_ROOT),
        "script09_reports": str(SCRIPT09_REPORTS),
        "script09_charts": str(SCRIPT09_CHARTS),
        "script10_package_root": str(SCRIPT10_PACKAGE_ROOT),
        "input_files": {
            "final_combined_tray_synthesis": str(FINAL_COMBINED_TRAY_SYNTHESIS),
            "final_group_synthesis": str(FINAL_GROUP_SYNTHESIS),
            "final_inside_outside_synthesis": str(FINAL_INSIDE_OUTSIDE_SYNTHESIS),
            "input_file_inventory": str(INPUT_FILE_INVENTORY),
            "package_manifest": str(PACKAGE_MANIFEST),
            "package_summary": str(PACKAGE_SUMMARY),
            "package_zip": str(PACKAGE_ZIP),
        },
        "outputs": {
            "word_report": str(WORD_REPORT),
            "key_findings_csv": str(KEY_FINDINGS_CSV),
            "settings_json": str(SETTINGS_JSON),
            "charts_used_folder": str(CHARTS_USED_ROOT),
        },
        "key_findings_count": int(len(key_findings)),
        "day7_policy": (
            "Observed Day 7 values are preserved. Adjusted Day 7 values are kept separately "
            "for likely bug-eaten cells and are not treated as direct observations."
        ),
        "interpretation_warning": (
            "The final combined score is descriptive and not a formal statistical test."
        ),
    }

    SETTINGS_JSON.write_text(
        json.dumps(
            settings,
            indent=2,
        ),
        encoding="utf-8",
    )

    return SETTINGS_JSON


# ============================================================
# 11) MAIN
# ============================================================

def main() -> int:
    print(
        "\nSCRIPT 11 — THIRD TRIAL SHORT FINAL SUMMARY REPORT"
    )

    print(
        "=" * 78
    )

    print(
        f"Script 09 reports:\n{SCRIPT09_REPORTS}"
    )

    print(
        f"\nOutput folder:\n{OUTPUT_ROOT}"
    )

    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run this first:\n"
            '& "C:\\Users\\tshib\\anaconda3\\python.exe" -m pip install python-docx'
        )

    REPORTS_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    CHARTS_USED_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    combined = read_csv_required(
        FINAL_COMBINED_TRAY_SYNTHESIS,
        "Script 09 final combined tray synthesis",
    )

    group_synthesis = read_csv_required(
        FINAL_GROUP_SYNTHESIS,
        "Script 09 final group synthesis",
    )

    inside_outside = read_csv_required(
        FINAL_INSIDE_OUTSIDE_SYNTHESIS,
        "Script 09 final Inside vs Outside synthesis",
    )

    input_inventory = read_csv_optional(
        INPUT_FILE_INVENTORY
    )

    package_manifest = read_csv_optional(
        PACKAGE_MANIFEST
    )

    package_summary = read_csv_optional(
        PACKAGE_SUMMARY
    )

    key_findings = create_key_findings(
        combined,
        group_synthesis,
        inside_outside,
        package_manifest,
    )

    key_findings.to_csv(
        KEY_FINDINGS_CSV,
        index=False,
    )

    report_path = create_word_report(
        combined,
        group_synthesis,
        inside_outside,
        input_inventory,
        package_manifest,
        key_findings,
    )

    settings_path = save_settings(
        key_findings
    )

    print(
        "\n"
        + "=" * 78
    )

    print(
        "SCRIPT 11 FINISHED"
    )

    print(
        "=" * 78
    )

    print(
        f"Combined tray rows read: {len(combined)}"
    )

    print(
        f"Group synthesis rows read: {len(group_synthesis)}"
    )

    print(
        f"Inside/Outside rows read: {len(inside_outside)}"
    )

    print(
        f"Key findings written: {len(key_findings)}"
    )

    if not package_manifest.empty:
        print(
            f"Package manifest rows read: {len(package_manifest)}"
        )

    print(
        f"\nWord report:\n{report_path}"
    )

    print(
        f"\nKey findings CSV:\n{KEY_FINDINGS_CSV}"
    )

    print(
        f"\nSettings:\n{settings_path}"
    )

    print(
        f"\nCharts copied to:\n{CHARTS_USED_ROOT}"
    )

    return 0


if __name__ == "__main__":
    raise SystemExit(
        main()
    )